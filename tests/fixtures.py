"""Shared, synthetic test data: DOCX builders and a matching `MasterResume`.

Every function here is either fabricated content or mechanically derived from it — no
personal data, and no dependence on `data/master_resume.json` or
`templates/main_template.docx`, which are a developer's own upload and therefore not
present (or not in any particular state) on a clean checkout or in CI. See
`conftest.py`'s `built_template` fixture, which builds the tagged template these DOCX
builders describe exactly once per test session.

The low-level builders (`_add_bullet_numbering`, `_make_bullet`, `_docx_bytes`,
`_add_hyperlink`) and the four resume-shaped builders below were originally duplicated
across `test_template_analyze.py` and `test_template_build.py` (the latter importing
from the former via `from tests.test_template_analyze import ...`); moved here once that
cross-suite import became indistinguishable from a real shared-fixtures module.
"""

from __future__ import annotations

import io

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from resume_tailor.data import (
    Bullet,
    Contact,
    Education,
    EducationSection,
    Experience,
    ExperienceSection,
    MasterResume,
    Project,
    ProjectSection,
    SkillGroup,
    SkillsSection,
)

# --------------------------------------------------------------------------------------
# Low-level DOCX construction helpers
# --------------------------------------------------------------------------------------


def _add_bullet_numbering(document) -> str:
    """Install a minimal bullet abstract + num instance; return the numId string."""
    # python-docx creates numbering part lazily; force it via a style-based list then
    # fall back to a hand-rolled abstract if needed.
    numbering = document.part.numbering_part
    root = numbering.element

    abstract_id = "10"
    # Skip if we already added one in this document.
    for anum in root.findall(qn("w:abstractNum")):
        if anum.get(qn("w:abstractNumId")) == abstract_id:
            break
    else:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), abstract_id)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "bullet")
        lvl.append(num_fmt)
        abstract.append(lvl)
        # abstractNum elements must precede w:num
        first_num = root.find(qn("w:num"))
        if first_num is not None:
            first_num.addprevious(abstract)
        else:
            root.append(abstract)

    num_id = "20"
    for num in root.findall(qn("w:num")):
        if num.get(qn("w:numId")) == num_id:
            return num_id
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), num_id)
    abs_el = OxmlElement("w:abstractNumId")
    abs_el.set(qn("w:val"), abstract_id)
    num.append(abs_el)
    root.append(num)
    return num_id


def _make_bullet(document, text: str, num_id: str):
    """Append a list paragraph with the given numId."""
    paragraph = document.add_paragraph(text)
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    nid = OxmlElement("w:numId")
    nid.set(qn("w:val"), num_id)
    numPr.append(ilvl)
    numPr.append(nid)
    pPr.append(numPr)
    return paragraph


def _docx_bytes(build) -> bytes:
    """Run `build(document)` and return the saved .docx bytes."""
    document = docx.Document()
    build(document)
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    """Append a real w:hyperlink run so analyze's has_hyperlink path fires."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# --------------------------------------------------------------------------------------
# Resume-shaped DOCX builders
# --------------------------------------------------------------------------------------


def _standard_resume(document) -> None:
    """Minimal single-column resume matching the legacy heading contract."""
    num_id = _add_bullet_numbering(document)
    document.add_paragraph("Ada Lovelace")
    document.add_paragraph("London • ada@example.com • LinkedIn")
    document.add_paragraph("EDUCATION")
    document.add_paragraph("University of London | UK\t2018 - 2022")
    _make_bullet(document, "BSc Computer Science | GPA: 3.9", num_id)
    _make_bullet(document, "Relevant Coursework: Algorithms, Databases", num_id)
    document.add_paragraph("WORK EXPERIENCES")
    document.add_paragraph("Analytical Engines | London\t2022 - Present")
    document.add_paragraph("Software Engineer")
    _make_bullet(document, "Built numerical engines in Python.", num_id)
    document.add_paragraph("PROJECTS")
    document.add_paragraph("Note Engine | Python, FastAPI\t2024")
    _make_bullet(document, "Indexed research notes with embeddings.", num_id)
    document.add_paragraph("SKILLS")
    p = document.add_paragraph()
    run = p.add_run("Languages:")
    run.bold = True
    p.add_run(" Python, SQL")


def _multi_section_resume(document) -> None:
    """Mirrors the real-world motivating case: three experience-shaped sections, one of
    them ("OTHER ACTIVITIES") unrecognized by any heading alias, plus a single-line
    (non-bulleted) education entry — both structural-detection edge cases."""
    num_id = _add_bullet_numbering(document)
    document.add_paragraph("Nina Dao")
    document.add_paragraph("nina@example.com")
    document.add_paragraph("EDUCATION")
    document.add_paragraph("UC Irvine\tExpected June 2027")
    document.add_paragraph("B.A. in Business Administration")  # plain, not a Word bullet
    document.add_paragraph("WORK EXPERIENCE")
    document.add_paragraph("Langmaster JSC\tAug 2025 - Present")
    document.add_paragraph("Online Tutor")
    _make_bullet(document, "Tutored students in English.", num_id)
    document.add_paragraph("LEADERSHIP EXPERIENCE")
    document.add_paragraph("In the Green at UCI\tMay 2025 - Present")
    document.add_paragraph("Co-president")
    _make_bullet(document, "Oversaw club operations.", num_id)
    document.add_paragraph("OTHER ACTIVITIES")
    document.add_paragraph("Heartbeat Bazaar\tMar 2022 - Jun 2022")
    document.add_paragraph("Organizer")
    _make_bullet(document, "Directed fundraising events.", num_id)
    document.add_paragraph("SKILLS")
    sk = document.add_paragraph()
    run = sk.add_run("Languages:")
    run.bold = True
    sk.add_run(" Python, SQL")


def _spacer_multi_section_resume(document) -> None:
    """Mirrors the real-world motivating document: a blank paragraph before every
    heading, one right after every heading, and one between entries within a body that
    has more than one entry — exactly the pattern `_detect_spacing` looks for. Unlike
    `_multi_section_resume`/`_standard_resume` (both used by
    `tests/test_template_build.py` and asserted spacer-free elsewhere), this fixture
    exists solely to exercise spacer detection."""
    num_id = _add_bullet_numbering(document)
    document.add_paragraph("Nina Dao")
    document.add_paragraph("nina@example.com")
    document.add_paragraph("")
    document.add_paragraph("EDUCATION")
    document.add_paragraph("")
    document.add_paragraph("UC Irvine\tExpected June 2027")
    _make_bullet(document, "B.A. in Business Administration", num_id)
    document.add_paragraph("")
    document.add_paragraph("WORK EXPERIENCE")
    document.add_paragraph("")
    document.add_paragraph("Langmaster JSC\tAug 2025 - Present")
    document.add_paragraph("Online Tutor")
    _make_bullet(document, "Tutored students in English.", num_id)
    document.add_paragraph("")
    document.add_paragraph("Garin JSC\tNov 2023 - May 2024")
    document.add_paragraph("Logistics Intern")
    _make_bullet(document, "Managed international logistics.", num_id)
    document.add_paragraph("")
    document.add_paragraph("LEADERSHIP EXPERIENCE")
    document.add_paragraph("")
    document.add_paragraph("In the Green at UCI\tMay 2025 - Present")
    document.add_paragraph("Co-president")
    _make_bullet(document, "Oversaw club operations.", num_id)
    document.add_paragraph("")
    document.add_paragraph("SKILLS")
    document.add_paragraph("")
    _make_bullet(document, "Languages: Python, SQL", num_id)


def _rule_separated_resume(document) -> None:
    """Mirrors the real installed baseline: a horizontal rule *and* a blank under every
    heading, and two consecutive blanks at one entry boundary. Both shapes defeated
    single-paragraph spacer detection — the rule is chrome but not blank, and with two
    adjacent blanks neither one's neighbour is the bullet/header pair a boundary test
    looks for."""
    num_id = _add_bullet_numbering(document)
    rule = "_" * 60
    document.add_paragraph("Nina Dao")
    document.add_paragraph("nina@example.com")
    document.add_paragraph("")
    document.add_paragraph("EDUCATION")
    document.add_paragraph(rule)
    document.add_paragraph("")
    document.add_paragraph("UC Irvine\tExpected June 2028")
    _make_bullet(document, "B.A. in Business Administration", num_id)
    document.add_paragraph(rule)
    document.add_paragraph("")
    document.add_paragraph("WORK EXPERIENCE")
    document.add_paragraph(rule)
    document.add_paragraph("")
    document.add_paragraph("Langmaster JSC\tAug 2025 - Present")
    document.add_paragraph("Online Tutor")
    _make_bullet(document, "Tutored students in English.", num_id)
    document.add_paragraph("")
    document.add_paragraph("")
    document.add_paragraph("Garin JSC\tNov 2023 - May 2024")
    document.add_paragraph("Logistics Intern")
    _make_bullet(document, "Managed international logistics.", num_id)
    document.add_paragraph(rule)
    document.add_paragraph("")
    document.add_paragraph("SKILLS")
    document.add_paragraph(rule)
    document.add_paragraph("")
    _make_bullet(document, "Languages: Python, SQL", num_id)


def _full_featured_resume(document) -> None:
    """One entry per fixed-mode kind, deliberately touching every feature
    `test_render.py` needs to check without a developer's own upload: a hyperlinked
    project, education with GPA + coursework, an ampersand in both a skills group and a
    bullet, and three location strings ("Irvine, CA" / "Boston, MA" / "Remote") chosen
    to never be substrings of one another, so a contact-field-override test can assert
    one is absent without a false negative from an unrelated section using the same
    text. `synthetic_resume()` below is the matching `MasterResume` — kept in
    section-kind-and-order lockstep, though under `section_mode="fixed"` (what this
    fixture builds, since it has exactly one heading per kind) only THIS docx's own
    heading text and paragraph shapes affect what renders; `synthetic_resume()`'s
    section titles are inert prototypes, not templated in.
    """
    num_id = _add_bullet_numbering(document)
    document.add_paragraph("Jordan Rivera")
    document.add_paragraph(
        "Irvine, CA • jordan@example.com • (555) 123-4567 • LinkedIn • GitHub"
    )
    document.add_paragraph("EDUCATION")
    edu = document.add_paragraph()
    school_run = edu.add_run("State University | ")
    school_run.bold = True
    edu.add_run("Boston, MA")
    edu_tab = edu.add_run()
    edu_tab._r.append(OxmlElement("w:tab"))
    edu.add_run("2019 - 2023")
    _make_bullet(document, "BS Computer Science", num_id)
    _make_bullet(document, "Relevant Coursework: Algorithms, Databases", num_id)
    document.add_paragraph("WORK EXPERIENCES")
    exp = document.add_paragraph()
    company_run = exp.add_run("Example Corp | ")
    company_run.bold = True
    exp.add_run("Remote")
    exp_tab = exp.add_run()
    exp_tab._r.append(OxmlElement("w:tab"))
    exp.add_run("2023 - Present")
    document.add_paragraph("Software Engineer")
    _make_bullet(
        document, "Improved reliability & throughput for production services.", num_id
    )
    document.add_paragraph("PROJECTS")
    proj = document.add_paragraph()
    proj.add_run("Note Engine | ")
    proj.add_run("Python, FastAPI | ")
    _add_hyperlink(proj, "Github", "https://github.com/jordanrivera/note-engine")
    tab_run = proj.add_run()
    tab_run._r.append(OxmlElement("w:tab"))
    proj.add_run("2024")
    _make_bullet(document, "Indexed research notes with embeddings.", num_id)
    document.add_paragraph("SKILLS")
    sk = document.add_paragraph()
    run = sk.add_run("Tools:")
    run.bold = True
    sk.add_run(" Python, Data & Analytics")


def synthetic_resume() -> MasterResume:
    """A `MasterResume` matching `_full_featured_resume`'s shape and kind order.

    Under `section_mode="fixed"` (what `_full_featured_resume` builds) the tagged
    template's headings and prototype text come entirely from the uploaded docx, not
    from this resume — so nothing here needs to textually match the docx builder above
    except the *kind* and *count* of sections (one experience entry, one project, one
    education entry, one skills group), which is what determines whether a bullet has
    anywhere to render into.
    """
    return MasterResume(
        contact=Contact(
            name="Jordan Rivera",
            email="jordan@example.com",
            phone="(555) 123-4567",
            location="Irvine, CA",
            linkedin="https://linkedin.com/in/jordanrivera",
            github="https://github.com/jordanrivera",
        ),
        sections=[
            EducationSection(
                id="education",
                title="Education",
                entries=[
                    Education(
                        school="State University",
                        degree="BS Computer Science",
                        dates="2019 - 2023",
                        location="Boston, MA",
                        coursework=["Algorithms", "Databases"],
                        gpa="3.8",
                        show_gpa=False,
                    )
                ],
            ),
            ExperienceSection(
                id="experience",
                title="Work Experience",
                entries=[
                    Experience(
                        id="example-corp",
                        company="Example Corp",
                        title="Software Engineer",
                        location="Remote",
                        start="2023-01",
                        end="present",
                        bullets=[
                            Bullet(
                                id="exp_b1",
                                text=(
                                    "Improved reliability & throughput for "
                                    "production services."
                                ),
                                # "python" specifically (not e.g. "backend") so a
                                # consumer can write a plain `Keyword(canonical="python")`
                                # requirement and get a real, non-trivial match — the
                                # convention most other fixtures in this suite already
                                # use for their one "obviously covered" tag.
                                tags=["python"],
                            ),
                        ],
                    )
                ],
            ),
            ProjectSection(
                id="projects",
                title="Projects",
                entries=[
                    Project(
                        id="note-engine",
                        name="Note Engine",
                        tech=["Python", "FastAPI"],
                        date="2024",
                        link="Repo",
                        url="https://github.com/jordanrivera/note-engine",
                        bullets=[
                            Bullet(
                                id="proj_b1",
                                text="Indexed research notes with embeddings.",
                                tags=["search"],
                            ),
                        ],
                    )
                ],
            ),
            SkillsSection(
                id="skills",
                title="Skills",
                entries=[
                    SkillGroup(label="Tools", items=["Python", "Data & Analytics"]),
                ],
            ),
        ],
    )


# --------------------------------------------------------------------------------------
# Table-layout DOCX builder — content lives inside one invisible layout table (used to
# right-align dates/locations without tab stops), the shape the driving real-world
# document ("Nina Dao - aug.docx") turned out to have. See CLAUDE.md's "Template
# generation" section for `layout="table"` and its "Non-obvious gotchas" for the
# `{%tr %}` marker-row idiom this shape requires at build time.
# --------------------------------------------------------------------------------------


def _make_cell_bullet(cell, text: str, num_id: str):
    """Append a list paragraph with the given numId inside a table cell."""
    paragraph = cell.add_paragraph(text)
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    nid = OxmlElement("w:numId")
    nid.set(qn("w:val"), num_id)
    numPr.append(ilvl)
    numPr.append(nid)
    pPr.append(numPr)
    return paragraph


def _shape_row(table, row_idx: int, spans: list[int]) -> list:
    """Reshape row `row_idx`'s physical `<w:tc>` cells into the given colspans
    (summing to the table's declared column count), returning the resulting `_Cell`
    objects in order.

    Raw `w:gridSpan` rather than `_Cell.merge` — merge operates on the table's
    *logical* grid and renumbers `Table.cell(r, c)`, so a fixture that needs an exact
    physical `<w:tc>` count (3+1 vs 2+2 vs 1+3 vs 4 — the driving document uses all
    four) cannot express itself through it.
    """
    from docx.table import _Cell

    row = table.rows[row_idx]
    tr = row._tr
    tcs = tr.findall(qn("w:tc"))
    if sum(spans) != len(tcs):
        raise ValueError(f"spans {spans} must sum to {len(tcs)} physical cells, row {row_idx}")

    kept: list = []
    idx = 0
    for span in spans:
        keep_tc = tcs[idx]
        for absorbed in tcs[idx + 1 : idx + span]:
            tr.remove(absorbed)
        if span > 1:
            tcPr = keep_tc.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr")
                keep_tc.insert(0, tcPr)
            grid_span_el = OxmlElement("w:gridSpan")
            grid_span_el.set(qn("w:val"), str(span))
            tcPr.append(grid_span_el)
        kept.append(keep_tc)
        idx += span
    return [_Cell(tc, table) for tc in kept]


def _cell_text(cell, first: str, *rest: str) -> None:
    """Set a cell's first (already-existing) paragraph text, then append one
    paragraph per remaining arg. Plain (no run formatting) — fine for content that
    `_split_entries` never scans (contact block, skills grid); an entry's own header
    fields must use `_cell_para` instead, see its docstring."""
    cell.paragraphs[0].text = first
    for text in rest:
        cell.add_paragraph(text)


def _cell_para(
    cell, index: int, text: str, *, bold: bool = False, italic: bool = False, align_right: bool = False
):
    """Set text and run formatting on the cell's paragraph at `index` (0 is the
    cell's already-existing first paragraph; higher indices are appended).

    Formatting here is not cosmetic: `template_analyze._split_entries` re-splits a
    bootstrapped entry list by formatting fingerprint when a majority of entries'
    header paragraphs share one — real resumes naturally satisfy this (a bold company
    name differs from an italic, right-aligned date), but an unformatted fixture
    (every paragraph identically plain) makes EVERY paragraph share one fingerprint,
    so each wrongly starts its own "entry". `align_right` matters here too, not just
    `bold`/`italic`: the real driving document right-aligns location/dates exactly
    this way, and `_Fingerprint` includes alignment — without it, a bold *location*
    paragraph collides with a bold *company* paragraph's fingerprint (both simply
    "bold"), which mis-splits the entry the same way unformatted text does.
    """
    para = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
    if align_right:
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    return para


def _table_resume(document) -> None:
    """Reproduces the driving real-world document's shape: name/address/email/phone
    spread across paragraphs in a 4-column layout table; a `3+1` gridSpan for
    EDUCATION/WORK EXPERIENCE's heading and entry-header rows; a `4`-span heading and
    `2+2` entry-header rows for LEADERSHIP (the gridSpan inconsistency the real
    document has between its two experience-kind sections); a `1+3` label/value grid
    for ADDITIONAL INFORMATION's skills; full-width (`4`-span) bullet rows; no tabs
    anywhere; no hyperlinks. Two experience-kind headings makes this `needs_generic`
    on its own, matching the real document. Entry headers use bold company/school
    names and italic titles/dates — see `_cell_para` for why that is load-bearing,
    not decorative.
    """
    num_id = _add_bullet_numbering(document)
    table = document.add_table(rows=14, cols=4)

    name_row = _shape_row(table, 0, [3, 1])
    _cell_text(name_row[0], "JORDAN RIVERA")
    _cell_text(name_row[1], "123 Example St.", "Springfield, IL 60000")

    contact_row = _shape_row(table, 1, [3, 1])
    _cell_text(contact_row[0], "jordan@example.com")
    _cell_text(contact_row[1], "(555) 123-4567")

    _shape_row(table, 2, [3, 1])  # spacer row, both cells left blank

    edu_heading = _shape_row(table, 3, [3, 1])
    _cell_text(edu_heading[0], "EDUCATION")

    edu_header = _shape_row(table, 4, [3, 1])
    _cell_para(edu_header[0], 0, "State University", bold=True)
    _cell_para(edu_header[0], 1, "BS Computer Science", italic=True)
    _cell_para(edu_header[1], 0, "IL", bold=True, align_right=True)
    _cell_para(edu_header[1], 1, "Class of 2024", italic=True, align_right=True)

    edu_bullets = _shape_row(table, 5, [4])
    _make_cell_bullet(edu_bullets[0], "Cumulative GPA: 3.8/4.0", num_id)
    _make_cell_bullet(edu_bullets[0], "Relevant Coursework: Algorithms, Databases", num_id)

    exp_heading = _shape_row(table, 6, [3, 1])
    _cell_text(exp_heading[0], "WORK EXPERIENCE")

    exp_header = _shape_row(table, 7, [3, 1])
    _cell_para(exp_header[0], 0, "Example Corp", bold=True)
    _cell_para(exp_header[0], 1, "Software Engineer", italic=True)
    _cell_para(exp_header[1], 0, "Remote", bold=True, align_right=True)
    _cell_para(exp_header[1], 1, "Jan 2023 - Present", italic=True, align_right=True)

    exp_bullets = _shape_row(table, 8, [4])
    _make_cell_bullet(
        exp_bullets[0], "Improved reliability & throughput for production services.", num_id
    )

    lead_heading = _shape_row(table, 9, [4])
    _cell_text(lead_heading[0], "LEADERSHIP")

    lead_header = _shape_row(table, 10, [2, 2])
    _cell_para(lead_header[0], 0, "Campus Club", bold=True)
    _cell_para(lead_header[0], 1, "President", italic=True)
    _cell_para(lead_header[1], 0, "Springfield, IL", bold=True, align_right=True)
    _cell_para(lead_header[1], 1, "2022 - 2023", italic=True, align_right=True)

    lead_bullets = _shape_row(table, 11, [4])
    _make_cell_bullet(lead_bullets[0], "Grew membership by 30%.", num_id)

    info_heading = _shape_row(table, 12, [4])
    _cell_text(info_heading[0], "ADDITIONAL INFORMATION")

    info_grid = _shape_row(table, 13, [1, 3])
    _cell_text(info_grid[0], "Languages:", "Skills:")
    _cell_text(info_grid[1], "Fluent in English and Spanish", "Python, SQL, Excel")


def _sidebar_table_resume(document) -> None:
    """Negative fixture: same table shape as `_table_resume`, but the work-experience
    bullets sit in the row's SECOND cell alongside its own header text — the
    structural signature of two parallel reading columns (a sidebar), which
    `template_analyze.classify_table_layout` must reject via `table_sidebar_bullets`.
    """
    num_id = _add_bullet_numbering(document)
    table = document.add_table(rows=8, cols=4)

    name_row = _shape_row(table, 0, [3, 1])
    _cell_text(name_row[0], "JORDAN RIVERA")
    _cell_text(name_row[1], "jordan@example.com")

    _shape_row(table, 1, [3, 1])  # spacer

    exp_heading = _shape_row(table, 2, [3, 1])
    _cell_text(exp_heading[0], "WORK EXPERIENCE")

    exp_header = _shape_row(table, 3, [3, 1])
    _cell_text(exp_header[0], "Example Corp", "Software Engineer")
    exp_header[1].paragraphs[0].text = "Remote"
    _make_cell_bullet(exp_header[1], "Improved reliability for production services.", num_id)
