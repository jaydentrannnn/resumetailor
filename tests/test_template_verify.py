"""Tests for `template_verify.py` — no Word, no LibreOffice.

Two kinds of case: a known-good build must verify clean (regression guard against
`expected_tags`/`verify_tagged`/`verify_roundtrip` themselves drifting out of sync with
what `template_build.py` actually emits), and a deliberately corrupted build must be
caught (proof the checks actually catch something, not just pass vacuously). Corruption
is injected by editing a *built* template's XML directly after a real, successful build
— simulating a tagging bug without needing to first reproduce a specific analyzer bug
that happens to produce it.
"""

from __future__ import annotations

import tempfile
from pathlib import Path

import docx

from resume_tailor import template_analyze, template_build, template_verify
from tests.fixtures import _docx_bytes, _full_featured_resume, _multi_section_resume, synthetic_resume


def _build(builder) -> tuple[Path, object]:
    """Analyze + build `builder`'s docx into a fresh temp dir. Returns (tagged_path, profile)."""
    tmp = Path(tempfile.mkdtemp())
    src = tmp / "original_export.docx"
    src.write_bytes(_docx_bytes(builder))
    result = template_analyze.analyze_docx(raw=src.read_bytes())
    assert result.ready, f"fixture failed to analyze cleanly: {result.issues}"
    dst = tmp / "main_template.docx"
    template_build.build_from_profile(src, dst, result.suggested_profile)
    return dst, result.suggested_profile


def _remove_paragraph_containing(doc_path: Path, needle: str) -> None:
    """Delete the first paragraph containing `needle` — simulates a build that dropped
    a tag/control-tag it should have emitted."""
    doc = docx.Document(str(doc_path))
    for p in doc.paragraphs:
        if needle in p.text:
            p._p.getparent().remove(p._p)
            doc.save(str(doc_path))
            return
    raise AssertionError(f"no paragraph containing {needle!r} to corrupt")


def _replace_run_text(doc_path: Path, old: str, new: str) -> None:
    """Replace the first run whose text equals `old` with `new` — simulates a tag
    pointed at the wrong field (a mapping bug rebuild would not otherwise catch)."""
    doc = docx.Document(str(doc_path))
    for p in doc.paragraphs:
        for r in p.runs:
            if r.text == old:
                r.text = new
                doc.save(str(doc_path))
                return
    raise AssertionError(f"no run with text {old!r} to corrupt")


# --------------------------------------------------------------------------------------
# expected_tags
# --------------------------------------------------------------------------------------


def test_expected_tags_always_includes_name_contact_and_bullet():
    _, profile = _build(_full_featured_resume)
    tags = template_verify.expected_tags(profile)
    assert "{{ name }}" in tags
    assert "{{r contact }}" in tags
    assert "{{ bullet }}" in tags


def test_expected_tags_omits_an_unmapped_optional_field():
    _, profile = _build(_full_featured_resume)
    # `_full_featured_resume`'s experience header has no explicit location omission —
    # flip one off to prove expected_tags actually reads field.present, not just kind.
    fields = dict(profile.experience.header.fields)
    fields["location"] = fields["location"].model_copy(update={"present": False, "span": None})
    header = profile.experience.header.model_copy(update={"fields": fields})
    mutated = profile.model_copy(update={"experience": profile.experience.model_copy(update={"header": header})})
    tags = template_verify.expected_tags(mutated)
    assert "{{ job.location }}" not in tags
    assert "{{ job.company }}" in tags  # untouched fields still expected


def test_expected_tags_omits_a_disabled_kind():
    _, profile = _build(_full_featured_resume)
    mutated = profile.model_copy(
        update={
            "enabled": profile.enabled.model_copy(update={"skills": False}),
            "skills": None,
        }
    )
    tags = template_verify.expected_tags(mutated)
    assert "{{ group.label }}" not in tags
    assert "{{ group.entries }}" not in tags


def test_expected_tags_includes_section_title_only_under_generic_mode():
    _, fixed_profile = _build(_full_featured_resume)
    assert fixed_profile.section_mode == "fixed"
    assert "{{ section.title }}" not in template_verify.expected_tags(fixed_profile)

    _, generic_profile = _build(_multi_section_resume)
    assert generic_profile.section_mode == "generic"
    assert "{{ section.title }}" in template_verify.expected_tags(generic_profile)


# --------------------------------------------------------------------------------------
# verify_tagged — known-good builds verify clean
# --------------------------------------------------------------------------------------


def test_verify_tagged_clean_on_fixed_mode_build():
    tagged, profile = _build(_full_featured_resume)
    assert template_verify.verify_tagged(tagged, profile) == []


def test_verify_tagged_clean_on_generic_mode_build():
    tagged, profile = _build(_multi_section_resume)
    assert template_verify.verify_tagged(tagged, profile) == []


# --------------------------------------------------------------------------------------
# verify_tagged — catches injected defects
# --------------------------------------------------------------------------------------


def test_verify_tagged_catches_a_missing_expected_tag():
    """The exact class of bug that motivated this module: a field the profile marked
    present never actually reaching the built template."""
    tagged, profile = _build(_full_featured_resume)
    _remove_paragraph_containing(tagged, "{{ job.dates }}")

    issues = template_verify.verify_tagged(tagged, profile)
    codes = {i.code for i in issues}
    assert "tag_missing" in codes
    assert any("job.dates" in i.message for i in issues)


def test_verify_tagged_catches_an_unclosed_for_loop():
    tagged, profile = _build(_full_featured_resume)
    _remove_paragraph_containing(tagged, "{%p endfor %}")

    issues = template_verify.verify_tagged(tagged, profile)
    assert any(i.code == "unbalanced_control_tags" for i in issues)


def test_verify_tagged_catches_a_missing_generic_section_loop():
    tagged, profile = _build(_multi_section_resume)
    _remove_paragraph_containing(tagged, "{%p for section in sections %}")

    issues = template_verify.verify_tagged(tagged, profile)
    codes = {i.code for i in issues}
    # Removing the loop-open paragraph also drops every tag nested inside it, so this
    # trips both the missing-tag check and the section-loop-count check — assert the
    # one this test is actually about.
    assert "section_loop_count" in codes


def test_verify_tagged_catches_a_leftover_hyperlink():
    """The project header's baked-in hyperlink must never survive tagging — every
    per-project URL comes from render-time RichText instead."""
    tagged, profile = _build(_full_featured_resume)
    doc = docx.Document(str(tagged))
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    hyperlink = OxmlElement("w:hyperlink")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "stray link"
    run.append(text)
    hyperlink.append(run)
    doc.paragraphs[0]._p.append(hyperlink)
    doc.save(str(tagged))

    issues = template_verify.verify_tagged(tagged, profile)
    assert any(i.code == "leftover_hyperlink" for i in issues)


# --------------------------------------------------------------------------------------
# verify_roundtrip
# --------------------------------------------------------------------------------------


def test_verify_roundtrip_clean_on_a_known_good_build():
    tagged, profile = _build(_full_featured_resume)
    assert template_verify.verify_roundtrip(tagged, profile, synthetic_resume()) == []


def test_verify_roundtrip_catches_a_field_pointed_at_the_wrong_value():
    """A tag that is present (passes verify_tagged) but tags the wrong thing — the
    class of bug verify_tagged cannot see, since the *tag* is exactly where it should
    be, just wired to the wrong field."""
    tagged, profile = _build(_full_featured_resume)
    # Swap the company tag for the location tag: `{{ job.company }}` disappears from
    # the template entirely, so the company value can never reach the output.
    _replace_run_text(tagged, "{{ job.company }}", "{{ job.location }}")

    issues = template_verify.verify_roundtrip(tagged, profile, synthetic_resume())
    assert any(i.code == "roundtrip_missing" and "company" in i.message.lower() for i in issues)


def test_verify_roundtrip_ignores_entries_with_no_surviving_bullets():
    """Matches render.build_context's own filtering — an entry with zero bullets never
    renders at all, so its fields have nothing to check against."""
    resume = synthetic_resume()
    for section in resume.sections:
        if section.kind == "experience":
            section.entries[0].bullets = []
    tagged, profile = _build(_full_featured_resume)
    # The now-bulletless entry's company must not be demanded of the rendered output.
    assert template_verify.verify_roundtrip(tagged, profile, resume) == []
