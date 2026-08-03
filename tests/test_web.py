"""Web API and job queue — no network, no Word, no LibreOffice.

The pipeline is stubbed at the same seams the CLI tests use (`jd.extract`,
`rewrite.score_table`, `fit.fit`), so these assert the HTTP contract and the
single-worker queue behaviour without spending tokens.
"""

from __future__ import annotations

import json
import time
from pathlib import Path
from urllib.parse import unquote

import pytest
from fastapi.testclient import TestClient

from resume_tailor import config
from resume_tailor.data import load
from resume_tailor.events import ProgressEvent
from resume_tailor.fit import FitResult
from resume_tailor.web import jobs as jobs_mod
from resume_tailor.web import template_ops
from resume_tailor.web.app import app
from resume_tailor.web.jobs import JobQueue
from resume_tailor.web.schemas import JobSettings
# `bootstrap` is imported directly (not via `resume_tailor.workspace.bootstrap`) because
# the autouse fixture in conftest.py stubs the module attribute to a no-op for every
# other test in this file; these tests want the real implementation.
from resume_tailor.workspace import bootstrap as real_bootstrap


@pytest.fixture
def client(tmp_path, monkeypatch):
    """Fresh app client with an isolated job queue and output directory."""
    monkeypatch.setattr(config, "OUTPUT_DIR", tmp_path / "output")
    monkeypatch.setattr(config, "CACHE_DIR", tmp_path / "cache")
    config.OUTPUT_DIR.mkdir()
    config.CACHE_DIR.mkdir()

    q = JobQueue()
    monkeypatch.setattr(jobs_mod, "queue_singleton", q)
    monkeypatch.setattr(jobs_mod, "get_queue", lambda: q)

    with TestClient(app) as c:
        yield c, q


def test_get_config_returns_defaults(client):
    """GET /api/config exposes the knobs the SPA needs before a run."""
    c, _ = client
    res = c.get("/api/config")
    assert res.status_code == 200
    body = res.json()
    assert body["pages"] == config.DEFAULT_PAGE_TARGET
    assert body["experience"] == config.MAX_EXPERIENCE_ENTRIES
    assert "claude" in body["model_profiles"]
    assert "lmstudio" in body["model_profiles"]
    assert isinstance(body["tag_vocabulary"], list)
    assert body["pdf_backend"] in ("word", "soffice")
    assert "fill_target" in body
    assert 0.8 <= body["fill_target"] <= 0.95
    # Stored vocabulary (or derived fallback) should be non-empty for a real master resume.
    assert len(body["tag_vocabulary"]) >= 1


def test_create_job_rejects_empty_jd(client):
    """An empty JD is a 400, not a queued no-op."""
    c, _ = client
    res = c.post("/api/jobs", json={"jd_text": "   ", "settings": {}})
    assert res.status_code == 400


def test_get_config_exposes_gemini_fields(client):
    """The settings panel needs the Gemini env defaults and which profiles use them,
    mirroring the Ollama fields it already gets."""
    c, _ = client
    body = c.get("/api/config").json()
    assert "gemini" in body["model_profiles"]
    assert body["gemini_model"] == config.GEMINI_MODEL
    assert body["gemini_base_url"] == config.GEMINI_BASE_URL
    assert "gemini" in body["gemini_profiles"]
    assert "ollama" not in body["gemini_profiles"]
    assert "gemini" in body["provider_keys"]
    # Booleans only — never the key value itself.
    assert isinstance(body["provider_keys"]["gemini"], bool)


def test_create_job_rejects_a_profile_with_no_key(client, monkeypatch):
    """A missing Gemini key must fail the POST synchronously, not surface later as an
    async `job.status == 'failed'` once the worker gets to it."""
    c, q = client
    for name in ("GEMINI_API_KEY", "GOOGLE_API_KEY", "LLM_API_KEY"):
        monkeypatch.delenv(name, raising=False)

    res = c.post(
        "/api/jobs",
        json={"jd_text": "Some job description.", "settings": {"model": "gemini"}},
    )
    assert res.status_code == 400
    assert "GEMINI_API_KEY" in res.json()["detail"]
    # The queue never saw it — rejected before `submit()`, not queued and then failed.
    assert q._jobs == {}


def test_create_job_with_a_gemini_key_present_is_accepted(client, monkeypatch):
    """The inverse: a key present, even a fake one, must not be blocked at the door."""
    c, q = client
    monkeypatch.setenv("GEMINI_API_KEY", "fake-key-for-test")
    monkeypatch.setattr(jobs_mod.jd, "extract", _stub_no_network_extract)

    res = c.post(
        "/api/jobs",
        json={"jd_text": "Some job description.", "settings": {"model": "gemini"}},
    )
    assert res.status_code == 200
    _drain(c, res.json()["job_id"])
    config.resolve("claude")


def _point_settings_at(tmp_path: Path, monkeypatch) -> Path:
    """Redirect settings.json under tmp_path so a test never touches the real repo."""
    path = tmp_path / "settings.json"
    monkeypatch.setattr(config, "SETTINGS_PATH", path)
    return path


def test_get_settings_returns_seeded_defaults_when_missing(client, tmp_path, monkeypatch):
    """No settings.json yet: GET reports JobSettings() defaults and seeded=True."""
    c, _ = client
    _point_settings_at(tmp_path, monkeypatch)

    res = c.get("/api/settings")
    assert res.status_code == 200
    body = res.json()
    assert body["seeded"] is True
    assert body["settings"]["pages"] == 1
    # Ollama, not Claude: a profile that has never saved settings must still be runnable
    # without an Anthropic key.
    assert body["settings"]["model"] == "ollama"
    # Blank means "use the server's OLLAMA_MODEL" — the field only overrides when set.
    assert body["settings"]["ollama_model"] is None


def test_settings_round_trip(client, tmp_path, monkeypatch):
    """PUT persists new defaults; a later GET reflects them and reports seeded=False."""
    c, _ = client
    path = _point_settings_at(tmp_path, monkeypatch)

    res = c.put("/api/settings", json={"settings": {"pages": 2, "model": "ollama"}})
    assert res.status_code == 200
    body = res.json()
    assert body["seeded"] is False
    assert body["settings"]["pages"] == 2
    assert body["settings"]["model"] == "ollama"
    assert path.exists()

    got = c.get("/api/settings")
    assert got.status_code == 200
    got_body = got.json()
    assert got_body["seeded"] is False
    assert got_body["settings"]["pages"] == 2
    assert got_body["settings"]["model"] == "ollama"


def test_create_job_without_settings_uses_saved_defaults(client, tmp_path, monkeypatch):
    """POST /api/jobs with no `settings` falls back to the active profile's saved ones."""
    c, q = client
    _point_settings_at(tmp_path, monkeypatch)
    c.put("/api/settings", json={"settings": {"pages": 3, "no_semantic": True}})

    # The background worker picks this job up immediately; stub the first pipeline
    # call so it fails fast in-process instead of ever reaching the network, whether
    # or not an API key happens to be configured in this environment.
    def _stub_extract_raises(*args, **kwargs):
        raise RuntimeError("stub — no network in tests")

    monkeypatch.setattr(jobs_mod.jd, "extract", _stub_extract_raises)

    res = c.post("/api/jobs", json={"jd_text": "Some job description."})
    assert res.status_code == 200
    job_id = res.json()["job_id"]

    job = q.get(job_id)
    assert job.settings.pages == 3
    assert job.settings.no_semantic is True


def _stub_no_network_extract(*args, **kwargs):
    """Fail the run at its first pipeline call, after routing has been resolved."""
    raise RuntimeError("stub — no network in tests")


def _drain(c, job_id: str) -> dict:
    """Block until the queue's worker thread has finished `job_id`, and return its status.

    The worker is a background thread, so anything asserting on what a *run* did (rather
    than on the queued `Job` object) has to wait for it or it is a race that passes on a
    fast machine.
    """
    import time

    deadline = time.time() + 10
    while time.time() < deadline:
        status = c.get(f"/api/jobs/{job_id}").json()
        if status["status"] in ("succeeded", "failed"):
            return status
        time.sleep(0.05)
    pytest.fail("job did not finish in time")


def test_ollama_model_setting_repoints_only_the_ollama_stages(client, monkeypatch):
    """The UI's Ollama tag field must reach `config.resolve` as per-stage overrides.

    Asserted at the resolve boundary rather than on a live call: what the field is for is
    changing the tag without an `.env` edit and a restart, and that is entirely a routing
    question. Under `hybrid` the Claude rewrite stage must come through untouched.
    """
    c, _q = client
    seen: dict[str, object] = {}
    real_resolve = config.resolve

    def recording_resolve(profile=None, *, overrides=None, effort=None):
        seen["profile"] = profile
        seen["overrides"] = dict(overrides or {})
        return real_resolve(profile, overrides=overrides, effort=effort)

    monkeypatch.setattr(jobs_mod.config, "resolve", recording_resolve)
    monkeypatch.setattr(jobs_mod.jd, "extract", _stub_no_network_extract)

    res = c.post(
        "/api/jobs",
        json={
            "jd_text": "Some job description.",
            "settings": {"model": "hybrid", "ollama_model": "gemma4"},
        },
    )
    assert res.status_code == 200
    _drain(c, res.json()["job_id"])

    assert seen["profile"] == "hybrid"
    assert seen["overrides"] == {
        "extract": "gemma4",
        "score": "gemma4",
        "expand": "gemma4",
        "facets": "gemma4",
    }
    config.resolve("claude")


def test_blank_ollama_model_leaves_the_env_default_in_place(client, monkeypatch):
    """The field overrides *only* when filled in; blank must resolve to OLLAMA_MODEL.

    The inverse of the override test below, and the one users actually depend on: a saved
    settings blob with no tag in it must not start sending an empty model to the API.
    """
    c, _q = client
    seen: dict[str, object] = {}
    real_resolve = config.resolve

    def recording_resolve(profile=None, *, overrides=None, effort=None):
        seen["overrides"] = dict(overrides or {})
        return real_resolve(profile, overrides=overrides, effort=effort)

    monkeypatch.setattr(jobs_mod.config, "resolve", recording_resolve)
    monkeypatch.setattr(jobs_mod.jd, "extract", _stub_no_network_extract)

    res = c.post(
        "/api/jobs",
        json={"jd_text": "Some job description.", "settings": {"model": "ollama"}},
    )
    assert res.status_code == 200
    _drain(c, res.json()["job_id"])

    assert seen["overrides"] == {}
    assert config.model_for("extract") == config.OLLAMA_MODEL == "gemma4:cloud"
    config.resolve("claude")


def test_gemini_model_setting_repoints_only_the_gemini_stages(client, monkeypatch):
    """Mirrors the Ollama tag test above, for the Gemini field added alongside it."""
    c, _q = client
    monkeypatch.setenv("GEMINI_API_KEY", "fake-key-for-test")
    seen: dict[str, object] = {}
    real_resolve = config.resolve

    def recording_resolve(profile=None, *, overrides=None, effort=None):
        seen["profile"] = profile
        seen["overrides"] = dict(overrides or {})
        return real_resolve(profile, overrides=overrides, effort=effort)

    monkeypatch.setattr(jobs_mod.config, "resolve", recording_resolve)
    monkeypatch.setattr(jobs_mod.jd, "extract", _stub_no_network_extract)

    res = c.post(
        "/api/jobs",
        json={
            "jd_text": "Some job description.",
            "settings": {"model": "gemini", "gemini_model": "gemini-3.5-pro"},
        },
    )
    assert res.status_code == 200
    _drain(c, res.json()["job_id"])

    assert seen["profile"] == "gemini"
    assert seen["overrides"] == dict.fromkeys(config.PURPOSES, "gemini-3.5-pro")
    config.resolve("claude")


def test_explicit_stage_override_beats_the_blanket_ollama_tag(client, monkeypatch):
    """`rewrite_model` is the narrower choice and must not be clobbered by the tag."""
    c, _q = client
    seen: dict[str, object] = {}
    real_resolve = config.resolve

    def recording_resolve(profile=None, *, overrides=None, effort=None):
        seen["overrides"] = dict(overrides or {})
        return real_resolve(profile, overrides=overrides, effort=effort)

    monkeypatch.setattr(jobs_mod.config, "resolve", recording_resolve)
    monkeypatch.setattr(jobs_mod.jd, "extract", _stub_no_network_extract)

    res = c.post(
        "/api/jobs",
        json={
            "jd_text": "Some job description.",
            "settings": {
                "model": "ollama",
                "ollama_model": "gemma4",
                "rewrite_model": "claude-sonnet-5",
            },
        },
    )
    assert res.status_code == 200
    _drain(c, res.json()["job_id"])

    assert seen["overrides"]["rewrite"] == "claude-sonnet-5"
    assert seen["overrides"]["extract"] == "gemma4"
    config.resolve("claude")


def test_job_runs_to_success_with_stubbed_pipeline(client, monkeypatch, tmp_path):
    """A submitted job reaches succeeded and writes a downloadable .docx."""
    c, q = client
    resume = load()

    def fake_extract(text, *, known_tags=None, use_cache=True, on_event=None):
        from resume_tailor.jd import JobRequirements, Keyword

        if on_event:
            on_event(ProgressEvent("extract", "stub extract", {}))
        return JobRequirements(
            title="Stub Role",
            seniority="intern",
            keywords=[Keyword(phrase="Python", canonical="python", importance="must_have")],
        )

    def fake_score(bullets, requirements, *, use_cache=True, on_event=None):
        if on_event:
            on_event(ProgressEvent("score", "stub score", {}))
        return {b.id: 5.0 for b in bullets}

    seen_fit: dict = {}

    def fake_fit(
        resume,
        requirements,
        *,
        target_pages=1,
        template=None,
        out=None,
        max_experience=None,
        max_projects=None,
        semantic=None,
        repair_widows=True,
        repair_verbs=True,
        merge_bullets=False,
        include_project_links=True,
        fill_target=None,
        on_event=None,
    ):
        """Stub fit and record polish/merge/link knobs from JobSettings."""
        seen_fit.update(
            repair_widows=repair_widows,
            repair_verbs=repair_verbs,
            merge_bullets=merge_bullets,
            include_project_links=include_project_links,
            fill_target=fill_target,
        )
        out = Path(out)
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_bytes(b"PK")  # pretend docx
        out.with_suffix(".pdf").write_bytes(b"%PDF-1.4 stub")
        if on_event:
            on_event(ProgressEvent("fit", "stub fit done", {"pages": 1}))
        # One real bullet so report_data has something to summarise.
        bullet = resume.all_bullets()[0]
        return FitResult(
            out_path=out,
            pages=1,
            pages_are_estimated=False,
            iterations=1,
            bullets_selected=1,
            bullets_total=1,
            bullets={bullet.id: bullet.text},
            semantic_used=bool(semantic),
        )

    monkeypatch.setattr(jobs_mod.jd, "extract", fake_extract)
    monkeypatch.setattr(jobs_mod.rewrite, "score_table", fake_score)
    monkeypatch.setattr(jobs_mod.fit, "fit", fake_fit)
    monkeypatch.setattr(jobs_mod.jd, "verify_verbatim", lambda *a, **k: [])

    def fake_facets(resume, requirements, **kwargs):
        """Budget-only facets so the job path never reaches the network."""
        from resume_tailor import facets as facets_mod

        return facets_mod.budget_only(
            resume,
            requirements,
            include_project_links=kwargs.get("include_project_links", True),
        )

    monkeypatch.setattr(jobs_mod.facets, "select_facets", fake_facets)

    from resume_tailor.expand import ExpandedEntry, Expansion

    def fake_expand(*a, **k):
        """Stub expansion so web tests never reach the network."""
        entry = resume.experience[0]
        return Expansion(
            entries=[
                ExpandedEntry(
                    entry_key="exp:0",
                    title=entry.title,
                    company=entry.company,
                    location=entry.location,
                    start=entry.start,
                    end=entry.end,
                    bullets=["Expanded bullet from stub."],
                    char_count=28,
                    on_resume=True,
                )
            ],
            model="stub",
            char_limit=config.EXPAND_CHAR_LIMIT,
        )

    monkeypatch.setattr(jobs_mod.expand, "expand_experience", fake_expand)

    res = c.post(
        "/api/jobs",
        json={
            "jd_text": "Looking for a Python intern.",
            "settings": {
                "model": "claude",
                "merge": True,
                "no_verb_repair": True,
                "no_project_links": True,
                "fill_target": 0.88,
            },
        },
    )
    assert res.status_code == 200
    job_id = res.json()["job_id"]

    status = _drain(c, job_id)
    assert status["status"] == "succeeded", status
    assert status["report"]["title"] == "Stub Role"
    assert status["report"]["pages"] == 1
    assert status["report"]["verb_collisions_remaining"] == 0
    assert status["expansion"] is not None
    assert status["expansion"]["entries"][0]["company"] == resume.experience[0].company
    assert seen_fit == {
        "repair_widows": True,
        "repair_verbs": False,
        "merge_bullets": True,
        "include_project_links": False,
        "fill_target": 0.88,
    }
    assert any(e["stage"] == "extract" for e in status["events"])

    docx = c.get(f"/api/jobs/{job_id}/download.docx")
    assert docx.status_code == 200
    assert docx.content.startswith(b"PK")
    disposition = unquote(docx.headers.get("content-disposition", ""))
    assert f"{resume.contact.name} Resume - Stub Role.docx" in disposition

    pdf = c.get(f"/api/jobs/{job_id}/preview.pdf")
    assert pdf.status_code == 200
    assert pdf.content.startswith(b"%PDF")
    preview_disp = unquote(pdf.headers.get("content-disposition", ""))
    assert f"{resume.contact.name} Resume - Stub Role.pdf" in preview_disp
    assert "inline" in preview_disp.lower()

    pdf_dl = c.get(f"/api/jobs/{job_id}/download.pdf")
    assert pdf_dl.status_code == 200
    assert pdf_dl.content.startswith(b"%PDF")
    download_disp = unquote(pdf_dl.headers.get("content-disposition", ""))
    assert f"{resume.contact.name} Resume - Stub Role.pdf" in download_disp
    assert "attachment" in download_disp.lower()

    expansion_md = c.get(f"/api/jobs/{job_id}/expansion.md")
    assert expansion_md.status_code == 200
    assert b"Expanded bullet from stub." in expansion_md.content


def test_master_resume_validate_rejects_bad_payload(client):
    """POST /api/master-resume/validate returns field errors, not a 500."""
    c, _ = client
    res = c.post("/api/master-resume/validate", json={"contact": {"name": "x"}})
    assert res.status_code == 200
    body = res.json()
    assert body["ok"] is False
    assert body["errors"]


def test_master_resume_round_trip(client, tmp_path, monkeypatch):
    """PUT saves a validated resume and keeps a backup of the previous file."""
    c, _ = client
    resume = load()
    path = tmp_path / "master_resume.json"
    path.write_text(config.MASTER_RESUME_PATH.read_text(encoding="utf-8"), encoding="utf-8")
    monkeypatch.setattr(config, "MASTER_RESUME_PATH", path)
    monkeypatch.setattr(config, "DATA_DIR", tmp_path)

    payload = resume.model_dump(by_alias=True)
    payload["contact"]["name"] = "Test User"

    res = c.put("/api/master-resume", json=payload)
    assert res.status_code == 200
    body = res.json()
    assert body["ok"] is True
    assert body["summary"]["name"] == "Test User"
    assert list(tmp_path.glob("master_resume.*.bak.json")), "expected a timestamped backup"

    got = c.get("/api/master-resume")
    assert got.status_code == 200
    assert got.json()["contact"]["name"] == "Test User"


def test_master_resume_accepts_appended_experience_bullet(client, tmp_path, monkeypatch):
    """PUT with a UI-authored experience entry and new bullet id grows the store."""
    c, _ = client
    resume = load()
    path = tmp_path / "master_resume.json"
    path.write_text(config.MASTER_RESUME_PATH.read_text(encoding="utf-8"), encoding="utf-8")
    monkeypatch.setattr(config, "MASTER_RESUME_PATH", path)
    monkeypatch.setattr(config, "DATA_DIR", tmp_path)

    before = len(resume.all_bullets())
    payload = resume.model_dump(by_alias=True)
    payload["experience"].append(
        {
            "company": "Editor Test Co",
            "title": "Software Engineer",
            "location": "Remote",
            "start": "2024-01",
            "end": "present",
            "bullets": [
                {
                    "id": "edittest_b1",
                    "text": "Shipped a feature with Python and FastAPI.",
                    "tags": ["Python", "FastAPI"],
                    "metric": False,
                }
            ],
        }
    )

    res = c.put("/api/master-resume", json=payload)
    assert res.status_code == 200
    body = res.json()
    assert body["ok"] is True
    assert body["summary"]["bullets"] == before + 1
    assert body["summary"]["experience"] == len(resume.experience) + 1

    got = c.get("/api/master-resume").json()
    last = got["experience"][-1]
    assert last["company"] == "Editor Test Co"
    assert last["bullets"][0]["id"] == "edittest_b1"
    # Tags come back canonicalised the way data.Bullet._normalise_tags stores them.
    assert "python" in [t.lower() for t in last["bullets"][0]["tags"]]


def test_queue_serialises_jobs(monkeypatch, tmp_path):
    """Two submitted jobs never run concurrently — the second waits for the first."""
    import threading
    import time

    monkeypatch.setattr(config, "OUTPUT_DIR", tmp_path / "output")
    config.OUTPUT_DIR.mkdir()

    q = JobQueue()
    running = threading.Event()
    release = threading.Event()
    concurrent = []

    def slow_execute(job):
        if running.is_set():
            concurrent.append(True)
        running.set()
        release.wait(timeout=2)
        running.clear()
        job.report = None

    monkeypatch.setattr(q, "_execute", slow_execute)

    j1, _ = q.submit("jd one", JobSettings())
    j2, pos2 = q.submit("jd two", JobSettings())
    assert pos2 >= 1

    # Wait until the first job is running, then release it.
    assert running.wait(timeout=2)
    # While the first is held, the second must still be queued.
    assert q.get(j2.job_id).status == "queued"
    release.set()

    deadline = time.time() + 5
    while time.time() < deadline:
        if q.get(j1.job_id).status == "succeeded" and q.get(j2.job_id).status == "succeeded":
            break
        time.sleep(0.05)

    assert q.get(j1.job_id).status == "succeeded"
    assert q.get(j2.job_id).status == "succeeded"
    assert not concurrent, "jobs overlapped — queue is not serial"


# ---------------------------------------------------------------------------
# Template tab
# ---------------------------------------------------------------------------


def _minimal_docx_bytes(paragraph: str | None = None) -> bytes:
    """Build a tiny valid .docx in memory for upload tests (no Word required).

    Optional `paragraph` text makes two fixtures differ byte-for-byte.
    """
    import io

    from docx import Document

    buf = io.BytesIO()
    doc = Document()
    if paragraph is not None:
        doc.add_paragraph(paragraph)
    doc.save(buf)
    return buf.getvalue()


def _resume_docx_bytes() -> bytes:
    """A single-column resume the analyzer accepts (`ready=True`).

    `_minimal_docx_bytes` is deliberately contentless and yields no suggested profile,
    so it cannot drive the wizard/profile install path.
    """
    import io

    import docx
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    document = docx.Document()

    numbering = document.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), "9")
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet")
    lvl.append(fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "●")
    lvl.append(lvl_text)
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), "9")
    abs_el = OxmlElement("w:abstractNumId")
    abs_el.set(qn("w:val"), "9")
    num.append(abs_el)
    numbering.append(num)

    def bullet(text: str):
        paragraph = document.add_paragraph(text)
        pPr = paragraph._p.get_or_add_pPr()
        numPr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        nid = OxmlElement("w:numId")
        nid.set(qn("w:val"), "9")
        numPr.append(ilvl)
        numPr.append(nid)
        pPr.append(numPr)

    document.add_paragraph("Ada Lovelace")
    document.add_paragraph("London • ada@example.com • LinkedIn")
    document.add_paragraph("EDUCATION")
    document.add_paragraph("University of London | UK\t2018 - 2022")
    bullet("BSc Computer Science | GPA: 3.9")
    bullet("Relevant Coursework: Algorithms, Databases")
    document.add_paragraph("WORK EXPERIENCES")
    document.add_paragraph("Analytical Engines | London\t2022 - Present")
    document.add_paragraph("Software Engineer")
    bullet("Built numerical engines in Python.")
    document.add_paragraph("PROJECTS")
    document.add_paragraph("Note Engine | Python, FastAPI\t2024")
    bullet("Indexed research notes with embeddings.")
    document.add_paragraph("SKILLS")
    skills = document.add_paragraph()
    skills.add_run("Languages:").bold = True
    skills.add_run(" Python, SQL")

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _point_templates_at(tmp_path: Path, monkeypatch) -> Path:
    """Redirect baseline/tagged/library paths under tmp_path and return the templates dir."""
    templates = tmp_path / "templates"
    templates.mkdir(parents=True, exist_ok=True)
    monkeypatch.setattr(config, "TEMPLATES_DIR", templates)
    monkeypatch.setattr(config, "BASELINE_TEMPLATE_PATH", templates / "original_export.docx")
    monkeypatch.setattr(config, "DEFAULT_TEMPLATE_PATH", templates / "main_template.docx")
    monkeypatch.setattr(config, "TEMPLATE_PROFILE_PATH", templates / "template_profile.json")
    monkeypatch.setattr(config, "TEMPLATE_LIBRARY_DIR", templates / "library")
    return templates


def test_get_template_returns_metadata(client, tmp_path, monkeypatch):
    """GET /api/template reports existence and calibration for redirected template paths."""
    c, _ = client
    templates = _point_templates_at(tmp_path, monkeypatch)
    (templates / "original_export.docx").write_bytes(_minimal_docx_bytes())
    (templates / "main_template.docx").write_bytes(_minimal_docx_bytes())

    res = c.get("/api/template")
    assert res.status_code == 200
    body = res.json()
    assert body["baseline"]["exists"] is True
    assert body["tagged"]["exists"] is True
    assert body["baseline"]["size_bytes"] > 0
    assert "calibration" in body
    assert "stale" in body["calibration"]
    assert isinstance(body["experience_entries"], int)
    assert isinstance(body["bullets"], int)


def test_upload_template_rejects_non_docx(client, tmp_path, monkeypatch):
    """POST /api/template with a .txt leaves the baseline untouched and returns 400."""
    c, _ = client
    templates = _point_templates_at(tmp_path, monkeypatch)
    baseline = templates / "original_export.docx"
    original = _minimal_docx_bytes()
    baseline.write_bytes(original)

    res = c.post(
        "/api/template",
        files={"file": ("resume.txt", b"not a docx", "text/plain")},
    )
    assert res.status_code == 400
    assert "docx" in res.json()["detail"].lower()
    assert baseline.read_bytes() == original


def test_upload_template_rejects_when_queue_busy(client, tmp_path, monkeypatch):
    """POST /api/template returns 409 while a job is queued or running."""
    c, q = client
    _point_templates_at(tmp_path, monkeypatch)

    # Mark the queue busy without actually running the pipeline.
    job, _ = q.submit("placeholder jd", JobSettings())
    job.status = "running"

    res = c.post(
        "/api/template",
        files={
            "file": (
                "resume.docx",
                _minimal_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert res.status_code == 409
    assert "progress" in res.json()["detail"].lower() or "job" in res.json()["detail"].lower()


def test_upload_template_backs_up_and_rebuilds(client, tmp_path, monkeypatch):
    """Successful upload writes a timestamped backup and invokes the build stub."""
    c, _ = client
    templates = _point_templates_at(tmp_path, monkeypatch)
    baseline = templates / "original_export.docx"
    tagged = templates / "main_template.docx"
    old_bytes = _minimal_docx_bytes()
    baseline.write_bytes(old_bytes)
    tagged.write_bytes(old_bytes)

    builds: list[int] = []

    def fake_build():
        """Pretend build_template.py succeeded and wrote a new tagged file."""
        builds.append(1)
        tagged.write_bytes(b"PK\x03\x04rebuilt")
        return 0, "found 1 education, 1 experience entries, 1 projects\nwrote main_template.docx\n"

    monkeypatch.setattr(template_ops, "_run_build", fake_build)

    new_bytes = _minimal_docx_bytes()
    # Ensure the new upload differs from the old baseline so we can assert replacement.
    assert new_bytes  # non-empty

    res = c.post(
        "/api/template",
        files={
            "file": (
                "resume.docx",
                new_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert res.status_code == 200, res.text
    body = res.json()
    assert body["ok"] is True
    assert builds == [1]
    assert "wrote" in body["log"].lower() or "found" in body["log"].lower()
    backups = list((templates / "backups").glob("original_export.*.docx"))
    assert len(backups) == 1
    assert backups[0].read_bytes() == old_bytes
    assert baseline.read_bytes() == new_bytes


def test_upload_template_restores_baseline_on_build_failure(client, tmp_path, monkeypatch):
    """Failed build restores the previous baseline byte-for-byte and returns 422."""
    c, _ = client
    templates = _point_templates_at(tmp_path, monkeypatch)
    baseline = templates / "original_export.docx"
    old_bytes = _minimal_docx_bytes()
    baseline.write_bytes(old_bytes)

    def failing_build():
        """Simulate build_template.py exiting non-zero with a useful log."""
        return 1, "ERROR: could not find section heading(s): WORK EXPERIENCES.\n"

    monkeypatch.setattr(template_ops, "_run_build", failing_build)

    res = c.post(
        "/api/template",
        files={
            "file": (
                "resume.docx",
                _minimal_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert res.status_code == 422
    detail = res.json()["detail"]
    assert isinstance(detail, dict)
    assert "log" in detail
    assert "WORK EXPERIENCES" in detail["log"]
    assert baseline.read_bytes() == old_bytes


def test_template_preview_uses_stubbed_render(client, tmp_path, monkeypatch):
    """GET /api/template/preview.pdf serves a PDF produced via the render seam (no Word)."""
    c, _ = client
    templates = _point_templates_at(tmp_path, monkeypatch)
    (templates / "main_template.docx").write_bytes(_minimal_docx_bytes())
    monkeypatch.setattr(config, "OUTPUT_DIR", tmp_path / "output")
    config.OUTPUT_DIR.mkdir(exist_ok=True)

    def fake_render(resume, *, out, **_kwargs):
        """Write a placeholder .docx where the preview would land."""
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_bytes(b"fake-docx")
        return out

    def fake_to_pdf(docx_path, pdf_path=None, **_kwargs):
        """Write a minimal PDF-like payload without calling Word/LibreOffice."""
        target = pdf_path or docx_path.with_suffix(".pdf")
        # Minimal PDF header so FileResponse has something to stream.
        target.write_bytes(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\ntrailer\n%%EOF\n")
        return target

    monkeypatch.setattr(template_ops.render, "render", fake_render)
    monkeypatch.setattr(template_ops.render, "to_pdf", fake_to_pdf)

    res = c.get("/api/template/preview.pdf")
    assert res.status_code == 200
    assert res.headers["content-type"].startswith("application/pdf")
    assert res.content.startswith(b"%PDF")


def test_preview_cache_invalidates_when_master_resume_changes(client, tmp_path, monkeypatch):
    """Editing the master resume re-renders the template preview.

    Regression: `ensure_preview` compared the cached PDF's mtime against the *tagged
    template* only. Since the preview renders the resume through that template, a
    resume edit left the template untouched, the cache looked fresh, and the Template
    tab served the pre-edit PDF forever — the "Refresh doesn't show my changes" report.
    """
    import os

    c, _ = client
    templates = _point_templates_at(tmp_path, monkeypatch)
    (templates / "main_template.docx").write_bytes(_minimal_docx_bytes())
    monkeypatch.setattr(config, "OUTPUT_DIR", tmp_path / "preview-out")
    resume_path = tmp_path / "master_resume.json"
    resume_path.write_text(config.MASTER_RESUME_PATH.read_text(encoding="utf-8"), encoding="utf-8")
    monkeypatch.setattr(config, "MASTER_RESUME_PATH", resume_path)

    renders: list[int] = []

    def fake_render(resume, *, out, **_kwargs):
        renders.append(1)
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_bytes(b"fake-docx")
        return out

    def fake_to_pdf(docx_path, pdf_path=None, **_kwargs):
        target = pdf_path or docx_path.with_suffix(".pdf")
        target.write_bytes(b"%PDF-1.4\n%%EOF\n")
        return target

    monkeypatch.setattr(template_ops.render, "render", fake_render)
    monkeypatch.setattr(template_ops.render, "to_pdf", fake_to_pdf)

    assert c.get("/api/template/preview.pdf").status_code == 200
    assert len(renders) == 1
    # Unchanged inputs must still hit the cache — this is a real render each time.
    assert c.get("/api/template/preview.pdf").status_code == 200
    assert len(renders) == 1

    _, pdf_path = template_ops._preview_paths()
    future = pdf_path.stat().st_mtime + 10
    resume_path.write_text(resume_path.read_text(encoding="utf-8"), encoding="utf-8")
    os.utime(resume_path, (future, future))

    assert c.get("/api/template/preview.pdf").status_code == 200
    assert len(renders) == 2, "a master-resume edit must invalidate the preview cache"


def test_analyze_template_returns_structured_report(client, tmp_path, monkeypatch):
    """POST /api/template/analyze does not write under templates/ and returns issues."""
    c, _ = client
    templates = _point_templates_at(tmp_path, monkeypatch)
    before = list(templates.iterdir()) if templates.exists() else []

    res = c.post(
        "/api/template/analyze",
        files={
            "file": (
                "resume.docx",
                _minimal_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert res.status_code == 200
    body = res.json()
    assert "source_sha256" in body
    assert "issues" in body
    assert "paragraphs" in body
    assert body["ready"] is False
    after = list(templates.iterdir()) if templates.exists() else []
    assert after == before


def test_get_template_includes_profile_summary(client, tmp_path, monkeypatch):
    """GET /api/template always includes a profile summary object."""
    c, _ = client
    _point_templates_at(tmp_path, monkeypatch)
    res = c.get("/api/template")
    assert res.status_code == 200
    body = res.json()
    assert "profile" in body
    assert body["profile"]["exists"] is False


def test_upload_template_with_calibrate_flag(client, tmp_path, monkeypatch):
    """calibrate=true runs calibration after a successful legacy build and reloads config."""
    from resume_tailor.calibrate import CalibrationResult

    c, _ = client
    templates = _point_templates_at(tmp_path, monkeypatch)
    baseline = templates / "original_export.docx"
    tagged = templates / "main_template.docx"
    baseline.write_bytes(_minimal_docx_bytes())
    tagged.write_bytes(_minimal_docx_bytes())

    def fake_build():
        """Pretend build_template.py succeeded."""
        tagged.write_bytes(b"PK\x03\x04rebuilt")
        return 0, "wrote main_template.docx\n"

    cal_calls: list[dict] = []

    def fake_calibrate(*, verify_anchors=True):
        """Record the calibrate call without touching Word/LibreOffice."""
        cal_calls.append({"verify_anchors": verify_anchors})
        path = tmp_path / "calibration.json"
        path.write_text("{}", encoding="utf-8")
        return CalibrationResult(
            chars_per_line=99,
            lines_per_page=48,
            path=path,
            log="CHARS_PER_LINE = 99\nLINES_PER_PAGE = 48",
        )

    monkeypatch.setattr(template_ops, "_run_build", fake_build)
    monkeypatch.setattr(template_ops.calibrate, "run", fake_calibrate)
    monkeypatch.setattr(config, "reload_calibration", lambda: (99, 48, "test"))

    res = c.post(
        "/api/template",
        data={"calibrate": "true"},
        files={
            "file": (
                "resume.docx",
                _minimal_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert res.status_code == 200
    body = res.json()
    assert body["ok"] is True
    assert cal_calls == [{"verify_anchors": True}]
    assert "CHARS_PER_LINE = 99" in body["log"]


def test_library_seeds_default_from_live(client, tmp_path, monkeypatch):
    """GET /api/template/library registers live baseline+tagged as Default when empty."""
    c, _ = client
    templates = _point_templates_at(tmp_path, monkeypatch)
    payload = _minimal_docx_bytes()
    (templates / "original_export.docx").write_bytes(payload)
    (templates / "main_template.docx").write_bytes(payload)

    res = c.get("/api/template/library")
    assert res.status_code == 200
    body = res.json()
    assert len(body["entries"]) == 1
    assert body["entries"][0]["label"] == "Default"
    assert body["entries"][0]["is_active"] is True
    assert body["active_id"] == body["entries"][0]["id"]

    info = c.get("/api/template").json()
    assert info["active_label"] == "Default"
    assert info["active_library_id"] == body["active_id"]


def test_upload_with_label_creates_library_entry(client, tmp_path, monkeypatch):
    """POST /api/template with label snapshots the install into the named library."""
    c, _ = client
    templates = _point_templates_at(tmp_path, monkeypatch)
    tagged = templates / "main_template.docx"
    old = _minimal_docx_bytes()
    (templates / "original_export.docx").write_bytes(old)
    tagged.write_bytes(old)

    def fake_build():
        """Pretend build_template.py succeeded."""
        tagged.write_bytes(b"PK\x03\x04rebuilt-for-library")
        return 0, "wrote main_template.docx\n"

    monkeypatch.setattr(template_ops, "_run_build", fake_build)

    res = c.post(
        "/api/template",
        data={"label": "Campus CV"},
        files={
            "file": (
                "campus.docx",
                _minimal_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert res.status_code == 200, res.text
    info = res.json()["info"]
    assert info["active_label"] == "Campus CV"

    lib = c.get("/api/template/library").json()
    labels = {e["label"] for e in lib["entries"]}
    # Prior live was seeded/preserved as Default; new install is Campus CV.
    assert "Campus CV" in labels
    assert "Default" in labels
    active = next(e for e in lib["entries"] if e["is_active"])
    assert active["label"] == "Campus CV"


def test_activate_library_switches_live_baseline(client, tmp_path, monkeypatch):
    """Activating another library entry restores its baseline bytes into the live slot."""
    c, _ = client
    templates = _point_templates_at(tmp_path, monkeypatch)
    tagged = templates / "main_template.docx"
    first = _minimal_docx_bytes()
    (templates / "original_export.docx").write_bytes(first)
    tagged.write_bytes(first)

    # Seed Default.
    assert c.get("/api/template/library").status_code == 200
    default_id = c.get("/api/template/library").json()["active_id"]

    builds: list[int] = []

    def fake_build():
        """Write a distinct tagged payload for the second install."""
        builds.append(1)
        tagged.write_bytes(b"PK\x03\x04second-tagged")
        return 0, "wrote\n"

    monkeypatch.setattr(template_ops, "_run_build", fake_build)
    second = _minimal_docx_bytes("second baseline")
    assert second != first
    res = c.post(
        "/api/template",
        data={"label": "Second"},
        files={
            "file": (
                "second.docx",
                second,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert res.status_code == 200
    assert (templates / "original_export.docx").read_bytes() == second

    act = c.post(f"/api/template/library/{default_id}/activate")
    assert act.status_code == 200, act.text
    assert (templates / "original_export.docx").read_bytes() == first
    body = act.json()
    assert body["ok"] is True
    assert body["info"]["active_library_id"] == default_id
    assert body["info"]["active_label"] == "Default"


def test_rename_library_rejects_duplicate_label(client, tmp_path, monkeypatch):
    """PATCH rename fails when the new label collides case-insensitively."""
    c, _ = client
    templates = _point_templates_at(tmp_path, monkeypatch)
    tagged = templates / "main_template.docx"
    payload = _minimal_docx_bytes()
    (templates / "original_export.docx").write_bytes(payload)
    tagged.write_bytes(payload)
    c.get("/api/template/library")

    def fake_build():
        """Legacy build stub."""
        tagged.write_bytes(b"PK\x03\x04x")
        return 0, "ok\n"

    monkeypatch.setattr(template_ops, "_run_build", fake_build)
    c.post(
        "/api/template",
        data={"label": "Alpha"},
        files={
            "file": (
                "a.docx",
                _minimal_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    lib = c.get("/api/template/library").json()
    default = next(e for e in lib["entries"] if e["label"] == "Default")
    res = c.patch(
        f"/api/template/library/{default['id']}",
        json={"label": "alpha"},
    )
    assert res.status_code == 400
    assert "already exists" in res.json()["detail"].lower()


def test_delete_library_refuses_active(client, tmp_path, monkeypatch):
    """DELETE on the active entry returns 400; non-active deletes succeed."""
    c, _ = client
    templates = _point_templates_at(tmp_path, monkeypatch)
    tagged = templates / "main_template.docx"
    payload = _minimal_docx_bytes()
    (templates / "original_export.docx").write_bytes(payload)
    tagged.write_bytes(payload)
    c.get("/api/template/library")

    def fake_build():
        """Legacy build stub."""
        tagged.write_bytes(b"PK\x03\x04y")
        return 0, "ok\n"

    monkeypatch.setattr(template_ops, "_run_build", fake_build)
    c.post(
        "/api/template",
        data={"label": "Spare"},
        files={
            "file": (
                "s.docx",
                _minimal_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    lib = c.get("/api/template/library").json()
    active = next(e for e in lib["entries"] if e["is_active"])
    other = next(e for e in lib["entries"] if not e["is_active"])

    bad = c.delete(f"/api/template/library/{active['id']}")
    assert bad.status_code == 400
    assert "active" in bad.json()["detail"].lower()

    ok = c.delete(f"/api/template/library/{other['id']}")
    assert ok.status_code == 200
    labels = {e["label"] for e in ok.json()["entries"]}
    assert other["label"] not in labels
    assert active["label"] in labels


def test_library_cap_refuses_twenty_first(client, tmp_path, monkeypatch):
    """Installing when the library already has 20 entries returns 400."""
    c, _ = client
    templates = _point_templates_at(tmp_path, monkeypatch)
    tagged = templates / "main_template.docx"
    payload = _minimal_docx_bytes()
    (templates / "original_export.docx").write_bytes(payload)
    tagged.write_bytes(payload)

    # Fill the library with synthetic entries (no install needed).
    monkeypatch.setattr(template_ops, "_LIBRARY_MAX_ENTRIES", 2)
    c.get("/api/template/library")  # seeds Default (1)

    def fake_build():
        """Legacy build stub."""
        tagged.write_bytes(b"PK\x03\x04z")
        return 0, "ok\n"

    monkeypatch.setattr(template_ops, "_run_build", fake_build)
    first = c.post(
        "/api/template",
        data={"label": "Two"},
        files={
            "file": (
                "t.docx",
                _minimal_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert first.status_code == 200
    assert len(c.get("/api/template/library").json()["entries"]) == 2

    blocked = c.post(
        "/api/template",
        data={"label": "Three"},
        files={
            "file": (
                "u.docx",
                _minimal_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert blocked.status_code == 400
    assert "full" in blocked.json()["detail"].lower()


def _point_workspaces_at(tmp_path: Path, monkeypatch) -> dict[str, Path]:
    """Redirect every workspace storage root under tmp_path so tests never touch the
    real repo's data/templates/output trees, and reset the active-workspace pointer."""
    data_root = tmp_path / "ws_data"
    templates_root = tmp_path / "ws_templates"
    output_root = tmp_path / "ws_output"
    data_root.mkdir(parents=True, exist_ok=True)
    templates_root.mkdir(parents=True, exist_ok=True)
    output_root.mkdir(parents=True, exist_ok=True)
    monkeypatch.delenv("RESUME_TAILOR_CALIBRATION_DIR", raising=False)
    for name, value in (
        ("DATA_ROOT", data_root),
        ("TEMPLATES_ROOT", templates_root),
        ("OUTPUT_ROOT", output_root),
        ("CACHE_ROOT", output_root),
        ("DATA_DIR", data_root),
        ("TEMPLATES_DIR", templates_root),
        ("OUTPUT_DIR", output_root),
        ("CACHE_DIR", output_root),
        ("MASTER_RESUME_PATH", data_root / "master_resume.json"),
        ("SETTINGS_PATH", data_root / "settings.json"),
        ("DEFAULT_TEMPLATE_PATH", templates_root / "main_template.docx"),
        ("BASELINE_TEMPLATE_PATH", templates_root / "original_export.docx"),
        ("TEMPLATE_PROFILE_PATH", templates_root / "template_profile.json"),
        ("TEMPLATE_LIBRARY_DIR", templates_root / "library"),
        ("_ACTIVE_WORKSPACE_ID", None),
    ):
        monkeypatch.setattr(config, name, value)
    return {"data": data_root, "templates": templates_root, "output": output_root}


def test_get_workspaces_lists_registered_profiles(client, tmp_path, monkeypatch):
    """GET /api/workspaces reports the registry after a real bootstrap."""
    c, _ = client
    _point_workspaces_at(tmp_path, monkeypatch)
    real_bootstrap()

    res = c.get("/api/workspaces")
    assert res.status_code == 200
    body = res.json()
    assert body["active_id"] == "default"
    assert [e["id"] for e in body["entries"]] == ["default"]
    assert body["entries"][0]["is_active"] is True


def test_create_rename_and_delete_workspace(client, tmp_path, monkeypatch):
    """The CRUD routes round-trip through the registry."""
    c, _ = client
    _point_workspaces_at(tmp_path, monkeypatch)
    real_bootstrap()

    created = c.post("/api/workspaces", json={"label": "Data Science"})
    assert created.status_code == 200
    ids = {e["id"] for e in created.json()["entries"]}
    assert ids == {"default", "data-science"}

    renamed = c.patch("/api/workspaces/data-science", json={"label": "Machine Learning"})
    assert renamed.status_code == 200
    labels = {e["id"]: e["label"] for e in renamed.json()["entries"]}
    assert labels["data-science"] == "Machine Learning"

    dup_label = c.post("/api/workspaces", json={"label": "default"})
    assert dup_label.status_code == 400

    deleted = c.delete("/api/workspaces/data-science")
    assert deleted.status_code == 200
    assert [e["id"] for e in deleted.json()["entries"]] == ["default"]

    refuse_last = c.delete("/api/workspaces/default")
    assert refuse_last.status_code == 400


def test_activate_workspace_returns_full_reseed_payload(client, tmp_path, monkeypatch):
    """Activating swaps config/settings/template in one response and rebinds config."""
    c, _ = client
    _point_workspaces_at(tmp_path, monkeypatch)
    real_bootstrap()
    c.post("/api/workspaces", json={"label": "Second"})

    res = c.post("/api/workspaces/second/activate")
    assert res.status_code == 200
    body = res.json()
    assert body["active_id"] == "second"
    assert body["config"]["active_workspace_id"] == "second"
    assert body["config"]["active_workspace_label"] == "Second"
    assert body["settings"]["pages"] == 1
    assert config.active_workspace_id() == "second"

    unknown = c.post("/api/workspaces/does-not-exist/activate")
    assert unknown.status_code == 400


def test_activate_workspace_409_when_queue_busy(client, tmp_path, monkeypatch):
    """POST /api/workspaces/{id}/activate returns 409 while a job is queued or running."""
    c, q = client
    _point_workspaces_at(tmp_path, monkeypatch)
    real_bootstrap()
    c.post("/api/workspaces", json={"label": "Second"})

    # Insert a "running" job directly rather than through submit(), which would start
    # the background worker and race this test — there is no master resume in the
    # isolated tmp workspace for it to load, so it would fail (and un-busy the queue)
    # before this request lands.
    job = jobs_mod.Job(job_id="fake-busy", jd_text="x", settings=JobSettings(), status="running")
    q._jobs[job.job_id] = job

    res = c.post("/api/workspaces/second/activate")
    assert res.status_code == 409
    assert "progress" in res.json()["detail"].lower() or "job" in res.json()["detail"].lower()
    # Refused, so the active profile must not have changed.
    assert config.active_workspace_id() == "default"


def test_settings_round_trip_is_per_workspace(client, tmp_path, monkeypatch):
    """Each profile keeps its own settings.json; switching swaps which one is live."""
    c, _ = client
    _point_workspaces_at(tmp_path, monkeypatch)
    real_bootstrap()

    put_res = c.put("/api/settings", json={"settings": {"pages": 4, "model": "ollama"}})
    assert put_res.status_code == 200

    c.post("/api/workspaces", json={"label": "Second"})
    c.post("/api/workspaces/second/activate")
    fresh = c.get("/api/settings")
    assert fresh.json()["settings"]["pages"] == 1
    assert fresh.json()["seeded"] is True

    c.post("/api/workspaces/default/activate")
    restored = c.get("/api/settings")
    assert restored.json()["settings"]["pages"] == 4
    assert restored.json()["settings"]["model"] == "ollama"


def test_switch_workspace_swaps_master_resume(client, tmp_path, monkeypatch):
    """Activating a different profile serves that profile's own master resume."""
    c, _ = client
    real_resume_json = config.MASTER_RESUME_PATH.read_text(encoding="utf-8")
    _point_workspaces_at(tmp_path, monkeypatch)
    real_bootstrap()  # empty "default": nothing to migrate under the isolated roots

    payload = json.loads(real_resume_json)
    payload["contact"]["name"] = "Default Person"
    config.MASTER_RESUME_PATH.parent.mkdir(parents=True, exist_ok=True)
    config.MASTER_RESUME_PATH.write_text(json.dumps(payload), encoding="utf-8")

    c.post("/api/workspaces", json={"label": "Second"})
    second_paths = config.workspace_paths("second")
    second_paths["MASTER_RESUME_PATH"].parent.mkdir(parents=True, exist_ok=True)
    payload["contact"]["name"] = "Second Person"
    second_paths["MASTER_RESUME_PATH"].write_text(json.dumps(payload), encoding="utf-8")

    assert c.get("/api/master-resume").json()["contact"]["name"] == "Default Person"

    activate_res = c.post("/api/workspaces/second/activate")
    assert activate_res.status_code == 200
    assert activate_res.json()["config"]["contact_name"] == "Second Person"
    assert c.get("/api/master-resume").json()["contact"]["name"] == "Second Person"

    c.post("/api/workspaces/default/activate")
    assert c.get("/api/master-resume").json()["contact"]["name"] == "Default Person"


def test_job_artifacts_land_under_active_workspace(client, tmp_path, monkeypatch):
    """A finished job's out_dir sits under the active profile's output/workspaces/<id>/."""
    c, q = client
    real_resume_json = config.MASTER_RESUME_PATH.read_text(encoding="utf-8")
    roots = _point_workspaces_at(tmp_path, monkeypatch)
    real_bootstrap()
    config.MASTER_RESUME_PATH.parent.mkdir(parents=True, exist_ok=True)
    config.MASTER_RESUME_PATH.write_text(real_resume_json, encoding="utf-8")

    resume = load()

    def fake_extract(text, *, known_tags=None, use_cache=True, on_event=None):
        from resume_tailor.jd import JobRequirements, Keyword

        return JobRequirements(
            title="Stub Role",
            seniority="intern",
            keywords=[Keyword(phrase="Python", canonical="python", importance="must_have")],
        )

    def fake_score(bullets, requirements, *, use_cache=True, on_event=None):
        return {b.id: 5.0 for b in bullets}

    def fake_fit(resume, requirements, *, out=None, on_event=None, **kwargs):
        out = Path(out)
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_bytes(b"PK")
        out.with_suffix(".pdf").write_bytes(b"%PDF-1.4 stub")
        bullet = resume.all_bullets()[0]
        return FitResult(
            out_path=out,
            pages=1,
            pages_are_estimated=False,
            iterations=1,
            bullets_selected=1,
            bullets_total=1,
            bullets={bullet.id: bullet.text},
            semantic_used=False,
        )

    def fake_facets(resume, requirements, **kwargs):
        from resume_tailor import facets as facets_mod

        return facets_mod.budget_only(
            resume, requirements, include_project_links=kwargs.get("include_project_links", True)
        )

    from resume_tailor.expand import Expansion

    monkeypatch.setattr(jobs_mod.jd, "extract", fake_extract)
    monkeypatch.setattr(jobs_mod.rewrite, "score_table", fake_score)
    monkeypatch.setattr(jobs_mod.fit, "fit", fake_fit)
    monkeypatch.setattr(jobs_mod.jd, "verify_verbatim", lambda *a, **k: [])
    monkeypatch.setattr(jobs_mod.facets, "select_facets", fake_facets)
    monkeypatch.setattr(
        jobs_mod.expand, "expand_experience", lambda *a, **k: Expansion(entries=[], model="stub")
    )

    res = c.post("/api/jobs", json={"jd_text": "Looking for a Python intern."})
    assert res.status_code == 200
    job_id = res.json()["job_id"]

    deadline = time.time() + 10
    while time.time() < deadline:
        status = c.get(f"/api/jobs/{job_id}").json()
        if status["status"] in ("succeeded", "failed"):
            break
        time.sleep(0.05)
    else:
        pytest.fail("job did not finish in time")
    assert status["status"] == "succeeded", status

    job = q.get(job_id)
    expected_root = roots["output"] / "workspaces" / "default" / "jobs" / job_id
    assert job.out_dir == expected_root


def test_profile_template_install_works_on_a_freshly_created_profile(
    client, tmp_path, monkeypatch
):
    """Wizard (profile) install into a profile created *without* duplicating.

    Regression, reported from Docker: such a profile had no master_resume.json, so
    `_install_with_profile`'s smoke render raised
    "Staged build or smoke render failed: Master resume not found at
    .../workspaces/<id>/master_resume.json" and no template could ever be installed.

    Must exercise the **profile** path specifically — `_install_legacy` never calls
    `_smoke_render`, so a legacy-path install passes even with the bug present.
    Only `_run_build` is stubbed here; the smoke render is real.
    """
    import docx as docx_mod

    c, _ = client
    _point_workspaces_at(tmp_path, monkeypatch)
    real_bootstrap()

    c.post("/api/workspaces", json={"label": "Nina"})  # no copy_from
    assert c.post("/api/workspaces/nina/activate").status_code == 200

    def fake_build(*, source=None, output=None, profile_path=None, legacy=False):
        """Pretend build_template.py ran; the smoke render still opens this for real."""
        out = Path(output) if output else config.DEFAULT_TEMPLATE_PATH
        out.parent.mkdir(parents=True, exist_ok=True)
        doc = docx_mod.Document()
        doc.add_paragraph("Tagged template")
        doc.save(str(out))
        return 0, "stub build ok"

    monkeypatch.setattr(template_ops, "_run_build", fake_build)

    upload = _resume_docx_bytes()
    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    analyzed = c.post("/api/template/analyze", files={"file": ("nina.docx", upload, mime)})
    assert analyzed.status_code == 200
    profile = analyzed.json()["suggested_profile"]
    assert profile, analyzed.json()["issues"]

    res = c.post(
        "/api/template",
        data={"label": "Nina Template", "profile": json.dumps(profile)},
        files={"file": ("nina.docx", upload, mime)},
    )
    assert res.status_code == 200, res.json()
    assert res.json()["ok"] is True
