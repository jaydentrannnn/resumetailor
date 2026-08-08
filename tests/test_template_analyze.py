"""Unit tests for deterministic DOCX template analysis (no Word / no network)."""

from __future__ import annotations

import io

import docx
import pytest
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from resume_tailor import template_analyze
from tests.fixtures import _sidebar_table_resume, _table_resume


def _add_bullet_numbering(document) -> str:
    """Install a minimal bullet abstract + num instance; return the numId string."""
    # python-docx creates numbering part lazily; force it via a style-based list then
    # fall back to a hand-rolled abstract if needed.
    numbering = document.part.numbering_part
    root = numbering.element

    abstract_id = "10"
    # Skip if we already added one in this document.
    for anum in root.findall(qn("w:abstractNum")):
        if anum.get(qn("w:abstractNumId")) == abstract_id:
            break
    else:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), abstract_id)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "bullet")
        lvl.append(num_fmt)
        abstract.append(lvl)
        # abstractNum elements must precede w:num
        first_num = root.find(qn("w:num"))
        if first_num is not None:
            first_num.addprevious(abstract)
        else:
            root.append(abstract)

    num_id = "20"
    for num in root.findall(qn("w:num")):
        if num.get(qn("w:numId")) == num_id:
            return num_id
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), num_id)
    abs_el = OxmlElement("w:abstractNumId")
    abs_el.set(qn("w:val"), abstract_id)
    num.append(abs_el)
    root.append(num)
    return num_id


def _make_bullet(document, text: str, num_id: str):
    """Append a list paragraph with the given numId."""
    paragraph = document.add_paragraph(text)
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    nid = OxmlElement("w:numId")
    nid.set(qn("w:val"), num_id)
    numPr.append(ilvl)
    numPr.append(nid)
    pPr.append(numPr)
    return paragraph


def _make_cell_bullet(cell, text: str, num_id: str):
    """Append a list paragraph with the given numId inside a table cell."""
    paragraph = cell.add_paragraph(text)
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    nid = OxmlElement("w:numId")
    nid.set(qn("w:val"), num_id)
    numPr.append(ilvl)
    numPr.append(nid)
    pPr.append(numPr)
    return paragraph


def _docx_bytes(build) -> bytes:
    """Run `build(document)` and return the saved .docx bytes."""
    document = docx.Document()
    build(document)
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _standard_resume(document) -> None:
    """Minimal single-column resume matching the legacy heading contract."""
    num_id = _add_bullet_numbering(document)
    document.add_paragraph("Ada Lovelace")
    document.add_paragraph("London • ada@example.com • LinkedIn")
    document.add_paragraph("EDUCATION")
    document.add_paragraph("University of London | UK\t2018 - 2022")
    _make_bullet(document, "BSc Computer Science | GPA: 3.9", num_id)
    _make_bullet(document, "Relevant Coursework: Algorithms, Databases", num_id)
    document.add_paragraph("WORK EXPERIENCES")
    document.add_paragraph("Analytical Engines | London\t2022 - Present")
    document.add_paragraph("Software Engineer")
    _make_bullet(document, "Built numerical engines in Python.", num_id)
    document.add_paragraph("PROJECTS")
    document.add_paragraph("Note Engine | Python, FastAPI\t2024")
    _make_bullet(document, "Indexed research notes with embeddings.", num_id)
    document.add_paragraph("SKILLS")
    p = document.add_paragraph()
    run = p.add_run("Languages:")
    run.bold = True
    p.add_run(" Python, SQL")


def test_analyze_standard_resume_is_ready():
    """A well-formed single-column resume yields a suggested profile."""
    raw = _docx_bytes(_standard_resume)
    result = template_analyze.analyze_docx(raw=raw)
    assert result.ready is True
    assert result.suggested_profile is not None
    assert result.suggested_profile.enabled.experience is True
    assert result.suggested_profile.source_sha256 == template_analyze.sha256_bytes(raw)
    keys = {s.key for s in result.sections}
    assert keys == {"education", "experience", "projects", "skills"}


def _multi_section_resume(document) -> None:
    """Mirrors the real-world motivating case: three experience-shaped sections, one of
    them ("OTHER ACTIVITIES") unrecognized by any heading alias, plus a single-line
    (non-bulleted) education entry — both structural-detection edge cases."""
    num_id = _add_bullet_numbering(document)
    document.add_paragraph("Nina Dao")
    document.add_paragraph("nina@example.com")
    document.add_paragraph("EDUCATION")
    document.add_paragraph("UC Irvine\tExpected June 2027")
    document.add_paragraph("B.A. in Business Administration")  # plain, not a Word bullet
    document.add_paragraph("WORK EXPERIENCE")
    document.add_paragraph("Langmaster JSC\tAug 2025 - Present")
    document.add_paragraph("Online Tutor")
    _make_bullet(document, "Tutored students in English.", num_id)
    document.add_paragraph("LEADERSHIP EXPERIENCE")
    document.add_paragraph("In the Green at UCI\tMay 2025 - Present")
    document.add_paragraph("Co-president")
    _make_bullet(document, "Oversaw club operations.", num_id)
    document.add_paragraph("OTHER ACTIVITIES")
    document.add_paragraph("Heartbeat Bazaar\tMar 2022 - Jun 2022")
    document.add_paragraph("Organizer")
    _make_bullet(document, "Directed fundraising events.", num_id)
    document.add_paragraph("SKILLS")
    sk = document.add_paragraph()
    run = sk.add_run("Languages:")
    run.bold = True
    sk.add_run(" Python, SQL")


def test_structural_fallback_detects_unaliased_experience_heading():
    """'OTHER ACTIVITIES' matches no known alias but is still recognised, structurally,
    as its own experience-shaped section rather than being absorbed into the section
    above it."""
    result = template_analyze.analyze_docx(raw=_docx_bytes(_multi_section_resume))
    other = next((s for s in result.sections if s.heading_text == "OTHER ACTIVITIES"), None)
    assert other is not None
    assert other.key == "experience"
    assert other.entry_count == 1
    assert other.bullet_count == 1


def test_all_three_experience_shaped_sections_detected_separately():
    """WORK EXPERIENCE, LEADERSHIP EXPERIENCE, and OTHER ACTIVITIES each become their own
    candidate — none absorbed into another."""
    result = template_analyze.analyze_docx(raw=_docx_bytes(_multi_section_resume))
    experience_headings = {s.heading_text for s in result.sections if s.key == "experience"}
    assert experience_headings == {"WORK EXPERIENCE", "LEADERSHIP EXPERIENCE", "OTHER ACTIVITIES"}


def test_generic_mode_suggested_when_multiple_same_kind_sections_found():
    """Three experience-kind headings cannot all keep their own title under one
    hard-coded fixed-mode heading, so the suggestion switches to generic mode."""
    result = template_analyze.analyze_docx(raw=_docx_bytes(_multi_section_resume))
    assert result.ready is True
    profile = result.suggested_profile
    assert profile is not None
    assert profile.section_mode == "generic"
    assert profile.heading_prototype is not None
    section_titles = {s.title for s in profile.sections}
    assert section_titles == {
        "EDUCATION", "WORK EXPERIENCE", "LEADERSHIP EXPERIENCE", "OTHER ACTIVITIES", "SKILLS",
    }


def test_fixed_mode_unchanged_for_a_single_heading_per_kind():
    """The ordinary case — one heading per kind — still suggests fixed mode, byte-for-byte
    the same profile shape as before generic mode existed."""
    result = template_analyze.analyze_docx(raw=_docx_bytes(_standard_resume))
    profile = result.suggested_profile
    assert profile is not None
    assert profile.section_mode == "fixed"
    assert profile.sections == []
    assert profile.heading_prototype is None


def test_plain_education_line_is_a_non_blocking_warning():
    """A degree line with no Word list bullet no longer blocks the install — `retarget_
    bullet` can convert it into one at build time — it only warns."""
    result = template_analyze.analyze_docx(raw=_docx_bytes(_multi_section_resume))
    assert result.ready is True
    assert not any(i.code == "no_education_bullets" for i in result.issues)
    assert any(
        i.code == "education_bullets_not_list" and not i.blocking for i in result.issues
    )


def test_structural_fallback_does_not_misclassify_the_name_line():
    """An all-caps name (a common style) must never be read as a section heading."""
    def build(document):
        num_id = _add_bullet_numbering(document)
        document.add_paragraph("JANE APPLESEED")
        document.add_paragraph("jane@example.com")
        document.add_paragraph("EXPERIENCE")
        document.add_paragraph("Acme | Remote\t2020 - Present")
        document.add_paragraph("Engineer")
        _make_bullet(document, "Shipped features.", num_id)

    result = template_analyze.analyze_docx(raw=_docx_bytes(build))
    assert {s.heading_text for s in result.sections} == {"EXPERIENCE"}
    assert result.suggested_profile is not None
    assert result.suggested_profile.name_paragraph_id == 0


def _spacer_multi_section_resume(document) -> None:
    """Mirrors the real-world motivating document: a blank paragraph before every
    heading, one right after every heading, and one between entries within a body that
    has more than one entry — exactly the pattern `_detect_spacing` looks for. Unlike
    `_multi_section_resume`/`_standard_resume` (both imported by
    `tests/test_template_build.py` and asserted spacer-free elsewhere), this fixture
    exists solely to exercise spacer detection."""
    num_id = _add_bullet_numbering(document)
    document.add_paragraph("Nina Dao")
    document.add_paragraph("nina@example.com")
    document.add_paragraph("")
    document.add_paragraph("EDUCATION")
    document.add_paragraph("")
    document.add_paragraph("UC Irvine\tExpected June 2027")
    _make_bullet(document, "B.A. in Business Administration", num_id)
    document.add_paragraph("")
    document.add_paragraph("WORK EXPERIENCE")
    document.add_paragraph("")
    document.add_paragraph("Langmaster JSC\tAug 2025 - Present")
    document.add_paragraph("Online Tutor")
    _make_bullet(document, "Tutored students in English.", num_id)
    document.add_paragraph("")
    document.add_paragraph("Garin JSC\tNov 2023 - May 2024")
    document.add_paragraph("Logistics Intern")
    _make_bullet(document, "Managed international logistics.", num_id)
    document.add_paragraph("")
    document.add_paragraph("LEADERSHIP EXPERIENCE")
    document.add_paragraph("")
    document.add_paragraph("In the Green at UCI\tMay 2025 - Present")
    document.add_paragraph("Co-president")
    _make_bullet(document, "Oversaw club operations.", num_id)
    document.add_paragraph("")
    document.add_paragraph("SKILLS")
    document.add_paragraph("")
    _make_bullet(document, "Languages: Python, SQL", num_id)


def test_spacer_donors_detected_on_a_blank_separated_document():
    """A majority of detected sections show a blank before, after, and (where more than
    one entry exists) between entries — all three runs get recorded, each one paragraph
    long since this fixture uses a single blank at every slot."""
    result = template_analyze.analyze_docx(raw=_docx_bytes(_spacer_multi_section_resume))
    profile = result.suggested_profile
    assert profile is not None
    assert profile.section_mode == "generic"
    assert len(profile.spacing.before_heading) == 1
    assert len(profile.spacing.after_heading) == 1
    assert len(profile.spacing.between_entries) == 1

    paras = {p.id: p.text for p in result.paragraphs}
    for run in (
        profile.spacing.before_heading,
        profile.spacing.after_heading,
        profile.spacing.between_entries,
    ):
        for pid in run:
            assert paras[pid].strip() == ""


def _rule_separated_resume(document) -> None:
    """Mirrors the real installed baseline: a horizontal rule *and* a blank under every
    heading, and two consecutive blanks at one entry boundary. Both shapes defeated
    single-paragraph spacer detection — the rule is chrome but not blank, and with two
    adjacent blanks neither one's neighbour is the bullet/header pair a boundary test
    looks for."""
    num_id = _add_bullet_numbering(document)
    rule = "_" * 60
    document.add_paragraph("Nina Dao")
    document.add_paragraph("nina@example.com")
    document.add_paragraph("")
    document.add_paragraph("EDUCATION")
    document.add_paragraph(rule)
    document.add_paragraph("")
    document.add_paragraph("UC Irvine\tExpected June 2028")
    _make_bullet(document, "B.A. in Business Administration", num_id)
    document.add_paragraph("")
    document.add_paragraph("EXPERIENCE")
    document.add_paragraph(rule)
    document.add_paragraph("")
    document.add_paragraph("Yellow Daisy\tAug 2022 - Jun 2024")
    document.add_paragraph("Co-founder")
    _make_bullet(document, "Raised $1,200 for surgery.", num_id)
    document.add_paragraph("")
    document.add_paragraph("")  # two blanks at this boundary, not one
    document.add_paragraph("Youth Opportunity\tSep 2021 - Oct 2021")
    document.add_paragraph("Ambassador")
    _make_bullet(document, "Designed social posts.", num_id)
    document.add_paragraph("")
    document.add_paragraph("INTERNSHIPS & PROGRAMS")
    document.add_paragraph(rule)
    document.add_paragraph("")
    document.add_paragraph("Deloitte\tFeb 2025 - May 2025")
    document.add_paragraph("Mentee")
    _make_bullet(document, "Exposure to financial analysis.", num_id)
    document.add_paragraph("")
    document.add_paragraph("Garin JSC\tNov 2023 - May 2024")
    document.add_paragraph("Audit Assistant")
    _make_bullet(document, "Tracked logistics.", num_id)
    document.add_paragraph("")
    document.add_paragraph("SKILLS")
    document.add_paragraph(rule)
    document.add_paragraph("")
    _make_bullet(document, "Technical: Canva, Excel", num_id)


def test_after_heading_run_captures_a_rule_plus_blank():
    """The paragraph at `body_start` is the horizontal rule, which is chrome but not
    blank. Detecting only a single *blank* found nothing here, and the rule was then
    deleted with the rest of the body — the section heading lost its underline."""
    result = template_analyze.analyze_docx(raw=_docx_bytes(_rule_separated_resume))
    profile = result.suggested_profile
    assert profile is not None
    assert profile.section_mode == "generic"

    paras = {p.id: p.text for p in result.paragraphs}
    run = profile.spacing.after_heading
    assert len(run) == 2
    assert set(paras[run[0]].strip()) == {"_"}  # the rule
    assert paras[run[1]].strip() == ""  # then the blank


def test_between_entries_run_survives_two_consecutive_blanks():
    """One boundary in this fixture uses two blanks and two use one. The modal length
    wins, so the document-wide gap stays one paragraph rather than being inflated by the
    single outlier — and, critically, is detected at all."""
    result = template_analyze.analyze_docx(raw=_docx_bytes(_rule_separated_resume))
    profile = result.suggested_profile
    assert profile is not None

    paras = {p.id: p.text for p in result.paragraphs}
    run = profile.spacing.between_entries
    assert len(run) == 1
    assert paras[run[0]].strip() == ""


def test_spacer_donors_absent_on_the_existing_multi_section_fixture():
    """`_multi_section_resume` has no blank paragraphs at all — every run stays empty,
    pinning that an ordinary (non-blank-separated) upload is unaffected by this feature."""
    result = template_analyze.analyze_docx(raw=_docx_bytes(_multi_section_resume))
    profile = result.suggested_profile
    assert profile is not None
    assert profile.section_mode == "generic"
    assert profile.spacing.before_heading == []
    assert profile.spacing.after_heading == []
    assert profile.spacing.between_entries == []


def test_spacer_donors_absent_under_fixed_mode():
    """`_standard_resume` has one heading per kind, so it stays on fixed mode — spacing
    detection never runs there (`build_generic` is never used to consume it)."""
    result = template_analyze.analyze_docx(raw=_docx_bytes(_standard_resume))
    profile = result.suggested_profile
    assert profile is not None
    assert profile.section_mode == "fixed"
    assert profile.spacing.before_heading == []
    assert profile.spacing.after_heading == []
    assert profile.spacing.between_entries == []


def test_spacing_profile_accepts_the_legacy_single_id_shape():
    """A `template_profile.json` written by the first cut of this feature stored one int
    (or null) per slot. It must still load rather than 500 the whole app."""
    from resume_tailor.template_profile import SpacingProfile

    legacy = SpacingProfile.model_validate(
        {"before_heading": 2, "after_heading": None, "between_entries": 13}
    )
    assert legacy.before_heading == [2]
    assert legacy.after_heading == []
    assert legacy.between_entries == [13]


def test_detected_section_uses_singular_project_kind():
    """`DetectedSection.kind` uses `GenericSectionKind`'s singular `"project"`, not
    `SectionCandidate.key`'s legacy plural `"projects"` — a resume with two
    projects-shaped headings must not crash `analyze_docx` with a pydantic
    `ValidationError` the first time this path is exercised."""
    def build(document):
        num_id = _add_bullet_numbering(document)
        document.add_paragraph("Name")
        document.add_paragraph("email@example.com")
        document.add_paragraph("EXPERIENCE")
        document.add_paragraph("Acme | Remote\t2020 - Present")
        document.add_paragraph("Engineer")
        _make_bullet(document, "Shipped features.", num_id)
        document.add_paragraph("SELECTED PROJECTS")
        document.add_paragraph("Note Engine\t2024")
        _make_bullet(document, "Indexed notes.", num_id)
        document.add_paragraph("PERSONAL PROJECTS")
        document.add_paragraph("Side App\t2023")
        _make_bullet(document, "Shipped solo.", num_id)

    result = template_analyze.analyze_docx(raw=_docx_bytes(build))
    assert result.ready is True
    profile = result.suggested_profile
    assert profile is not None
    assert profile.section_mode == "generic"
    project_sections = [s for s in profile.sections if s.kind == "project"]
    assert len(project_sections) == 2


def test_analyze_renamed_experience_heading():
    """'PROFESSIONAL EXPERIENCE' maps to the experience section."""
    def build(document):
        num_id = _add_bullet_numbering(document)
        document.add_paragraph("Name")
        document.add_paragraph("email@example.com")
        document.add_paragraph("PROFESSIONAL EXPERIENCE")
        document.add_paragraph("Acme | Remote\t2020 - Present")
        document.add_paragraph("Engineer")
        _make_bullet(document, "Shipped features.", num_id)

    result = template_analyze.analyze_docx(raw=_docx_bytes(build))
    assert any(s.key == "experience" for s in result.sections)
    assert result.suggested_profile is not None
    assert result.suggested_profile.enabled.education is False
    assert result.suggested_profile.enabled.projects is False


def test_analyze_missing_experience_is_blocking():
    """No experience heading → ready false with a blocker."""
    def build(document):
        document.add_paragraph("Name")
        document.add_paragraph("email@example.com")
        document.add_paragraph("SKILLS")
        document.add_paragraph("Languages: Python")

    result = template_analyze.analyze_docx(raw=_docx_bytes(build))
    assert result.ready is False
    assert any(i.code == "missing_experience" and i.blocking for i in result.issues)


#: Every blocking issue code `classify_table_layout` can emit — used to assert a
#: linear table isn't rejected by the table gate itself (whatever else may still be
#: incomplete about its mapping).
_TABLE_LAYOUT_BLOCKING_CODES = {
    "nested_tables",
    "multiple_tables",
    "table_vertical_merge",
    "table_parallel_columns",
    "table_sidebar_bullets",
    "table_sidebar_headings",
    "table_no_headings",
}


def test_analyze_sidebar_table_is_blocking():
    """A table with bulleted content in a second (sidebar) column is rejected — that
    is the structural signature of two parallel reading columns, not an invisible
    single-column layout grid."""
    def build(document):
        document.add_paragraph("Name")
        document.add_paragraph("email@example.com")
        num_id = _add_bullet_numbering(document)
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "WORK EXPERIENCE"
        _make_cell_bullet(table.cell(0, 1), "Python", num_id)

    result = template_analyze.analyze_docx(raw=_docx_bytes(build))
    assert any(i.code == "table_sidebar_bullets" and i.blocking for i in result.issues)
    assert result.ready is False


def test_analyze_linear_table_is_not_blocked_by_table_classification():
    """A table used only as an invisible single-column layout grid — a heading alone
    in its row, an ordinary two-cell entry-header row below it — passes the table
    classifier itself. (Whether the rest of the mapping is complete enough for
    `ready=True` is a separate, later concern.)"""
    def build(document):
        document.add_paragraph("Name")
        document.add_paragraph("email@example.com")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "WORK EXPERIENCE"
        table.cell(0, 1).text = ""
        table.cell(1, 0).text = "Acme"
        table.cell(1, 1).text = "Engineer"

    result = template_analyze.analyze_docx(raw=_docx_bytes(build))
    assert not any(i.code in _TABLE_LAYOUT_BLOCKING_CODES for i in result.issues)


def test_analyze_table_with_three_populated_cells_is_blocking():
    """A row with three independently populated cells reads as more than two parallel
    columns — never produced by a table used only to right-align dates."""
    def build(document):
        document.add_paragraph("Name")
        document.add_paragraph("email@example.com")
        table = document.add_table(rows=1, cols=3)
        table.cell(0, 0).text = "Skills"
        table.cell(0, 1).text = "Languages"
        table.cell(0, 2).text = "Interests"

    result = template_analyze.analyze_docx(raw=_docx_bytes(build))
    assert any(i.code == "table_parallel_columns" and i.blocking for i in result.issues)


def test_analyze_table_with_no_headings_is_blocking():
    """A table with no recognizable section heading anywhere reads as a data table,
    not a resume layout grid."""
    def build(document):
        document.add_paragraph("Name")
        document.add_paragraph("email@example.com")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Acme"
        table.cell(0, 1).text = "Engineer"

    result = template_analyze.analyze_docx(raw=_docx_bytes(build))
    assert any(i.code == "table_no_headings" and i.blocking for i in result.issues)


def test_analyze_table_vertical_merge_is_blocking():
    """A vertically merged cell is how a sidebar column spans several rows — the
    defining structural feature of a layout the table-classifier does not support."""
    def build(document):
        document.add_paragraph("Name")
        document.add_paragraph("email@example.com")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "EXPERIENCE"
        table.cell(1, 0).text = "Acme"
        table.cell(0, 1).merge(table.cell(1, 1))
        table.cell(0, 1).text = "Sidebar spanning two rows"

    result = template_analyze.analyze_docx(raw=_docx_bytes(build))
    assert any(i.code == "table_vertical_merge" and i.blocking for i in result.issues)


def test_analyze_multiple_tables_is_blocking():
    """More than one top-level table has no well-defined single reading order."""
    def build(document):
        document.add_paragraph("Name")
        document.add_paragraph("email@example.com")
        t1 = document.add_table(rows=1, cols=1)
        t1.cell(0, 0).text = "EXPERIENCE"
        document.add_paragraph("")
        t2 = document.add_table(rows=1, cols=1)
        t2.cell(0, 0).text = "Acme"

    result = template_analyze.analyze_docx(raw=_docx_bytes(build))
    assert any(i.code == "multiple_tables" and i.blocking for i in result.issues)


def test_validate_profile_hash_mismatch():
    """Installing with a profile from a different file is rejected."""
    raw_a = _docx_bytes(_standard_resume)
    result = template_analyze.analyze_docx(raw=raw_a)
    assert result.suggested_profile is not None
    profile = result.suggested_profile

    raw_b = _docx_bytes(lambda d: (d.add_paragraph("Other"), d.add_paragraph("x@y.z")))
    issues = template_analyze.validate_profile_against_doc(profile, raw=raw_b)
    assert any(i.code == "hash_mismatch" and i.blocking for i in issues)


def test_contact_separator_detection():
    """Pipe-separated contact lines become ' | ' separators."""
    def build(document):
        num_id = _add_bullet_numbering(document)
        document.add_paragraph("Name")
        document.add_paragraph("City | a@b.co | LinkedIn")
        document.add_paragraph("WORK EXPERIENCES")
        document.add_paragraph("Acme | Remote\t2020 - Present")
        document.add_paragraph("Engineer")
        _make_bullet(document, "Did work.", num_id)

    result = template_analyze.analyze_docx(raw=_docx_bytes(build))
    assert result.suggested_profile is not None
    assert "|" in result.suggested_profile.contact.separator


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    """Append a real w:hyperlink run so analyze's has_hyperlink path fires."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def test_project_header_name_tech_github_no_overlap():
    """name | tech | Github\\tdate must not fold Github into the tech span."""
    def build(document):
        num_id = _add_bullet_numbering(document)
        document.add_paragraph("Name")
        document.add_paragraph("email@example.com")
        document.add_paragraph("WORK EXPERIENCES")
        document.add_paragraph("Acme | Remote\t2020 - Present")
        document.add_paragraph("Engineer")
        _make_bullet(document, "Shipped features.", num_id)
        document.add_paragraph("PROJECTS")
        # Mirror the live resume: name | tech | Github\\tdates with a real hyperlink.
        p = document.add_paragraph(
            "Text-to-SQL Reward Fine-Tuning via RL | "
            "GRPO, Deeplearn, HuggingFace, SQL | "
        )
        _add_hyperlink(p, "Github", "https://github.com/example/repo")
        # Tab + dates as a following run (python-docx paragraph.text joins them).
        run = p.add_run()
        run._r.append(OxmlElement("w:tab"))
        p.add_run("Jan 2026 - Mar 2026")
        _make_bullet(document, "Built reward functions.", num_id)

    raw = _docx_bytes(build)
    result = template_analyze.analyze_docx(raw=raw)
    assert result.ready is True
    assert result.suggested_profile is not None
    projects = result.suggested_profile.projects
    assert projects is not None
    tech = projects.header.fields.get("tech")
    assert tech is not None and tech.present and tech.span is not None
    header_text = next(
        p.text for p in result.paragraphs if p.id == projects.header.header_paragraph_id
    )
    tech_preview = header_text[tech.span.start : tech.span.end]
    assert "Github" not in tech_preview
    assert "GRPO" in tech_preview
    assert projects.link.present is True
    assert projects.link.span is not None
    link_preview = header_text[projects.link.span.start : projects.link.span.end]
    assert link_preview == "Github"
    assert tech.span.end <= projects.link.span.start


def test_project_link_label_comes_from_hyperlink_text():
    """The link span is derived from the hyperlink's own text, not a fixed word list."""
    def build(document):
        num_id = _add_bullet_numbering(document)
        document.add_paragraph("Name")
        document.add_paragraph("email@example.com")
        document.add_paragraph("WORK EXPERIENCES")
        document.add_paragraph("Acme | Remote\t2020 - Present")
        document.add_paragraph("Engineer")
        _make_bullet(document, "Shipped features.", num_id)
        document.add_paragraph("PROJECTS")
        p = document.add_paragraph("Note Engine | Python, FastAPI | ")
        _add_hyperlink(p, "Source", "https://example.com/source")
        run = p.add_run()
        run._r.append(OxmlElement("w:tab"))
        p.add_run("2024")
        _make_bullet(document, "Indexed notes.", num_id)

    result = template_analyze.analyze_docx(raw=_docx_bytes(build))
    assert result.suggested_profile is not None
    projects = result.suggested_profile.projects
    assert projects is not None
    assert projects.link.present is True
    assert projects.link.span is not None
    header_text = next(
        p.text for p in result.paragraphs if p.id == projects.header.header_paragraph_id
    )
    assert header_text[projects.link.span.start : projects.link.span.end] == "Source"


def test_project_header_without_tech_does_not_duplicate_link_span():
    """'Name | Github\\tdate' maps only the link, not an overlapping tech span."""
    def build(document):
        num_id = _add_bullet_numbering(document)
        document.add_paragraph("Name")
        document.add_paragraph("email@example.com")
        document.add_paragraph("WORK EXPERIENCES")
        document.add_paragraph("Acme | Remote\t2020 - Present")
        document.add_paragraph("Engineer")
        _make_bullet(document, "Shipped features.", num_id)
        document.add_paragraph("PROJECTS")
        p = document.add_paragraph("Solo Proj | ")
        _add_hyperlink(p, "Github", "https://github.com/example/solo")
        run = p.add_run()
        run._r.append(OxmlElement("w:tab"))
        p.add_run("Jan 2026")
        _make_bullet(document, "Shipped solo.", num_id)

    result = template_analyze.analyze_docx(raw=_docx_bytes(build))
    assert result.suggested_profile is not None
    projects = result.suggested_profile.projects
    assert projects is not None
    tech = projects.header.fields.get("tech")
    assert tech is None or not tech.present
    assert projects.link.present is True
    assert projects.link.span is not None
    header_text = next(
        p.text for p in result.paragraphs if p.id == projects.header.header_paragraph_id
    )
    assert header_text[projects.link.span.start : projects.link.span.end] == "Github"


def test_validate_rejects_tab_inside_span():
    """A mapped span that straddles a real tab is rejected as blocking."""
    raw = _docx_bytes(_standard_resume)
    result = template_analyze.analyze_docx(raw=raw)
    assert result.suggested_profile is not None
    profile = result.suggested_profile
    exp = profile.experience
    dates_span = exp.header.fields["dates"].span
    assert dates_span is not None
    bad_span = dates_span.model_copy(update={"start": dates_span.start - 1})
    bad_fields = dict(exp.header.fields)
    bad_fields["dates"] = bad_fields["dates"].model_copy(update={"span": bad_span})
    bad_profile = profile.model_copy(
        update={
            "experience": exp.model_copy(
                update={"header": exp.header.model_copy(update={"fields": bad_fields})}
            )
        }
    )
    issues = template_analyze.validate_profile_against_doc(bad_profile, raw=raw)
    assert any(i.code == "span_has_tab" and i.blocking for i in issues)


def test_validate_rejects_overlapping_spans():
    """Two mapped fields claiming the same text on one paragraph are rejected."""
    raw = _docx_bytes(_standard_resume)
    result = template_analyze.analyze_docx(raw=raw)
    assert result.suggested_profile is not None
    profile = result.suggested_profile
    exp = profile.experience
    company_span = exp.header.fields["company"].span
    assert company_span is not None
    bad_fields = dict(exp.header.fields)
    bad_fields["location"] = bad_fields["location"].model_copy(
        update={"span": company_span}
    )
    bad_profile = profile.model_copy(
        update={
            "experience": exp.model_copy(
                update={"header": exp.header.model_copy(update={"fields": bad_fields})}
            )
        }
    )
    issues = template_analyze.validate_profile_against_doc(bad_profile, raw=raw)
    assert any(i.code == "overlapping_spans" and i.blocking for i in issues)


# --- Phase 4: structure-first heading detection -----------------------------------


def _bold_heading(document, text: str):
    """Append a bold-run heading paragraph, so real headings form their own formatting
    class distinct from plain body/entry text — the shape most real resume exports use
    and the one `_heading_classes` is designed to corroborate against."""
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p


def test_summary_section_is_excluded_when_a_real_heading_follows():
    """'PROFESSIONAL SUMMARY' followed only by a paragraph of prose — never a bullet or
    a tab-aligned entry header — introduces nothing and must not become a section, even
    though it is short, unaliased, and all-caps enough to pass `_looks_like_heading`."""
    def build(document):
        num_id = _add_bullet_numbering(document)
        document.add_paragraph("Jordan Ray")
        document.add_paragraph("jordan@example.com")
        _bold_heading(document, "PROFESSIONAL SUMMARY")
        document.add_paragraph(
            "Detail-oriented engineer with five years of experience shipping products."
        )
        _bold_heading(document, "EDUCATION")
        document.add_paragraph("State University\t2016 - 2020")
        _make_bullet(document, "B.S. Computer Science", num_id)
        _bold_heading(document, "WORK EXPERIENCE")
        document.add_paragraph("Acme Corp | Remote\t2020 - Present")
        document.add_paragraph("Engineer")
        _make_bullet(document, "Built things.", num_id)

    result = template_analyze.analyze_docx(raw=_docx_bytes(build))
    keys = {s.key for s in result.sections}
    assert keys == {"education", "experience"}
    assert "PROFESSIONAL SUMMARY" not in {s.heading_text for s in result.sections}
    assert result.ready is True


def test_summary_only_resume_is_blocking():
    """A resume with nothing but a professional summary — no real section heading at
    all — must fail loudly (`missing_experience`), not report `ready: true` with the
    summary mis-mapped as an experience section."""
    def build(document):
        document.add_paragraph("Jordan Ray")
        document.add_paragraph("jordan@example.com")
        _bold_heading(document, "PROFESSIONAL SUMMARY")
        document.add_paragraph(
            "Detail-oriented engineer with five years of experience shipping products."
        )

    result = template_analyze.analyze_docx(raw=_docx_bytes(build))
    assert result.ready is False
    assert any(i.code == "missing_experience" and i.blocking for i in result.issues)
    assert "PROFESSIONAL SUMMARY" not in {s.heading_text for s in result.sections}


def test_all_caps_entry_line_not_corroborated_as_heading():
    """An all-caps company name with no distinguishing formatting from body text is
    structurally plausible on text alone (`_looks_like_heading` returns True for it in
    isolation) but must not become a spurious heading when the document's real headings
    are styled distinctly (bold here) and this line is not."""
    def build(document):
        num_id = _add_bullet_numbering(document)
        document.add_paragraph("Jamie Lee")
        document.add_paragraph("jamie@example.com")
        _bold_heading(document, "EDUCATION")
        document.add_paragraph("State University\t2016 - 2020")
        _make_bullet(document, "B.S. Computer Science", num_id)
        _bold_heading(document, "WORK EXPERIENCE")
        document.add_paragraph("AMAZON WEB SERVICES")
        document.add_paragraph("Software Engineer Intern\tJune 2022 - Present")
        _make_bullet(document, "Shipped a feature.", num_id)

    result = template_analyze.analyze_docx(raw=_docx_bytes(build))
    keys = {s.key for s in result.sections}
    assert keys == {"education", "experience"}
    assert "AMAZON WEB SERVICES" not in {s.heading_text for s in result.sections}


def test_job_title_containing_a_section_keyword_not_misdetected_as_heading():
    """A job title like 'Experience Designer' sitting right under its own entry's
    company/dates header line must never be read as a new 'experience' section, even
    though the word 'experience' gives it a nonzero-confidence text match — the
    <=0.6-confidence tiers have no case requirement at all. Regression test for the
    false positive found while validating the structural-fallback corroboration work:
    fingerprint corroboration alone can't screen this out, since a plain title line's
    formatting can coincidentally land in the same class as the real headings.
    """
    def build(document):
        num_id = _add_bullet_numbering(document)
        document.add_paragraph("Alex Kim")
        document.add_paragraph("alex@example.com")
        _bold_heading(document, "EDUCATION")
        document.add_paragraph("State University\t2016 - 2020")
        _make_bullet(document, "B.S. Computer Science", num_id)
        _bold_heading(document, "WORK EXPERIENCE")
        document.add_paragraph("Acme Corp | Remote\tJan 2023 - Present")
        document.add_paragraph("Experience Designer")
        _make_bullet(document, "Redesigned the onboarding flow.", num_id)

    result = template_analyze.analyze_docx(raw=_docx_bytes(build))
    keys = {s.key for s in result.sections}
    assert keys == {"education", "experience"}
    assert "Experience Designer" not in {s.heading_text for s in result.sections}


def test_entry_header_with_keyword_and_date_not_misdetected_as_heading():
    """A later entry's own tab-aligned 'Name\\tDate' header line must never be read as
    a new section just because its text happens to contain another section's keyword
    ('Education' inside 'Advocate of Sexual Education in School'). Regression test for
    the real-world case found in the `nina` workspace's uploaded resume: a second,
    bulletless-header entry inside an unaliased 'OTHER ACTIVITIES' section, whose own
    header line carries a trailing tab-aligned date exactly like a real section heading
    never does.
    """
    def build(document):
        num_id = _add_bullet_numbering(document)
        document.add_paragraph("Nina Dao")
        document.add_paragraph("nina@example.com")
        _bold_heading(document, "EDUCATION")
        document.add_paragraph("UC Irvine\t2027")
        _make_bullet(document, "B.A. in Business Administration", num_id)
        _bold_heading(document, "OTHER ACTIVITIES")
        document.add_paragraph("Heartbeat Bazaar\tMar 2022 - Jun 2022")
        document.add_paragraph("Organizer")
        _make_bullet(document, "Directed fundraising events.", num_id)
        document.add_paragraph("Advocate of Sexual Education in School\t2022")
        _make_bullet(document, "Issued a petition.", num_id)

    result = template_analyze.analyze_docx(raw=_docx_bytes(build))
    keys = {s.key for s in result.sections}
    assert keys == {"education", "experience"}
    assert "Advocate of Sexual Education in School\t2022" not in {
        s.heading_text for s in result.sections
    }


def test_experience_section_with_no_dates_is_blocking():
    """Every experience entry header lacking a detectable date must block install
    (`experience_dates_not_detected`), not silently drop dates from every rendered job
    with `ready: true`."""
    def build(document):
        num_id = _add_bullet_numbering(document)
        document.add_paragraph("Sam Rivers")
        document.add_paragraph("sam@example.com")
        _bold_heading(document, "EDUCATION")
        document.add_paragraph("State University\t2016 - 2020")
        _make_bullet(document, "B.S. Computer Science", num_id)
        _bold_heading(document, "WORK EXPERIENCE")
        document.add_paragraph("Acme Corp | Remote")
        document.add_paragraph("Engineer")
        _make_bullet(document, "Built things.", num_id)

    result = template_analyze.analyze_docx(raw=_docx_bytes(build))
    assert result.ready is False
    assert any(
        i.code == "experience_dates_not_detected" and i.blocking for i in result.issues
    )


def test_experience_partial_dates_is_a_non_blocking_warning():
    """A minority of experience entries missing a date is a warning, not a blocker —
    only a majority-missing date is a `experience_dates_not_detected` blocker."""
    def build(document):
        num_id = _add_bullet_numbering(document)
        document.add_paragraph("Sam Rivers")
        document.add_paragraph("sam@example.com")
        _bold_heading(document, "EDUCATION")
        document.add_paragraph("State University\t2016 - 2020")
        _make_bullet(document, "B.S. Computer Science", num_id)
        _bold_heading(document, "WORK EXPERIENCE")
        document.add_paragraph("Acme Corp | Remote\t2022 - Present")
        document.add_paragraph("Engineer")
        _make_bullet(document, "Built things.", num_id)
        document.add_paragraph("Beta LLC | Remote\t2020 - 2022")
        document.add_paragraph("Analyst")
        _make_bullet(document, "Analyzed things.", num_id)
        document.add_paragraph("Gamma Inc | Remote")
        document.add_paragraph("Intern")
        _make_bullet(document, "Interned.", num_id)

    result = template_analyze.analyze_docx(raw=_docx_bytes(build))
    assert any(
        i.code == "experience_dates_partial" and not i.blocking for i in result.issues
    )
    assert not any(i.code == "experience_dates_not_detected" for i in result.issues)


def test_project_section_with_no_dates_is_blocking():
    """Same guarantee as experience, for projects: no detected date on any project
    entry blocks install instead of silently dropping every project's date."""
    def build(document):
        num_id = _add_bullet_numbering(document)
        document.add_paragraph("Sam Rivers")
        document.add_paragraph("sam@example.com")
        _bold_heading(document, "WORK EXPERIENCE")
        document.add_paragraph("Acme Corp | Remote\t2022 - Present")
        document.add_paragraph("Engineer")
        _make_bullet(document, "Built things.", num_id)
        _bold_heading(document, "PROJECTS")
        document.add_paragraph("Note Engine | Python, FastAPI")
        _make_bullet(document, "Indexed research notes with embeddings.", num_id)

    result = template_analyze.analyze_docx(raw=_docx_bytes(build))
    assert result.ready is False
    assert any(
        i.code == "project_dates_not_detected" and i.blocking for i in result.issues
    )


def test_contact_field_order_does_not_read_a_date_range_as_a_phone_number():
    """A bare year range ('2021 - 2025') is digits-space-punctuation-digits just like a
    phone number to `_PHONE_RE` alone; `_contact_field_order` must exclude anything that
    also looks like a date before trusting it as a phone field. `_contact_field_order`
    always returns the full field set (order is a rendering preference, not a presence
    filter — see its own "always allow the full set" comment), so the observable is
    which field the text's own content is attributed to first, not membership."""
    order = template_analyze._contact_field_order("2021 - 2025")
    assert order[0] == "location"  # not "phone" — a date range is not a phone number

    order_with_real_phone = template_analyze._contact_field_order("555-123-4567")
    assert order_with_real_phone[0] == "phone"


# --------------------------------------------------------------------------------------
# Table-layout resume: content lives inside one invisible layout table. See
# `tests/fixtures.py::_table_resume` for the exact shape (mirrors the real-world
# document that motivated this — gridSpan inconsistency between its two
# experience-kind sections, a cross-cell skills grid, a multi-paragraph contact block).
# --------------------------------------------------------------------------------------


def test_analyze_table_resume_is_ready():
    """The driving document's shape analyzes cleanly: `ready`, `layout == "table"`,
    `section_mode == "generic"` (two experience-kind headings forces it), and all four
    sections detected in document order with the right kind."""
    result = template_analyze.analyze_docx(raw=_docx_bytes(_table_resume))
    assert result.ready is True
    profile = result.suggested_profile
    assert profile is not None
    assert profile.layout == "table"
    assert profile.section_mode == "generic"
    assert [(s.title, s.kind) for s in profile.sections] == [
        ("EDUCATION", "education"),
        ("WORK EXPERIENCE", "experience"),
        ("LEADERSHIP", "experience"),
        ("ADDITIONAL INFORMATION", "skills"),
    ]


def test_analyze_table_resume_cross_cell_dates_independent_of_gridspan():
    """Experience dates are detected via the row's second cell regardless of whether
    that row is a 3+1 split (WORK EXPERIENCE) or a 2+2 split (LEADERSHIP) — pins that
    the cross-cell detector keys on "a second populated cell", never a fixed gridSpan."""
    result = template_analyze.analyze_docx(raw=_docx_bytes(_table_resume))
    profile = result.suggested_profile
    assert profile is not None
    header = profile.experience.header
    dates = header.fields["dates"]
    assert dates.present is True
    assert header.date_alignment == "separate_paragraph"
    assert dates.span.paragraph_id != header.header_paragraph_id


def test_analyze_table_resume_contact_slots_and_unmapped_address():
    """A multi-paragraph contact block (name row's second cell holds a street address
    AND a city/state/zip line) splits into slots — location/email/phone recognized,
    the street address left unmapped with a warning rather than silently guessed at or
    silently dropped."""
    result = template_analyze.analyze_docx(raw=_docx_bytes(_table_resume))
    profile = result.suggested_profile
    assert profile is not None
    slot_fields = [s.fields for s in profile.contact.slots]
    assert slot_fields == [["location"], ["email"], ["phone"]]
    assert any(i.code == "contact_unmapped_paragraph" and not i.blocking for i in result.issues)


def test_analyze_table_resume_skills_cross_cell_pairing():
    """ADDITIONAL INFORMATION's label column ('Languages:'/'Skills:') and value column
    live in different paragraphs — `label_span`/`body_span` must point at different
    paragraph ids, unlike the single-paragraph 'Label: item, item' shape."""
    result = template_analyze.analyze_docx(raw=_docx_bytes(_table_resume))
    profile = result.suggested_profile
    assert profile is not None
    skills = profile.skills
    assert skills is not None
    assert skills.label_span.paragraph_id != skills.body_span.paragraph_id


def test_table_resume_single_contact_paragraph_still_uses_slots_empty():
    """A table-layout document whose contact block happens to be ONE paragraph (name
    row's second cell holds just one line, not a multi-paragraph spread) must still
    produce `slots == []` and the ordinary joined-line contract — a table layout does
    not, by itself, force slot-based contact mapping.

    Exercises `_detect_name_and_contact` directly rather than through the full
    `analyze_docx` pipeline, since the point here is contact detection in isolation,
    not a fully mappable resume (this minimal fixture has no title/dates line, which
    would fail for unrelated reasons through the full pipeline)."""
    from tests.fixtures import _add_bullet_numbering, _cell_text, _shape_row

    def build(document):
        _add_bullet_numbering(document)
        table = document.add_table(rows=2, cols=4)
        name_row = _shape_row(table, 0, [4])
        _cell_text(name_row[0], "JORDAN RIVERA")
        contact_row = _shape_row(table, 1, [4])
        _cell_text(contact_row[0], "jordan@example.com • (555) 123-4567")

    raw = _docx_bytes(build)
    doc = docx.Document(io.BytesIO(raw))
    paras = template_analyze._load_paras(doc)
    name_id, contact_para, slots, unmapped = template_analyze._detect_name_and_contact(
        paras, first_heading_id=None
    )
    assert slots == []
    assert unmapped == []
    assert contact_para is not None
    assert "jordan@example.com" in contact_para.text


def test_analyze_sidebar_table_resume_is_blocking():
    """The negative-fixture sidebar table (bullets in the row's second cell) is
    rejected by the table classifier itself."""
    result = template_analyze.analyze_docx(raw=_docx_bytes(_sidebar_table_resume))
    assert result.ready is False
    assert any(i.code == "table_sidebar_bullets" and i.blocking for i in result.issues)


def _candidates_for(result, heading_paragraph_id: int) -> dict[str, template_analyze.FieldCandidate]:
    """Best (highest-confidence) candidate per field for one section, keyed by field
    name — mirrors the wizard's own `SectionFieldRows` de-duplication."""
    best: dict[str, template_analyze.FieldCandidate] = {}
    for c in result.field_candidates:
        if c.section_heading_paragraph_id != heading_paragraph_id:
            continue
        if c.field not in best or c.confidence > best[c.field].confidence:
            best[c.field] = c
    return best


def test_every_experience_section_gets_its_own_field_candidates():
    """Regression test for the reported bug: a second (or third) same-kind section must
    not read as "not detected" merely because the kind's single installed prototype
    came from a different section. `_multi_section_resume` has three experience-kind
    sections (WORK EXPERIENCE, LEADERSHIP EXPERIENCE, OTHER ACTIVITIES); every one of
    them has its own company and dates, and each must get its own candidates tagged
    with its own `heading_paragraph_id` — not just whichever section the pooled
    kind-wide prototype happened to come from."""
    result = template_analyze.analyze_docx(raw=_docx_bytes(_multi_section_resume))
    experience_sections = [s for s in result.sections if s.key == "experience"]
    assert len(experience_sections) == 3
    for sec in experience_sections:
        cands = _candidates_for(result, sec.heading_paragraph_id)
        assert "company" in cands, f"no company candidate for {sec.heading_text!r}"
        assert "dates" in cands, f"no dates candidate for {sec.heading_text!r}"
        assert cands["company"].confidence == 1.0
        assert cands["dates"].confidence == 1.0

    # Every candidate is attributed to some section — nothing falls through the old
    # range-based fallback that motivated this fix in the first place.
    assert all(c.section_heading_paragraph_id is not None for c in result.field_candidates)


def test_table_resume_second_experience_section_has_candidates():
    """The real driving document's shape: WORK EXPERIENCE (3+1 gridSpan header rows)
    and LEADERSHIP (2+2 gridSpan header rows) are both experience-kind sections: the
    second one (LEADERSHIP) must get its own company/dates/location candidates from
    its own cross-cell entries, not the pooled kind-wide prototype's."""
    result = template_analyze.analyze_docx(raw=_docx_bytes(_table_resume))
    experience_sections = {s.heading_text: s for s in result.sections if s.key == "experience"}
    assert set(experience_sections) == {"WORK EXPERIENCE", "LEADERSHIP"}

    leadership = experience_sections["LEADERSHIP"]
    cands = _candidates_for(result, leadership.heading_paragraph_id)
    assert cands["company"].preview == "Campus Club"
    assert cands["location"].preview == "Springfield, IL"
    assert cands["dates"].preview == "2022 - 2023"

    work = experience_sections["WORK EXPERIENCE"]
    work_cands = _candidates_for(result, work.heading_paragraph_id)
    assert work_cands["company"].preview == "Example Corp"
    assert work_cands["dates"].preview == "Jan 2023 - Present"


def test_single_section_field_candidates_unchanged():
    """No-regression gate for the common case: a document with exactly one section per
    kind gets the same fields, at the same confidences, as before this change — the
    per-section helper degenerates to "one section" without altering behavior."""
    result = template_analyze.analyze_docx(raw=_docx_bytes(_standard_resume))
    exp = next(s for s in result.sections if s.key == "experience")
    cands = _candidates_for(result, exp.heading_paragraph_id)
    assert cands["company"].preview == "Analytical Engines"
    assert cands["company"].confidence == 1.0
    assert cands["dates"].preview == "2022 - Present"
    assert cands["dates"].confidence == 1.0
    assert cands["title"].preview == "Software Engineer"
    assert cands["title"].confidence == 0.9

    edu = next(s for s in result.sections if s.key == "education")
    edu_cands = _candidates_for(result, edu.heading_paragraph_id)
    assert edu_cands["school"].preview == "University of London"
    assert edu_cands["dates"].preview == "2018 - 2022"


def _two_education_sections_resume(document) -> None:
    """Two education-kind headings, each with its own detectable entry — exercises
    per-section field candidates for the `education` kind specifically, since
    `_multi_section_resume` only ever has one. "ACADEMIC BACKGROUND" reaches
    education-kind via `_classify_heading`'s "academic" keyword heuristic (0.6
    confidence), not an exact alias — a second, differently-worded education-kind
    heading, the same way a real resume's second section rarely repeats the first
    section's exact title."""
    num_id = _add_bullet_numbering(document)
    document.add_paragraph("Nina Dao")
    document.add_paragraph("nina@example.com")
    document.add_paragraph("EDUCATION")
    document.add_paragraph("UC Irvine\tExpected June 2027")
    _make_bullet(document, "B.A. in Business Administration", num_id)
    document.add_paragraph("ACADEMIC BACKGROUND")
    document.add_paragraph("Coursera\tSummer 2024")
    _make_bullet(document, "Data Analytics Certificate", num_id)
    document.add_paragraph("WORK EXPERIENCE")
    document.add_paragraph("Langmaster JSC\tAug 2025 - Present")
    document.add_paragraph("Online Tutor")
    _make_bullet(document, "Tutored students in English.", num_id)


def test_education_sections_each_get_candidates():
    """Two education-kind sections (a real "EDUCATION" plus an unaliased structural
    match) must each surface their own school/dates candidates."""
    result = template_analyze.analyze_docx(raw=_docx_bytes(_two_education_sections_resume))
    education_sections = [s for s in result.sections if s.key == "education"]
    assert len(education_sections) == 2
    for sec in education_sections:
        cands = _candidates_for(result, sec.heading_paragraph_id)
        assert "school" in cands, f"no school candidate for {sec.heading_text!r}"
        assert "dates" in cands, f"no dates candidate for {sec.heading_text!r}"


def test_entry_header_fields_resolves_cross_cell_project_entry():
    """`_entry_header_fields` (the dispatcher the `:2042` bonus fix now routes through
    for the installed Projects prototype, same as experience/education already did)
    correctly reads a table-layout project entry whose name sits in the main cell and
    whose tech + date sit in the row's other cell (the same main/side shape experience
    and education already use) — independent of `_table_resume`'s own paragraph ids,
    which every other table test pins against."""
    from tests.fixtures import _cell_para, _shape_row

    document = docx.Document()
    _add_bullet_numbering(document)
    table = document.add_table(rows=1, cols=4)
    row = _shape_row(table, 0, [3, 1])
    _cell_para(row[0], 0, "Note Engine", bold=True)
    _cell_para(row[1], 0, "Python, FastAPI", italic=True, align_right=True)
    _cell_para(row[1], 1, "2024", italic=True, align_right=True)

    buf = io.BytesIO()
    document.save(buf)
    doc = docx.Document(io.BytesIO(buf.getvalue()))
    paras = template_analyze._load_paras(doc)
    entry = [p for p in paras if p.text.strip()]

    header, candidates = template_analyze._entry_header_fields(
        entry, primary="name", secondary="tech", date_field="date"
    )
    assert header.fields["name"].present is True
    assert header.fields["tech"].present is True
    assert header.fields["date"].present is True
    assert header.date_alignment == "separate_paragraph"
    by_field = {c.field: c for c in candidates}
    assert by_field["name"].preview == "Note Engine"
    assert by_field["tech"].preview == "Python, FastAPI"
    assert by_field["date"].preview == "2024"
