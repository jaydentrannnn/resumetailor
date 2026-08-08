"""Tests for profile-driven template tagging (no Word / no network)."""

from __future__ import annotations

import re
from pathlib import Path

import docx
import pytest
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docxtpl import DocxTemplate

from resume_tailor import template_analyze, template_build, template_profile
from resume_tailor.data import (
    Bullet,
    Experience,
    ExperienceSection,
    ListItem,
    ListSection,
    MasterResume,
)
from resume_tailor.template_profile import DetectedSection, HeadingPrototype
from resume_tailor.render import build_context
from tests.test_template_analyze import (
    _add_bullet_numbering,
    _add_hyperlink,
    _docx_bytes,
    _make_bullet,
    _standard_resume,
)
from tests.fixtures import _table_resume


def _run_bold(run) -> bool:
    """Whether a python-docx `Run` is explicitly bold."""
    rpr = run._r.find(qn("w:rPr"))
    if rpr is None:
        return False
    b = rpr.find(qn("w:b"))
    if b is None:
        return False
    val = b.get(qn("w:val"))
    return val is None or val not in ("0", "false")


def _run_has_tab(run) -> bool:
    """Whether a python-docx `Run` carries a real `<w:tab/>` element."""
    return run._r.find(qn("w:tab")) is not None


def _paragraph_containing(doc, needle: str):
    """The single paragraph whose text contains `needle`."""
    matches = [p for p in doc.paragraphs if needle in p.text]
    assert len(matches) == 1, (
        f"expected exactly one paragraph containing {needle!r}, got {len(matches)}: "
        f"{[p.text for p in matches]}"
    )
    return matches[0]


def _run_with_text(paragraph, text: str):
    """The single run in `paragraph` whose text is exactly `text`."""
    matches = [r for r in paragraph.runs if r.text == text]
    assert len(matches) == 1, (
        f"expected exactly one run with text {text!r} in {paragraph.text!r}, "
        f"got {len(matches)}"
    )
    return matches[0]


def _project_resume_with_link(document, *, tech: bool = True) -> None:
    """PROJECTS resume with a bold name run, plain tech run, and a real hyperlink+tab date.

    Mirrors the run structure of an actual Google Docs export: the name and its
    trailing " | " share one bold run, tech is a separate plain run, the link is a real
    `w:hyperlink`, and the date follows a real `<w:tab/>` element.
    """
    num_id = _add_bullet_numbering(document)
    document.add_paragraph("Name")
    document.add_paragraph("email@example.com")
    document.add_paragraph("EDUCATION")
    edu = document.add_paragraph()
    school = edu.add_run("State University | ")
    school.bold = True
    edu.add_run("Remote")
    edu_tab = edu.add_run()
    edu_tab._r.append(OxmlElement("w:tab"))
    edu.add_run("2018 - 2022")
    _make_bullet(document, "BSc Computer Science", num_id)
    document.add_paragraph("WORK EXPERIENCES")
    exp = document.add_paragraph()
    company = exp.add_run("Acme | ")
    company.bold = True
    exp.add_run("Remote")
    exp_tab = exp.add_run()
    exp_tab._r.append(OxmlElement("w:tab"))
    exp.add_run("2020 - Present")
    document.add_paragraph("Engineer")
    _make_bullet(document, "Shipped features.", num_id)
    document.add_paragraph("PROJECTS")
    p = document.add_paragraph()
    name_run = p.add_run("Text-to-SQL Reward Fine-Tuning via RL | ")
    name_run.bold = True
    if tech:
        p.add_run("GRPO, Deeplearn, HuggingFace, SQL | ")
    _add_hyperlink(p, "Github", "https://github.com/example/repo")
    tab_run = p.add_run()
    tab_run._r.append(OxmlElement("w:tab"))
    p.add_run("Jan 2026 - Mar 2026")
    _make_bullet(document, "Built reward functions.", num_id)
    document.add_paragraph("SKILLS")
    sk = document.add_paragraph()
    bold = sk.add_run("Languages:")
    bold.bold = True
    sk.add_run(" Python, SQL")


def test_build_from_profile_inserts_jinja_tags(tmp_path: Path):
    """Profile build produces experience/project loop tags and name/contact placeholders."""
    raw = _docx_bytes(_standard_resume)
    analysis = template_analyze.analyze_docx(raw=raw)
    assert analysis.suggested_profile is not None

    src = tmp_path / "baseline.docx"
    dst = tmp_path / "main_template.docx"
    src.write_bytes(raw)

    template_build.build_from_profile(src, dst, analysis.suggested_profile)
    assert dst.exists()

    doc = docx.Document(str(dst))
    texts = [p.text for p in doc.paragraphs]
    joined = "\n".join(texts)
    assert "{{ name }}" in joined
    assert "{{r contact }}" in joined
    assert any("{%p for job in experience %}" in t for t in texts)
    assert any("{{ job.company }}" in t or "{{ job.title }}" in t for t in texts)
    assert any("{%p for bullet in job.bullets %}" in t for t in texts)
    assert any("{%p for proj in projects %}" in t for t in texts)
    assert any("{%p for group in skills %}" in t for t in texts)


def test_build_omits_disabled_projects(tmp_path: Path):
    """Disabling projects removes that section's loop from the tagged template."""
    raw = _docx_bytes(_standard_resume)
    analysis = template_analyze.analyze_docx(raw=raw)
    profile = analysis.suggested_profile
    assert profile is not None

    profile = profile.model_copy(
        update={
            "enabled": profile.enabled.model_copy(update={"projects": False}),
            "projects": None,
        }
    )

    src = tmp_path / "baseline.docx"
    dst = tmp_path / "main_template.docx"
    src.write_bytes(raw)
    template_build.build_from_profile(src, dst, profile)

    joined = "\n".join(p.text for p in docx.Document(str(dst)).paragraphs)
    assert "{%p for job in experience %}" in joined
    assert "{%p for proj in projects %}" not in joined


def test_build_project_header_with_github_link(tmp_path: Path):
    """Three-part project headers (name | tech | Github\\tdate) tag without overlap errors."""
    from tests.test_template_analyze import (
        _add_bullet_numbering,
        _add_hyperlink,
        _make_bullet,
    )
    from docx.oxml import OxmlElement

    def build(document):
        num_id = _add_bullet_numbering(document)
        document.add_paragraph("Name")
        document.add_paragraph("email@example.com")
        document.add_paragraph("WORK EXPERIENCES")
        document.add_paragraph("Acme | Remote\t2020 - Present")
        document.add_paragraph("Engineer")
        _make_bullet(document, "Shipped features.", num_id)
        document.add_paragraph("PROJECTS")
        p = document.add_paragraph(
            "Text-to-SQL Reward Fine-Tuning via RL | "
            "GRPO, Deeplearn, HuggingFace, SQL | "
        )
        _add_hyperlink(p, "Github", "https://github.com/example/repo")
        run = p.add_run()
        run._r.append(OxmlElement("w:tab"))
        p.add_run("Jan 2026 - Mar 2026")
        _make_bullet(document, "Built reward functions.", num_id)
        document.add_paragraph("SKILLS")
        sk = document.add_paragraph()
        bold = sk.add_run("Languages:")
        bold.bold = True
        sk.add_run(" Python, SQL")

    raw = _docx_bytes(build)
    analysis = template_analyze.analyze_docx(raw=raw)
    assert analysis.suggested_profile is not None

    src = tmp_path / "baseline.docx"
    dst = tmp_path / "main_template.docx"
    src.write_bytes(raw)
    template_build.build_from_profile(src, dst, analysis.suggested_profile)

    joined = "\n".join(p.text for p in docx.Document(str(dst)).paragraphs)
    assert "{{ proj.name }}" in joined
    assert "{{ proj.tech }}" in joined
    assert "{{ proj.date }}" in joined
    assert "{{r proj.link }}" in joined


def _build_project_template(tmp_path: Path, *, tech: bool = True, keep_link: bool = True):
    """Build the tagged template for `_project_resume_with_link` and return the doc."""
    raw = _docx_bytes(lambda d: _project_resume_with_link(d, tech=tech))
    analysis = template_analyze.analyze_docx(raw=raw)
    assert analysis.suggested_profile is not None
    profile = analysis.suggested_profile
    if not keep_link:
        assert profile.projects is not None
        profile = profile.model_copy(
            update={
                "projects": profile.projects.model_copy(
                    update={"link": template_analyze.OptionalSpan(present=False)}
                )
            }
        )

    src = tmp_path / "baseline.docx"
    dst = tmp_path / "main_template.docx"
    src.write_bytes(raw)
    template_build.build_from_profile(src, dst, profile)
    return docx.Document(str(dst))


def test_project_header_keeps_bold_name_and_plain_tech(tmp_path: Path):
    """The project name tag stays bold; tech, link, and date tags come out plain."""
    doc = _build_project_template(tmp_path)
    header = _paragraph_containing(doc, "{{ proj.name }}")
    assert _run_bold(_run_with_text(header, "{{ proj.name }}"))
    assert not _run_bold(_run_with_text(header, "{{ proj.tech }}"))
    assert not _run_bold(_run_with_text(header, "{{r proj.link }}"))
    assert not _run_bold(_run_with_text(header, "{{ proj.date }}"))


def test_project_header_date_keeps_tab_element(tmp_path: Path):
    """A real <w:tab/> survives immediately before the project date tag."""
    doc = _build_project_template(tmp_path)
    header = _paragraph_containing(doc, "{{ proj.name }}")
    runs = header.runs
    date_idx = next(i for i, r in enumerate(runs) if r.text == "{{ proj.date }}")
    assert any(_run_has_tab(r) for r in runs[:date_idx])
    assert "\t{{ proj.date }}" in header.text or (
        date_idx > 0 and _run_has_tab(runs[date_idx - 1])
    )


def test_no_separator_before_project_link_tag(tmp_path: Path):
    """No doubled/dangling separator sits between the tech tag and the link tag."""
    doc = _build_project_template(tmp_path)
    header = _paragraph_containing(doc, "{{ proj.name }}")
    text = header.text
    assert "{{ proj.tech }}{{r proj.link }}" in text
    assert re.search(r"[|•·–—/]\s*\{\{r proj\.link \}\}", text) is None


def test_project_header_without_tech_maps_link_only(tmp_path: Path):
    """'Name | Github\\tdate' builds without an overlap error and has no dangling separator."""
    doc = _build_project_template(tmp_path, tech=False)
    header = _paragraph_containing(doc, "{{ proj.name }}")
    text = header.text
    assert "{{ proj.tech }}" not in text
    assert "{{ proj.name }}{{r proj.link }}" in text
    assert re.search(r"[|•·–—/]\s*\{\{r proj\.link \}\}", text) is None


def test_unmapped_hyperlink_label_is_dropped(tmp_path: Path):
    """A hyperlink label with no link mapping is dropped, and its separator does not dangle."""
    doc = _build_project_template(tmp_path, keep_link=False)
    header = _paragraph_containing(doc, "{{ proj.name }}")
    text = header.text
    assert "Github" not in text
    assert "{{r proj.link }}" not in text
    assert re.search(r"[|•·–—/]\s*\t", text) is None


def test_skills_prototype_keeps_plain_body_run(tmp_path: Path):
    """The skills label tag stays bold; the entries tag is plain, with the gap preserved."""
    doc = _build_project_template(tmp_path)
    header = _paragraph_containing(doc, "{{ group.label }}")
    assert _run_bold(_run_with_text(header, "{{ group.label }}"))
    assert not _run_bold(_run_with_text(header, "{{ group.entries }}"))
    assert "{{ group.label }}: {{ group.entries }}" in header.text


def test_experience_header_keeps_bold_company_and_plain_location(tmp_path: Path):
    """The experience company tag stays bold; location and dates come out plain."""
    doc = _build_project_template(tmp_path)
    header = _paragraph_containing(doc, "{{ job.company }}")
    assert _run_bold(_run_with_text(header, "{{ job.company }}"))
    assert not _run_bold(_run_with_text(header, "{{ job.location }}"))
    assert not _run_bold(_run_with_text(header, "{{ job.dates }}"))
    assert any(_run_has_tab(r) for r in header.runs)


def test_span_past_paragraph_end_raises():
    """A field span extending past the paragraph length fails loudly, naming the field."""
    raw = _docx_bytes(_standard_resume)
    from io import BytesIO

    doc = docx.Document(BytesIO(raw))
    para = next(p for p in doc.paragraphs if "Analytical Engines" in p.text)
    with pytest.raises(RuntimeError, match="dates"):
        template_build.replace_span_with_tag(
            para,
            template_analyze.CharSpan(
                paragraph_id=0, start=0, end=len(para.text) + 5
            ),
            "{{ job.dates }}",
            field="dates",
        )


def test_tab_inside_span_raises():
    """A span that straddles a real tab fails loudly instead of dropping the tab."""
    raw = _docx_bytes(_standard_resume)
    from io import BytesIO

    doc = docx.Document(BytesIO(raw))
    para = next(p for p in doc.paragraphs if "Analytical Engines" in p.text)
    tab_idx = para.text.index("\t")
    with pytest.raises(RuntimeError, match="tab"):
        template_build.replace_span_with_tag(
            para,
            template_analyze.CharSpan(
                paragraph_id=0, start=tab_idx - 1, end=tab_idx + 2
            ),
            "{{ job.dates }}",
            field="dates",
        )


def test_overlapping_spans_name_both_fields():
    """Overlapping mapped spans raise, naming both offending fields."""
    raw = _docx_bytes(_standard_resume)
    from io import BytesIO

    doc = docx.Document(BytesIO(raw))
    para = next(p for p in doc.paragraphs if "Analytical Engines" in p.text)
    slices = template_build.docx_text.paragraph_run_slices(para)
    text = template_build.docx_text.paragraph_text(para)
    with pytest.raises(RuntimeError, match=r"(?s)company.*location|location.*company"):
        template_build.build_segments(
            text,
            slices,
            [
                (0, 12, "{{ job.company }}", "company"),
                (8, 20, "{{ job.location }}", "location"),
            ],
        )


# ----------------------------------------------------------------------------------------
# Generic-mode build (section_mode="generic"): one shared block instead of one loop per
# kind, driven by `data.MasterResume.sections` at render time. Phase 4's analyzer is what
# will eventually populate `TemplateProfile.sections`/`heading_prototype` from a real
# upload; these tests hand-convert the analyzer's own fixed-mode output (same document,
# same prototype spans) so the build/render path is exercised independently of detection.
# ----------------------------------------------------------------------------------------


def _to_generic(profile: template_profile.TemplateProfile) -> template_profile.TemplateProfile:
    """Fixed-mode profile -> bare generic-mode profile, for testing `build_generic`."""
    return profile.model_copy(
        update={
            "section_mode": "generic",
            "heading_prototype": HeadingPrototype(
                paragraph_id=profile.experience.heading_paragraph_id
            ),
        }
    )


def test_build_generic_inserts_one_shared_block(tmp_path: Path):
    """Generic mode tags one `{%p for section in sections %}` block, not one loop per
    kind — no `{%p for job in experience %}`-style fixed-mode tag survives."""
    raw = _docx_bytes(_standard_resume)
    analysis = template_analyze.analyze_docx(raw=raw)
    assert analysis.suggested_profile is not None
    profile = _to_generic(analysis.suggested_profile)

    src = tmp_path / "baseline.docx"
    dst = tmp_path / "main_template.docx"
    src.write_bytes(raw)
    template_build.build_from_profile(src, dst, profile)

    joined = "\n".join(p.text for p in docx.Document(str(dst)).paragraphs)

    assert "{%p for section in sections %}" in joined
    assert "{{ section.title }}" in joined
    assert "{%p if section.kind == 'experience' %}" in joined
    assert "{%p if section.kind == 'project' %}" in joined
    assert "{%p if section.kind == 'skills' %}" in joined
    assert "{%p if section.kind == 'education' %}" in joined
    assert "{%p if section.kind == 'list' %}" not in joined  # no list_section mapping

    assert "{%p for job in section.entries %}" in joined
    assert "{{ job.company }}" in joined
    assert "{%p for proj in section.entries %}" in joined
    assert "{{ proj.name }}" in joined
    assert "{%p for group in section.entries %}" in joined
    assert "{{ group.label }}" in joined
    assert "{%p for edu in section.entries %}" in joined
    assert "{{ edu.school }}" in joined

    # No fixed-mode loop tags leaked through.
    assert "{%p for job in experience %}" not in joined
    assert "{%p for proj in projects %}" not in joined
    assert "{%p for group in skills %}" not in joined
    assert "{%p for edu in education %}" not in joined
    # The original heading text is gone — replaced by the section.title tag.
    assert "WORK EXPERIENCES" not in joined


def test_build_generic_renders_multiple_sections_of_the_same_kind(tmp_path: Path):
    """The whole point: N experience-kind sections, each with its own custom title,
    render correctly from a template tagged exactly once — no rebuild needed to add,
    rename, or reorder a section. Mirrors the motivating case (WORK EXPERIENCE /
    LEADERSHIP EXPERIENCE / OTHER ACTIVITIES all being experience-shaped)."""
    raw = _docx_bytes(_standard_resume)
    analysis = template_analyze.analyze_docx(raw=raw)
    assert analysis.suggested_profile is not None
    profile = _to_generic(analysis.suggested_profile)

    src = tmp_path / "baseline.docx"
    dst = tmp_path / "main_template.docx"
    src.write_bytes(raw)
    template_build.build_from_profile(src, dst, profile)

    resume = MasterResume(
        contact={"name": "Nina Dao", "email": "nina@example.com"},
        sections=[
            ExperienceSection(
                id="work",
                title="WORK EXPERIENCE",
                entries=[
                    Experience(
                        company="Langmaster JSC",
                        title="Online Tutor",
                        start="2025",
                        end="Present",
                        bullets=[
                            Bullet(id="w1", text="Tutored students in English.", tags=["teaching"])
                        ],
                    )
                ],
            ),
            ExperienceSection(
                id="leadership",
                title="LEADERSHIP EXPERIENCE",
                entries=[
                    Experience(
                        company="In the Green at UCI",
                        title="Co-president",
                        start="2025",
                        end="Present",
                        bullets=[
                            Bullet(id="l1", text="Oversaw club operations.", tags=["leadership"])
                        ],
                    )
                ],
            ),
            ExperienceSection(
                id="other",
                title="OTHER ACTIVITIES",
                entries=[
                    Experience(
                        company="Heartbeat Bazaar",
                        title="Organizer",
                        start="2022",
                        end="2022",
                        bullets=[
                            Bullet(id="o1", text="Directed a fundraising event.", tags=["events"])
                        ],
                    )
                ],
            ),
        ],
    )

    tpl = DocxTemplate(str(dst))
    layout = template_profile.active_layout(profile=profile)
    ctx = build_context(resume, tpl, layout=layout)
    tpl.render(ctx, autoescape=True)
    out = tmp_path / "out.docx"
    tpl.save(str(out))

    texts = [p.text for p in docx.Document(str(out)).paragraphs]
    joined = "\n".join(texts)

    assert (
        texts.index("WORK EXPERIENCE")
        < texts.index("LEADERSHIP EXPERIENCE")
        < texts.index("OTHER ACTIVITIES")
    )
    assert "Langmaster JSC" in joined and "Tutored students in English." in joined
    assert "In the Green at UCI" in joined and "Oversaw club operations." in joined
    assert "Heartbeat Bazaar" in joined and "Directed a fundraising event." in joined
    # No leftover control tags or original upload heading text.
    assert "{%p" not in joined
    assert "{{" not in joined
    assert "WORK EXPERIENCES" not in texts


def test_build_context_omits_a_section_kind_with_no_enabled_prototype(tmp_path: Path):
    """A resume section whose kind the active layout has no prototype for (the default
    for `list_section`) is silently omitted from `sections` — the render-side half of
    the "skip it and warn loudly" policy; `fit.fit` supplies the warning half."""
    resume = MasterResume(
        contact={"name": "N", "email": "n@example.com"},
        sections=[
            ListSection(
                id="certs",
                title="CERTIFICATIONS",
                entries=[ListItem(id="c1", text="AWS Certified Cloud Practitioner")],
            )
        ],
    )
    blank = tmp_path / "blank.docx"
    blank.write_bytes(_docx_bytes(lambda d: d.add_paragraph("placeholder")))
    tpl = DocxTemplate(str(blank))

    ctx = build_context(resume, tpl, layout=template_profile.legacy_defaults())
    assert ctx["sections"] == []


def test_build_context_includes_list_section_when_enabled(tmp_path: Path):
    """The same list-kind section renders once the layout's `list_section` prototype is
    marked enabled — confirming the omission above is the enabled-flag gate, not a bug."""
    resume = MasterResume(
        contact={"name": "N", "email": "n@example.com"},
        sections=[
            ListSection(
                id="certs",
                title="CERTIFICATIONS",
                entries=[ListItem(id="c1", text="AWS Certified Cloud Practitioner")],
            )
        ],
    )
    blank = tmp_path / "blank.docx"
    blank.write_bytes(_docx_bytes(lambda d: d.add_paragraph("placeholder")))
    tpl = DocxTemplate(str(blank))

    layout = dict(template_profile.legacy_defaults())
    layout["enabled"] = {**layout["enabled"], "list_section": True}
    ctx = build_context(resume, tpl, layout=layout)
    assert ctx["sections"] == [
        {
            "id": "certs",
            "title": "CERTIFICATIONS",
            "kind": "list",
            "entries": ["AWS Certified Cloud Practitioner"],
        }
    ]


def test_build_generic_survives_a_paragraph_insertion_during_tagging(tmp_path: Path):
    """Regression: `_tag_education_prototype` clones the degree paragraph in place
    (`addnext`) when a single-line entry has no separate detail bullet, which shifts
    every later paragraph's index. A doc with EDUCATION physically above SKILLS used to
    corrupt the skills prototype's span once education's insertion ran first; processing
    kinds bottom-up (by document position, not by `_GENERIC_KIND_ORDER`) fixes it."""
    from tests.test_template_analyze import _multi_section_resume

    raw = _docx_bytes(_multi_section_resume)
    analysis = template_analyze.analyze_docx(raw=raw)
    profile = analysis.suggested_profile
    assert profile is not None
    assert profile.section_mode == "generic"  # this fixture has 3 experience-kind headings

    src = tmp_path / "baseline.docx"
    dst = tmp_path / "main_template.docx"
    src.write_bytes(raw)
    template_build.build_from_profile(src, dst, profile)  # must not raise

    joined = "\n".join(p.text for p in docx.Document(str(dst)).paragraphs)
    assert "{{ group.label }}" in joined
    assert "{{ group.entries }}" in joined
    assert "{{ edu.degree_line }}" in joined


def test_build_generic_default_spacing_emits_no_loop_first_guards(tmp_path: Path):
    """All-`None` spacing (today's default) inserts no spacer paragraphs and no `loop.
    first` guards — a template with no spacer donors builds byte-identically to before
    this feature existed."""
    raw = _docx_bytes(_standard_resume)
    analysis = template_analyze.analyze_docx(raw=raw)
    assert analysis.suggested_profile is not None
    profile = _to_generic(analysis.suggested_profile)
    assert profile.spacing == template_profile.SpacingProfile()

    src = tmp_path / "baseline.docx"
    dst = tmp_path / "main_template.docx"
    src.write_bytes(raw)
    template_build.build_from_profile(src, dst, profile)

    joined = "\n".join(p.text for p in docx.Document(str(dst)).paragraphs)
    assert "loop.first" not in joined


def test_build_generic_with_spacing_emits_loop_first_guards(tmp_path: Path):
    """A profile with all three spacer donors set emits one `{%p if not loop.first %}`
    guard around the section heading, plus one more per entry-shaped kind present
    (experience and education here — no project or list section in this fixture)."""
    from tests.test_template_analyze import _spacer_multi_section_resume

    raw = _docx_bytes(_spacer_multi_section_resume)
    analysis = template_analyze.analyze_docx(raw=raw)
    profile = analysis.suggested_profile
    assert profile is not None
    assert profile.section_mode == "generic"
    assert profile.spacing.before_heading
    assert profile.spacing.after_heading
    assert profile.spacing.between_entries

    src = tmp_path / "baseline.docx"
    dst = tmp_path / "main_template.docx"
    src.write_bytes(raw)
    template_build.build_from_profile(src, dst, profile)

    joined = "\n".join(p.text for p in docx.Document(str(dst)).paragraphs)
    # One "before heading" guard (shared across all sections) plus one "between
    # entries" guard per entry-shaped kind present: experience and education.
    assert joined.count("{%p if not loop.first %}") == 3
    assert "{%p" not in joined.replace("{%p if not loop.first %}", "").replace(
        "{%p endif %}", ""
    ).replace("{%p endfor %}", "").replace("{%p for", "").replace(
        "{%p if section.kind", ""
    )


def _line_spacing(paragraph) -> str | None:
    """The paragraph's explicit `w:line` value, or None when unset."""
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        return None
    spacing = pPr.find(qn("w:spacing"))
    return None if spacing is None else spacing.get(qn("w:line"))


def _line_rule(paragraph) -> str | None:
    """The paragraph's explicit `w:lineRule` value, or None when unset."""
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        return None
    spacing = pPr.find(qn("w:spacing"))
    return None if spacing is None else spacing.get(qn("w:lineRule"))


def test_sub_single_text_paragraphs_are_normalized_to_single(tmp_path: Path):
    """A text line set below single (`120` = 0.5 line, `72` = 0.3) only renders as its
    author intended in Word, which honours proportional spacing literally. LibreOffice —
    what the container actually renders with — refuses to compress below the glyph height
    and floors the same paragraph ~3x taller, so the two disagree wildly. Text is
    therefore always normalised up to single, where both agree and nothing can overlap."""
    def build(document):
        num_id = _add_bullet_numbering(document)
        document.add_paragraph("Name")
        document.add_paragraph("email@example.com")
        heading = document.add_paragraph("EXPERIENCE")
        heading.paragraph_format.line_spacing = 0.5
        header = document.add_paragraph("Acme | Remote\t2020 - Present")
        header.paragraph_format.line_spacing = 0.3
        document.add_paragraph("Engineer")
        _make_bullet(document, "Shipped features.", num_id)  # spacing left unset

    raw = _docx_bytes(build)
    src = tmp_path / "baseline.docx"
    src.write_bytes(raw)
    doc = docx.Document(str(src))
    template_build.normalize_single_spacing(doc)
    out = tmp_path / "normalized.docx"
    doc.save(str(out))

    paras = {p.text: p for p in docx.Document(str(out)).paragraphs}
    for text in ("EXPERIENCE", "Acme | Remote\t2020 - Present", "Engineer"):
        assert _line_spacing(paras[text]) == "240", text
        assert _line_rule(paras[text]) == "auto", text
    assert _line_spacing(paras["Shipped features."]) == "240"


def test_bullets_always_get_the_exact_line_rule(tmp_path: Path):
    """Regression: a bullet already carrying `line="240" lineRule="auto"` reads as
    "already single" to the only-ever-tighten rule, so it was skipped and kept `auto`.
    Under `auto` a taller substitute bullet glyph (LibreOffice without Noto) inflates the
    line box past single — measured ~15.7pt for 10pt text — which cost a whole bullet per
    rendered page. `exact` is a measurement guarantee and applies to every bullet."""
    def build(document):
        num_id = _add_bullet_numbering(document)
        document.add_paragraph("Name")
        document.add_paragraph("email@example.com")
        document.add_paragraph("EXPERIENCE")
        document.add_paragraph("Acme | Remote\t2020 - Present")
        document.add_paragraph("Engineer")
        # Exactly single, expressed as auto — the shape a real Google Docs export uses.
        b = _make_bullet(document, "Shipped features.", num_id)
        b.paragraph_format.line_spacing = 1.0

    raw = _docx_bytes(build)
    src = tmp_path / "baseline.docx"
    src.write_bytes(raw)
    doc = docx.Document(str(src))
    assert _line_rule(doc.paragraphs[-1]) == "auto"  # sanity: starts as auto

    template_build.normalize_single_spacing(doc)
    out = tmp_path / "normalized.docx"
    doc.save(str(out))

    bullet = docx.Document(str(out)).paragraphs[-1]
    assert bullet.text == "Shipped features."
    assert _line_rule(bullet) == "exact"
    assert _line_spacing(bullet) == "240"


def test_chrome_keeps_its_height_but_is_pinned_to_exact(tmp_path: Path):
    """Chrome carries no text, so "one line of text should be one line tall" says nothing
    about it — its authored height *is* its content and must survive. But a sub-single
    `auto` value is the non-portable one, so it is re-expressed as the same number of
    twips under `exact`, which Word and LibreOffice draw identically. Values at or above
    single are already consistent between the two and are left alone."""
    def build(document):
        document.add_paragraph("Name")
        document.add_paragraph("email@example.com")
        loose = document.add_paragraph("")
        loose.paragraph_format.line_spacing = 1.5  # looser than single, still chrome
        thin = document.add_paragraph("")
        thin.paragraph_format.line_spacing = 0.3  # the non-portable case
        document.add_paragraph("_" * 40)

    raw = _docx_bytes(build)
    src = tmp_path / "baseline.docx"
    src.write_bytes(raw)
    doc = docx.Document(str(src))
    template_build.normalize_single_spacing(doc)
    out = tmp_path / "normalized.docx"
    doc.save(str(out))

    blanks = [p for p in docx.Document(str(out)).paragraphs if not p.text.strip()]
    loose, thin = blanks[0], blanks[1]
    # Looser than single: already portable, untouched.
    assert (_line_spacing(loose), _line_rule(loose)) == ("360", "auto")
    # Tighter than single: same height, now pinned so LibreOffice cannot inflate it.
    assert (_line_spacing(thin), _line_rule(thin)) == ("72", "exact")
    rule = next(p for p in docx.Document(str(out)).paragraphs if set(p.text.strip()) == {"_"})
    assert _line_spacing(rule) is None  # unset, left exactly as authored


def test_no_built_template_paragraph_is_sub_single_auto(tmp_path: Path):
    """The invariant the renderer-portability fix establishes: after tagging, nothing in
    the document may carry `lineRule="auto"` with `w:line` below single. Word and
    LibreOffice diverge by ~3x on that combination — Word compresses it literally, real
    LibreOffice floors it at the glyph height — so it must never reach a built template,
    whether the source paragraph was normalised (text) or pinned to `exact` (chrome).

    Uses `_rule_separated_resume`, which mirrors the real motivating document: rules and
    blanks at 0.3 of a line, sub-single headings and entry headers, plus a two-blank
    entry boundary — the exact combination that exposed this bug."""
    from tests.test_template_analyze import _rule_separated_resume

    raw = _docx_bytes(_rule_separated_resume)
    analysis = template_analyze.analyze_docx(raw=raw)
    assert analysis.suggested_profile is not None

    src = tmp_path / "baseline.docx"
    dst = tmp_path / "main_template.docx"
    src.write_bytes(raw)
    template_build.build_from_profile(src, dst, analysis.suggested_profile)

    for p in docx.Document(str(dst)).paragraphs:
        if _line_rule(p) != "auto":
            continue
        line = _line_spacing(p)
        if line is None:
            continue
        assert int(line) >= 240, f"{p.text!r} is sub-single auto: line={line}"


def test_clamp_tab_stops_uses_paragraph_right_edge(tmp_path: Path):
    """Google Docs exports can carry a right tab stop (`w:tab/@w:pos`, measured from the
    left text margin) that overruns the paragraph's own right boundary
    (`text_width - w:ind/@w:right`) and even the page's text column. Word silently clamps
    such a stop back to the paragraph's right edge when it renders; LibreOffice — what the
    container actually renders with — honours it literally and lets the tab-aligned date
    hang past the page margin. `clamp_tab_stops` bakes Word's clamp into the XML.

    Also pins that the section-geometry reader tolerates fractional twips: Google Docs
    writes non-integer `w:pgMar` values (e.g. `w:left="1417.3228346456694"`), which raise
    `ValueError` through python-docx's typed `Section.left_margin`/`.right_margin`."""

    def build(document):
        over = document.add_paragraph("Company")
        over.paragraph_format.right_indent = Pt(10)  # 200 twips
        over.paragraph_format.tab_stops.add_tab_stop(Pt(500), WD_TAB_ALIGNMENT.RIGHT)
        within = document.add_paragraph("In range")
        within.paragraph_format.tab_stops.add_tab_stop(Pt(100), WD_TAB_ALIGNMENT.RIGHT)

    raw = _docx_bytes(build)
    src = tmp_path / "baseline.docx"
    src.write_bytes(raw)
    doc = docx.Document(str(src))

    # Real Google Docs fractional-twip margins — the shape that breaks typed accessors.
    sectPr = doc.element.body.find(qn("w:sectPr"))
    pgMar = sectPr.find(qn("w:pgMar"))
    pgMar.set(qn("w:left"), "1417.3228346456694")
    pgMar.set(qn("w:right"), "708.5433070866151")
    pgSz = sectPr.find(qn("w:pgSz"))
    text_width = float(pgSz.get(qn("w:w"))) - 1417.3228346456694 - 708.5433070866151

    template_build.clamp_tab_stops(doc)
    out = tmp_path / "clamped.docx"
    doc.save(str(out))

    paras = {p.text: p for p in docx.Document(str(out)).paragraphs}

    over_tab = paras["Company"]._p.find(qn("w:pPr")).find(qn("w:tabs")).find(qn("w:tab"))
    limit = text_width - 200  # the paragraph's own right edge (right_indent = 200 twips)
    assert float(over_tab.get(qn("w:pos"))) == round(limit)

    within_tab = paras["In range"]._p.find(qn("w:pPr")).find(qn("w:tabs")).find(qn("w:tab"))
    assert float(within_tab.get(qn("w:pos"))) == 2000  # Pt(100), untouched


def test_no_built_template_tab_stop_exceeds_its_paragraph(tmp_path: Path):
    """Companion to the sub-single-auto invariant above: after tagging, no explicit tab
    stop may sit beyond its own paragraph's right edge either. Reproduces the real Nina
    defect end to end — an entry header's date tab stop overruns the paragraph, which
    Word silently clamps on render and LibreOffice does not, letting the date hang past
    the page margin."""
    from io import BytesIO

    from tests.test_template_analyze import _rule_separated_resume

    raw = _docx_bytes(_rule_separated_resume)
    doc = docx.Document(BytesIO(raw))
    for p in doc.paragraphs:
        if "\t" not in p.text:
            continue
        p.paragraph_format.right_indent = Pt(10)
        p.paragraph_format.tab_stops.add_tab_stop(Pt(500), WD_TAB_ALIGNMENT.RIGHT)
    buf = BytesIO()
    doc.save(buf)
    raw = buf.getvalue()

    analysis = template_analyze.analyze_docx(raw=raw)
    assert analysis.suggested_profile is not None

    src = tmp_path / "baseline.docx"
    dst = tmp_path / "main_template.docx"
    src.write_bytes(raw)
    template_build.build_from_profile(src, dst, analysis.suggested_profile)

    built = docx.Document(str(dst))
    text_width = template_build._section_text_width(built)
    assert text_width is not None

    found_tabs = 0
    for p in built.paragraphs:
        pPr = p._p.find(qn("w:pPr"))
        if pPr is None:
            continue
        tabs = pPr.find(qn("w:tabs"))
        if tabs is None:
            continue
        ind = pPr.find(qn("w:ind"))
        right = (
            float(ind.get(qn("w:right")))
            if ind is not None and ind.get(qn("w:right")) is not None
            else 0.0
        )
        limit = text_width - right
        for tab in tabs.findall(qn("w:tab")):
            pos = tab.get(qn("w:pos"))
            if pos is None:
                continue
            found_tabs += 1
            assert float(pos) <= limit + 0.5, (
                f"{p.text!r} tab pos={pos} exceeds paragraph edge {limit}"
            )
    assert found_tabs > 0  # sanity: the fixture actually exercised the clamp


def test_build_generic_reproduces_a_rule_under_every_heading(tmp_path: Path):
    """Regression: the horizontal rule under each heading is a *body* paragraph, so it
    was swept into `all_victims` and deleted, leaving every rendered heading without its
    underline. It is chrome (not blank), so single-blank spacer detection never caught
    it; the `after_heading` run does."""
    from tests.test_template_analyze import _rule_separated_resume

    raw = _docx_bytes(_rule_separated_resume)
    analysis = template_analyze.analyze_docx(raw=raw)
    profile = analysis.suggested_profile
    assert profile is not None
    assert len(profile.spacing.after_heading) == 2  # rule + blank

    src = tmp_path / "baseline.docx"
    dst = tmp_path / "main_template.docx"
    src.write_bytes(raw)
    template_build.build_from_profile(src, dst, profile)

    resume = MasterResume(
        contact={"name": "Nina", "email": "n@example.com"},
        sections=[
            ExperienceSection(
                id="work",
                title="EXPERIENCE",
                entries=[
                    Experience(
                        company="Yellow Daisy", title="Co-founder", start="2022", end="2024",
                        bullets=[Bullet(id="a", text="Raised funds.", tags=["x"])],
                    ),
                    Experience(
                        company="Youth Opportunity", title="Ambassador", start="2021", end="2021",
                        bullets=[Bullet(id="b", text="Designed posts.", tags=["x"])],
                    ),
                ],
            ),
        ],
    )

    tpl = DocxTemplate(str(dst))
    ctx = build_context(resume, tpl, layout=template_profile.active_layout(profile=profile))
    tpl.render(ctx, autoescape=True)
    out = tmp_path / "out.docx"
    tpl.save(str(out))

    texts = [p.text for p in docx.Document(str(out)).paragraphs]
    head = texts.index("EXPERIENCE")
    # Heading, then the rule, then a blank — the upload's own shape.
    assert set(texts[head + 1].strip()) == {"_"}
    assert texts[head + 2].strip() == ""
    # And the entries are still separated by exactly one blank.
    daisy = next(i for i, t in enumerate(texts) if "Yellow Daisy" in t)
    youth = next(i for i, t in enumerate(texts) if "Youth Opportunity" in t)
    assert texts[youth - 1].strip() == ""
    assert texts[daisy - 1].strip() == ""  # the after-heading blank
    assert youth - daisy == 4  # header, title, bullet, blank


def test_build_generic_renders_spacers_at_the_right_positions(tmp_path: Path):
    """End-to-end: a 2-section (education, experience), 2-entry-in-experience resume
    renders exactly one blank line after each heading, one before the second section,
    and one between the two experience entries — none before the very first section or
    before the first entry of either section."""
    from tests.test_template_analyze import _spacer_multi_section_resume

    raw = _docx_bytes(_spacer_multi_section_resume)
    analysis = template_analyze.analyze_docx(raw=raw)
    profile = analysis.suggested_profile
    assert profile is not None

    src = tmp_path / "baseline.docx"
    dst = tmp_path / "main_template.docx"
    src.write_bytes(raw)
    template_build.build_from_profile(src, dst, profile)

    from resume_tailor.data import EducationSection

    resume = MasterResume(
        contact={"name": "Nina Dao", "email": "nina@example.com"},
        sections=[
            EducationSection(
                id="edu",
                title="EDUCATION",
                entries=[
                    {
                        "school": "UC Irvine",
                        "degree": "B.A. in Business Administration",
                        "dates": "Expected June 2027",
                    }
                ],
            ),
            ExperienceSection(
                id="work",
                title="WORK EXPERIENCE",
                entries=[
                    Experience(
                        company="Langmaster JSC",
                        title="Online Tutor",
                        start="2025",
                        end="Present",
                        bullets=[Bullet(id="w1", text="Tutored students.", tags=["teaching"])],
                    ),
                    Experience(
                        company="Garin JSC",
                        title="Logistics Intern",
                        start="2023",
                        end="2024",
                        bullets=[Bullet(id="w2", text="Managed logistics.", tags=["logistics"])],
                    ),
                ],
            ),
        ],
    )

    tpl = DocxTemplate(str(dst))
    layout = template_profile.active_layout(profile=profile)
    ctx = build_context(resume, tpl, layout=layout)
    tpl.render(ctx, autoescape=True)
    out = tmp_path / "out.docx"
    tpl.save(str(out))

    texts = [p.text for p in docx.Document(str(out)).paragraphs]

    blank_indices = [i for i, t in enumerate(texts) if t.strip() == ""]
    edu_idx = texts.index("EDUCATION")
    work_idx = texts.index("WORK EXPERIENCE")
    tutor_idx = next(i for i, t in enumerate(texts) if "Langmaster JSC" in t)
    garin_idx = next(i for i, t in enumerate(texts) if "Garin JSC" in t)

    assert edu_idx < work_idx < garin_idx
    assert tutor_idx < garin_idx

    # Exactly one blank precedes the very first section's heading too — but it comes
    # from the untouched original paragraph sitting above the whole `{%p for section in
    # sections %}` block, not from the `before_heading` spacer (which is guarded on
    # `not loop.first` specifically so it does not double up with that survivor).
    assert edu_idx - 1 in blank_indices
    assert edu_idx - 2 not in blank_indices
    # One blank right after the EDUCATION heading.
    assert edu_idx + 1 in blank_indices
    # One blank right before WORK EXPERIENCE (not the first section) and right after it
    # — the same blank that immediately precedes the first entry (Langmaster), since the
    # `between_entries` spacer is suppressed before a section's first entry.
    assert work_idx - 1 in blank_indices
    assert work_idx + 1 in blank_indices
    assert tutor_idx - 1 == work_idx + 1
    # One blank between the two entries (Langmaster, then Garin), and none anywhere
    # else inside the experience body.
    assert garin_idx - 1 in blank_indices
    assert blank_indices == [edu_idx - 1, edu_idx + 1, work_idx - 1, work_idx + 1, garin_idx - 1]
    # No leftover control tags.
    assert "{%p" not in "\n".join(texts)
    assert "{{" not in "\n".join(texts)


# ----------------------------------------------------------------------------------------
# Bullet marker shrink: glyph-aware and idempotent.
# ----------------------------------------------------------------------------------------


def _marker_size(paragraph) -> str | None:
    """The paragraph mark's own `w:sz` (governs the bullet glyph size), or None."""
    pPr = paragraph._p.find(qn("w:pPr"))
    rPr = pPr.find(qn("w:rPr")) if pPr is not None else None
    sz = rPr.find(qn("w:sz")) if rPr is not None else None
    return sz.get(qn("w:val")) if sz is not None else None


def _bullet_resume_with_glyph(document, glyph: str):
    """A minimal bulleted resume whose lvl0 marker uses `glyph`. Returns the bullet
    paragraph. Builds on `_add_bullet_numbering`'s abstract/num ids (10/20) rather than
    inventing new ones — a blank `docx.Document()` already ships built-in numbering
    definitions, and low hand-picked ids like "1" collide with those, resolving to the
    wrong glyph entirely."""
    num_id = _add_bullet_numbering(document)
    root = document.part.numbering_part.element
    abstract_id = template_build._num_to_abstract(root)[num_id]
    for anum in root.findall(qn("w:abstractNum")):
        if anum.get(qn("w:abstractNumId")) != abstract_id:
            continue
        for lvl in anum.findall(qn("w:lvl")):
            if lvl.get(qn("w:ilvl")) != "0":
                continue
            lvl_text = lvl.find(qn("w:lvlText"))
            if lvl_text is None:
                lvl_text = OxmlElement("w:lvlText")
                lvl.append(lvl_text)
            lvl_text.set(qn("w:val"), glyph)

    document.add_paragraph("Name")
    document.add_paragraph("email@example.com")
    bullet = _make_bullet(document, "Shipped features.", num_id)
    for run in bullet.runs:
        run.font.size = Pt(10)  # w:sz 20 half-points
    pPr = bullet._p.get_or_add_pPr()
    rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "20")
    rPr.append(sz)
    pPr.append(rPr)
    return bullet


def test_shrink_bullet_marker_leaves_a_dash_at_body_size(tmp_path: Path):
    """A `-` glyph (Jayden's export uses this in every list definition) is already
    small and thin at body size — shrinking it the same ~53% that tames a large round
    dot renders it as a near-invisible hairline, so it must be left alone."""
    doc = docx.Document()
    bullet = _bullet_resume_with_glyph(doc, "-")
    template_build.shrink_bullet_marker(doc, bullet)
    assert _marker_size(bullet) == "20"


def test_shrink_bullet_marker_still_shrinks_a_round_dot():
    """`●` renders as a near-full-em disc in most fonts — the shape the ratio was
    tuned against — and must still shrink."""
    doc = docx.Document()
    bullet = _bullet_resume_with_glyph(doc, "●")
    template_build.shrink_bullet_marker(doc, bullet)
    assert _marker_size(bullet) == "11"


def test_shrink_bullet_marker_is_idempotent():
    """Applying the shrink twice must not halve an already-shrunk marker — the target
    is derived from the body run's own size, not the marker's current size."""
    doc = docx.Document()
    bullet = _bullet_resume_with_glyph(doc, "●")
    template_build.shrink_bullet_marker(doc, bullet)
    once = _marker_size(bullet)
    template_build.shrink_bullet_marker(doc, bullet)
    assert _marker_size(bullet) == once


def test_entry_bullet_matching_another_headings_text_does_not_truncate_body(tmp_path: Path):
    """A bullet whose own text happens to exactly equal another section's heading text
    ("SKILLS") must not be mistaken for that heading and truncate the experience
    section early — regression test for `_section_body_paragraphs`'s move from text
    matching to paragraph object identity as the stop condition.
    """
    def build(document):
        num_id = _add_bullet_numbering(document)
        document.add_paragraph("Name")
        document.add_paragraph("email@example.com")
        document.add_paragraph("WORK EXPERIENCES")
        document.add_paragraph("Acme | Remote\t2020 - Present")
        document.add_paragraph("Engineer")
        _make_bullet(document, "Built things.", num_id)
        _make_bullet(document, "SKILLS", num_id)
        _make_bullet(document, "Led a project after acquiring new skills.", num_id)
        document.add_paragraph("SKILLS")
        sk = document.add_paragraph()
        run = sk.add_run("Languages:")
        run.bold = True
        sk.add_run(" Python, SQL")

    raw = _docx_bytes(build)
    result = template_analyze.analyze_docx(raw=raw)
    assert result.suggested_profile is not None, result.issues

    src = tmp_path / "baseline.docx"
    dst = tmp_path / "out.docx"
    src.write_bytes(raw)
    template_build.build_from_profile(src, dst, result.suggested_profile)

    doc = docx.Document(str(dst))
    texts = [p.text for p in doc.paragraphs]
    # Both bullets after (and including) the "SKILLS"-text one must have been swept up
    # as experience victims and deleted — under the old text-matching stop condition,
    # the walk would have given up at the "SKILLS" bullet, leaving both this bullet and
    # everything after it as stray, undeleted paragraphs in the built template.
    assert "Built things." not in texts
    assert "Led a project after acquiring new skills." not in texts
    # The real SKILLS heading (fixed mode keeps it verbatim) and skills prototype tag
    # still made it through untouched.
    assert "SKILLS" in texts
    assert any("{{ group.label }}" in t for t in texts)


# --------------------------------------------------------------------------------------
# Table-layout build (`build_generic_table`): the shared block repeats table ROWS, not
# paragraphs, via `{%tr %}` marker rows. See `tests/fixtures.py::_table_resume`.
# --------------------------------------------------------------------------------------


def _build_table_resume(tmp_path: Path):
    """Analyze + build `_table_resume`, returning `(profile, built_docx_path)`."""
    raw = _docx_bytes(_table_resume)
    analysis = template_analyze.analyze_docx(raw=raw)
    assert analysis.suggested_profile is not None, analysis.issues
    profile = analysis.suggested_profile
    assert profile.layout == "table"

    src = tmp_path / "baseline.docx"
    dst = tmp_path / "main_template.docx"
    src.write_bytes(raw)
    template_build.build_from_profile(src, dst, profile)
    return profile, dst


def test_build_generic_table_produces_no_empty_cells(tmp_path: Path):
    """A `<w:tc>` with zero `<w:p>` children is invalid OOXML — Word refuses to open
    such a file, but `python-docx` opens it happily, so this must be checked
    explicitly rather than relying on the build simply not raising."""
    _profile, dst = _build_table_resume(tmp_path)
    doc = docx.Document(str(dst))
    assert len(doc.tables) == 1
    for row in doc.tables[0].rows:
        for tc in row._tr.findall(qn("w:tc")):
            assert tc.find(qn("w:p")) is not None


def test_build_generic_table_row_loop_structure(tmp_path: Path):
    """Exactly one row-level section loop; per-kind row-level branches for the two
    experience-kind sections' shared "experience" branch, education, and skills; no
    paragraph-level `SECTION_LOOP_OPEN` (that's the non-table generic-mode tag)."""
    _profile, dst = _build_table_resume(tmp_path)
    doc = docx.Document(str(dst))
    NS = qn("w:p")
    texts = []
    for table in doc.tables:
        for row in table.rows:
            for tc in row._tr.findall(qn("w:tc")):
                texts.extend(p.text for p in tc.findall(NS))
    joined = "\n".join(texts)

    assert texts.count("{%tr for section in sections %}") == 1
    assert "{{ section.title }}" in joined
    assert "{%tr if section.kind == 'experience' %}" in joined
    assert "{%tr if section.kind == 'education' %}" in joined
    assert "{%tr if section.kind == 'skills' %}" in joined
    assert "{%tr for job in section.entries %}" in joined
    assert "{{ job.company }}" in joined
    assert "{%p for bullet in job.bullets %}" in joined
    assert "{%tr for edu in section.entries %}" in joined
    assert "{{ edu.school }}" in joined
    assert "{%p for detail in edu.details %}" in joined
    assert "{%p for group in section.entries %}" in joined
    assert "{{ group.label }}" in joined

    # Paragraph-level generic-mode tags must not leak into a table-layout build.
    assert "{%p for section in sections %}" not in joined


def test_build_generic_table_no_leftover_sibling_bullets(tmp_path: Path):
    """A row's OTHER bullets (siblings of the one chosen as the loop's tagged
    prototype, stacked in the same cell) must not survive verbatim beside the
    `{%p for bullet in job.bullets %}` loop — see `_wrap_cell_loop`. Only ONE bullet
    paragraph should remain per bullet-loop cell in the tagged template."""
    _profile, dst = _build_table_resume(tmp_path)
    doc = docx.Document(str(dst))
    for table in doc.tables:
        for row in table.rows:
            for tc in row._tr.findall(qn("w:tc")):
                para_texts = [p.text for p in tc.findall(qn("w:p"))]
                if "{%p for bullet in job.bullets %}" in para_texts:
                    assert para_texts.count("{{ bullet }}") == 1
                if "{%p for detail in edu.details %}" in para_texts:
                    assert para_texts.count("{{ detail }}") == 1


def test_build_generic_table_verify_tagged_clean(tmp_path: Path):
    """The full `verify_tagged` check (flattened-walk aware, table-mode loop count,
    balanced `{%tr %}`/`{%p %}` control tags, no empty cells) passes with zero issues."""
    from resume_tailor import template_verify

    profile, dst = _build_table_resume(tmp_path)
    issues = template_verify.verify_tagged(dst, profile)
    assert issues == []


def test_build_generic_table_verify_roundtrip_clean(tmp_path: Path):
    """A real render through the table-layout template reaches every mapped field's
    actual value — the full round-trip signal, not just that tagging looks right."""
    from resume_tailor import resume_import, template_verify

    raw = _docx_bytes(_table_resume)
    analysis = template_analyze.analyze_docx(raw=raw)
    assert analysis.suggested_profile is not None
    profile = analysis.suggested_profile

    src = tmp_path / "baseline.docx"
    dst = tmp_path / "main_template.docx"
    src.write_bytes(raw)
    template_build.build_from_profile(src, dst, profile)

    doc = docx.Document(str(src))
    imported = resume_import.import_from_analysis(analysis, doc)

    issues = template_verify.verify_roundtrip(dst, profile, imported.resume)
    assert issues == []
