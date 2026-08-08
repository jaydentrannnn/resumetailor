"""Tests for the deterministic DOCX -> MasterResume importer (no LLM, no network)."""

from __future__ import annotations

import io

import docx

from resume_tailor import config, resume_import, template_analyze
from resume_tailor.data import (
    Bullet,
    Contact,
    Education,
    EducationSection,
    Experience,
    ExperienceSection,
    ListItem,
    ListSection,
    MasterResume,
    SkillGroup,
    SkillsSection,
    SummaryVariant,
)
from tests.fixtures import (
    _add_bullet_numbering,
    _add_hyperlink,
    _docx_bytes,
    _full_featured_resume,
    _make_bullet,
    _multi_section_resume,
    _standard_resume,
    _table_resume,
)


def _import(build) -> resume_import.ImportedResume:
    """Analyze + import a synthetic docx built by `build`."""
    raw = _docx_bytes(build)
    result = template_analyze.analyze_docx(raw=raw)
    doc = docx.Document(io.BytesIO(raw))
    return resume_import.import_from_analysis(result, doc)


def test_full_featured_resume_round_trips_every_field():
    """Every field the source docx actually carries reaches the draft, and the draft
    is a fully valid MasterResume (round-trips through JSON, not just construction)."""
    imported = _import(_full_featured_resume)
    resume = imported.resume

    # Round-trips through full JSON serialization, not just in-memory construction.
    import json

    reloaded = MasterResume.model_validate(json.loads(json.dumps(resume.model_dump(mode="json"))))
    assert reloaded == resume

    assert resume.contact.name == "Jordan Rivera"
    assert resume.contact.email == "jordan@example.com"
    assert resume.contact.phone == "(555) 123-4567"
    assert resume.contact.location == "Irvine, CA"

    edu = resume.education[0]
    assert edu.school == "State University"
    assert edu.location == "Boston, MA"
    assert edu.dates == "2019 - 2023"
    assert edu.degree == "BS Computer Science"
    assert edu.coursework == ["Algorithms", "Databases"]

    exp = resume.experience[0]
    assert exp.company == "Example Corp"
    assert exp.location == "Remote"
    assert exp.title == "Software Engineer"
    assert exp.start == "2023"
    assert exp.end == "Present"
    assert len(exp.bullets) == 1
    assert exp.bullets[0].text == "Improved reliability & throughput for production services."

    proj = resume.projects[0]
    assert proj.name == "Note Engine"
    assert proj.tech == ["Python", "FastAPI"]
    assert proj.date == "2024"
    assert proj.link == "Github"
    assert proj.url == "https://github.com/jordanrivera/note-engine"
    assert len(proj.bullets) == 1

    skills = resume.skills[0]
    assert skills.label == "Tools"
    assert skills.items == ["Python", "Data & Analytics"]


def test_bullet_ids_are_entry_scoped_and_sequential():
    imported = _import(_full_featured_resume)
    exp_bullet = imported.resume.experience[0].bullets[0]
    proj_bullet = imported.resume.projects[0].bullets[0]
    assert exp_bullet.id == f"{imported.resume.experience[0].id}_b1"
    assert proj_bullet.id == f"{imported.resume.projects[0].id}_b1"


def test_known_tag_is_seeded_deterministically():
    """A bullet whose text names a known tag (e.g. "embeddings") gets it without any
    LLM call — the deterministic pass this whole module is built around."""
    imported = _import(_full_featured_resume)
    proj_bullet = imported.resume.projects[0].bullets[0]
    assert "embeddings" in proj_bullet.tags


def test_unmatched_bullet_gets_the_untagged_sentinel_and_is_counted():
    imported = _import(_full_featured_resume)
    exp_bullet = imported.resume.experience[0].bullets[0]
    assert exp_bullet.tags == [resume_import.UNTAGGED]
    assert imported.untagged_bullet_count >= 1
    assert any("untagged" in w.lower() for w in imported.warnings)


def test_entry_ids_collide_safely_across_kinds():
    """Two entries that would slugify to the same base (one experience, one project,
    both named 'Acme') must not collide — the second gets a `-2` suffix, the same
    collision suffixing `MasterResume._fill_entry_ids` already uses."""

    def build(document):
        num_id = _add_bullet_numbering(document)
        document.add_paragraph("Sam Rivers")
        document.add_paragraph("sam@example.com")
        document.add_paragraph("WORK EXPERIENCE")
        document.add_paragraph("Acme | Remote\t2022 - Present")
        document.add_paragraph("Engineer")
        _make_bullet(document, "Built things.", num_id)
        document.add_paragraph("PROJECTS")
        document.add_paragraph("Acme | Python\t2023")
        _make_bullet(document, "Built a different thing.", num_id)

    imported = _import(build)
    exp_id = imported.resume.experience[0].id
    proj_id = imported.resume.projects[0].id
    assert exp_id != proj_id
    assert {exp_id, proj_id} == {"acme", "acme-2"}


def test_multi_section_resume_imports_every_experience_shaped_section():
    """Mirrors the analyzer's own motivating multi-section fixture: three
    experience-kind sections (one unaliased, 'OTHER ACTIVITIES') must all import,
    not just the first one found."""
    imported = _import(_multi_section_resume)
    experience_sections = [s for s in imported.resume.sections if s.kind == "experience"]
    assert len(experience_sections) == 3
    titles = {s.title for s in experience_sections}
    assert titles == {"WORK EXPERIENCE", "LEADERSHIP EXPERIENCE", "OTHER ACTIVITIES"}
    for s in experience_sections:
        assert len(s.entries) == 1
        assert len(s.entries[0].bullets) == 1


def test_standard_resume_imports_cleanly_with_no_warnings():
    imported = _import(_standard_resume)
    assert imported.warnings == []
    assert imported.resume.experience[0].company == "Analytical Engines"
    assert imported.resume.education[0].school == "University of London"


def test_contact_hyperlinks_are_extracted_when_real():
    """A real `w:hyperlink` (not just visible "LinkedIn"/"GitHub" text) is resolved to
    its actual target URL via `docx_text.hyperlink_target`."""

    def build(document):
        document.add_paragraph("Jordan Rivera")
        p = document.add_paragraph("Irvine, CA | jordan@example.com | ")
        _add_hyperlink(p, "LinkedIn", "https://linkedin.com/in/jordanrivera")

    imported = _import(build)
    assert imported.resume.contact.linkedin == "https://linkedin.com/in/jordanrivera"


def test_contact_plain_text_profile_url_is_recovered():
    """A profile URL typed as plain text (no real hyperlink — common when an export
    loses them) is still recognized, and a bare "/" inside it is not mistaken for a
    field separator."""

    def build(document):
        document.add_paragraph("Nina Dao")
        document.add_paragraph(
            "nina@example.com • +1 (626) 206-6947 • www.linkedin.com/in/ngocdao2006"
        )

    imported = _import(build)
    assert imported.resume.contact.linkedin == "www.linkedin.com/in/ngocdao2006"
    assert imported.resume.contact.location == ""


def test_phone_with_leading_parenthesis_is_not_truncated():
    """`_PHONE_RE` requires its match to start on a digit, so "(555) 123-4567" would
    lose its opening paren without the explicit restore."""

    def build(document):
        document.add_paragraph("Jordan Rivera")
        document.add_paragraph("Irvine, CA • jordan@example.com • (555) 123-4567")

    imported = _import(build)
    assert imported.resume.contact.phone == "(555) 123-4567"


def test_two_line_header_title_does_not_leak_its_own_trailing_date():
    """A title line that itself carries a tab-aligned date ('Organizer and Marketing
    Member\\tJan. 2023 - Feb. 2023') must contribute only its text to `title`, using
    its date as a fallback only when the entry header itself had none."""

    def build(document):
        num_id = _add_bullet_numbering(document)
        document.add_paragraph("Nina Dao")
        document.add_paragraph("nina@example.com")
        document.add_paragraph("OTHER ACTIVITIES")
        document.add_paragraph("Heartbeat Bazaar\tMar 2022 - Jun 2022")
        document.add_paragraph("Organizer\tJan 2023 - Feb 2023")
        _make_bullet(document, "Directed fundraising events.", num_id)

    imported = _import(build)
    entry = imported.resume.experience[0]
    assert entry.title == "Organizer"
    assert "\t" not in entry.title
    # The primary header's own dates win over the title line's fallback date.
    assert (entry.start, entry.end) == ("2022-03", "2022-06")


def test_education_gpa_and_coursework_parsed_from_the_degree_line():
    def build(document):
        num_id = _add_bullet_numbering(document)
        document.add_paragraph("Ada Lovelace")
        document.add_paragraph("ada@example.com")
        document.add_paragraph("EDUCATION")
        document.add_paragraph("State University\t2018 - 2022")
        _make_bullet(document, "BSc Computer Science | GPA: 3.9", num_id)
        _make_bullet(document, "Relevant Coursework: Algorithms, Databases", num_id)
        document.add_paragraph("WORK EXPERIENCES")
        document.add_paragraph("Acme | Remote\t2022 - Present")
        document.add_paragraph("Engineer")
        _make_bullet(document, "Built things.", num_id)

    imported = _import(build)
    edu = imported.resume.education[0]
    assert edu.degree == "BSc Computer Science"
    assert edu.gpa == "3.9"
    assert edu.show_gpa is True
    assert edu.coursework == ["Algorithms", "Databases"]


def test_table_resume_imports_cross_cell_fields_correctly():
    """The driving document's shape: contact split across cells, entry headers with
    location/dates in a sibling cell (independent of 3+1 vs 2+2 gridSpan), an
    education entry whose plain-prose degree line is distinct from its real bulleted
    details, and a cross-cell skills grid — every field lands where it belongs, not
    just "some string ends up somewhere"."""
    imported = _import(_table_resume)
    resume = imported.resume

    assert resume.contact.name == "JORDAN RIVERA"
    assert resume.contact.email == "jordan@example.com"
    assert resume.contact.phone == "(555) 123-4567"
    assert resume.contact.location == "Springfield, IL 60000"

    assert len(resume.experience) == 2
    by_company = {e.company: e for e in resume.experience}
    assert by_company["Example Corp"].title == "Software Engineer"
    assert by_company["Example Corp"].location == "Remote"
    assert by_company["Example Corp"].start == "2023-01"
    # LEADERSHIP's 2+2 gridSpan entry must resolve dates identically to WORK
    # EXPERIENCE's 3+1 one — the cross-cell detector keys on cell population, not span.
    assert by_company["Campus Club"].location == "Springfield, IL"
    assert by_company["Campus Club"].start == "2022"

    edu = resume.education[0]
    assert edu.school == "State University"
    assert edu.degree == "BS Computer Science"
    assert edu.location == "IL"
    assert edu.dates == "Class of 2024"
    # The cross-cell location/dates paragraphs must never leak into details.
    assert "IL" not in edu.details
    assert "Class of 2024" not in edu.details

    groups = {g.label: g.items for g in resume.skills}
    assert groups["Languages"] == ["Fluent in English and Spanish"]
    assert groups["Skills"] == ["Python", "SQL", "Excel"]


# ----------------------------------------------------------------------------------
# merge_into — folding an imported draft into an existing master resume (pure, no I/O)
# ----------------------------------------------------------------------------------


def _contact(**overrides) -> Contact:
    base: dict = {"name": "Test User", "email": "test@example.com"}
    base.update(overrides)
    return Contact(**base)


def _resume(sections=None, **overrides) -> MasterResume:
    base: dict = {"contact": _contact(), "sections": sections or []}
    base.update(overrides)
    return MasterResume(**base)


def test_merge_experience_match_is_case_and_punctuation_insensitive():
    existing = _resume(
        sections=[
            ExperienceSection(
                id="exp",
                title="Experience",
                entries=[
                    Experience(
                        id="acme-id",
                        company="Acme, Inc.",
                        title="Old Title",
                        location="Old Loc",
                        start="2020-01",
                        end="2020-06",
                        bullets=[Bullet(id="acme-id_b1", text="Old bullet.", tags=["x"])],
                    ),
                ],
            ),
        ]
    )
    incoming = _resume(
        sections=[
            ExperienceSection(
                id="exp2",
                title="Experience",
                entries=[
                    Experience(
                        company="ACME INC",  # different case/punctuation, same identity
                        title="New Title",
                        location="New Loc",
                        start="2021-01",
                        end="2021-06",
                        bullets=[
                            Bullet(id="whatever_b1", text="New bullet one.", tags=["y"]),
                            Bullet(id="whatever_b2", text="New bullet two.", tags=["y"]),
                        ],
                    ),
                ],
            ),
        ]
    )

    merged, stats = resume_import.merge_into(existing, incoming)

    assert stats.updated == ["ACME INC"]
    assert stats.added == []
    entries = merged.experience
    assert len(entries) == 1
    entry = entries[0]
    assert entry.id == "acme-id"  # kept the existing id
    assert entry.company == "ACME INC"  # refreshed from incoming
    assert entry.title == "New Title"
    assert entry.location == "New Loc"
    # bullets re-minted under the SURVIVING (existing) id, not the incoming entry's own
    assert [b.id for b in entry.bullets] == ["acme-id_b1", "acme-id_b2"]
    assert [b.text for b in entry.bullets] == ["New bullet one.", "New bullet two."]


def test_merge_experience_matches_globally_not_by_section_title():
    """An entry living in 'LEADERSHIP EXPERIENCE' must update in place when the
    incoming section is titled just 'LEADERSHIP' — not duplicate into a new section.
    This is the exact real-world case that motivated global (not section-scoped)
    matching: Nina Dao's export titles this section 'LEADERSHIP', not
    'LEADERSHIP EXPERIENCE'."""
    existing = _resume(
        sections=[
            ExperienceSection(id="work", title="WORK EXPERIENCE", entries=[]),
            ExperienceSection(
                id="leadership-exp",
                title="LEADERSHIP EXPERIENCE",
                entries=[
                    Experience(
                        id="itg-uci",
                        company="In the Green at UCI",
                        title="Co-President",
                        location="Irvine, CA",
                        start="2025-05",
                        end="Present",
                        bullets=[Bullet(id="itg-uci_b1", text="Old.", tags=["x"])],
                    ),
                ],
            ),
        ]
    )
    incoming = _resume(
        sections=[
            ExperienceSection(
                id="leadership",
                title="LEADERSHIP",
                entries=[
                    Experience(
                        company="In the Green at UCI",
                        title="Co-President",
                        location="Irvine, CA",
                        start="2025-05",
                        end="Present",
                        bullets=[Bullet(id="x_b1", text="Refreshed bullet.", tags=["y"])],
                    ),
                ],
            ),
        ]
    )

    merged, stats = resume_import.merge_into(existing, incoming)

    assert stats.updated == ["In the Green at UCI"]
    assert stats.added == []
    assert stats.added_sections == []  # no new "LEADERSHIP" section created
    titles = [s.title for s in merged.sections if s.kind == "experience"]
    assert titles == ["WORK EXPERIENCE", "LEADERSHIP EXPERIENCE"]
    entry = merged.experience[0]
    assert entry.id == "itg-uci"
    assert entry.bullets[0].text == "Refreshed bullet."


def test_merge_experience_does_not_confuse_companies_sharing_a_slugify_prefix():
    """`config.slugify` truncates at 40 chars; `_match_key` must not, or two distinct
    long company names would slugify to the same string and one would silently
    overwrite the other."""
    name_a = "Global Technology Solutions and Consulting Group North"
    name_b = "Global Technology Solutions and Consulting Group South"
    assert config.slugify(name_a) == config.slugify(name_b)  # the failure mode this guards against

    existing = _resume(
        sections=[
            ExperienceSection(
                id="exp",
                title="Experience",
                entries=[
                    Experience(
                        id="a",
                        company=name_a,
                        title="T",
                        start="2020-01",
                        end="2020-06",
                        bullets=[Bullet(id="a_b1", text="A bullet.", tags=["x"])],
                    ),
                ],
            ),
        ]
    )
    incoming = _resume(
        sections=[
            ExperienceSection(
                id="exp2",
                title="Experience",
                entries=[
                    Experience(
                        company=name_b,
                        title="T2",
                        start="2021-01",
                        end="2021-06",
                        bullets=[Bullet(id="b_b1", text="B bullet.", tags=["y"])],
                    ),
                ],
            ),
        ]
    )

    merged, stats = resume_import.merge_into(existing, incoming)

    assert stats.added == [name_b]
    assert stats.updated == []
    companies = {e.company for e in merged.experience}
    assert companies == {name_a, name_b}


def test_merge_leaves_unmatched_existing_entries_completely_untouched():
    existing_entry = Experience(
        id="foo-inc",
        company="Foo Inc",
        title="Analyst",
        location="NYC",
        start="2019-01",
        end="2020-01",
        bullets=[Bullet(id="foo-inc_b1", text="Did analyst things.", tags=["excel"], metric=True)],
    )
    existing = _resume(sections=[ExperienceSection(id="exp", title="Experience", entries=[existing_entry])])
    incoming = _resume(
        sections=[
            ExperienceSection(
                id="exp2",
                title="Experience",
                entries=[
                    Experience(
                        company="New Co",
                        title="Engineer",
                        start="2022-01",
                        end="2023-01",
                        bullets=[Bullet(id="z_b1", text="New bullet.", tags=["y"])],
                    ),
                ],
            ),
        ]
    )

    merged, stats = resume_import.merge_into(existing, incoming)

    assert stats.added == ["New Co"]
    assert stats.updated == []
    by_company = {e.company: e for e in merged.experience}
    assert by_company["Foo Inc"] == existing_entry  # byte-identical, untouched
    new_entry = by_company["New Co"]
    assert new_entry.id and new_entry.id != "foo-inc"
    assert new_entry.bullets[0].id == f"{new_entry.id}_b1"


def test_merge_adopts_an_empty_existing_section_of_the_right_kind():
    """A fresh workspace's starter resume has an empty default-titled section per
    kind — an unmatched incoming entry should land there, renamed, rather than a
    second same-kind section being created beside it."""
    existing = _resume(sections=[ExperienceSection(id="experience", title="Experience", entries=[])])
    incoming = _resume(
        sections=[
            ExperienceSection(
                id="work",
                title="WORK EXPERIENCE",
                entries=[
                    Experience(
                        company="Acme",
                        title="Engineer",
                        start="2022-01",
                        end="2023-01",
                        bullets=[Bullet(id="a_b1", text="Bullet.", tags=["x"])],
                    ),
                ],
            ),
        ]
    )

    merged, stats = resume_import.merge_into(existing, incoming)

    assert stats.added_sections == []  # adopted, not created
    exp_sections = [s for s in merged.sections if s.kind == "experience"]
    assert len(exp_sections) == 1
    assert exp_sections[0].title == "WORK EXPERIENCE"  # renamed
    assert exp_sections[0].id == "experience"  # the same section, not a new one
    assert len(exp_sections[0].entries) == 1


def test_merge_uses_the_sole_existing_section_of_a_kind_when_no_title_matches():
    existing = _resume(
        sections=[
            SkillsSection(
                id="skills", title="SKILLS", entries=[SkillGroup(label="Technical", items=["Python"])]
            ),
        ]
    )
    incoming = _resume(
        sections=[
            SkillsSection(
                id="info",
                title="ADDITIONAL INFORMATION",
                entries=[SkillGroup(label="Interests", items=["Chess"])],
            ),
        ]
    )

    merged, stats = resume_import.merge_into(existing, incoming)

    assert stats.added_sections == []
    skills_sections = [s for s in merged.sections if s.kind == "skills"]
    assert len(skills_sections) == 1
    assert skills_sections[0].title == "SKILLS"  # not renamed — rule 3, not rule 2
    labels = {g.label for g in skills_sections[0].entries}
    assert labels == {"Technical", "Interests"}


def test_merge_creates_a_new_section_only_when_no_targeting_rule_resolves():
    existing = _resume(
        sections=[
            ExperienceSection(
                id="work",
                title="WORK EXPERIENCE",
                entries=[
                    Experience(
                        id="a", company="A", title="T", start="2020-01", end="2020-06",
                        bullets=[Bullet(id="a_b1", text="X.", tags=["x"])],
                    ),
                ],
            ),
            ExperienceSection(
                id="other",
                title="OTHER ACTIVITIES",
                entries=[
                    Experience(
                        id="b", company="B", title="T", start="2020-01", end="2020-06",
                        bullets=[Bullet(id="b_b1", text="X.", tags=["x"])],
                    ),
                ],
            ),
        ]
    )
    incoming = _resume(
        sections=[
            ExperienceSection(
                id="lead",
                title="LEADERSHIP",
                entries=[
                    Experience(
                        company="C", title="T", start="2021-01", end="2021-06",
                        bullets=[Bullet(id="c_b1", text="Y.", tags=["y"])],
                    ),
                ],
            ),
        ]
    )

    merged, stats = resume_import.merge_into(existing, incoming)

    assert stats.added_sections == ["LEADERSHIP"]
    assert stats.added == ["C"]
    titles = [s.title for s in merged.sections if s.kind == "experience"]
    assert titles == ["WORK EXPERIENCE", "OTHER ACTIVITIES", "LEADERSHIP"]


def test_merge_pairs_duplicate_match_keys_in_document_order():
    """Two existing entries at the same company (two stints) pair with two incoming
    entries of the same company in document order, not by any other heuristic."""
    existing = _resume(
        sections=[
            ExperienceSection(
                id="exp",
                title="Experience",
                entries=[
                    Experience(
                        id="acme-1", company="Acme", title="Intern", start="2018-01", end="2018-06",
                        bullets=[Bullet(id="acme-1_b1", text="First stint.", tags=["x"])],
                    ),
                    Experience(
                        id="acme-2", company="Acme", title="Engineer", start="2020-01", end="2021-01",
                        bullets=[Bullet(id="acme-2_b1", text="Second stint.", tags=["x"])],
                    ),
                ],
            ),
        ]
    )
    incoming = _resume(
        sections=[
            ExperienceSection(
                id="exp2",
                title="Experience",
                entries=[
                    Experience(
                        company="Acme", title="Intern (updated)", start="2018-01", end="2018-06",
                        bullets=[Bullet(id="z1_b1", text="First stint refreshed.", tags=["y"])],
                    ),
                    Experience(
                        company="Acme", title="Engineer (updated)", start="2020-01", end="2021-01",
                        bullets=[Bullet(id="z2_b1", text="Second stint refreshed.", tags=["y"])],
                    ),
                ],
            ),
        ]
    )

    merged, stats = resume_import.merge_into(existing, incoming)

    assert stats.updated == ["Acme", "Acme"]
    assert stats.added == []
    entries = merged.experience
    assert len(entries) == 2
    assert entries[0].id == "acme-1"
    assert entries[0].title == "Intern (updated)"
    assert entries[1].id == "acme-2"
    assert entries[1].title == "Engineer (updated)"


def test_merge_skills_group_items_are_replaced_on_label_match():
    existing = _resume(
        sections=[
            SkillsSection(
                id="skills", title="Skills", entries=[SkillGroup(label="Languages", items=["English"])]
            ),
        ]
    )
    incoming = _resume(
        sections=[
            SkillsSection(
                id="skills2",
                title="Skills",
                entries=[SkillGroup(label="languages", items=["English", "Spanish"])],  # case differs
            ),
        ]
    )

    merged, stats = resume_import.merge_into(existing, incoming)

    assert stats.updated == ["Languages"]  # existing casing is what's reported
    group = merged.skills[0]
    assert group.label == "Languages"  # existing label wording preserved
    assert group.items == ["English", "Spanish"]  # items refreshed


def test_merge_list_items_dedupe_by_exact_text():
    existing = _resume(
        sections=[
            ListSection(
                id="certs", title="Certifications", entries=[ListItem(id="cert-1", text="AWS Certified", tags=["aws"])]
            ),
        ]
    )
    incoming = _resume(
        sections=[
            ListSection(
                id="certs2",
                title="Certifications",
                entries=[
                    ListItem(id="whatever", text="AWS Certified", tags=["aws"]),  # duplicate
                    ListItem(id="whatever2", text="PMP Certified", tags=["pmp"]),  # new
                ],
            ),
        ]
    )

    merged, stats = resume_import.merge_into(existing, incoming)

    assert stats.added == ["PMP Certified"]
    entries = merged.sections[0].entries
    assert [i.text for i in entries] == ["AWS Certified", "PMP Certified"]
    assert len(entries) == 2  # no duplicate


def test_merge_contact_only_overwrites_fields_incoming_actually_has():
    existing = _resume(
        contact=Contact(
            name="Old Name",
            email="old@example.com",
            phone="555-1234",
            location="Old City",
            linkedin="https://linkedin.com/in/old",
            github="",
        )
    )
    incoming = _resume(
        contact=Contact(
            name="New Name",
            email="new@example.com",
            phone="",  # nothing detected
            location="",  # nothing detected
            linkedin="",  # nothing detected
            github="",
        )
    )

    merged, _stats = resume_import.merge_into(existing, incoming)

    assert merged.contact.name == "New Name"
    assert merged.contact.email == "new@example.com"
    assert merged.contact.phone == "555-1234"  # kept — incoming had none
    assert merged.contact.location == "Old City"  # kept
    assert merged.contact.linkedin == "https://linkedin.com/in/old"  # kept


def test_merge_preserves_summary_variants_and_comment_and_unions_tag_vocabulary():
    existing = MasterResume(
        comment="hand-maintained notes",
        contact=_contact(),
        summary_variants=[SummaryVariant(id="ic", text="IC framing", tags=["zzzic"])],
        sections=[
            ExperienceSection(
                id="exp",
                title="Experience",
                entries=[
                    Experience(
                        id="a", company="A", title="T", start="2020-01", end="2020-06",
                        bullets=[Bullet(id="a_b1", text="X.", tags=["zzzpython"])],
                    ),
                ],
            ),
        ],
        tag_vocabulary=["zzzpython", "zzzsql"],
    )
    incoming = _resume(
        sections=[
            ExperienceSection(
                id="exp2",
                title="Experience",
                entries=[
                    Experience(
                        company="B", title="T", start="2021-01", end="2021-06",
                        bullets=[Bullet(id="b_b1", text="Y.", tags=["zzzexcel"])],
                    ),
                ],
            ),
        ],
        tag_vocabulary=["zzzexcel", "zzzpython"],
    )

    merged, _stats = resume_import.merge_into(existing, incoming)

    assert merged.comment == "hand-maintained notes"
    assert merged.summary_variants == existing.summary_variants
    assert merged.tag_vocabulary == ["zzzexcel", "zzzpython", "zzzsql"]


def test_merge_education_near_miss_school_updates_in_place_with_coursework():
    """The real-world case that motivated near-miss matching: an incoming school string
    that is the existing one plus a school/college suffix matches and updates in place —
    coursework the existing entry lacked comes through, and the curated GPA the .docx
    could not express survives (see `_merge_education_entry`)."""
    existing = _resume(
        sections=[
            EducationSection(
                id="edu",
                title="EDUCATION",
                entries=[
                    Education(
                        school="University of California, Irvine",
                        degree="B.A. in Business Administration",
                        dates="Expected graduation: June 2027",
                        location="Irvine, CA",
                        gpa="3.92",
                        show_gpa=True,
                        details=["Dean's List"],
                    ),
                ],
            ),
        ]
    )
    incoming = _resume(
        sections=[
            EducationSection(
                id="edu2",
                title="EDUCATION",
                entries=[
                    Education(
                        school="University of California, Irvine --- Paul Merage School of Business",
                        degree="Bachelor of Arts in Business Administration",
                        dates="Class of 2027",
                        location="CA",
                        coursework=["Financial Accounting", "Management Science", "Intro to Consulting"],
                        details=["Cumulative GPA: 3.9/4.0; Major GPA: 3.88/4.0"],
                    ),
                ],
            ),
        ]
    )

    merged, stats = resume_import.merge_into(existing, incoming)

    assert stats.added == []
    assert stats.updated == ["University of California, Irvine --- Paul Merage School of Business"]
    assert len(merged.education) == 1
    entry = merged.education[0]
    assert entry.school == "University of California, Irvine --- Paul Merage School of Business"
    assert entry.coursework == ["Financial Accounting", "Management Science", "Intro to Consulting"]
    # Curated GPA survives: the incoming entry has no `gpa`, since the .docx wrote it as
    # a free-text detail line, not the `| GPA: …` form `_GPA_RE` recognizes.
    assert entry.gpa == "3.92"
    assert entry.show_gpa is True
    assert any("near-miss" in w for w in stats.warnings)


def test_merge_education_empty_incoming_field_never_blanks_a_populated_one():
    existing = _resume(
        sections=[
            EducationSection(
                id="edu",
                title="EDUCATION",
                entries=[
                    Education(
                        school="Acme University",
                        degree="B.S. Computer Science",
                        dates="2020-2024",
                        location="Springfield",
                        details=["Honors"],
                    ),
                ],
            ),
        ]
    )
    incoming = _resume(
        sections=[
            EducationSection(
                id="edu2",
                title="EDUCATION",
                entries=[Education(school="Acme University", degree="", dates="", location="")],
            ),
        ]
    )

    merged, stats = resume_import.merge_into(existing, incoming)

    entry = merged.education[0]
    assert entry.degree == "B.S. Computer Science"
    assert entry.dates == "2020-2024"
    assert entry.location == "Springfield"
    assert entry.details == ["Honors"]


def test_merge_education_exact_beats_near_miss():
    """When the existing resume already holds both the short and the long spelling of
    one school, an incoming long spelling must claim its exact twin — not fuzzily match
    the short entry too, which would leave one of the two existing entries orphaned."""
    existing = _resume(
        sections=[
            EducationSection(
                id="edu",
                title="EDUCATION",
                entries=[
                    Education(school="University of California, Irvine", degree="BA", dates="2024"),
                    Education(
                        school="University of California, Irvine --- Paul Merage School of Business",
                        degree="Old Degree Text",
                        dates="2024",
                    ),
                ],
            ),
        ]
    )
    incoming = _resume(
        sections=[
            EducationSection(
                id="edu2",
                title="EDUCATION",
                entries=[
                    Education(
                        school="University of California, Irvine --- Paul Merage School of Business",
                        degree="New Degree Text",
                        dates="2027",
                    ),
                ],
            ),
        ]
    )

    merged, stats = resume_import.merge_into(existing, incoming)

    assert stats.added == []
    assert stats.updated == ["University of California, Irvine --- Paul Merage School of Business"]
    assert not any("near-miss" in w for w in stats.warnings)
    by_school = {e.school: e for e in merged.education}
    assert by_school["University of California, Irvine"].degree == "BA"
    assert by_school["University of California, Irvine --- Paul Merage School of Business"].degree == (
        "New Degree Text"
    )


def test_merge_education_ambiguous_near_miss_is_added_not_guessed():
    """Two existing entries both near-miss one incoming school — refuse to guess which
    one it means; add it as new and warn, rather than silently picking one."""
    existing = _resume(
        sections=[
            EducationSection(
                id="edu",
                title="EDUCATION",
                entries=[
                    Education(school="Acme University North Campus", degree="BA", dates="2020"),
                    Education(school="Acme University South Campus", degree="BS", dates="2021"),
                ],
            ),
        ]
    )
    incoming = _resume(
        sections=[
            EducationSection(
                id="edu2",
                title="EDUCATION",
                entries=[Education(school="Acme University", degree="MS", dates="2027")],
            ),
        ]
    )

    merged, stats = resume_import.merge_into(existing, incoming)

    assert stats.updated == []
    assert stats.added == ["Acme University"]
    assert len(merged.education) == 3
    assert any("matched" in w.lower() and "existing entries" in w for w in stats.warnings)


def test_merge_education_genuinely_distinct_schools_stay_two_entries():
    existing = _resume(
        sections=[
            EducationSection(
                id="edu",
                title="EDUCATION",
                entries=[
                    Education(school="University of California, Irvine", degree="BA", dates="2024"),
                ],
            ),
        ]
    )
    incoming = _resume(
        sections=[
            EducationSection(
                id="edu2",
                title="EDUCATION",
                entries=[
                    Education(school="University of California, Los Angeles", degree="BS", dates="2027"),
                ],
            ),
        ]
    )

    merged, stats = resume_import.merge_into(existing, incoming)

    assert stats.updated == []
    assert stats.added == ["University of California, Los Angeles"]
    schools = {e.school for e in merged.education}
    assert schools == {
        "University of California, Irvine",
        "University of California, Los Angeles",
    }


def test_merge_experience_near_miss_scoping_does_not_apply_to_companies():
    """Near-miss matching is education-only — a company name differing by the same kind
    of suffix must still be treated as a distinct entry, not fuzzily merged."""
    existing = _resume(
        sections=[
            ExperienceSection(
                id="exp",
                title="EXPERIENCE",
                entries=[
                    Experience(
                        id="acme-id",
                        company="Acme Corp",
                        title="Engineer",
                        start="2020",
                        end="2021",
                    ),
                ],
            ),
        ]
    )
    incoming = _resume(
        sections=[
            ExperienceSection(
                id="exp2",
                title="EXPERIENCE",
                entries=[
                    Experience(
                        id="acme2-id",
                        company="Acme Corp Consulting Division",
                        title="Consultant",
                        start="2022",
                        end="2023",
                    ),
                ],
            ),
        ]
    )

    merged, stats = resume_import.merge_into(existing, incoming)

    assert stats.updated == []
    assert stats.added == ["Acme Corp Consulting Division"]
    companies = {e.company for e in merged.experience}
    assert companies == {"Acme Corp", "Acme Corp Consulting Division"}
