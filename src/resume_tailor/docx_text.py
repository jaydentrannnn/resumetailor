"""Map `Paragraph.text` character offsets onto the runs that produced them.

`python-docx`'s `Paragraph.text` (`CT_P.text`) is built by walking `w:r | w:hyperlink`
children and concatenating their `.text` — a hyperlink's visible text is included even
though `Paragraph.runs` (`self._p.r_lst`) only lists top-level `w:r` elements and skips
runs nested inside `w:hyperlink`. Any code that measures character spans against
`paragraph.text` but then edits via `paragraph.runs` silently mis-slices everything
after the first hyperlink — the project date's tab stop was lost exactly this way,
because a span computed before `strip_hyperlinks()` ran was applied after it, once the
hyperlink's characters (and therefore every later offset) no longer existed.

This module is the single place that reconciles the two: every offset here is measured
in the same space as `paragraph.text`, including hyperlink text.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from typing import Iterator

from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

_R = qn("w:r")
_HYPERLINK = qn("w:hyperlink")
_P_TAG = qn("w:p")
_TBL_TAG = qn("w:tbl")
_TC_TAG = qn("w:tc")
_TR_TAG = qn("w:tr")

#: Blank, or a decorative rule made only of underscores/dashes/dots/middots.
_CHROME_RE = re.compile(r"[\s_\-–—=·.]+")


def is_chrome_text(text: str) -> bool:
    """True for a paragraph carrying no content: blank, or a decorative rule line.

    Lives here rather than in `template_analyze` because `template_build` needs the
    identical judgement — the analyzer records which chrome paragraphs to reproduce as
    spacers, and the builder clones exactly those. Two copies of this rule would drift,
    and the failure mode is silent (a rule paragraph read as an entry header, or a
    recorded spacer donor the builder then refuses to clone).
    """
    stripped = text.strip()
    return not stripped or bool(_CHROME_RE.fullmatch(stripped))


@dataclass(frozen=True)
class RunSlice:
    """One `w:r` element and the half-open `paragraph.text` range it produced.

    `element` is the raw `w:r` oxml element (not a `docx.text.run.Run`) so callers can
    read or clone its `w:rPr` directly. `hyperlink` is the enclosing `w:hyperlink`
    element when this run lives inside one, else `None`.
    """

    element: object
    start: int
    end: int
    hyperlink: object | None = None

    @property
    def in_hyperlink(self) -> bool:
        """True when this run's text is part of a hyperlink's visible text."""
        return self.hyperlink is not None

    @property
    def is_empty(self) -> bool:
        """True for a zero-length run (e.g. a bare `<w:tab/>` run with no `w:t`)."""
        return self.start == self.end


def paragraph_text(paragraph: Paragraph) -> str:
    """`paragraph.text` — the offset space every `CharSpan` is measured in.

    Named explicitly to flag the contract at every call site: this INCLUDES hyperlink
    visible text. Pair it with `paragraph_run_slices`, never with `paragraph.runs`.
    """
    return paragraph.text or ""


def paragraph_run_slices(paragraph: Paragraph) -> list[RunSlice]:
    """Map every run — including ones nested in a hyperlink — onto `paragraph_text` offsets.

    Walks `./w:r | ./w:hyperlink` in document order, the same xpath `CT_P.text` uses,
    descending into each hyperlink's own `./w:r` children (mirroring `CT_Hyperlink.text`).
    Cumulative offsets therefore line up with `paragraph_text` exactly, tabs and other
    special children counting as however many characters `w:r.text` renders them as.
    """
    slices: list[RunSlice] = []
    cursor = 0
    for child in paragraph._p.xpath("w:r | w:hyperlink"):
        if child.tag == _R:
            length = len(child.text or "")
            slices.append(RunSlice(element=child, start=cursor, end=cursor + length))
            cursor += length
        elif child.tag == _HYPERLINK:
            for run_el in child.xpath("w:r"):
                length = len(run_el.text or "")
                slices.append(
                    RunSlice(
                        element=run_el,
                        start=cursor,
                        end=cursor + length,
                        hyperlink=child,
                    )
                )
                cursor += length
    return slices


def hyperlink_char_spans(paragraph: Paragraph) -> list[tuple[int, int]]:
    """Char ranges of each `w:hyperlink`'s visible text within `paragraph_text`.

    Multiple runs inside one hyperlink are merged into a single contiguous span.
    """
    spans: list[tuple[int, int]] = []
    current: object | None = None
    start = end = 0
    for s in paragraph_run_slices(paragraph):
        if s.hyperlink is current and current is not None:
            end = s.end
            continue
        if current is not None:
            spans.append((start, end))
        current = s.hyperlink
        if current is not None:
            start, end = s.start, s.end
    if current is not None:
        spans.append((start, end))
    return spans


def hyperlink_target(paragraph: Paragraph, hyperlink_element) -> str | None:
    """Resolve a `w:hyperlink` element's `r:id` to its target URL via the paragraph
    part's relationships. Returns `None` for an internal bookmark reference (`w:anchor`,
    no `r:id`) or a dangling relationship id — a caller wants the label text either way,
    so a missing target should degrade to "no link", not raise.
    """
    r_id = hyperlink_element.get(qn("r:id"))
    if not r_id:
        return None
    try:
        return paragraph.part.rels[r_id].target_ref
    except KeyError:
        return None


def slice_at(
    slices: list[RunSlice], offset: int, *, skip_hyperlinks: bool = False
) -> RunSlice | None:
    """Return the run that should donate formatting for text at `offset`.

    Prefers a non-empty slice actually containing `offset`; falls back to the last slice
    starting at or before it (formatting for a position right at a run boundary, or past
    the end of the paragraph, comes from what precedes it); falls back to the first slice
    when nothing starts before `offset` (position 0 of an empty-first-run paragraph).
    Zero-length slices never win a lookup — they exist only to carry structural elements
    like a bare `<w:tab/>` and have no formatting decision to make.

    `skip_hyperlinks=True` ignores hyperlink-nested runs entirely, falling back to the
    nearest preceding non-hyperlink run: a Jinja tag must never inherit a link's blue/
    underline styling, since the real per-item link is reinstated separately at render
    time as a `RichText`.
    """
    candidates = [s for s in slices if s.hyperlink is None] if skip_hyperlinks else slices
    if not candidates:
        candidates = slices
    if not candidates:
        return None

    for s in candidates:
        if s.start <= offset < s.end and not s.is_empty:
            return s

    before = [s for s in candidates if s.start <= offset]
    if before:
        return max(before, key=lambda s: s.start)

    return candidates[0]


@dataclass(frozen=True)
class ParaLocation:
    """Where a paragraph sits inside a table, in *physical* cells.

    `cell` indexes `<w:tc>` children of the row directly — NOT `Row.cells`, which
    python-docx expands over `gridSpan` so a 3-span-plus-1 row reports as four cells,
    two of them the same object. Every consumer of this dataclass asks "does this row
    have a second populated cell", which only the physical count answers correctly.
    """

    table: int
    row: int
    cell: int
    para: int
    row_cells: int
    row_content_cells: int


def has_nested_tables(doc) -> bool:
    """True when any top-level table contains a `w:tbl` inside one of its cells.

    A resume built this way cannot be read as a single reading column with any
    confidence, so callers use this to block rather than attempt to flatten it.
    """
    for table in doc.tables:
        if table._tbl.find(f".//{_TC_TAG}/{_TBL_TAG}") is not None:
            return True
    return False


def iter_document_paragraphs(
    doc,
) -> Iterator[tuple[Paragraph, ParaLocation | None]]:
    """Every paragraph in the document body, in reading order.

    Walks `w:body`'s direct children in document order; a `w:p` is yielded as-is
    (`location=None`), a `w:tbl` is descended into as rows -> physical cells ->
    paragraphs. This is THE invariant the whole template system rests on:
    `template_analyze._load_paras` and `template_build._para_by_id` must both use this
    walk and therefore enumerate identically, because a `CharSpan.paragraph_id` is an
    index into this sequence and nothing else.

    A nested table (a `w:tbl` inside a `w:tc`) is not descended into — callers must
    check `has_nested_tables` first and block rather than silently skip its content.
    """
    body = doc.element.body
    table_idx = 0
    for child in body.iterchildren():
        if child.tag == _P_TAG:
            yield Paragraph(child, doc), None
        elif child.tag == _TBL_TAG:
            table = Table(child, doc)
            rows = child.findall(_TR_TAG)
            for row_idx, tr in enumerate(rows):
                tc_list = tr.findall(_TC_TAG)
                row_content_cells = sum(
                    1
                    for tc in tc_list
                    if any(
                        not is_chrome_text(p.text or "")
                        for p in _Cell(tc, table).paragraphs
                    )
                )
                for cell_idx, tc in enumerate(tc_list):
                    for para_idx, paragraph in enumerate(_Cell(tc, table).paragraphs):
                        yield paragraph, ParaLocation(
                            table=table_idx,
                            row=row_idx,
                            cell=cell_idx,
                            para=para_idx,
                            row_cells=len(tc_list),
                            row_content_cells=row_content_cells,
                        )
            table_idx += 1
