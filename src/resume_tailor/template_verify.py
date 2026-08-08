"""Post-build verification: does a tagged template actually contain what tagging
claims to have produced, and does a real render's content actually reach the output?

Two independent checks, both pure functions of their explicit arguments (a built
`.docx` path, a confirmed `TemplateProfile`, and — for the roundtrip check — a
`MasterResume`) — no dependency on `config`'s path globals, no mutation.
`web/template_ops.py` runs both after `_smoke_render` and before a staged install
replaces the live template; a blocking issue from either aborts the install the same
way a smoke-render failure already does.

`verify_tagged` alone would have caught the bug that motivated this module: a resume
with a `PROFESSIONAL SUMMARY` heading made the analyzer anchor the experience prototype
on the summary sentence instead of a real job entry, and the mapping it produced had
`dates`/`location` both `present=False` on every entry. A template built from that
mapping is syntactically fine and renders without error — it just silently drops every
job's dates. `expected_tags` would have flagged `{{ job.dates }}` as missing before the
bad template ever reached the live slot.
"""

from __future__ import annotations

import tempfile
from pathlib import Path

import docx
from docx.oxml.ns import qn

from . import docx_text
from . import render as render_mod
from . import template_profile
from .data import MasterResume
from .template_analyze import Issue
from .template_build import (
    BULLET_TAG,
    CONTACT_SLOT_TAG_FMT,
    CONTACT_TAG,
    EDUCATION_DEGREE_TAG,
    EDUCATION_DETAIL_TAG,
    EDUCATION_HEADER_TAGS,
    EXPERIENCE_HEADER_TAGS,
    EXPERIENCE_TITLE_TAG,
    LIST_ITEM_TAG,
    NAME_TAG,
    PROJECT_HEADER_TAGS,
    PROJECT_LINK_TAG,
    SECTION_LOOP_OPEN,
    SECTION_LOOP_OPEN_TR,
    SECTION_TITLE_TAG,
    SKILLS_BODY_TAG,
    SKILLS_LABEL_TAG,
)
from .template_profile import TemplateProfile

_TC_TAG = qn("w:tc")
_P_TAG = qn("w:p")


def expected_tags(profile: TemplateProfile) -> set[str]:
    """Every Jinja tag a correctly-built template must contain for `profile`.

    Mirrors exactly what `template_build`'s `_tag_*_prototype` functions emit, reading
    the same module constants those functions tag with, so this can never drift into a
    second, hand-maintained copy of the list. Only fields the profile actually mapped
    (`OptionalSpan.present`) are expected — an unmapped field (no `dates` detected in
    the upload) correctly has no corresponding tag to look for; whether *that* is
    itself a bug is `template_analyze`'s job, not this module's. This module only holds
    tagging accountable to whatever mapping it was actually given.
    """
    if profile.contact.slots:
        tags = {NAME_TAG} | {CONTACT_SLOT_TAG_FMT % i for i in range(len(profile.contact.slots))}
    else:
        tags = {NAME_TAG, CONTACT_TAG}

    exp = profile.experience
    for name, field in exp.header.fields.items():
        if field.present and name in EXPERIENCE_HEADER_TAGS:
            tags.add(EXPERIENCE_HEADER_TAGS[name])
    if exp.title.present:
        tags.add(EXPERIENCE_TITLE_TAG)
    tags.add(BULLET_TAG)

    if profile.enabled.education and profile.education is not None:
        edu = profile.education
        for name, field in edu.header.fields.items():
            if field.present and name in EDUCATION_HEADER_TAGS:
                tags.add(EDUCATION_HEADER_TAGS[name])
        tags.add(EDUCATION_DEGREE_TAG)
        tags.add(EDUCATION_DETAIL_TAG)

    if profile.enabled.projects and profile.projects is not None:
        proj = profile.projects
        for name, field in proj.header.fields.items():
            if field.present and name in PROJECT_HEADER_TAGS:
                tags.add(PROJECT_HEADER_TAGS[name])
        if proj.link.present:
            tags.add(PROJECT_LINK_TAG)
        tags.add(BULLET_TAG)

    if profile.enabled.skills and profile.skills is not None:
        tags.add(SKILLS_LABEL_TAG)
        tags.add(SKILLS_BODY_TAG)

    if profile.enabled.list_section and profile.list_section is not None:
        tags.add(LIST_ITEM_TAG)

    if profile.section_mode == "generic":
        tags.add(SECTION_TITLE_TAG)

    return tags


def _check_control_tags_balanced(texts: list[str]) -> list[Issue]:
    """Walk paragraph texts as one token stream and verify every `{%p for %}`/
    `{%p if %}` (and, under a table layout, `{%tr for %}`/`{%tr if %}`) closes before
    an ancestor's close does.

    A plain open/close *count* match would still pass a misnested (as opposed to
    merely incomplete) structure; walking a stack catches both. An imbalance here means
    invalid Jinja that would fail at render time, or worse, silently mis-nest.

    One stack for both `{%p %}` and `{%tr %}` levels is correct, not a simplification:
    `texts` is already the flattened reading-order walk (rows in order, cells
    left-to-right, paragraphs in order within a cell), so a `{%p for bullet %}` /
    `{%p endfor %}` pair opens and closes entirely inside one cell — strictly between
    its own row's `{%tr for job %}` and `{%tr endfor %}` markers. The two levels are
    never actually interleaved in the token stream, only nested.
    """
    stack: list[tuple[str, str]] = []
    issues: list[Issue] = []
    for t in texts:
        stripped = t.strip()
        if stripped.startswith(("{%p for", "{%tr for")):
            stack.append(("for", stripped))
        elif stripped.startswith(("{%p if", "{%tr if")):
            stack.append(("if", stripped))
        elif stripped in ("{%p endfor %}", "{%tr endfor %}"):
            if stack and stack[-1][0] == "for":
                stack.pop()
            else:
                issues.append(
                    Issue(
                        code="unbalanced_control_tags",
                        message=(
                            f"Found {stripped!r} with no matching open for-loop "
                            f"(open stack: {[s[1] for s in stack]})."
                        ),
                        blocking=True,
                    )
                )
        elif stripped in ("{%p endif %}", "{%tr endif %}"):
            if stack and stack[-1][0] == "if":
                stack.pop()
            else:
                issues.append(
                    Issue(
                        code="unbalanced_control_tags",
                        message=(
                            f"Found {stripped!r} with no matching open if-block "
                            f"(open stack: {[s[1] for s in stack]})."
                        ),
                        blocking=True,
                    )
                )
    if stack:
        issues.append(
            Issue(
                code="unbalanced_control_tags",
                message=f"{len(stack)} control tag(s) never closed: {[s[1] for s in stack]}.",
                blocking=True,
            )
        )
    return issues


def _check_no_empty_cells(doc) -> list[Issue]:
    """A `<w:tc>` with no `<w:p>` child is invalid OOXML — Word refuses such a file
    with a generic "cannot open" error, but `python-docx` opens it happily, making this
    exactly the class of bug that would otherwise reach a user unnoticed. The two ways
    to produce one in a table-layout build: blanking a cloned row's cell by removing
    its last paragraph instead of just its text, or a cell whose only content is a
    `{%p for %}` loop that renders zero iterations.
    """
    issues: list[Issue] = []
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            for tc in row._tr.findall(_TC_TAG):
                if tc.find(_P_TAG) is None:
                    issues.append(
                        Issue(
                            code="empty_table_cell",
                            message=(
                                f"A table cell in row {row_idx} has no paragraphs at "
                                "all, which is invalid OOXML (Word will refuse to open "
                                "the file)."
                            ),
                            blocking=True,
                        )
                    )
    return issues


def verify_tagged(tagged: Path, profile: TemplateProfile) -> list[Issue]:
    """Re-open a built template and check it against what `profile` demands.

    All-blocking checks: every tag `expected_tags` names is present somewhere (as a
    substring of some paragraph's text — a tag may share a paragraph with literal
    text); `{%p for %}`/`{%p if %}` control tags are balanced and correctly nested;
    under generic mode, exactly one outer `{%p for section in sections %}` (zero means
    nothing will ever render; more than one is a structural bug this module has no
    other way to catch); and no baked-in `w:hyperlink` survives anywhere — every
    tagging path that ever touches one (`build_contact_profile`, `_tag_mapped_header`
    on a project header) strips it, since the real per-item URL is reinstated only at
    render time as a `RichText`. A leftover would bake one person's link into every
    rendered document.
    """
    doc = docx.Document(str(tagged))
    texts = [p.text for p, _location in docx_text.iter_document_paragraphs(doc)]
    joined = "\n".join(texts)

    issues: list[Issue] = []

    for tag in sorted(expected_tags(profile)):
        if tag not in joined:
            issues.append(
                Issue(
                    code="tag_missing",
                    message=(
                        f"Expected tag {tag!r} does not appear anywhere in the built "
                        "template. A field the profile marked present was not tagged."
                    ),
                    blocking=True,
                )
            )

    issues.extend(_check_control_tags_balanced(texts))

    if profile.layout == "table":
        issues.extend(_check_no_empty_cells(doc))

    if profile.section_mode == "generic":
        expected_loop_open = SECTION_LOOP_OPEN_TR if profile.layout == "table" else SECTION_LOOP_OPEN
        loop_opens = sum(1 for t in texts if t.strip() == expected_loop_open)
        if loop_opens != 1:
            issues.append(
                Issue(
                    code="section_loop_count",
                    message=f"Expected exactly one {expected_loop_open!r}, found {loop_opens}.",
                    blocking=True,
                )
            )

    hyperlink_count = len(list(doc.element.body.iter(qn("w:hyperlink"))))
    if hyperlink_count:
        issues.append(
            Issue(
                code="leftover_hyperlink",
                message=(
                    f"{hyperlink_count} baked-in hyperlink(s) survive in the built "
                    "template; every per-item link must come from render-time "
                    "RichText, not a literal hyperlink left over from tagging."
                ),
                blocking=True,
            )
        )

    return issues


def verify_roundtrip(
    tagged: Path, profile: TemplateProfile, resume: MasterResume
) -> list[Issue]:
    """Render `resume` through `tagged` and confirm each mapped field's *value* reaches
    the output — catches a mapping that tags correctly (passing `verify_tagged`) but
    points at the wrong span, which builds and renders without error, just with the
    wrong text sitting where a field should be.

    Only checks entries that will actually render (at least one bullet, matching
    `render.build_context`'s own filtering) and only fields the profile mapped as
    present — an unmapped field has no expected value to look for.

    Renders with `layout=template_profile.active_layout(profile)` explicitly, not
    `render.render`'s default disk-read: `profile` here is the one being verified,
    which may not be (or not yet be) the one installed at `config.
    TEMPLATE_PROFILE_PATH` — a staged install verifying itself before commit must not
    silently pick up whatever profile happens to already be live.
    """
    layout = template_profile.active_layout(profile)
    with tempfile.TemporaryDirectory() as tmp:
        out = Path(tmp) / "roundtrip.docx"
        render_mod.render(resume, template=tagged, out=out, layout=layout)
        doc = docx.Document(str(out))
        body_text = "\n".join(p.text for p, _location in docx_text.iter_document_paragraphs(doc))

    issues: list[Issue] = []
    if profile.layout == "table":
        issues.extend(_check_no_empty_cells(doc))

    def _check(value: str, label: str) -> None:
        if value and value not in body_text:
            issues.append(
                Issue(
                    code="roundtrip_missing",
                    message=f"{label} {value!r} did not reach the rendered output.",
                    blocking=True,
                )
            )

    exp_fields = profile.experience.header.fields
    for entry in resume.experience:
        if not entry.bullets:
            continue
        _check(entry.company, f"Experience company for {entry.company!r}")
        if profile.experience.title.present:
            _check(entry.title, f"Experience title for {entry.company!r}")
        location = exp_fields.get("location")
        if location is not None and location.present:
            _check(entry.location, f"Experience location for {entry.company!r}")
        dates = exp_fields.get("dates")
        if dates is not None and dates.present:
            _check(
                render_mod.format_range(entry.start, entry.end),
                f"Experience dates for {entry.company!r}",
            )

    if profile.enabled.education and profile.education is not None:
        edu_fields = profile.education.header.fields
        for edu in resume.education:
            _check(edu.school, f"Education school for {edu.school!r}")
            dates = edu_fields.get("dates")
            if dates is not None and dates.present:
                _check(edu.dates, f"Education dates for {edu.school!r}")

    if profile.enabled.projects and profile.projects is not None:
        proj_fields = profile.projects.header.fields
        for proj in resume.projects:
            if not proj.bullets:
                continue
            _check(proj.name, f"Project name for {proj.name!r}")
            date = proj_fields.get("date")
            if date is not None and date.present:
                _check(proj.date, f"Project date for {proj.name!r}")
            if profile.projects.link.present:
                _check(proj.link, f"Project link label for {proj.name!r}")

    return issues
