"""Generate a tagged `main_template.docx` from a baseline export.

This module is the ONLY producer of the tagged template (see CLAUDE.md). The CLI
`scripts/build_template.py` is a thin wrapper.

Two modes:
- **Legacy** (no profile): exact all-caps headings and the original Google Docs layout.
- **Profile** (confirmed mapping): tags only mapped character spans so separators and
  tabs from the upload survive; optional sections may be omitted.
"""

from __future__ import annotations

import copy
from dataclasses import dataclass
from pathlib import Path
from typing import Literal

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from . import config, docx_text
from .docx_text import RunSlice
from .template_profile import CharSpan, OptionalSpan, TemplateProfile

W = qn("w:p")
NOTO_MARKER_FONT = "Noto Sans Symbols"

# --------------------------------------------------------------------------------------
# Jinja tag strings
#
# Hoisted to module constants — not just inlined at each tagging call site — so
# `template_verify.py` can assert a built template actually contains what tagging
# claims to have produced, reading the exact same source rather than a second,
# driftable copy of this list.
# --------------------------------------------------------------------------------------

NAME_TAG = "{{ name }}"
CONTACT_TAG = "{{r contact }}"
#: One RichText per table-layout contact slot — named keys (`contact_slot_0`, …), not
#: a subscript into a list: a subscript past the end of a short list renders as a
#: silent empty string under Jinja's `Undefined`, where a missing named key would at
#: least be a KeyError during development. See `TemplateProfile.contact.slots`.
CONTACT_SLOT_TAG_FMT = "{{r contact_slot_%d }}"

EXPERIENCE_HEADER_TAGS: dict[str, str] = {
    "company": "{{ job.company }}",
    "location": "{{ job.location }}",
    "dates": "{{ job.dates }}",
}
EXPERIENCE_TITLE_TAG = "{{ job.title }}"
#: Shared with `PROJECT_BULLET_TAG` deliberately — each is a different `{%p for bullet
#: in ... %}` loop's own variable, so the identical tag *text* is not evidence of one
#: loop leaking into another; `expected_tags` treats it as presence-only, never an
#: exactly-once count, for this reason.
BULLET_TAG = "{{ bullet }}"

EDUCATION_HEADER_TAGS: dict[str, str] = {
    "school": "{{ edu.school }}",
    "location": "{{ edu.location }}",
    "dates": "{{ edu.dates }}",
}
EDUCATION_DEGREE_TAG = "{{ edu.degree_line }}"
EDUCATION_DETAIL_TAG = "{{ detail }}"

PROJECT_HEADER_TAGS: dict[str, str] = {
    "name": "{{ proj.name }}",
    "tech": "{{ proj.tech }}",
    "date": "{{ proj.date }}",
}
PROJECT_LINK_TAG = "{{r proj.link }}"
PROJECT_BULLET_TAG = BULLET_TAG

SKILLS_LABEL_TAG = "{{ group.label }}"
SKILLS_BODY_TAG = "{{ group.entries }}"

LIST_ITEM_TAG = "{{ item }}"

#: Generic-mode only: the shared heading clone's tag, and the outer per-section loop.
SECTION_TITLE_TAG = "{{ section.title }}"
SECTION_LOOP_OPEN = "{%p for section in sections %}"
#: Table-layout counterpart of `SECTION_LOOP_OPEN` — a ROW-level loop, since a table
#: layout's shared block repeats table rows, not paragraphs. See `build_generic_table`.
SECTION_LOOP_OPEN_TR = "{%tr for section in sections %}"


# --------------------------------------------------------------------------------------
# Low-level XML helpers
# --------------------------------------------------------------------------------------


def make_para(text: str) -> OxmlElement:
    """Build a bare paragraph carrying a `{%p %}` control tag.

    The `p` suffix is load-bearing, not cosmetic. A plain `{% for %}` is substituted
    inline, so the loop body starts mid-paragraph and each iteration re-emits the
    surrounding `<w:p>` fragments — nesting a paragraph inside a paragraph. The result
    is well-formed XML that violates the OOXML schema, and Word refuses to open it with
    a generic "Word experienced an error trying to open the file". `{%p %}` makes
    docxtpl drop the whole enclosing paragraph instead, so loops span entire paragraphs.

    Deliberately style-less: these paragraphs are deleted at render time, so formatting
    on them would be wasted, and inheriting a list style would leave a stray bullet
    behind if a tag were ever mistyped.
    """
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    p.append(r)
    return p


def delete(paragraph: Paragraph) -> None:
    paragraph._p.getparent().remove(paragraph._p)


def strip_hyperlinks(paragraph: Paragraph) -> list[str]:
    """Remove `w:hyperlink` wrappers, returning the link texts that were dropped.

    Each project carries its own URL, so a hyperlink baked into the template would point
    every rendered project at the prototype's target. The link is reinstated at render
    time as a docxtpl `RichText`, which supplies the correct per-project URL.
    """
    dropped = []
    for link in paragraph._p.findall(qn("w:hyperlink")):
        dropped.append("".join(t.text or "" for t in link.iter(qn("w:t"))))
        paragraph._p.remove(link)
    return dropped


def has_tab_element(run) -> bool:
    """Whether the run's tab is a `<w:tab/>` element rather than a literal character.

    Word writes tab stops as elements, which `run.text` renders as "\\t". Treating that
    rendered form as if it were real text and rewriting it produced a duplicated tab —
    the element survived the rewrite and a literal tab was added alongside it.
    """
    return run._r.find(qn("w:tab")) is not None


def set_run_text(run, text: str, *, keep_tabs: bool = False) -> None:
    """Replace a run's text, preserving significant whitespace.

    With `keep_tabs`, any `<w:tab/>` elements are left in place and `text` is appended
    after them — use it for the run that holds a header's alignment tab.
    """
    for t in run._r.findall(qn("w:t")):
        run._r.remove(t)
    if not keep_tabs:
        for tab in run._r.findall(qn("w:tab")):
            run._r.remove(tab)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run._r.append(t)


def clone_run_after(run, text: str) -> Run:
    """Duplicate a run (inheriting its formatting) with new text, placed just after it."""
    new_r = copy.deepcopy(run._r)
    run._r.addnext(new_r)
    cloned = Run(new_r, run._parent)
    set_run_text(cloned, text)
    return cloned


def strip_bold(run: Run) -> None:
    """Remove bold from a run's character properties."""
    rPr = run._r.find(qn("w:rPr"))
    if rPr is None:
        return
    for tag in ("w:b", "w:bCs"):
        el = rPr.find(qn(tag))
        if el is not None:
            rPr.remove(el)


def tab_run_index(paragraph: Paragraph) -> int:
    """Return the index of the run holding the header's date-alignment tab."""
    for i, run in enumerate(paragraph.runs):
        if "\t" in run.text or has_tab_element(run):
            return i
    return len(paragraph.runs) - 1


def set_num_id(paragraph: Paragraph, num_id: str) -> None:
    """Point a list paragraph at a different numbering definition instance."""
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        paragraph._p.insert(0, pPr)
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        numPr = OxmlElement("w:numPr")
        pPr.insert(0, numPr)
    num_id_el = numPr.find(qn("w:numId"))
    if num_id_el is None:
        num_id_el = OxmlElement("w:numId")
        numPr.append(num_id_el)
    num_id_el.set(qn("w:val"), num_id)


def strip_right_indent(paragraph: Paragraph) -> None:
    """Drop a paragraph's right indent — it narrows the column and inflates line count."""
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        return
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        return
    if ind.get(qn("w:right")) is not None:
        del ind.attrib[qn("w:right")]


# Word single spacing under lineRule=auto is 240 (240ths of a line).
SINGLE_LINE = "240"


def set_single_spacing(paragraph: Paragraph, *, exact: bool = False) -> None:
    """Force Word single line spacing on a paragraph, preserving before/after gaps.

    `lineRule=auto` is Word's "Single". Bullets also get `exact` so a taller substitute
    bullet font (common under LibreOffice when Noto is missing) cannot inflate the
    auto line box past single — measured wraps were ~15.7pt for 10pt body text.
    """
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        paragraph._p.insert(0, pPr)
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:line"), SINGLE_LINE)
    # exact locks height in twips (240 = 12pt); auto is Word UI "Single".
    spacing.set(qn("w:lineRule"), "exact" if exact else "auto")


def pin_sub_single_to_exact(paragraph: Paragraph) -> bool:
    """Rewrite a sub-single `lineRule="auto"` paragraph to the equivalent `"exact"` height.

    Returns whether anything was changed.

    Proportional (`auto`) line spacing below 100% is **not portable**. Word honours it
    literally — `w:line="72"` is 30% of a line, ~3.5pt — while LibreOffice refuses to
    compress a line below its glyph height and floors the same paragraph at ~9.6pt, 2.8x
    taller. A source whose blank filler paragraphs are invisible in Word therefore grows a
    full blank line per filler once rendered through the container's LibreOffice.

    `exact` is honoured identically by both, and the twip number carries over directly:
    `w:line="72"` as `exact` is 72 twips = 3.6pt, within a rounding error of what Word
    already drew. This is the same reasoning that already pins bullets to `exact` in
    `set_single_spacing`.

    Only safe for paragraphs holding no text — an exact 3.6pt box would clip a 10pt glyph.
    Text paragraphs are normalised up to single instead; see `normalize_single_spacing`.
    """
    pPr = paragraph._p.find(qn("w:pPr"))
    spacing = pPr.find(qn("w:spacing")) if pPr is not None else None
    if spacing is None or spacing.get(qn("w:lineRule")) != "auto":
        return False
    line = spacing.get(qn("w:line"))
    if line is None:
        return False
    try:
        if int(line) >= int(SINGLE_LINE):
            return False
    except ValueError:
        return False
    spacing.set(qn("w:lineRule"), "exact")
    return True


def normalize_single_spacing(doc) -> None:
    """Make every body paragraph's line height renderer-independent.

    Google Docs exports can leave spacing unset (inherits a looser default) or carry
    Multiple > 1 on some paragraphs. Prototypes alone cannot fix education bullets or
    the name/contact lines, so this runs over the whole document after tagging.

    The governing rule is that a sub-100% `lineRule="auto"` value must never survive into
    a built template, because Word and LibreOffice disagree about it by ~3x (see
    `pin_sub_single_to_exact`). Two treatments, split on whether the paragraph holds text:

    - **Text** is normalised to single. Anything tighter only renders as the author
      intended in Word; under LibreOffice it floors out and merely looks cramped, or
      overlaps the line above.
    - **Chrome** (blank spacers, horizontal rules) keeps its authored height but is pinned
      to `exact`, so the thin gaps a source encodes as fractional filler paragraphs come
      out the same in both renderers instead of ballooning to a full blank line.
    """
    for paragraph, _location in docx_text.iter_document_paragraphs(doc):
        text = (paragraph.text or "").strip()
        # Control-tag paragraphs are dropped at render time; leave them bare.
        if text.startswith("{%p ") or text.startswith("{%tr "):
            continue
        if docx_text.is_chrome_text(text):
            pin_sub_single_to_exact(paragraph)
            continue
        set_single_spacing(paragraph, exact=is_bullet(paragraph))


def _section_text_width(doc) -> float | None:
    """Body text column width in twips, or `None` if the section geometry is missing.

    Deliberately reads `w:pgSz`/`w:pgMar` from the raw XML as `float` rather than using
    python-docx's typed `Section.page_width`/`.left_margin`/`.right_margin`: Google Docs
    exports non-integer twips (e.g. `w:left="1417.3228346456694"`), and those accessors
    raise `ValueError` on such a file.
    """
    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))
    if sectPr is None:
        return None
    pgSz = sectPr.find(qn("w:pgSz"))
    pgMar = sectPr.find(qn("w:pgMar"))
    if pgSz is None or pgMar is None:
        return None
    try:
        width = float(pgSz.get(qn("w:w")))
        left = float(pgMar.get(qn("w:left")))
        right = float(pgMar.get(qn("w:right")))
    except (TypeError, ValueError):
        return None
    return width - left - right


def clamp_tab_stops(doc) -> None:
    """Clamp every tab stop to its own paragraph's right edge, as Word silently does.

    A Google Docs export can carry a right tab stop (`w:tab/@w:pos`, measured from the
    left text margin) that overruns the paragraph's own right boundary — `text_width -
    w:ind/@w:right` — and even the page's text column. Word clamps such a stop back to the
    paragraph's right edge when it renders; LibreOffice honours it literally and lets the
    tab-aligned text hang past the margin. Baking Word's clamp into the XML makes both
    renderers agree, the same reasoning `pin_sub_single_to_exact` applies to line spacing.
    """
    text_width = _section_text_width(doc)
    if text_width is None:
        return
    for paragraph, _location in docx_text.iter_document_paragraphs(doc):
        pPr = paragraph._p.find(qn("w:pPr"))
        if pPr is None:
            continue
        tabs = pPr.find(qn("w:tabs"))
        if tabs is None:
            continue
        ind = pPr.find(qn("w:ind"))
        try:
            right_indent = float(ind.get(qn("w:right"))) if ind is not None else 0.0
        except (TypeError, ValueError):
            right_indent = 0.0
        limit = text_width - right_indent
        for tab in tabs.findall(qn("w:tab")):
            pos = tab.get(qn("w:pos"))
            if pos is None:
                continue
            try:
                pos_val = float(pos)
            except ValueError:
                continue
            if pos_val > limit:
                tab.set(qn("w:pos"), str(round(limit)))


# --------------------------------------------------------------------------------------
# Document structure
# --------------------------------------------------------------------------------------

SECTIONS = ("EDUCATION", "WORK EXPERIENCES", "PROJECTS", "SKILLS")


def is_bullet(p: Paragraph) -> bool:
    """A real list paragraph, i.e. one carrying numbering properties."""
    return p._p.find(f".//{qn('w:numPr')}") is not None


def find_sections(doc) -> dict[str, tuple[int, int]]:
    """Map each section heading to the (start, end) paragraph range of its body."""
    heads: dict[str, int] = {}
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text in SECTIONS:
            heads[text] = i

    bounds: dict[str, tuple[int, int]] = {}
    ordered = sorted(heads.items(), key=lambda kv: kv[1])
    for n, (name, idx) in enumerate(ordered):
        end = ordered[n + 1][1] if n + 1 < len(ordered) else len(doc.paragraphs)
        bounds[name] = (idx + 1, end)
    return bounds


def _num_to_abstract(numbering_root) -> dict[str, str]:
    """Map numbering instance ids to their abstract definition ids."""
    mapping: dict[str, str] = {}
    for num in numbering_root.findall(qn("w:num")):
        nid = num.get(qn("w:numId"))
        abs_el = num.find(qn("w:abstractNumId"))
        if nid is not None and abs_el is not None:
            mapping[nid] = abs_el.get(qn("w:val"), "")
    return mapping


def _abstract_has_noto_marker(numbering_root, abstract_id: str) -> bool:
    """Return whether an abstract numbering definition pins Noto for its lvl0 bullet."""
    for anum in numbering_root.findall(qn("w:abstractNum")):
        if anum.get(qn("w:abstractNumId")) != abstract_id:
            continue
        for lvl in anum.findall(qn("w:lvl")):
            if lvl.get(qn("w:ilvl")) != "0":
                continue
            num_fmt = lvl.find(qn("w:numFmt"))
            if num_fmt is None or num_fmt.get(qn("w:val")) != "bullet":
                continue
            rPr = lvl.find(qn("w:rPr"))
            rf = rPr.find(qn("w:rFonts")) if rPr is not None else None
            return rf is not None and rf.get(qn("w:ascii")) == NOTO_MARKER_FONT
    return False


def discover_noto_num_id(doc) -> str | None:
    """Find a list instance id whose lvl0 bullet marker uses Noto Sans Symbols.

    Prefer an id used by education bullets so experience/projects match that section.
    """
    numbering = doc.part.numbering_part
    root = numbering.element
    num_to_abs = _num_to_abstract(root)

    bounds = find_sections(doc)
    candidates: list[str] = []
    if "EDUCATION" in bounds:
        start, end = bounds["EDUCATION"]
        for paragraph in doc.paragraphs[start:end]:
            if not is_bullet(paragraph):
                continue
            pPr = paragraph._p.find(qn("w:pPr"))
            numPr = pPr.find(qn("w:numPr")) if pPr is not None else None
            if numPr is None:
                continue
            num_id_el = numPr.find(qn("w:numId"))
            if num_id_el is None:
                continue
            num_id = num_id_el.get(qn("w:val"))
            if num_id and _abstract_has_noto_marker(root, num_to_abs.get(num_id, "")):
                return num_id
            if num_id:
                candidates.append(num_id)

    for num_id, abstract_id in num_to_abs.items():
        if _abstract_has_noto_marker(root, abstract_id):
            return num_id
    return candidates[0] if candidates else None


def normalize_bullet_numbering(doc) -> None:
    """Pin Noto Sans Symbols on every lvl0 bullet marker in numbering.xml."""
    numbering = doc.part.numbering_part
    root = numbering.element

    ref_rfonts = None
    for anum in root.findall(qn("w:abstractNum")):
        for lvl in anum.findall(qn("w:lvl")):
            if lvl.get(qn("w:ilvl")) != "0":
                continue
            num_fmt = lvl.find(qn("w:numFmt"))
            if num_fmt is None or num_fmt.get(qn("w:val")) != "bullet":
                continue
            rPr = lvl.find(qn("w:rPr"))
            rf = rPr.find(qn("w:rFonts")) if rPr is not None else None
            if rf is not None and rf.get(qn("w:ascii")) == NOTO_MARKER_FONT:
                ref_rfonts = copy.deepcopy(rf)
                break
        if ref_rfonts is not None:
            break

    if ref_rfonts is None:
        ref_rfonts = OxmlElement("w:rFonts")
        for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
            ref_rfonts.set(qn(f"w:{attr}"), NOTO_MARKER_FONT)

    for anum in root.findall(qn("w:abstractNum")):
        for lvl in anum.findall(qn("w:lvl")):
            if lvl.get(qn("w:ilvl")) != "0":
                continue
            num_fmt = lvl.find(qn("w:numFmt"))
            if num_fmt is None or num_fmt.get(qn("w:val")) != "bullet":
                continue
            rPr = lvl.find(qn("w:rPr"))
            if rPr is None:
                rPr = OxmlElement("w:rPr")
                lvl.append(rPr)
            old_rf = rPr.find(qn("w:rFonts"))
            if old_rf is not None:
                rPr.remove(old_rf)
            rPr.insert(0, copy.deepcopy(ref_rfonts))


#: Most fonts (Times New Roman included) draw U+25CF "●" as a near-full-em disc, not
#: the small list dot a dedicated symbol font would give it. Pinning a symbol font at the
#: numbering level does not help: the paragraph mark's own direct rFonts (written by every
#: Google Docs export, on every bullet paragraph) takes precedence over the numbering
#: level's rPr, so the marker always renders in the body font regardless. Shrinking the
#: paragraph mark's own size sidesteps that precedence fight entirely, and works in
#: whatever font actually wins.
BULLET_MARKER_SIZE_RATIO = 0.8 * (2 / 3)  # ~0.53; a further 2/3 pass on top of the first cut
_MIN_MARKER_HALF_POINTS = 8  # 4pt floor so a tiny body font can't shrink the dot to nothing

#: Large geometric glyphs that render as a near-full-em disc/square in most fonts — the
#: shape `BULLET_MARKER_SIZE_RATIO` was tuned against. A hyphen, en-dash, small bullet
#: (U+2022, already text-sized), or asterisk is left alone: shrinking one of those by the
#: same ~53% that tames a "●" renders it as a near-invisible hairline, not a smaller dot.
_SHRINKABLE_MARKERS = frozenset("●⬤○◉■□◆◇")


def _marker_glyph(doc, paragraph: Paragraph) -> str | None:
    """The literal lvl0 bullet character governing `paragraph`'s numbering marker."""
    pPr = paragraph._p.find(qn("w:pPr"))
    numPr = pPr.find(qn("w:numPr")) if pPr is not None else None
    num_id_el = numPr.find(qn("w:numId")) if numPr is not None else None
    num_id = num_id_el.get(qn("w:val")) if num_id_el is not None else None
    if num_id is None:
        return None
    root = doc.part.numbering_part.element
    abstract_id = _num_to_abstract(root).get(num_id)
    if abstract_id is None:
        return None
    for anum in root.findall(qn("w:abstractNum")):
        if anum.get(qn("w:abstractNumId")) != abstract_id:
            continue
        for lvl in anum.findall(qn("w:lvl")):
            if lvl.get(qn("w:ilvl")) != "0":
                continue
            text = lvl.find(qn("w:lvlText"))
            return text.get(qn("w:val")) if text is not None else None
    return None


def _body_run_size(paragraph: Paragraph) -> int | None:
    """Font size in half-points of the paragraph's first run that sets one explicitly."""
    for run in paragraph.runs:
        rPr = run._r.find(qn("w:rPr"))
        sz = rPr.find(qn("w:sz")) if rPr is not None else None
        if sz is not None and sz.get(qn("w:val")):
            return int(sz.get(qn("w:val")))
    return None


def shrink_bullet_marker(doc, paragraph: Paragraph) -> None:
    """Scale down the paragraph-mark font size that governs the bullet glyph.

    The paragraph mark's `rPr` (inside `pPr`, not the visible runs) formats only the
    numbering symbol, so this changes the dot's size without touching the bullet's own
    text size. Only applies to `_SHRINKABLE_MARKERS` glyphs — see that constant.

    Idempotent: the target size is derived from the paragraph's own body-run size (via
    `_body_run_size`) when one is set, not from the marker's current size, so calling
    this twice does not halve an already-shrunk marker.
    """
    if _marker_glyph(doc, paragraph) not in _SHRINKABLE_MARKERS:
        return
    pPr = paragraph._p.find(qn("w:pPr"))
    rPr = pPr.find(qn("w:rPr")) if pPr is not None else None
    sz = rPr.find(qn("w:sz")) if rPr is not None else None
    if sz is None or not sz.get(qn("w:val")):
        return
    body_size = _body_run_size(paragraph) or int(sz.get(qn("w:val")))
    marker_size = str(max(_MIN_MARKER_HALF_POINTS, round(body_size * BULLET_MARKER_SIZE_RATIO)))
    sz.set(qn("w:val"), marker_size)
    sz_cs = rPr.find(qn("w:szCs"))
    if sz_cs is not None:
        sz_cs.set(qn("w:val"), marker_size)


def retarget_bullet(doc, paragraph: Paragraph, num_id: str | None) -> None:
    """Use the canonical small-marker list id and drop any right-indent inflation."""
    if num_id is not None:
        set_num_id(paragraph, num_id)
    strip_right_indent(paragraph)
    shrink_bullet_marker(doc, paragraph)


# --------------------------------------------------------------------------------------
# Tagging
# --------------------------------------------------------------------------------------


def tag_header(paragraph: Paragraph, fields: list[str], *, tail_field: str) -> None:
    """Tag a header line of the form `<bold> | <plain>\\t<date>`.

    Runs before the tab receive `fields` in order; the run holding the tab keeps it and
    the text after the tab becomes `tail_field`. Working run-by-run rather than rewriting
    the paragraph is what preserves the bold prefix and the right-aligned tab stop.
    """
    tab_idx = tab_run_index(paragraph)
    runs = paragraph.runs

    # Everything before the tab carries the leading fields.
    lead = runs[:tab_idx]
    for i, run in enumerate(lead):
        set_run_text(run, fields[i] if i < len(fields) else "")

    # More fields than runs: insert plain runs instead of merging into the bold run.
    if len(fields) > len(lead) and lead:
        anchor = lead[-1]
        for field in fields[len(lead) :]:
            plain = clone_run_after(anchor, field)
            strip_bold(plain)
            anchor = plain

    runs = paragraph.runs
    tab_idx = tab_run_index(paragraph)

    # The tab run keeps its tab; the date follows it.
    tab_run = runs[tab_idx]
    element_tab = has_tab_element(tab_run)
    after_tab = tab_run.text.split("\t", 1)[1] if "\t" in tab_run.text else ""

    if after_tab or tab_idx == len(runs) - 1:
        # Tab and date share one run.
        if element_tab:
            set_run_text(tab_run, tail_field, keep_tabs=True)
        else:
            set_run_text(tab_run, "\t" + tail_field)
        for extra in runs[tab_idx + 1 :]:
            paragraph._p.remove(extra._r)
    else:
        # Tab and date are in separate runs; keep that split so the date keeps its own
        # formatting (the tab run is sometimes bold, the date never is).
        if element_tab:
            set_run_text(tab_run, "", keep_tabs=True)
        else:
            set_run_text(tab_run, "\t")
        set_run_text(runs[tab_idx + 1], tail_field)
        for extra in runs[tab_idx + 2 :]:
            paragraph._p.remove(extra._r)


def tag_bullet(paragraph: Paragraph, expr: str) -> None:
    """Collapse a bullet to a single tagged run, keeping its list formatting."""
    runs = paragraph.runs
    set_run_text(runs[0], expr)
    for extra in runs[1:]:
        paragraph._p.remove(extra._r)


def build_loop(
    entry: list[Paragraph],
    anchor: Paragraph,
    *,
    outer: str,
    inner: str,
) -> None:
    """Insert a cloned, tagged entry wrapped in for-loops, before `anchor`."""
    anchor._p.addprevious(make_para(outer))
    for para in entry:
        anchor._p.addprevious(copy.deepcopy(para._p))
    anchor._p.addprevious(make_para("{%p endfor %}"))
    _ = inner


# --------------------------------------------------------------------------------------
# Span-aware tagging (profile mode)
# --------------------------------------------------------------------------------------


_R_TAG = qn("w:r")
_HYPERLINK_TAG = qn("w:hyperlink")

#: Characters that only ever appear as connective punctuation between header fields.
#: Character-class based rather than a fixed separator list, so " • ", " – ", " / " are
#: all recognised the same way " | " is — no per-template special-casing needed.
_SEPARATOR_CHARS = frozenset(" |•·–—/,;-")


@dataclass(frozen=True)
class _Segment:
    """One piece of a rebuilt paragraph, carrying the source offset it inherits from.

    `donor_offset` indexes into the ORIGINAL paragraph text (before any rebuild), and is
    resolved against a `RunSlice` list to decide which run's formatting this piece gets.
    """

    text: str
    donor_offset: int
    kind: Literal["text", "tag", "tab", "break"] = "text"
    field: str = ""


def _is_separator_literal(text: str) -> bool:
    """True when `text` is only whitespace and connective punctuation."""
    return bool(text) and all(ch in _SEPARATOR_CHARS for ch in text)


def _run_shell(donor: RunSlice | None):
    """Build an empty `w:r` carrying a deep copy of `donor`'s `w:rPr`, nothing else.

    Copying only the run properties (not the whole run, unlike `clone_run_after`) means
    cloning many segments from the same donor never duplicates non-text children such as
    a drawing or bookmark.
    """
    r = OxmlElement("w:r")
    if donor is not None:
        rpr = donor.element.find(qn("w:rPr"))
        if rpr is not None:
            r.append(copy.deepcopy(rpr))
    return r


def _append_text(run_el, text: str) -> None:
    """Append a `w:t` to a bare run element, preserving significant whitespace."""
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run_el.append(t)


def _split_literal(
    text: str, start: int, end: int, slices: list[RunSlice]
) -> list[_Segment]:
    """Split a literal gap ``[start:end)`` at source-run boundaries and tab/break chars.

    Splitting at run boundaries is what reproduces a document's own bold/plain split
    across a separator (e.g. a bold ``":"`` immediately followed by a plain space).
    Splitting at every ``\\t``/``\\n``/``\\r`` — regardless of whether it shares a run
    with surrounding text — is what turns a literal tab into a real ``<w:tab/>`` element
    instead of a character inside a `w:t`, which is what keeps a date right-aligned.
    """
    if start >= end:
        return []
    cut_points = {start, end}
    for s in slices:
        if start < s.start < end:
            cut_points.add(s.start)
        if start < s.end < end:
            cut_points.add(s.end)
    for i in range(start, end):
        if text[i] in "\t\n\r":
            cut_points.add(i)
            cut_points.add(i + 1)
    points = sorted(p for p in cut_points if start <= p <= end)

    segments: list[_Segment] = []
    for a, b in zip(points, points[1:]):
        if a == b:
            continue
        piece = text[a:b]
        if piece == "\t":
            segments.append(_Segment(text="", donor_offset=a, kind="tab"))
        elif piece in ("\n", "\r"):
            segments.append(_Segment(text="", donor_offset=a, kind="break"))
        else:
            segments.append(_Segment(text=piece, donor_offset=a, kind="text"))
    return segments


def build_segments(
    text: str,
    slices: list[RunSlice],
    items: list[tuple[int, int, str, str]],
    *,
    drop_hyperlink_literals: bool = True,
    render_owned_separator_before: frozenset[str] = frozenset(),
) -> list["_Segment"]:
    """Turn ordered field spans and the literal text between them into donor-tagged pieces.

    `items` is a list of ``(start, end, tag, field_name)``. Fails loudly rather than
    truncating: raises when a span exceeds the paragraph, contains a tab or line break,
    or overlaps another mapped span (naming both fields).

    `drop_hyperlink_literals` removes any literal piece whose source run sits inside a
    baked-in `w:hyperlink` — a per-project label like "Github" must never survive into
    the shared template. `render_owned_separator_before` names fields (e.g. ``"link"``)
    whose own separator is supplied by the render context rather than the template, so
    the literal immediately preceding that tag is dropped here to avoid a doubled or
    dangling separator. Either kind of removal can orphan its other neighbour (e.g. a
    ``" | "`` left with nothing after it but a tab); a cleanup pass drops any literal
    that is both a removal's neighbour and now dangling at a tab or a paragraph edge.
    """
    ordered = sorted(items, key=lambda it: it[0])

    for start, end, _tag, field in ordered:
        if end > len(text):
            raise RuntimeError(
                f"{field}: span [{start}:{end}] exceeds paragraph length {len(text)}: {text!r}"
            )
        chunk = text[start:end]
        if "\t" in chunk:
            raise RuntimeError(
                f"{field}: span [{start}:{end}] contains a tab: {chunk!r}. "
                "Map date fields after the tab, not through it."
            )
        if "\n" in chunk or "\r" in chunk:
            raise RuntimeError(
                f"{field}: span [{start}:{end}] contains a line break: {chunk!r}."
            )

    raw: list[_Segment] = []
    cursor = 0
    prev_field: str | None = None
    for start, end, tag, field in ordered:
        if start < cursor:
            lo, hi = min(start, cursor), max(end, cursor)
            raise RuntimeError(
                f"Overlapping header field spans: {prev_field!r} and {field!r} "
                f"both claim {text[lo:hi]!r}"
            )
        if start > cursor:
            raw.extend(_split_literal(text, cursor, start, slices))
        if tag or end > start:
            raw.append(_Segment(text=tag, donor_offset=start, kind="tag", field=field))
        cursor = max(cursor, end)
        prev_field = field
    if cursor < len(text):
        raw.extend(_split_literal(text, cursor, len(text), slices))

    removed: set[int] = set()
    if drop_hyperlink_literals:
        for i, seg in enumerate(raw):
            if seg.kind != "text":
                continue
            donor = docx_text.slice_at(slices, seg.donor_offset)
            if donor is not None and donor.in_hyperlink:
                removed.add(i)

    if render_owned_separator_before:
        for i, seg in enumerate(raw):
            if seg.kind != "tag" or seg.field not in render_owned_separator_before:
                continue
            j = i - 1
            while j in removed:
                j -= 1
            if j >= 0 and raw[j].kind == "text" and _is_separator_literal(raw[j].text):
                removed.add(j)

    kept = [(i, seg) for i, seg in enumerate(raw) if i not in removed]

    final: list[_Segment] = []
    for pos, (i, seg) in enumerate(kept):
        is_removal_neighbor = (i - 1) in removed or (i + 1) in removed
        if seg.kind == "text" and is_removal_neighbor and _is_separator_literal(seg.text):
            prev_seg = final[-1] if final else None
            next_seg = kept[pos + 1][1] if pos + 1 < len(kept) else None
            dangling = (
                prev_seg is None
                or prev_seg.kind in ("tab", "break")
                or next_seg is None
                or next_seg.kind in ("tab", "break")
            )
            if dangling:
                continue
        final.append(seg)

    return [seg for seg in final if seg.kind != "text" or seg.text]


def rebuild_paragraph(
    paragraph: Paragraph, segments: list["_Segment"], slices: list[RunSlice]
) -> None:
    """Rewrite a paragraph's inline content from `segments`, one run per segment.

    Removes every `w:r` and `w:hyperlink` child and inserts fresh runs where the first
    one stood (leaving `w:pPr`, bookmarks, and proofing marks alone). Because the
    replacement is total, a baked-in hyperlink disappears as a side effect — callers must
    NOT strip hyperlinks beforehand, or spans measured against the original text will no
    longer line up (this is exactly how the project date lost its tab stop).
    """
    p = paragraph._p
    old_children = [c for c in p if c.tag in (_R_TAG, _HYPERLINK_TAG)]
    if not old_children:
        if segments:
            raise RuntimeError("Cannot tag a paragraph with no runs.")
        return

    anchor = old_children[0]
    new_runs = []
    for seg in segments:
        donor = docx_text.slice_at(
            slices, seg.donor_offset, skip_hyperlinks=(seg.kind == "tag")
        )
        run_el = _run_shell(donor)
        if seg.kind == "tab":
            run_el.append(OxmlElement("w:tab"))
        elif seg.kind == "break":
            run_el.append(OxmlElement("w:br"))
        else:
            _append_text(run_el, seg.text)
        new_runs.append(run_el)

    for run_el in new_runs:
        anchor.addprevious(run_el)
    for child in old_children:
        p.remove(child)


def retag_paragraph(
    paragraph: Paragraph,
    items: list[tuple[int, int, str, str]],
    *,
    render_owned_separator_before: frozenset[str] = frozenset(),
) -> None:
    """Replace mapped character spans in `paragraph` with Jinja tags.

    Each surviving literal keeps the formatting of whichever source run covered it; each
    tag inherits formatting from the (non-hyperlink) run at its span's start. A
    `<w:tab/>` between spans is preserved as a real tab element, never a literal
    ``"\\t"`` inside a `w:t`. This is the one rebuild path both header tagging and
    skills tagging use, so the two can never diverge in how they treat formatting.
    """
    slices = docx_text.paragraph_run_slices(paragraph)
    text = docx_text.paragraph_text(paragraph)
    segments = build_segments(
        text, slices, items, render_owned_separator_before=render_owned_separator_before
    )
    rebuild_paragraph(paragraph, segments, slices)


def replace_span_with_tag(
    paragraph: Paragraph, span: CharSpan, tag: str, *, field: str = "field"
) -> None:
    """Replace characters ``[start:end)`` in the paragraph with a single-run Jinja tag.

    Surrounding text keeps its own runs' formatting; a `<w:tab/>` outside the span
    survives as a real tab element. A span containing a tab raises — map date fields
    after the tab, not through it.
    """
    if span.start == span.end and not tag:
        return
    retag_paragraph(paragraph, [(span.start, span.end, tag, field)])


def _para_by_id(doc, paragraph_id: int) -> Paragraph:
    """Return the body paragraph at `paragraph_id`, including one inside a table cell.

    Recomputed on every call, never cached: `template_analyze._load_paras` and this
    function must enumerate identically via `docx_text.iter_document_paragraphs`
    (`CharSpan.paragraph_id` is an index into that exact sequence), and several build
    steps insert or delete paragraphs mid-build, which shifts every later id — a cached
    list would silently return the wrong paragraph the moment that happens.
    """
    paras = [p for p, _location in docx_text.iter_document_paragraphs(doc)]
    if paragraph_id < 0 or paragraph_id >= len(paras):
        raise RuntimeError(f"Paragraph id {paragraph_id} out of range (0..{len(paras) - 1}).")
    return paras[paragraph_id]


def build_name_profile(doc, profile: TemplateProfile) -> None:
    """Tag the mapped name paragraph as ``{{ name }}``."""
    paragraph = _para_by_id(doc, profile.name_paragraph_id)
    runs = paragraph.runs
    if not runs:
        raise RuntimeError("Name line has no runs to tag.")
    set_run_text(runs[0], NAME_TAG)
    for extra in runs[1:]:
        paragraph._p.remove(extra._r)


def build_contact_profile(doc, profile: TemplateProfile) -> None:
    """Tag the mapped contact paragraph(s) as RichText placeholder(s).

    A table layout's `contact.slots` spreads the contact block across several
    paragraphs (name in one cell, address in another, email/phone in a third and
    fourth) — each slot gets its own `contact_slot_<i>` tag on its own paragraph,
    instead of collapsing everything onto one joined `{{r contact }}` line, which
    would discard the layout. Empty `slots` (every profile before this existed, and
    any single-paragraph contact block) keeps today's exact one-tag behaviour.
    """
    if profile.contact.slots:
        for i, slot in enumerate(profile.contact.slots):
            paragraph = _para_by_id(doc, slot.paragraph_id)
            strip_hyperlinks(paragraph)
            runs = paragraph.runs
            if not runs:
                raise RuntimeError(f"Contact slot {i} has no runs to tag.")
            set_run_text(runs[0], CONTACT_SLOT_TAG_FMT % i)
            for extra in runs[1:]:
                paragraph._p.remove(extra._r)
        return

    paragraph = _para_by_id(doc, profile.contact.paragraph_id)
    strip_hyperlinks(paragraph)
    runs = paragraph.runs
    if not runs:
        raise RuntimeError("Contact line has no runs to tag.")
    set_run_text(runs[0], CONTACT_TAG)
    for extra in runs[1:]:
        paragraph._p.remove(extra._r)


def _tag_mapped_header(
    doc,
    header_mapping,
    tag_for: dict[str, str],
    *,
    render_owned_separator_before: frozenset[str] = frozenset(),
) -> Paragraph:
    """Replace a header paragraph's mapped spans with tags, preserving runs and tabs.

    Spans are measured against the paragraph's text AS UPLOADED, including any baked-in
    hyperlink — callers must not strip hyperlinks before calling this. `retag_paragraph`
    removes them as a side effect of replacing every run, which is what keeps offsets
    valid (stripping first shortens the text the spans were measured against and slides
    every later span left — this is exactly how a project's date lost its tab stop).
    """
    paragraph = _para_by_id(doc, header_mapping.header_paragraph_id)

    items: list[tuple[int, int, str, str]] = []
    for name, field in header_mapping.fields.items():
        if not field.present or field.span is None or name not in tag_for:
            continue
        if field.span.paragraph_id != header_mapping.header_paragraph_id:
            other = _para_by_id(doc, field.span.paragraph_id)
            replace_span_with_tag(other, field.span, tag_for[name], field=name)
            continue
        items.append((field.span.start, field.span.end, tag_for[name], name))

    if not items:
        return paragraph

    retag_paragraph(
        paragraph, items, render_owned_separator_before=render_owned_separator_before
    )
    return paragraph


def _tag_experience_prototype(
    doc, mapping, *, noto_num_id: str | None
) -> tuple[Paragraph, Paragraph, Paragraph]:
    """Tag one experience header/title/bullet prototype in place. Pure tagging — no loop
    insertion, no deletion — so both `build_experience_profile` (fixed mode) and
    `build_generic` (generic mode) can wrap the result in whichever loop shape they need.
    Returns `(header, title, bullet)`, still at their original document positions.
    """
    header = _tag_mapped_header(doc, mapping.header, EXPERIENCE_HEADER_TAGS)
    if mapping.title.present and mapping.title.span is not None:
        title_para = _para_by_id(doc, mapping.title.span.paragraph_id)
        replace_span_with_tag(title_para, mapping.title.span, EXPERIENCE_TITLE_TAG)
    elif mapping.title_paragraph_id is not None:
        title_para = _para_by_id(doc, mapping.title_paragraph_id)
        if title_para.runs:
            set_run_text(title_para.runs[0], EXPERIENCE_TITLE_TAG)
            for extra in title_para.runs[1:]:
                title_para._p.remove(extra._r)
    else:
        raise RuntimeError("Experience mapping is missing a job title.")

    title_para = (
        _para_by_id(doc, mapping.title.span.paragraph_id)
        if mapping.title.present and mapping.title.span is not None
        else _para_by_id(doc, mapping.title_paragraph_id)  # type: ignore[arg-type]
    )
    bullet = _para_by_id(doc, mapping.bullet_paragraph_id)
    retarget_bullet(doc, bullet, noto_num_id)
    tag_bullet(bullet, BULLET_TAG)
    return header, title_para, bullet


def build_experience_profile(
    doc, profile: TemplateProfile, *, noto_num_id: str | None, other_headings: list[Paragraph]
) -> None:
    """Loop-tag the experience section from the confirmed mapping."""
    mapping = profile.experience
    victims = _section_body_paragraphs(doc, mapping.heading_paragraph_id, other_headings)
    header, title_para, bullet = _tag_experience_prototype(doc, mapping, noto_num_id=noto_num_id)

    # Insert clones before the (still-present) prototype header, then drop originals.
    anchor = header
    anchor._p.addprevious(make_para("{%p for job in experience %}"))
    anchor._p.addprevious(copy.deepcopy(header._p))
    anchor._p.addprevious(copy.deepcopy(title_para._p))
    anchor._p.addprevious(make_para("{%p for bullet in job.bullets %}"))
    anchor._p.addprevious(copy.deepcopy(bullet._p))
    anchor._p.addprevious(make_para("{%p endfor %}"))
    anchor._p.addprevious(make_para("{%p endfor %}"))

    _delete_paragraphs(victims)


def _tag_education_prototype(
    doc, mapping, *, noto_num_id: str | None
) -> tuple[Paragraph, Paragraph, Paragraph]:
    """Tag one education header/degree/detail prototype in place. Pure tagging, same
    split as `_tag_experience_prototype`. Returns `(header, degree, detail)`."""
    header = _tag_mapped_header(doc, mapping.header, EDUCATION_HEADER_TAGS)
    degree = _para_by_id(doc, mapping.degree_paragraph_id)
    retarget_bullet(doc, degree, noto_num_id)
    tag_bullet(degree, EDUCATION_DEGREE_TAG)

    if mapping.detail_paragraph_id is not None:
        detail = _para_by_id(doc, mapping.detail_paragraph_id)
        if detail._p is degree._p:
            detail_p = copy.deepcopy(degree._p)
            degree._p.addnext(detail_p)
            detail = Paragraph(detail_p, degree._parent)
            tag_bullet(detail, EDUCATION_DETAIL_TAG)
        else:
            retarget_bullet(doc, detail, noto_num_id)
            tag_bullet(detail, EDUCATION_DETAIL_TAG)
    else:
        detail_p = copy.deepcopy(degree._p)
        degree._p.addnext(detail_p)
        detail = Paragraph(detail_p, degree._parent)
        tag_bullet(detail, EDUCATION_DETAIL_TAG)
    return header, degree, detail


def build_education_profile(
    doc, profile: TemplateProfile, *, noto_num_id: str | None, other_headings: list[Paragraph]
) -> None:
    """Loop-tag the education section from the confirmed mapping."""
    mapping = profile.education
    assert mapping is not None
    victims = _section_body_paragraphs(doc, mapping.heading_paragraph_id, other_headings)
    header, degree, detail = _tag_education_prototype(doc, mapping, noto_num_id=noto_num_id)

    anchor = header
    anchor._p.addprevious(make_para("{%p for edu in education %}"))
    anchor._p.addprevious(copy.deepcopy(header._p))
    anchor._p.addprevious(copy.deepcopy(degree._p))
    anchor._p.addprevious(make_para("{%p for detail in edu.details %}"))
    anchor._p.addprevious(copy.deepcopy(detail._p))
    anchor._p.addprevious(make_para("{%p endfor %}"))
    anchor._p.addprevious(make_para("{%p endfor %}"))

    _delete_paragraphs(victims)
    if detail._p.getparent() is not None and detail not in victims:
        # Drop a temporary degree clone that was inserted beside the source bullet.
        delete(detail)


def _tag_project_prototype(doc, mapping, *, noto_num_id: str | None) -> tuple[Paragraph, Paragraph]:
    """Tag one project header/bullet prototype in place. Pure tagging, same split as
    `_tag_experience_prototype`. Returns `(header, bullet)`."""
    tag_for = dict(PROJECT_HEADER_TAGS)
    fields = dict(mapping.header.fields)
    if mapping.link.present and mapping.link.span is not None:
        fields["link"] = OptionalSpan(present=True, span=mapping.link.span)
        tag_for["link"] = PROJECT_LINK_TAG
    header_mapping = mapping.header.model_copy(update={"fields": fields})
    # The baked-in hyperlink, if any, is removed by `_tag_mapped_header`'s rebuild — not
    # stripped here first. `mapping.link.span` was measured against the paragraph WITH
    # the hyperlink still present; stripping it first would shorten the text and shift
    # every later offset left, which used to swallow the date's tab into the link tag.
    # `render_owned_separator_before` drops the literal " | " before the link tag,
    # because `render.py` supplies that separator itself, only when a link is present.
    header_para = _tag_mapped_header(
        doc, header_mapping, tag_for, render_owned_separator_before=frozenset({"link"})
    )

    bullet = _para_by_id(doc, mapping.bullet_paragraph_id)
    retarget_bullet(doc, bullet, noto_num_id)
    tag_bullet(bullet, PROJECT_BULLET_TAG)
    return header_para, bullet


def build_projects_profile(
    doc, profile: TemplateProfile, *, noto_num_id: str | None, other_headings: list[Paragraph]
) -> None:
    """Loop-tag the projects section from the confirmed mapping."""
    mapping = profile.projects
    assert mapping is not None
    victims = _section_body_paragraphs(doc, mapping.heading_paragraph_id, other_headings)
    header_para, bullet = _tag_project_prototype(doc, mapping, noto_num_id=noto_num_id)

    anchor = header_para
    anchor._p.addprevious(make_para("{%p for proj in projects %}"))
    anchor._p.addprevious(copy.deepcopy(header_para._p))
    anchor._p.addprevious(make_para("{%p for bullet in proj.bullets %}"))
    anchor._p.addprevious(copy.deepcopy(bullet._p))
    anchor._p.addprevious(make_para("{%p endfor %}"))
    anchor._p.addprevious(make_para("{%p endfor %}"))

    _delete_paragraphs(victims)


def _tag_skills_prototype(doc, mapping) -> tuple[Paragraph, Paragraph]:
    """Tag one skills label/entries prototype in place. Pure tagging. Returns
    `(label_paragraph, body_paragraph)` — the same paragraph twice when both spans
    share one paragraph (paragraph layout, today's exact behaviour), two different
    paragraphs when a table layout puts the label and its value in different cells
    (see `template_analyze._skills_pair_across_cells`) — mirrors how
    `_tag_mapped_header` already handles a cross-paragraph header field.
    """
    if mapping.label_span.paragraph_id == mapping.body_span.paragraph_id:
        prototype = _para_by_id(doc, mapping.label_span.paragraph_id)
        # A separate run per tag — rather than collapsing the whole line into one, as
        # the prior implementation did — is what keeps the label's bold and the
        # entries' plain formatting from the upload. Same tradeoff legacy
        # `build_skills` documents.
        items = sorted(
            [
                (
                    mapping.label_span.start,
                    mapping.label_span.end,
                    SKILLS_LABEL_TAG,
                    "skills_label",
                ),
                (
                    mapping.body_span.start,
                    mapping.body_span.end,
                    SKILLS_BODY_TAG,
                    "skills_body",
                ),
            ],
            key=lambda it: it[0],
        )
        retag_paragraph(prototype, items)
        shrink_bullet_marker(doc, prototype)
        return prototype, prototype

    label_para = _para_by_id(doc, mapping.label_span.paragraph_id)
    body_para = _para_by_id(doc, mapping.body_span.paragraph_id)
    replace_span_with_tag(label_para, mapping.label_span, SKILLS_LABEL_TAG, field="skills_label")
    replace_span_with_tag(body_para, mapping.body_span, SKILLS_BODY_TAG, field="skills_body")
    shrink_bullet_marker(doc, label_para)
    return label_para, body_para


def build_skills_profile(
    doc, profile: TemplateProfile, *, other_headings: list[Paragraph]
) -> None:
    """Loop-tag the skills section from the confirmed mapping."""
    mapping = profile.skills
    assert mapping is not None
    victims = _section_body_paragraphs(doc, mapping.heading_paragraph_id, other_headings)
    # Fixed mode is paragraph-layout only (a table layout always forces generic mode —
    # see `TemplateProfile.layout`), so label and body always share one paragraph here.
    prototype, _body_para = _tag_skills_prototype(doc, mapping)

    anchor = prototype
    anchor._p.addprevious(make_para("{%p for group in skills %}"))
    anchor._p.addprevious(copy.deepcopy(prototype._p))
    anchor._p.addprevious(make_para("{%p endfor %}"))

    _delete_paragraphs(victims)


def _tag_list_prototype(doc, mapping, *, noto_num_id: str | None) -> Paragraph:
    """Tag one plain-list bullet prototype in place. The simplest of the five kinds — no
    header, no entry structure, just a bullet loop. Returns the tagged bullet paragraph."""
    bullet = _para_by_id(doc, mapping.bullet_paragraph_id)
    retarget_bullet(doc, bullet, noto_num_id)
    tag_bullet(bullet, LIST_ITEM_TAG)
    return bullet


def _section_body_paragraphs(
    doc, heading_id: int, other_headings: list[Paragraph]
) -> list[Paragraph]:
    """Return current body paragraphs of a section (after heading, before next heading).

    Stops at the first *paragraph object* in `other_headings` — every other enabled
    kind's own heading, resolved by the caller once via `_para_by_id` before any
    section's body is touched, so the reference stays valid regardless of what gets
    inserted or deleted elsewhere in the document afterward (headings themselves are
    never moved or deleted mid-build; only the space between them is).

    Identity, not text: comparing walked-paragraph *text* against other headings' text
    (the previous approach) means an ordinary entry line that happens to read exactly
    "SKILLS" or "PROJECTS" — a bolded label inside a bullet, for instance — is
    indistinguishable from the real heading and silently truncates the body early.
    """
    heading = _para_by_id(doc, heading_id)
    other_elements = {p._p for p in other_headings}

    seen_heading = False
    body: list[Paragraph] = []
    for para in list(doc.paragraphs):
        if para._p is heading._p:
            seen_heading = True
            continue
        if not seen_heading:
            continue
        if para._p in other_elements:
            break
        body.append(para)
    return body


def _delete_paragraphs(paragraphs: list[Paragraph]) -> None:
    """Delete the given paragraph objects from the document."""
    for para in paragraphs:
        if para._p.getparent() is not None:
            delete(para)


#: Order the generic block's `{%p if section.kind == '...' %}` branches are emitted in.
#: Rendering order does not depend on this (each branch fires only for its own kind), but
#: a fixed order keeps the generated XML deterministic and readable.
_GENERIC_KIND_ORDER = ("experience", "project", "list", "education", "skills")

#: Entry-shaped kinds a `between_entries` spacer applies to — skills groups and list
#: items are consecutive bullet lines with no blank between them in any observed source.
_ENTRY_SHAPED_KINDS = ("experience", "project", "education")

_TR_TAG = qn("w:tr")
_TC_TAG = qn("w:tc")


def _row_of(paragraph: Paragraph):
    """The `<w:tr>` ancestor of `paragraph`, or `None` when it is body-level (not
    inside any table)."""
    el = paragraph._p.getparent()
    while el is not None and el.tag != _TR_TAG:
        el = el.getparent()
    return el


def _marker_row(text: str, grid_cols: int):
    """A disposable one-cell row carrying a `{%tr ... %}` row-level control tag.

    docxtpl's row pass (`DocxTemplate.patch_xml`, run for `y in ["tr","tc","p","r"]` —
    `tr` FIRST) replaces the ENTIRE `<w:tr>` containing a `{%tr %}` tag with the bare
    Jinja tag: the row is *consumed*, not repeated. Every row-level control tag
    therefore needs its own throwaway row — putting the tag inside a content row would
    delete that row's content along with it — and the rows meant to repeat sit between
    an open marker row and a matching close marker row.

    Must still be schema-valid OOXML on its own, since a user can open the tagged
    template in Word before it is ever rendered: at least one `<w:tc>`, with at least
    one `<w:p>`, spanning the table's declared column count so the grid isn't torn.
    """
    tr = OxmlElement("w:tr")
    tc = OxmlElement("w:tc")
    tcPr = OxmlElement("w:tcPr")
    grid_span = OxmlElement("w:gridSpan")
    grid_span.set(qn("w:val"), str(max(grid_cols, 1)))
    tcPr.append(grid_span)
    tc.append(tcPr)
    tc.append(make_para(text))
    tr.append(tc)
    return tr


def _wrap_cell_loop(paragraph: Paragraph, open_tag: str, close_tag: str = "{%p endfor %}") -> None:
    """Put a `{%p for %}` / `{%p endfor %}` pair around `paragraph`, inside its own
    cell, and remove every paragraph AFTER it within that same cell.

    A table cell commonly stacks several repeatable items as separate paragraphs
    (three bullets, three skill groups' labels), and only the *first* of them becomes
    the loop's tagged prototype — `template_analyze`'s own convention:
    `bullet_paragraph_id` / `detail_paragraph_id` / a skills `label_span.paragraph_id`
    always name the first paragraph of its kind in the entry. Left untouched, a later
    sibling would be duplicated verbatim into every rendered entry, since the whole
    ROW gets cloned once as the loop's per-iteration template
    (`build_generic_table`) — never a problem in a paragraph layout, where each bullet
    is an independent body-level paragraph the surrounding victim-deletion sweeps up
    regardless of this function.

    Never removes a paragraph BEFORE `paragraph` in the same cell — that is fixed
    content (a school/degree line the education fallback's synthetic single-paragraph
    detail clone sits *after*), not a repeatable sibling, and must survive untouched.
    """
    open_el = make_para(open_tag)
    close_el = make_para(close_tag)
    paragraph._p.addprevious(open_el)
    paragraph._p.addnext(close_el)

    tc = paragraph._p.getparent()
    while tc is not None and tc.tag != _TC_TAG:
        tc = tc.getparent()
    if tc is None:
        return
    siblings = tc.findall(W)
    idx = siblings.index(close_el)
    for extra in siblings[idx + 1 :]:
        tc.remove(extra)


def _section_body_rows(doc, heading_para: Paragraph, other_heading_paras: list[Paragraph]) -> list:
    """Rows after `heading_para`'s own row, up to (excluding) the next heading's row —
    the row-level analogue of `_section_body_paragraphs`, and for the same reason:
    matched by `<w:tr>` object identity, never by re-deriving an index or matching
    heading text, so an ordinary entry row is never mistaken for a heading row just
    because some paragraph inside it happens to equal another heading's text.
    """
    heading_row = _row_of(heading_para)
    if heading_row is None:
        raise RuntimeError(f"Heading paragraph {heading_para.text!r} is not inside a table row.")
    other_rows = {_row_of(p) for p in other_heading_paras}
    table_el = heading_row.getparent()
    rows = table_el.findall(_TR_TAG)
    start = rows.index(heading_row) + 1
    body: list = []
    for tr in rows[start:]:
        if tr in other_rows:
            break
        body.append(tr)
    return body


def _unique_rows(rows: list) -> list:
    """Consecutive-duplicate-collapsed `rows`, preserving order.

    A table layout often puts a header field and the paragraph right after it (a
    degree line, a synthesized single-paragraph detail) in the very same physical
    row — cloning that row twice would duplicate its content, rather than merging
    naturally the way sibling paragraphs within one cell already do once the row
    itself is cloned once.
    """
    out: list = []
    for r in rows:
        if r is None or (out and out[-1] is r):
            continue
        out.append(r)
    return out


def _spacer_donors(doc, paragraph_ids: list[int]) -> list:
    """Resolve a spacing donor *run* up front, before any tagging that could shift
    paragraph indices — same hazard `build_generic`'s descending processing order exists
    to dodge, sidestepped entirely by cloning while indices are still pristine.

    Returns the donors' deep-copied `<w:p>` elements in order. Any id that is out of
    range or no longer chrome drops out, and an empty result means "no spacer here" —
    degrading rather than failing the build, matching
    `template_analyze.validate_profile_against_doc`'s non-blocking `spacer_donor_missing`.
    """
    donors = []
    for paragraph_id in paragraph_ids:
        try:
            para = _para_by_id(doc, paragraph_id)
        except RuntimeError:
            continue
        if not docx_text.is_chrome_text(para.text):
            continue
        donors.append(copy.deepcopy(para._p))
    return donors


def _spacer_clones(donors: list) -> list:
    """Fresh copies of a resolved spacer run for one insertion point. Each physical spot
    in the tagged template needs its own elements — docxtpl's `{%p for %}` repeats
    paragraphs that appear once in the template, so this is called once per *position*
    (before-heading, after-heading, one per entry-shaped branch), never once per entry."""
    return [copy.deepcopy(el) for el in donors]


def build_generic(doc, profile: TemplateProfile) -> None:
    """Tag `doc` for `section_mode="generic"`: one shared block driven by the `sections`
    render-context key, so adding, renaming, or reordering a resume section never needs a
    template rebuild — see `data.MasterResume.sections` and `render.build_context`.

    Each enabled kind's prototype is tagged exactly the way `build_from_profile`'s
    per-kind builders already do (same helpers, same tag strings, same loop-variable
    names — `job`/`proj`/`edu`/`group`/`bullet`/`detail`, plus `item` for the new `list`
    kind), so a resume with exactly today's four sections renders identically to a
    fixed-mode build. What differs is the outer shape: one `{%p for section in sections %}`
    loop wraps a cloned, `{{ section.title }}`-tagged heading and one `{%p if
    section.kind == '<kind>' %}` branch per enabled kind, each with its own `{%p for
    <var> in section.entries %}` loop — so any number of sections of any kind, in any
    order, under any title, all render from the same tagged template.

    Deliberately independent `{%p if %}` blocks rather than an `elif` ladder: each kind's
    branch is then a self-contained unit that can be entirely omitted when that kind has
    no prototype, with no ladder ordering or dangling-`elif` bookkeeping to get right.
    """
    # Resolve spacer donors first, before any tagging — `_tag_education_prototype` can
    # insert a paragraph mid-build (see the descending-order comment below), which would
    # invalidate a donor id resolved afterward.
    before_heading_donors = _spacer_donors(doc, profile.spacing.before_heading)
    after_heading_donors = _spacer_donors(doc, profile.spacing.after_heading)
    between_entries_donors = _spacer_donors(doc, profile.spacing.between_entries)

    noto_num_id = discover_noto_num_id(doc) if profile.normalization.normalize_bullet_font else None

    build_name_profile(doc, profile)
    build_contact_profile(doc, profile)

    mappings = {
        "experience": profile.experience,
        "project": profile.projects,
        "list": profile.list_section,
        "education": profile.education,
        "skills": profile.skills,
    }
    enabled = {
        "experience": True,  # always required, same as fixed mode
        "project": profile.enabled.projects,
        "list": profile.enabled.list_section,
        "education": profile.enabled.education,
        "skills": profile.enabled.skills,
    }

    all_victims: list[Paragraph] = []
    heading_paragraphs: list[Paragraph] = []
    heading_ids: list[int] = []
    # kind -> ordered list of (paragraph-to-clone | literal control-tag string)
    branches: dict[str, list[Paragraph | str]] = {}

    # Process bottom-up (later headings first), same reasoning as `build_from_profile`'s
    # fixed-mode dispatch: `_tag_education_prototype` can INSERT a paragraph (cloning the
    # degree line when degree and detail share one paragraph — the common case for a
    # single-line education entry), which shifts every later paragraph's index. Tagging
    # a lower (later-in-document) kind first means that insertion can only affect
    # positions below kinds not yet processed, never above them, so an
    # already-resolved `heading_paragraph_id` for a still-pending kind never goes stale.
    processing_order = sorted(
        (
            (mappings[kind].heading_paragraph_id, kind)
            for kind in _GENERIC_KIND_ORDER
            if mappings[kind] is not None and enabled[kind]
        ),
        reverse=True,
    )

    # Resolve every enabled kind's heading paragraph *object* once, up front, before
    # any section's body is touched — same reasoning as `build_from_profile`'s
    # fixed-mode dispatch: `_section_body_paragraphs` stops on object identity, which
    # stays valid across the whole build, rather than re-deriving an index (which an
    # earlier insertion could shift) or matching heading text (which an ordinary entry
    # line could coincidentally equal).
    all_headings = {hid: _para_by_id(doc, hid) for hid, _kind in processing_order}

    for _heading_id, kind in processing_order:
        mapping = mappings[kind]
        other_headings = [
            p for h, p in all_headings.items() if h != mapping.heading_paragraph_id
        ]
        all_victims.extend(
            _section_body_paragraphs(doc, mapping.heading_paragraph_id, other_headings)
        )
        heading_paragraphs.append(_para_by_id(doc, mapping.heading_paragraph_id))
        heading_ids.append(mapping.heading_paragraph_id)

        if kind == "experience":
            header, title_para, bullet = _tag_experience_prototype(
                doc, mapping, noto_num_id=noto_num_id
            )
            branches[kind] = [
                "{%p for job in section.entries %}",
                header,
                title_para,
                "{%p for bullet in job.bullets %}",
                bullet,
                "{%p endfor %}",
                "{%p endfor %}",
            ]
        elif kind == "project":
            header, bullet = _tag_project_prototype(doc, mapping, noto_num_id=noto_num_id)
            branches[kind] = [
                "{%p for proj in section.entries %}",
                header,
                "{%p for bullet in proj.bullets %}",
                bullet,
                "{%p endfor %}",
                "{%p endfor %}",
            ]
        elif kind == "list":
            bullet = _tag_list_prototype(doc, mapping, noto_num_id=noto_num_id)
            branches[kind] = ["{%p for item in section.entries %}", bullet, "{%p endfor %}"]
        elif kind == "education":
            header, degree, detail = _tag_education_prototype(
                doc, mapping, noto_num_id=noto_num_id
            )
            branches[kind] = [
                "{%p for edu in section.entries %}",
                header,
                degree,
                "{%p for detail in edu.details %}",
                detail,
                "{%p endfor %}",
                "{%p endfor %}",
            ]
            # `_tag_education_prototype` clones the degree paragraph in place
            # (`addnext`) when degree/detail share one paragraph — the common
            # single-line-entry case. `branches[kind]` above already captured its own
            # independent copy for insertion, so the clone left sitting at its original
            # document position is now a duplicate; `_section_body_paragraphs` was
            # called (into `all_victims`) before this clone existed, so nothing else
            # will ever delete it. Same cleanup `build_education_profile` (fixed mode)
            # already does after its own call to this same tagging helper.
            if detail._p.getparent() is not None and detail not in all_victims:
                delete(detail)
        elif kind == "skills":
            group, _body_para = _tag_skills_prototype(doc, mapping)
            branches[kind] = ["{%p for group in section.entries %}", group, "{%p endfor %}"]

    if not branches:
        raise RuntimeError("Generic template build found no enabled sections to tag.")

    # Anchor the whole block at the topmost enabled heading; every other enabled heading
    # is deleted too (folded into `all_victims` below), since one shared block now covers
    # all of them. Compared by `heading_paragraph_id`, not `doc.paragraphs.index(...)` —
    # `Paragraph.__eq__` is identity, and `doc.paragraphs` builds fresh wrapper objects on
    # every access, so no paragraph fetched earlier is ever found in a freshly-built list.
    anchor = _para_by_id(doc, min(heading_ids))

    heading_source = _para_by_id(doc, profile.heading_prototype.paragraph_id)
    heading_clone_el = copy.deepcopy(heading_source._p)
    heading_clone = Paragraph(heading_clone_el, heading_source._parent)
    heading_runs = heading_clone.runs
    if not heading_runs:
        raise RuntimeError("Heading prototype paragraph has no runs to tag.")
    set_run_text(heading_runs[0], SECTION_TITLE_TAG)
    for extra in heading_runs[1:]:
        heading_clone._p.remove(extra._r)

    insertions: list = [make_para(SECTION_LOOP_OPEN)]
    if before_heading_donors:
        # `loop` here is the outer `for section in sections` loop — the only one open at
        # this point in the document — so `loop.first` means "first section", which is
        # what makes the gap precede every heading except the very first.
        insertions.append(make_para("{%p if not loop.first %}"))
        insertions.extend(_spacer_clones(before_heading_donors))
        insertions.append(make_para("{%p endif %}"))
    insertions.append(heading_clone_el)
    insertions.extend(_spacer_clones(after_heading_donors))
    for kind in _GENERIC_KIND_ORDER:
        body = branches.get(kind)
        if body is None:
            continue
        insertions.append(make_para(f"{{%p if section.kind == '{kind}' %}}"))
        for idx, item in enumerate(body):
            insertions.append(make_para(item) if isinstance(item, str) else copy.deepcopy(item._p))
            # `item` at index 0 of an entry-shaped branch is always its opening
            # `{%p for <var> in section.entries %}` tag (see the per-kind blocks above).
            # Inserting the spacer right after it, guarded on `loop.first` of that
            # now-innermost loop, means "first entry" — skipping the gap before the
            # first entry and placing it only between entries.
            if idx == 0 and between_entries_donors and kind in _ENTRY_SHAPED_KINDS:
                insertions.append(make_para("{%p if not loop.first %}"))
                insertions.extend(_spacer_clones(between_entries_donors))
                insertions.append(make_para("{%p endif %}"))
        insertions.append(make_para("{%p endif %}"))
    insertions.append(make_para("{%p endfor %}"))

    for element in insertions:
        anchor._p.addprevious(element)

    _delete_paragraphs(all_victims)
    for heading in heading_paragraphs:
        delete(heading)


def build_generic_table(doc, profile: TemplateProfile) -> None:
    """Table-layout counterpart of `build_generic`: one shared block whose repeating
    unit is table ROWS, not paragraphs — a table used purely as an invisible
    single-column layout grid (see `TemplateProfile.layout`). Reuses the exact same
    per-kind `_tag_*_prototype` functions `build_generic` does (pure tagging, no loop
    insertion or deletion — see their own docstrings) since a table layout tags the
    same character spans; what differs is that a `{%p for/if %}` construct cannot
    repeat a table ROW, so `{%tr %}` marker rows wrap whichever rows an entry spans,
    and bullet/detail loops live inside a cell as `{%p for %}` while entry/skills-group
    repetition happens at the row level via `{%tr for %}`.

    `layout="table"` always implies `section_mode="generic"` (see that field's
    docstring), so `build_from_profile` calls this instead of `build_generic` — never
    both.
    """
    noto_num_id = discover_noto_num_id(doc) if profile.normalization.normalize_bullet_font else None

    build_name_profile(doc, profile)
    build_contact_profile(doc, profile)

    mappings = {
        "experience": profile.experience,
        "project": profile.projects,
        "list": profile.list_section,
        "education": profile.education,
        "skills": profile.skills,
    }
    enabled = {
        "experience": True,  # always required, same as fixed mode / build_generic
        "project": profile.enabled.projects,
        "list": profile.enabled.list_section,
        "education": profile.enabled.education,
        "skills": profile.enabled.skills,
    }

    # Same bottom-up (later headings first) processing order `build_generic` uses, for
    # the same reason: `_tag_education_prototype` can insert a paragraph mid-build,
    # which must never invalidate an already-resolved heading id for a kind still
    # pending. Row insertion (below) happens only after every kind is fully tagged, so
    # that hazard never compounds with a second one.
    processing_order = sorted(
        (
            (mappings[kind].heading_paragraph_id, kind)
            for kind in _GENERIC_KIND_ORDER
            if mappings[kind] is not None and enabled[kind]
        ),
        reverse=True,
    )
    all_headings = {hid: _para_by_id(doc, hid) for hid, _kind in processing_order}

    all_victim_rows: list = []
    heading_rows: list = []
    heading_ids: list[int] = []
    branches: dict[str, list] = {}  # kind -> ordered list of (row element | control-tag string)

    for _heading_id, kind in processing_order:
        mapping = mappings[kind]
        heading_para = _para_by_id(doc, mapping.heading_paragraph_id)
        other_headings = [
            p for h, p in all_headings.items() if h != mapping.heading_paragraph_id
        ]
        all_victim_rows.extend(_section_body_rows(doc, heading_para, other_headings))
        heading_rows.append(_row_of(heading_para))
        heading_ids.append(mapping.heading_paragraph_id)

        if kind == "experience":
            header, _title_para, bullet = _tag_experience_prototype(
                doc, mapping, noto_num_id=noto_num_id
            )
            header_row, bullet_row = _row_of(header), _row_of(bullet)
            if header_row is None or bullet_row is None:
                raise RuntimeError("Table-layout experience prototype must live inside table rows.")
            _wrap_cell_loop(bullet, "{%p for bullet in job.bullets %}")
            rows = _unique_rows([header_row, bullet_row])
            body: list = ["{%tr for job in section.entries %}", rows[0]]
            if len(rows) > 1:
                body += ["{%tr if job.bullets %}", rows[1], "{%tr endif %}"]
            body.append("{%tr endfor %}")
            branches[kind] = body
        elif kind == "project":
            header, bullet = _tag_project_prototype(doc, mapping, noto_num_id=noto_num_id)
            header_row, bullet_row = _row_of(header), _row_of(bullet)
            if header_row is None or bullet_row is None:
                raise RuntimeError("Table-layout project prototype must live inside table rows.")
            _wrap_cell_loop(bullet, "{%p for bullet in proj.bullets %}")
            rows = _unique_rows([header_row, bullet_row])
            body = ["{%tr for proj in section.entries %}", rows[0]]
            if len(rows) > 1:
                body += ["{%tr if proj.bullets %}", rows[1], "{%tr endif %}"]
            body.append("{%tr endfor %}")
            branches[kind] = body
        elif kind == "list":
            bullet = _tag_list_prototype(doc, mapping, noto_num_id=noto_num_id)
            bullet_row = _row_of(bullet)
            if bullet_row is None:
                raise RuntimeError("Table-layout list prototype must live inside a table row.")
            _wrap_cell_loop(bullet, "{%p for item in section.entries %}")
            branches[kind] = ["{%tr for item in section.entries %}", bullet_row, "{%tr endfor %}"]
        elif kind == "education":
            header, degree, detail = _tag_education_prototype(
                doc, mapping, noto_num_id=noto_num_id
            )
            header_row, degree_row, detail_row = _row_of(header), _row_of(degree), _row_of(detail)
            if header_row is None:
                raise RuntimeError("Table-layout education prototype must live inside table rows.")
            _wrap_cell_loop(detail, "{%p for detail in edu.details %}")
            rows = _unique_rows([header_row, degree_row, detail_row])
            body = ["{%tr for edu in section.entries %}", rows[0]]
            body.extend(rows[1:-1] if len(rows) > 2 else [])
            if len(rows) > 1:
                body += ["{%tr if edu.details %}", rows[-1], "{%tr endif %}"]
            body.append("{%tr endfor %}")
            branches[kind] = body
            # `_tag_education_prototype`'s degree/detail-share-one-paragraph fallback
            # clones `degree` in place (`addnext`) when there is no separate detail
            # bullet — unlike `build_generic`'s paragraph-level cloning, that synthetic
            # clone is already inside whichever row got captured above (it shares
            # `degree`'s own cell), so deleting the victim rows below removes it too.
            # No separate cleanup needed here.
        elif kind == "skills":
            label_para, body_para = _tag_skills_prototype(doc, mapping)
            label_row, body_row = _row_of(label_para), _row_of(body_para)
            if label_row is None:
                raise RuntimeError("Table-layout skills prototype must live inside a table row.")
            _wrap_cell_loop(label_para, "{%p for group in section.entries %}")
            if body_para._p is not label_para._p:
                _wrap_cell_loop(body_para, "{%p for group in section.entries %}")
            rows = _unique_rows([label_row, body_row])
            # This row's cells hold ONLY the loop (a label cell, a value cell) — an
            # empty `section.entries` would otherwise leave both with zero `<w:p>`,
            # invalid OOXML. Removing the whole row instead is both valid and visually
            # correct: no skills, no row.
            branches[kind] = ["{%tr if section.entries %}", *rows, "{%tr endif %}"]

    if not branches:
        raise RuntimeError("Generic table template build found no enabled sections to tag.")

    table_el = heading_rows[0].getparent()
    grid_cols = len(table_el.find(qn("w:tblGrid")).findall(qn("w:gridCol"))) or 1

    anchor_row = _row_of(_para_by_id(doc, min(heading_ids)))
    if anchor_row is None:
        raise RuntimeError("Table-layout section heading is not inside a table row.")

    heading_source = _para_by_id(doc, profile.heading_prototype.paragraph_id)
    heading_row_source = _row_of(heading_source)
    if heading_row_source is None:
        raise RuntimeError("Table-layout heading prototype must live inside a table row.")
    heading_row_clone = copy.deepcopy(heading_row_source)
    # Tag the first non-blank paragraph in the cloned row as {{ section.title }};
    # blank any OTHER paragraph's text WITHOUT removing it — a <w:tc> must never end up
    # with zero <w:p>.
    heading_tagged = False
    for tc in heading_row_clone.findall(_TC_TAG):
        for p_el in tc.findall(W):
            p_obj = Paragraph(p_el, heading_source._parent)
            if not (p_obj.text or "").strip():
                continue
            if not heading_tagged:
                runs = p_obj.runs
                if runs:
                    set_run_text(runs[0], SECTION_TITLE_TAG)
                    for extra in runs[1:]:
                        p_obj._p.remove(extra._r)
                heading_tagged = True
            else:
                for run in list(p_obj.runs):
                    run._r.getparent().remove(run._r)
    if not heading_tagged:
        raise RuntimeError("Heading prototype row has no text to tag.")

    insertions: list = [_marker_row(SECTION_LOOP_OPEN_TR, grid_cols), heading_row_clone]
    for kind in _GENERIC_KIND_ORDER:
        body = branches.get(kind)
        if body is None:
            continue
        insertions.append(_marker_row(f"{{%tr if section.kind == '{kind}' %}}", grid_cols))
        for item in body:
            insertions.append(
                _marker_row(item, grid_cols) if isinstance(item, str) else copy.deepcopy(item)
            )
        insertions.append(_marker_row("{%tr endif %}", grid_cols))
    insertions.append(_marker_row("{%tr endfor %}", grid_cols))

    for element in insertions:
        anchor_row.addprevious(element)

    for tr in all_victim_rows:
        if tr.getparent() is not None:
            tr.getparent().remove(tr)
    for tr in heading_rows:
        if tr.getparent() is not None:
            tr.getparent().remove(tr)


def build_from_profile(
    src: Path,
    dst: Path,
    profile: TemplateProfile,
) -> None:
    """Tag `src` using `profile` and write the result to `dst`."""
    doc = docx.Document(str(src))

    if profile.paragraph_count is not None:
        actual = sum(1 for _ in docx_text.iter_document_paragraphs(doc))
        if actual != profile.paragraph_count:
            raise RuntimeError(
                f"Baseline paragraph count changed since analysis ({profile.paragraph_count} "
                f"-> {actual}); every mapped paragraph id would resolve to the wrong "
                "paragraph. Re-analyze the upload."
            )

    if profile.layout == "table":
        build_generic_table(doc, profile)
    elif profile.section_mode == "generic":
        build_generic(doc, profile)
    else:
        noto_num_id = None
        if profile.normalization.normalize_bullet_font:
            noto_num_id = discover_noto_num_id(doc)

        build_name_profile(doc, profile)
        build_contact_profile(doc, profile)

        # Build sections bottom-up (later headings first) so earlier paragraph indices
        # stay valid for still-pending sections. Deletion of a later section does not
        # shift earlier heading ids.
        builders: list[tuple[int, object]] = []
        builders.append((profile.experience.heading_paragraph_id, "experience"))
        if profile.enabled.education and profile.education is not None:
            builders.append((profile.education.heading_paragraph_id, "education"))
        if profile.enabled.projects and profile.projects is not None:
            builders.append((profile.projects.heading_paragraph_id, "projects"))
        if profile.enabled.skills and profile.skills is not None:
            builders.append((profile.skills.heading_paragraph_id, "skills"))
        builders.sort(key=lambda t: t[0], reverse=True)

        # Resolve every enabled kind's heading paragraph *object* once, before any
        # section's body is touched. `_section_body_paragraphs` then stops on object
        # identity rather than re-deriving indices (which insertions elsewhere would
        # go on to shift) or matching heading text (which an ordinary entry line could
        # coincidentally equal) — see its own docstring.
        heading_paragraphs = {hid: _para_by_id(doc, hid) for hid, _kind in builders}

        for hid, kind in builders:
            other_headings = [p for h, p in heading_paragraphs.items() if h != hid]
            if kind == "experience":
                build_experience_profile(
                    doc, profile, noto_num_id=noto_num_id, other_headings=other_headings
                )
            elif kind == "education":
                build_education_profile(
                    doc, profile, noto_num_id=noto_num_id, other_headings=other_headings
                )
            elif kind == "projects":
                build_projects_profile(
                    doc, profile, noto_num_id=noto_num_id, other_headings=other_headings
                )
            elif kind == "skills":
                build_skills_profile(doc, profile, other_headings=other_headings)

    if profile.normalization.normalize_bullet_font:
        normalize_bullet_numbering(doc)
    if profile.normalization.force_single_spacing:
        normalize_single_spacing(doc)
    clamp_tab_stops(doc)

    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst))


# --------------------------------------------------------------------------------------
# Entry points
# --------------------------------------------------------------------------------------


def build(
    src: Path | None = None,
    dst: Path | None = None,
    profile: TemplateProfile | None = None,
    profile_path: Path | None = None,
    verify: bool = True,
) -> int:
    """Build a tagged template from `profile`, or a readable profile file on disk.

    `verify` (default True) runs `template_verify.verify_tagged`/`verify_roundtrip`
    against the freshly-built template — the same check `web/template_ops.py`'s staged
    install runs before committing, exposed here so a CLI/scripted rebuild gets the
    same guarantee. Unlike the web install, this write is not staged/atomic — a
    verification failure is reported with a non-zero exit code, but the just-written
    `dst` is not rolled back; re-run after fixing the mapping.
    """
    from .template_profile import load_profile

    src = src or config.BASELINE_TEMPLATE_PATH
    dst = dst or config.DEFAULT_TEMPLATE_PATH
    if profile is None and profile_path is not None:
        profile = load_profile(profile_path)
    if profile is None:
        profile = load_profile()

    if profile is None:
        print(
            "ERROR: no template profile found (and none was passed in).\n"
            "Run the Template tab's analyze/confirm wizard, or pass --profile to "
            "point at a template_profile.json."
        )
        return 1

    try:
        build_from_profile(src, dst, profile)
    except Exception as exc:
        print(f"ERROR: profile build failed: {exc}")
        return 1

    if verify:
        from . import data, template_verify

        issues = template_verify.verify_tagged(dst, profile)
        try:
            resume = data.load()
        except (FileNotFoundError, ValueError) as exc:
            print(f"WARNING: could not load master resume to verify roundtrip: {exc}")
        else:
            issues += template_verify.verify_roundtrip(dst, profile, resume)
        blockers = [i for i in issues if i.blocking]
        if blockers:
            print(f"ERROR: build verification failed for {dst}:")
            for issue in blockers:
                print(f"  {issue.code}: {issue.message}")
            return 1

    print(
        f"wrote {dst} (profile v{profile.schema_version}; "
        f"sections: experience"
        f"{'+education' if profile.enabled.education else ''}"
        f"{'+projects' if profile.enabled.projects else ''}"
        f"{'+skills' if profile.enabled.skills else ''})"
    )
    return 0


def main(argv: list[str] | None = None) -> int:
    """CLI entry: optional ``--from``, ``--profile``, ``--out``, ``--workspace``."""
    import argparse
    import sys

    from . import workspace

    parser = argparse.ArgumentParser(description="Build main_template.docx from a baseline export.")
    parser.add_argument(
        "--from",
        dest="source",
        type=Path,
        default=None,
        help="Baseline DOCX (default: templates/original_export.docx).",
    )
    parser.add_argument(
        "--out",
        type=Path,
        default=None,
        help="Tagged output path (default: templates/main_template.docx).",
    )
    parser.add_argument(
        "--profile",
        type=Path,
        default=None,
        help="Template profile JSON (default: templates/template_profile.json if present).",
    )
    parser.add_argument(
        "--workspace",
        default=None,
        metavar="ID",
        help="Build against this profile instead of the active one (this invocation only).",
    )
    parser.add_argument(
        "--no-verify",
        dest="verify",
        action="store_false",
        help=(
            "Skip post-build verification (tag presence + a real render's field "
            "values) that runs by default for a profile-based build."
        ),
    )
    args = parser.parse_args(argv)
    try:
        workspace.bootstrap(workspace_id=args.workspace)
    except workspace.WorkspaceError as exc:
        print(f"error: {exc}", file=sys.stderr)
        return 1
    return build(args.source, args.out, profile_path=args.profile, verify=args.verify)


if __name__ == "__main__":
    raise SystemExit(main())
