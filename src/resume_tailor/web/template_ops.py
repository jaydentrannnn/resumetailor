"""Filesystem operations for the Template tab: inspect, preview, analyze, install.

All Word/LibreOffice work is serialised behind `LOCK` so a preview render and a
rebuild never overlap. `scripts/build_template.py` remains the sole CLI producer of
`main_template.docx`; this module shells out to it (or calls `template_build` for
staged profile installs).

Named library snapshots live under `templates/library/`; the live slot remains the
single `original_export.docx` / `main_template.docx` / `template_profile.json` trio.
"""

from __future__ import annotations

import hashlib
import json
import re
import secrets
import shutil
import subprocess
import sys
import tempfile
import threading
from datetime import datetime, timezone
from pathlib import Path

import docx

from resume_tailor import (
    calibrate,
    config,
    data,
    render,
    template_analyze,
    template_build,
    template_profile,
    template_verify,
)
from resume_tailor.labels import label_taken, normalize_label
from resume_tailor.template_profile import TemplateProfile, save_profile
from resume_tailor.web.schemas import (
    CalibrationInfo,
    TemplateAnalyzeResponse,
    TemplateBuildResponse,
    TemplateFieldCandidateOut,
    TemplateFileInfo,
    TemplateInfoResponse,
    TemplateIssueOut,
    TemplateLibraryEntry,
    TemplateLibraryResponse,
    TemplateParagraphOut,
    TemplateProfileSummary,
    TemplateSectionOut,
)

#: Reject uploads larger than this before writing anything under templates/.
_MAX_UPLOAD_BYTES = 10 * 1024 * 1024

#: Cap on named library entries (user must delete before adding more).
_LIBRARY_MAX_ENTRIES = 20

#: Max length for a user-facing library label.
_LIBRARY_LABEL_MAX = 80

#: Serialises preview render and baseline install so Word/LibreOffice is never concurrent.
LOCK = threading.Lock()


class TemplateValidationError(ValueError):
    """Raised when an upload fails a pre-install check (extension, size, OOXML)."""


class TemplateBuildError(RuntimeError):
    """Raised when template build fails; `.log` holds stdout/stderr or smoke-render detail."""

    def __init__(self, message: str, *, log: str = "") -> None:
        """Store the failure message and the captured build log."""
        super().__init__(message)
        self.log = log


def _file_info(path: Path) -> TemplateFileInfo:
    """Build a TemplateFileInfo for `path`, whether or not it exists."""
    if not path.exists():
        return TemplateFileInfo(exists=False, path=str(path))
    stat = path.stat()
    return TemplateFileInfo(
        exists=True,
        path=str(path),
        size_bytes=stat.st_size,
        modified_at=datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc).isoformat(
            timespec="seconds"
        ),
    )


def _other_backend_calibrations(*, exclude: str) -> list[str]:
    """Backend names (other than `exclude`) that already have a calibration file.

    Used to explain a "fallback" calibration state precisely: fit constants are not
    portable between Word and LibreOffice (they lay text out slightly differently), so
    a file present for the *other* backend cannot silently cover for a missing one —
    but a user seeing "fallback" with no further detail has no way to know that a real
    measurement exists at all, just for the wrong engine.
    """
    return [
        backend
        for backend in config.PDF_BACKENDS
        if backend != exclude and (config.CALIBRATION_DIR / f"{backend}.json").exists()
    ]


def _calibration_info(tagged: Path) -> CalibrationInfo:
    """Report fit-constant freshness relative to the tagged template's mtime."""
    cal_path = config.CALIBRATION_DIR / f"{config.PDF_BACKEND}.json"
    source = config.CALIBRATION_SOURCE
    rejection = config.CALIBRATION_REJECTION
    chars = config.CHARS_PER_LINE
    lines = config.LINES_PER_PAGE

    if not tagged.exists():
        return CalibrationInfo(
            source=source,
            chars_per_line=chars,
            lines_per_page=lines,
            stale=True,
            message="Tagged template is missing; generate it before calibrating.",
        )

    if rejection:
        # A file exists on disk but its numbers were implausible (see
        # `config.PLAUSIBLE_CHARS_PER_LINE`/`PLAUSIBLE_LINES_PER_PAGE`) — distinct from
        # "no file at all" below, and must be checked first: `cal_path.exists()` is
        # True here, so the next branch's "no calibration file" message would be wrong.
        return CalibrationInfo(
            source=source,
            chars_per_line=chars,
            lines_per_page=lines,
            stale=True,
            message=rejection,
        )

    if not cal_path.exists() or source == "fallback":
        other = _other_backend_calibrations(exclude=config.PDF_BACKEND)
        hint = (
            f" A calibration file exists for {', '.join(other)}, but fit constants are "
            "not portable between PDF backends — that measurement cannot be reused here."
            if other
            else ""
        )
        return CalibrationInfo(
            source=source,
            chars_per_line=chars,
            lines_per_page=lines,
            stale=True,
            message=(
                "No calibration file for this PDF backend. Re-upload with "
                "“Also calibrate fit constants”, or run `python scripts/calibrate.py`."
                + hint
            ),
        )

    # Module-level CHARS_PER_LINE / LINES_PER_PAGE were loaded at import time; if the
    # template is newer than the cal file, those numbers describe a previous layout.
    stale = tagged.stat().st_mtime > cal_path.stat().st_mtime
    return CalibrationInfo(
        source=source,
        chars_per_line=chars,
        lines_per_page=lines,
        stale=stale,
        message=(
            "Template is newer than calibration. Re-upload with "
            "“Also calibrate fit constants”, or run `python scripts/calibrate.py`."
            if stale
            else None
        ),
    )


def _profile_summary() -> TemplateProfileSummary:
    """Summarise the active template profile for the Template tab."""
    profile = template_profile.load_profile()
    if profile is None:
        return TemplateProfileSummary(exists=False)
    return TemplateProfileSummary(
        exists=True,
        schema_version=profile.schema_version,
        enabled=profile.enabled.model_dump(),
        warnings=list(profile.warnings),
        contact_separator=profile.contact.separator,
    )


def _preview_paths() -> tuple[Path, Path]:
    """Return `(preview.docx, preview.pdf)` under `output/template/`."""
    directory = config.OUTPUT_DIR / "template"
    return directory / "preview.docx", directory / "preview.pdf"


def info() -> TemplateInfoResponse:
    """Collect baseline/tagged metadata, master-resume counts, and calibration status."""
    baseline = _file_info(config.BASELINE_TEMPLATE_PATH)
    tagged = _file_info(config.DEFAULT_TEMPLATE_PATH)
    experience_entries = 0
    project_entries = 0
    bullets = 0
    try:
        resume = data.load()
        experience_entries = len(resume.experience)
        project_entries = len(resume.projects)
        bullets = len(resume.all_bullets())
    except (FileNotFoundError, ValueError):
        pass

    active_id, active_label = _library_active_meta()
    _, pdf_path = _preview_paths()
    return TemplateInfoResponse(
        baseline=baseline,
        tagged=tagged,
        experience_entries=experience_entries,
        project_entries=project_entries,
        bullets=bullets,
        calibration=_calibration_info(config.DEFAULT_TEMPLATE_PATH),
        preview_available=pdf_path.exists(),
        profile=_profile_summary(),
        active_library_id=active_id,
        active_label=active_label,
    )


# ---------------------------------------------------------------------------
# Named template library
# ---------------------------------------------------------------------------


def _library_root() -> Path:
    """Return the library directory (creates it on demand)."""
    root = getattr(config, "TEMPLATE_LIBRARY_DIR", None) or (
        config.TEMPLATES_DIR / "library"
    )
    root.mkdir(parents=True, exist_ok=True)
    return root


def _library_index_path() -> Path:
    """Path to the lightweight `{active_id}` index file."""
    return _library_root() / "index.json"


def _library_entry_dir(entry_id: str) -> Path:
    """Directory for one library snapshot."""
    return _library_root() / entry_id


def _read_library_index() -> dict:
    """Load `index.json`, or an empty active pointer when missing/corrupt."""
    path = _library_index_path()
    if not path.exists():
        return {"active_id": None}
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return {"active_id": None}
    if not isinstance(data, dict):
        return {"active_id": None}
    return {"active_id": data.get("active_id")}


def _write_library_index(*, active_id: str | None) -> None:
    """Persist the active library pointer."""
    path = _library_index_path()
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps({"active_id": active_id}, indent=2) + "\n",
        encoding="utf-8",
    )


def _normalize_library_label(label: str) -> str:
    """Strip and validate a library label; raise TemplateValidationError if invalid."""
    try:
        return normalize_label(label, max_len=_LIBRARY_LABEL_MAX, what="Template label")
    except ValueError as exc:
        raise TemplateValidationError(str(exc)) from exc


def _default_label_from_filename(filename: str) -> str:
    """Derive a library label from an upload filename stem."""
    stem = Path(filename).stem.strip() or "Untitled"
    # Collapse runs of whitespace / underscores for a readable default.
    stem = re.sub(r"[\s_]+", " ", stem).strip()
    return _normalize_library_label(stem[:_LIBRARY_LABEL_MAX])


def _sha256_file(path: Path) -> str | None:
    """Return hex SHA-256 of a file, or None if it does not exist."""
    if not path.exists():
        return None
    h = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def _new_library_id() -> str:
    """Allocate a filesystem-safe unique library entry id."""
    stamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    return f"{stamp}-{secrets.token_hex(2)}"


def _load_entry_meta(entry_id: str) -> dict | None:
    """Load `meta.json` for `entry_id`, or None if missing/invalid."""
    meta_path = _library_entry_dir(entry_id) / "meta.json"
    if not meta_path.exists():
        return None
    try:
        raw = json.loads(meta_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return None
    if not isinstance(raw, dict) or raw.get("id") != entry_id:
        return None
    return raw


def _iter_library_metas() -> list[dict]:
    """Return all valid entry metas, newest first."""
    root = _library_root()
    metas: list[dict] = []
    for child in root.iterdir():
        if not child.is_dir():
            continue
        meta = _load_entry_meta(child.name)
        if meta is not None:
            metas.append(meta)
    metas.sort(key=lambda m: m.get("created_at") or "", reverse=True)
    return metas


def _library_active_meta() -> tuple[str | None, str | None]:
    """Return `(active_id, active_label)` from the index, or `(None, None)`."""
    active_id = _read_library_index().get("active_id")
    if not active_id:
        return None, None
    meta = _load_entry_meta(str(active_id))
    if meta is None:
        return None, None
    return str(active_id), str(meta.get("label") or active_id)


def _label_taken(label: str, *, except_id: str | None = None) -> bool:
    """True when another entry already uses `label` (case-insensitive)."""
    others = (
        str(meta.get("label") or "")
        for meta in _iter_library_metas()
        if not (except_id and meta.get("id") == except_id)
    )
    return label_taken(label, others)


def _entry_to_schema(meta: dict, *, active_id: str | None) -> TemplateLibraryEntry:
    """Build an API entry from on-disk meta."""
    entry_id = str(meta["id"])
    entry_dir = _library_entry_dir(entry_id)
    baseline = entry_dir / "original_export.docx"
    size = baseline.stat().st_size if baseline.exists() else None
    return TemplateLibraryEntry(
        id=entry_id,
        label=str(meta.get("label") or entry_id),
        created_at=str(meta.get("created_at") or ""),
        source_filename=meta.get("source_filename"),
        size_bytes=size,
        has_profile=bool(meta.get("has_profile")),
        is_active=entry_id == active_id,
    )


def _snapshot_live_to_library(
    *,
    label: str,
    source_filename: str | None,
) -> dict:
    """Copy the live baseline/tagged/profile into a new library entry.

    Requires the live baseline and tagged template to exist. Caller must hold
    `LOCK` (or accept races) and have already validated label uniqueness / cap.
    """
    baseline = config.BASELINE_TEMPLATE_PATH
    tagged = config.DEFAULT_TEMPLATE_PATH
    if not baseline.exists():
        raise TemplateValidationError("No live baseline to save into the library.")
    if not tagged.exists():
        raise TemplateValidationError(
            "No tagged template to save; rebuild before adding to the library."
        )

    entry_id = _new_library_id()
    entry_dir = _library_entry_dir(entry_id)
    entry_dir.mkdir(parents=True, exist_ok=False)
    try:
        shutil.copy2(baseline, entry_dir / "original_export.docx")
        shutil.copy2(tagged, entry_dir / "main_template.docx")
        profile_file = template_profile.profile_path()
        has_profile = profile_file.exists()
        if has_profile:
            shutil.copy2(profile_file, entry_dir / "template_profile.json")
        sha = _sha256_file(entry_dir / "original_export.docx")
        meta = {
            "id": entry_id,
            "label": label,
            "created_at": datetime.now(timezone.utc).isoformat(timespec="seconds"),
            "source_filename": source_filename,
            "sha256": sha,
            "has_profile": has_profile,
        }
        (entry_dir / "meta.json").write_text(
            json.dumps(meta, indent=2) + "\n",
            encoding="utf-8",
        )
    except Exception:
        shutil.rmtree(entry_dir, ignore_errors=True)
        raise
    return meta


def _find_entry_by_sha(sha: str | None) -> dict | None:
    """Return the first library meta whose baseline sha256 matches `sha`."""
    if not sha:
        return None
    for meta in _iter_library_metas():
        if meta.get("sha256") == sha:
            return meta
    return None


def _library_seed_if_empty() -> None:
    """If the library is empty and a live baseline exists, register it as Default."""
    if _iter_library_metas():
        return
    if not config.BASELINE_TEMPLATE_PATH.exists():
        return
    if not config.DEFAULT_TEMPLATE_PATH.exists():
        return
    meta = _snapshot_live_to_library(label="Default", source_filename=None)
    _write_library_index(active_id=meta["id"])


def _library_ensure_room_for_new(*, label: str) -> None:
    """Raise if the library cannot accept a new entry with `label`."""
    label = _normalize_library_label(label)
    if _label_taken(label):
        raise TemplateValidationError(
            f"A saved template named “{label}” already exists. Choose another label."
        )
    if len(_iter_library_metas()) >= _LIBRARY_MAX_ENTRIES:
        raise TemplateValidationError(
            f"Template library is full ({_LIBRARY_MAX_ENTRIES} saved). "
            "Delete one before installing another."
        )


def _library_preserve_orphan_live() -> None:
    """If live baseline bytes are not in the library, snapshot them before overwrite.

    Uses label ``Default`` when the library is empty, otherwise an ``Autosaved …``
    stamp. No-op when live is missing or already matched by sha256. Raises when the
    library is at capacity and the orphan cannot be saved.
    """
    if not config.BASELINE_TEMPLATE_PATH.exists():
        return
    if not config.DEFAULT_TEMPLATE_PATH.exists():
        return
    sha = _sha256_file(config.BASELINE_TEMPLATE_PATH)
    if _find_entry_by_sha(sha) is not None:
        return
    metas = _iter_library_metas()
    if len(metas) >= _LIBRARY_MAX_ENTRIES:
        raise TemplateValidationError(
            f"Template library is full ({_LIBRARY_MAX_ENTRIES} saved) and the "
            "current live template is not saved. Delete one before replacing."
        )
    label = "Default" if not metas else (
        f"Autosaved {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M')} UTC"
    )
    # Disambiguate if the autosaved label somehow collides.
    base = label
    n = 2
    while _label_taken(label):
        label = f"{base} ({n})"
        n += 1
    # Do not change active_id — caller is about to replace live content.
    _snapshot_live_to_library(label=label, source_filename=None)


def list_library() -> TemplateLibraryResponse:
    """Seed Default if needed, then return all named library entries."""
    with LOCK:
        _library_seed_if_empty()
        active_id = _read_library_index().get("active_id")
        entries = [
            _entry_to_schema(m, active_id=active_id) for m in _iter_library_metas()
        ]
        # If index points nowhere but we have entries matching live sha, heal active.
        if active_id is None and entries:
            sha = _sha256_file(config.BASELINE_TEMPLATE_PATH)
            match = _find_entry_by_sha(sha)
            if match is not None:
                active_id = match["id"]
                _write_library_index(active_id=active_id)
                entries = [
                    _entry_to_schema(m, active_id=active_id)
                    for m in _iter_library_metas()
                ]
        return TemplateLibraryResponse(
            entries=entries,
            active_id=str(active_id) if active_id else None,
        )


def rename_library_entry(entry_id: str, label: str) -> TemplateLibraryResponse:
    """Rename a library entry; labels must stay unique (case-insensitive)."""
    cleaned = _normalize_library_label(label)
    with LOCK:
        meta = _load_entry_meta(entry_id)
        if meta is None:
            raise TemplateValidationError(f"Unknown template library id: {entry_id}")
        if _label_taken(cleaned, except_id=entry_id):
            raise TemplateValidationError(
                f"A saved template named “{cleaned}” already exists."
            )
        meta["label"] = cleaned
        meta_path = _library_entry_dir(entry_id) / "meta.json"
        meta_path.write_text(json.dumps(meta, indent=2) + "\n", encoding="utf-8")
    return list_library()


def delete_library_entry(entry_id: str) -> TemplateLibraryResponse:
    """Delete a non-active library entry from disk."""
    with LOCK:
        meta = _load_entry_meta(entry_id)
        if meta is None:
            raise TemplateValidationError(f"Unknown template library id: {entry_id}")
        active_id = _read_library_index().get("active_id")
        if entry_id == active_id:
            raise TemplateValidationError(
                "Cannot delete the active template. Activate another first."
            )
        shutil.rmtree(_library_entry_dir(entry_id), ignore_errors=False)
    return list_library()


def activate_library_entry(
    entry_id: str,
    *,
    do_calibrate: bool = False,
) -> TemplateBuildResponse:
    """Copy a library snapshot into the live slot; optionally recalibrate.

    Preserves the current live set as a new library entry when its baseline sha is
    not already present. Does not re-run `build_template` — tagged files travel with
    the snapshot.
    """
    with LOCK:
        meta = _load_entry_meta(entry_id)
        if meta is None:
            raise TemplateValidationError(f"Unknown template library id: {entry_id}")

        entry_dir = _library_entry_dir(entry_id)
        src_baseline = entry_dir / "original_export.docx"
        src_tagged = entry_dir / "main_template.docx"
        src_profile = entry_dir / "template_profile.json"
        if not src_baseline.exists() or not src_tagged.exists():
            raise TemplateValidationError(
                f"Library entry {entry_id} is incomplete (missing baseline or tagged)."
            )

        # Preserve orphan live content before overwrite (may raise if library full).
        _library_preserve_orphan_live()

        baseline = config.BASELINE_TEMPLATE_PATH
        tagged = config.DEFAULT_TEMPLATE_PATH
        profile_file = template_profile.profile_path()
        baseline.parent.mkdir(parents=True, exist_ok=True)

        previous_baseline = baseline.read_bytes() if baseline.exists() else None
        previous_tagged = tagged.read_bytes() if tagged.exists() else None
        previous_profile = profile_file.read_bytes() if profile_file.exists() else None

        def _restore() -> None:
            """Roll live files back after a failed activate commit."""
            if previous_baseline is not None:
                baseline.write_bytes(previous_baseline)
            elif baseline.exists():
                baseline.unlink()
            if previous_tagged is not None:
                tagged.write_bytes(previous_tagged)
            elif tagged.exists():
                tagged.unlink()
            if previous_profile is not None:
                profile_file.write_bytes(previous_profile)
            elif profile_file.exists():
                profile_file.unlink()

        try:
            # Smoke-open the tagged snapshot before committing.
            docx.Document(str(src_tagged))
            shutil.copy2(src_baseline, baseline)
            shutil.copy2(src_tagged, tagged)
            if src_profile.exists():
                shutil.copy2(src_profile, profile_file)
            elif profile_file.exists():
                profile_file.unlink()
        except Exception as exc:
            _restore()
            raise TemplateBuildError(
                f"Failed to activate library template: {exc}",
                log=str(exc),
            ) from exc

        _write_library_index(active_id=entry_id)
        invalidate_preview()
        log = f"Activated library template “{meta.get('label', entry_id)}” ({entry_id})."

    if do_calibrate:
        log = _maybe_calibrate(log, do_calibrate=True)
    return TemplateBuildResponse(ok=True, log=log, info=info())


def _library_record_after_install(*, label: str, source_filename: str) -> None:
    """Snapshot the just-installed live slot and mark it active."""
    cleaned = _normalize_library_label(label)
    # Cap / uniqueness were checked before install; re-check in case of races.
    if _label_taken(cleaned):
        # Install already committed — disambiguate rather than failing the install.
        base = cleaned
        n = 2
        while _label_taken(cleaned):
            cleaned = f"{base} ({n})"
            n += 1
    if len(_iter_library_metas()) >= _LIBRARY_MAX_ENTRIES:
        # Extremely unlikely after pre-check; skip snapshot rather than raise.
        return
    meta = _snapshot_live_to_library(label=cleaned, source_filename=source_filename)
    _write_library_index(active_id=meta["id"])


def ensure_preview() -> Path:
    """Render the full master resume through the tagged template and return its PDF path.

    Regenerates when the cached PDF is missing or older than **either** of its two
    inputs. Raises `RuntimeError` when PDF conversion is unavailable so the route can
    return 503 instead of a broken frame.
    """
    with LOCK:
        tagged = config.DEFAULT_TEMPLATE_PATH
        if not tagged.exists():
            raise FileNotFoundError(
                f"Tagged template not found: {tagged}. "
                "Upload a baseline or run `python scripts/build_template.py`."
            )

        docx_path, pdf_path = _preview_paths()
        pdf_path.parent.mkdir(parents=True, exist_ok=True)

        # The preview is the master resume rendered *through* the tagged template, so
        # both files are inputs. Keying staleness on the template alone meant editing
        # the resume never invalidated the cache — the Template tab kept serving the
        # pre-edit PDF indefinitely, including after an explicit Refresh. Checking
        # mtime (rather than having the editor call `invalidate_preview`) also covers
        # a master_resume.json edited by hand or by the CLI.
        newest_input = max(
            (p.stat().st_mtime for p in (tagged, config.MASTER_RESUME_PATH) if p.exists()),
            default=0.0,
        )
        needs_render = (
            not pdf_path.exists()
            or not docx_path.exists()
            or newest_input > pdf_path.stat().st_mtime
        )
        if not needs_render:
            return pdf_path

        resume = data.load()
        render.render(resume, out=docx_path)
        # Propagate RuntimeError from convert so the UI can say "no PDF backend".
        render.to_pdf(docx_path, pdf_path)
        return pdf_path


def invalidate_preview() -> None:
    """Delete the cached preview so the next `ensure_preview` regenerates it."""
    docx_path, pdf_path = _preview_paths()
    for path in (docx_path, pdf_path):
        if path.exists():
            path.unlink()


def _validate_upload_bytes(raw: bytes, filename: str) -> None:
    """Raise TemplateValidationError for bad extension, size, or empty payload."""
    name = Path(filename).name
    if not name.lower().endswith(".docx"):
        raise TemplateValidationError("Upload must be a .docx file.")
    if len(raw) > _MAX_UPLOAD_BYTES:
        raise TemplateValidationError(
            f"Upload is {len(raw)} bytes; maximum is {_MAX_UPLOAD_BYTES}."
        )
    if not raw:
        raise TemplateValidationError("Upload is empty.")


def _upload_cache_dir() -> Path:
    """Directory holding uploaded bytes for the wizard's remap/preview steps.

    Keyed by the upload's own sha256 so the remap and preview endpoints never need to
    re-upload the file — the wizard already sent it once, to `/analyze`. Lives under
    the active workspace's own `output/` root (via `config.OUTPUT_DIR`, one of the
    globals `set_active_workspace` rebinds), so switching profiles never leaks one
    workspace's in-progress upload into another's.
    """
    return config.OUTPUT_DIR / "template" / "uploads"


def _prune_upload_cache(*, max_age_seconds: float = 24 * 3600) -> None:
    """Delete cached uploads older than `max_age_seconds`.

    A wizard session that is abandoned mid-flow (tab closed after analyze, before
    install or reset) would otherwise leak its upload forever; this runs opportunistically
    whenever a new upload is cached, which is cheap and frequent enough that unbounded
    growth never accumulates.
    """
    directory = _upload_cache_dir()
    if not directory.exists():
        return
    cutoff = datetime.now(timezone.utc).timestamp() - max_age_seconds
    for entry in directory.glob("*.docx"):
        try:
            if entry.stat().st_mtime < cutoff:
                entry.unlink()
        except OSError:
            pass


def _cache_upload(raw: bytes, sha: str) -> None:
    """Persist an analyzed upload's bytes for later remap/preview calls."""
    _prune_upload_cache()
    directory = _upload_cache_dir()
    directory.mkdir(parents=True, exist_ok=True)
    (directory / f"{sha}.docx").write_bytes(raw)


def _load_cached_upload(sha: str) -> bytes:
    """Load a previously analyzed upload's bytes by its sha256.

    Raises `TemplateValidationError` (400, not 404/500) when the cache has expired or
    the wizard is pointing at a sha that was never analyzed — both are "start over",
    the same class of problem as a bad upload.
    """
    path = _upload_cache_dir() / f"{sha}.docx"
    if not path.exists():
        raise TemplateValidationError(
            "This upload is no longer available for preview/remap — analyze the file "
            "again."
        )
    return path.read_bytes()


def clear_upload_cache() -> None:
    """Drop every cached upload — called after a successful install, once the wizard's
    staged file has become the live template and has no further use as a cache entry."""
    directory = _upload_cache_dir()
    if directory.exists():
        shutil.rmtree(directory, ignore_errors=True)


def _analysis_to_response(result: template_analyze.AnalyzeResult) -> TemplateAnalyzeResponse:
    """Shared `AnalyzeResult` -> wire-shape conversion for `analyze_upload`/`remap_upload`."""
    return TemplateAnalyzeResponse(
        source_sha256=result.source_sha256,
        paragraphs=[
            TemplateParagraphOut(**p.model_dump()) for p in result.paragraphs
        ],
        sections=[TemplateSectionOut(**s.model_dump()) for s in result.sections],
        field_candidates=[
            TemplateFieldCandidateOut(
                field=c.field,
                paragraph_id=c.span.paragraph_id,
                start=c.span.start,
                end=c.span.end,
                confidence=c.confidence,
                preview=c.preview,
                section_heading_paragraph_id=c.section_heading_paragraph_id,
            )
            for c in result.field_candidates
        ],
        suggested_profile=(
            result.suggested_profile.model_dump(mode="json")
            if result.suggested_profile is not None
            else None
        ),
        issues=[TemplateIssueOut(**i.model_dump()) for i in result.issues],
        ready=result.ready,
    )


def analyze_upload(raw: bytes, filename: str) -> TemplateAnalyzeResponse:
    """Analyse an uploaded DOCX without writing under templates/."""
    _validate_upload_bytes(raw, filename)
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(raw)
        tmp_path = Path(tmp.name)
    try:
        try:
            docx.Document(str(tmp_path))
        except Exception as exc:
            raise TemplateValidationError(f"File is not a readable .docx: {exc}") from exc
        result = template_analyze.analyze_docx(raw=raw)
    finally:
        tmp_path.unlink(missing_ok=True)

    _cache_upload(raw, result.source_sha256)
    return _analysis_to_response(result)


def remap_upload(source_sha256: str, overrides: dict[int, str | None]) -> TemplateAnalyzeResponse:
    """Re-run analysis on a previously uploaded file with specific headings' kinds
    forced by the user, bypassing every heuristic gate for just those paragraphs (see
    `template_analyze._analyze_document`'s `overrides` parameter)."""
    raw = _load_cached_upload(source_sha256)
    result = template_analyze.analyze_docx(raw=raw, overrides=overrides)
    return _analysis_to_response(result)


def preview_source(source_sha256: str) -> Path:
    """Convert a previously uploaded (not-yet-installed) baseline to PDF, for the
    wizard's side-by-side comparison. Raises `RuntimeError` when no PDF backend is
    available, matching `ensure_preview`'s contract for the same 503 the route returns."""
    raw = _load_cached_upload(source_sha256)
    directory = _upload_cache_dir()
    docx_path = directory / f"{source_sha256}.source.docx"
    pdf_path = directory / f"{source_sha256}.source.pdf"
    # Keyed by the upload's own sha256, so the source bytes for a given path can never
    # change underneath it — an existing PDF here is never stale, unlike `ensure_preview`'s
    # mtime check (which exists only because *that* cache's docx input can be re-rendered
    # in place from an edited master resume).
    if pdf_path.exists():
        return pdf_path
    docx_path.write_bytes(raw)
    with LOCK:
        render.to_pdf(docx_path, pdf_path)
    return pdf_path


def preview_draft(source_sha256: str, profile: TemplateProfile | dict) -> Path:
    """Build a staged profile into a temp tagged template and render the master resume
    through it, without touching the live template slot — lets the wizard show what
    installing this exact mapping would produce before committing to it.

    Runs the same build `_install_with_profile` does, minus the atomic commit: no
    baseline/tagged/profile file under `templates/` is ever written or replaced.
    """
    confirmed = (
        profile if isinstance(profile, TemplateProfile) else TemplateProfile.model_validate(profile)
    )
    raw = _load_cached_upload(source_sha256)

    directory = _upload_cache_dir()
    directory.mkdir(parents=True, exist_ok=True)
    with LOCK:
        with tempfile.TemporaryDirectory() as stage_dir:
            stage = Path(stage_dir)
            staged_src = stage / "baseline.docx"
            staged_tagged = stage / "main_template.docx"
            staged_src.write_bytes(raw)
            template_build.build_from_profile(staged_src, staged_tagged, confirmed)

            docx_path = directory / f"{source_sha256}.draft.docx"
            pdf_path = directory / f"{source_sha256}.draft.pdf"
            render.render(
                data.load(),
                template=staged_tagged,
                out=docx_path,
                layout=template_profile.active_layout(confirmed),
            )
            render.to_pdf(docx_path, pdf_path)
    return pdf_path


def _run_build(
    *,
    source: Path | None = None,
    output: Path | None = None,
    profile_path: Path | None = None,
) -> tuple[int, str]:
    """Shell out to `scripts/build_template.py`. Returns `(exit_code, combined_log)`.

    Kept as a named function so tests can monkeypatch it without hitting Word or the
    real build script.
    """
    script = config.PROJECT_ROOT / "scripts" / "build_template.py"
    cmd = [sys.executable, str(script)]
    if source is not None:
        cmd.extend(["--from", str(source)])
    if output is not None:
        cmd.extend(["--out", str(output)])
    if profile_path is not None:
        cmd.extend(["--profile", str(profile_path)])
    # The subprocess re-imports config fresh, so without this it would build against
    # the legacy/default paths regardless of which workspace this process has active.
    active_workspace = config.active_workspace_id()
    if active_workspace is not None:
        cmd.extend(["--workspace", active_workspace])
    completed = subprocess.run(
        cmd,
        capture_output=True,
        text=True,
        cwd=str(config.PROJECT_ROOT),
        check=False,
    )
    log = (completed.stdout or "") + (completed.stderr or "")
    return completed.returncode, log


def _smoke_render(tagged: Path) -> None:
    """Render the master resume into a temp DOCX and reopen it (no PDF).

    Catches tag/XML mistakes before the staged files replace the live template.
    """
    resume = data.load()
    with tempfile.TemporaryDirectory() as tmp:
        out = Path(tmp) / "smoke.docx"
        render.render(resume, template=tagged, out=out)
        # Re-open: catches OOXML that python-docx/docxtpl wrote but Word would refuse.
        docx.Document(str(out))


def _verify_staged_build(tagged: Path, profile: TemplateProfile) -> None:
    """Hold a staged, profile-based build accountable to its own mapping before it can
    replace the live template — `_smoke_render` only proves the file opens; a template
    can smoke-render cleanly and still be missing a field entirely (an unmapped
    `dates`/`location` silently drops from every job) or point a mapped field at the
    wrong span (builds and renders fine, just with the wrong text). Both are
    `TemplateBuildError`s, the same as a smoke-render failure — a staged install that
    fails either check must never reach the live slot.
    """
    issues = template_verify.verify_tagged(tagged, profile)
    issues += template_verify.verify_roundtrip(tagged, profile, data.load())
    blockers = [i for i in issues if i.blocking]
    if blockers:
        detail = "; ".join(i.message for i in blockers)
        raise TemplateBuildError(
            f"Build verification failed: {detail}",
            log="\n".join(f"{i.code}: {i.message}" for i in blockers),
        )


def _maybe_calibrate(log: str, *, do_calibrate: bool) -> str:
    """Optionally measure fit constants and hot-reload them into `config`.

    Runs under `LOCK` so Word/LibreOffice is not concurrent with preview. Failures are
    appended to the log but do not undo a successful template install.
    """
    if not do_calibrate:
        return log
    with LOCK:
        try:
            # A fresh template install changes `template_sha256`, so the render-anchor
            # check re-baselines against this workspace's own resume+template rather
            # than warning — see `calibrate.check_render_anchors`.
            result = calibrate.run(verify_anchors=True)
            config.reload_calibration()
            combined = (log + "\n\n" + result.log).strip()
            if result.warnings:
                combined += "\n" + "\n".join(f"warning: {w}" for w in result.warnings)
            return combined
        except Exception as exc:
            return (log + f"\n\nCalibration failed (template install kept): {exc}").strip()


def install_baseline(
    raw: bytes,
    filename: str,
    *,
    profile: TemplateProfile | dict,
    do_calibrate: bool = False,
    label: str | None = None,
) -> TemplateBuildResponse:
    """Validate an uploaded .docx, rebuild the tagged template, then swap live files.

    Build + smoke-render happen in a temp directory and only then replace baseline,
    tagged template, and `template_profile.json` together.

    When `do_calibrate` is True, measure CHARS_PER_LINE / LINES_PER_PAGE after a successful
    install and reload them in-process (no server restart).

    After a successful install the live slot is snapshotted into the named library under
    `label` (default: upload filename stem).

    Raises `TemplateValidationError` for bad inputs and `TemplateBuildError` on build
    failure.
    """
    _validate_upload_bytes(raw, filename)
    resolved_label = _normalize_library_label(
        label
        if (label or "").strip()
        else _default_label_from_filename(filename)
    )

    with LOCK:
        _library_seed_if_empty()
        _library_preserve_orphan_live()
        _library_ensure_room_for_new(label=resolved_label)

    response = _install_with_profile(raw, filename, profile)

    with LOCK:
        _library_record_after_install(
            label=resolved_label,
            source_filename=Path(filename).name,
        )
    # The wizard's cached upload (for remap/preview) has now become the live template;
    # nothing further needs it, and keeping it around would just be dead weight until
    # `_prune_upload_cache`'s 24h sweep got to it.
    clear_upload_cache()

    if do_calibrate:
        new_log = _maybe_calibrate(response.log, do_calibrate=True)
        return TemplateBuildResponse(ok=True, log=new_log, info=info())
    return TemplateBuildResponse(ok=True, log=response.log, info=info())


def _install_with_profile(
    raw: bytes,
    filename: str,
    profile: TemplateProfile | dict,
) -> TemplateBuildResponse:
    """Staged profile install: build + smoke-render, then atomic commit."""
    confirmed = (
        profile
        if isinstance(profile, TemplateProfile)
        else TemplateProfile.model_validate(profile)
    )
    issues = template_analyze.validate_profile_against_doc(confirmed, raw=raw)
    blockers = [i for i in issues if i.blocking]
    if blockers:
        raise TemplateValidationError("; ".join(i.message for i in blockers))

    with LOCK:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(raw)
            tmp_path = Path(tmp.name)
        try:
            docx.Document(str(tmp_path))
        except Exception as exc:
            tmp_path.unlink(missing_ok=True)
            raise TemplateValidationError(
                f"File is not a readable .docx: {exc}"
            ) from exc

        baseline = config.BASELINE_TEMPLATE_PATH
        tagged = config.DEFAULT_TEMPLATE_PATH
        profile_file = template_profile.profile_path()
        baseline.parent.mkdir(parents=True, exist_ok=True)

        previous_baseline: bytes | None = (
            baseline.read_bytes() if baseline.exists() else None
        )
        previous_tagged: bytes | None = tagged.read_bytes() if tagged.exists() else None
        previous_profile: bytes | None = (
            profile_file.read_bytes() if profile_file.exists() else None
        )

        if previous_baseline is not None:
            stamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
            backup_dir = baseline.parent / "backups"
            backup_dir.mkdir(parents=True, exist_ok=True)
            shutil.copy2(baseline, backup_dir / f"original_export.{stamp}.docx")
            if previous_profile is not None:
                shutil.copy2(
                    profile_file, backup_dir / f"template_profile.{stamp}.json"
                )

        def _restore() -> None:
            """Roll baseline / tagged / profile back to their pre-install bytes."""
            if previous_baseline is not None:
                baseline.write_bytes(previous_baseline)
            elif baseline.exists():
                baseline.unlink()
            if previous_tagged is not None:
                tagged.write_bytes(previous_tagged)
            elif tagged.exists():
                tagged.unlink()
            if previous_profile is not None:
                profile_file.write_bytes(previous_profile)
            elif profile_file.exists():
                profile_file.unlink()

        with tempfile.TemporaryDirectory() as stage_dir:
            stage = Path(stage_dir)
            staged_src = stage / "baseline.docx"
            staged_out = stage / "main_template.docx"
            staged_profile = stage / "template_profile.json"
            shutil.copy2(tmp_path, staged_src)
            tmp_path.unlink(missing_ok=True)

            log = ""
            try:
                save_profile(confirmed, staged_profile)
                exit_code, log = _run_build(
                    source=staged_src,
                    output=staged_out,
                    profile_path=staged_profile,
                )
                if exit_code != 0 or not staged_out.exists():
                    # In-process fallback when the script seam fails or a test stub
                    # does not write the staged output path.
                    try:
                        template_build.build_from_profile(
                            staged_src, staged_out, confirmed
                        )
                        log = (
                            log + "\n(in-process profile build OK)\n"
                        ).strip()
                    except Exception as exc:
                        raise TemplateBuildError(
                            "Profile build failed during staged install.",
                            log=(log + f"\n{exc}").strip(),
                        ) from exc

                _smoke_render(staged_out)
                _verify_staged_build(staged_out, confirmed)
            except TemplateBuildError:
                raise
            except Exception as exc:
                raise TemplateBuildError(
                    f"Staged build or smoke render failed: {exc}",
                    log=(log + f"\n{exc}").strip(),
                ) from exc

            try:
                shutil.copy2(staged_src, baseline)
                shutil.copy2(staged_out, tagged)
                shutil.copy2(staged_profile, profile_file)
            except Exception as exc:
                _restore()
                raise TemplateBuildError(
                    f"Failed to commit staged template: {exc}",
                    log=log.strip(),
                ) from exc

        invalidate_preview()
        return TemplateBuildResponse(ok=True, log=log.strip(), info=info())
