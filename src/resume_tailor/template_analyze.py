"""Deterministic structural analysis of a single-column resume DOCX.

Produces ranked mapping suggestions and blocking compatibility issues. Never mutates
the document and never calls an LLM — the confirmed mapping is what `template_build`
consumes.
"""

from __future__ import annotations

import hashlib
import re
from collections import Counter
from collections.abc import Callable
from dataclasses import dataclass, field

import docx
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from pydantic import BaseModel, ConfigDict, Field

from . import config, docx_text
from .template_profile import (
    CharSpan,
    ContactField,
    ContactMapping,
    ContactSlot,
    DetectedSection,
    EducationMapping,
    EnabledSections,
    ExperienceMapping,
    HeaderFieldMapping,
    HeadingPrototype,
    ListMapping,
    NormalizationFlags,
    OptionalSpan,
    ProjectsMapping,
    SkillsMapping,
    SpacingProfile,
    TemplateProfile,
)

#: Canonical section keys the pipeline understands.
SECTION_KEYS = ("education", "experience", "projects", "skills")

#: `SectionCandidate.key`/`by_kind` use the legacy plural `"projects"` (matching
#: `SECTION_KEYS` and `EnabledSections.projects`); `template_profile.GenericSectionKind`
#: uses singular `"project"` (matching `data.Section`'s discriminator). Every other key is
#: spelled the same in both places. Without this remap, `DetectedSection(kind="projects")`
#: raises a `ValidationError` the first time a resume has two projects-shaped headings.
_TO_GENERIC_KIND: dict[str, str] = {"projects": "project"}

#: Heading aliases → canonical key. Exact match on stripped uppercase text first;
#: then substring heuristics below.
_HEADING_ALIASES: dict[str, str] = {
    "EDUCATION": "education",
    "ACADEMIC BACKGROUND": "education",
    "ACADEMICS": "education",
    "WORK EXPERIENCES": "experience",
    "WORK EXPERIENCE": "experience",
    "EXPERIENCE": "experience",
    "PROFESSIONAL EXPERIENCE": "experience",
    "EMPLOYMENT": "experience",
    "EMPLOYMENT HISTORY": "experience",
    "PROJECTS": "projects",
    "SELECTED PROJECTS": "projects",
    "PERSONAL PROJECTS": "projects",
    "SELECTED WORK": "projects",
    "SKILLS": "skills",
    "TECHNICAL SKILLS": "skills",
    "TECHNOLOGIES": "skills",
    "ADDITIONAL INFORMATION": "skills",
    "ADDITIONAL INFO": "skills",
}

_DATE_RE = re.compile(
    r"(?i)\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec|"
    r"january|february|march|april|june|july|august|september|october|november|december)"
    r"[a-z.]*\s+\d{4}"
    r"|(?:\d{4}\s*[-–—]\s*(?:\d{4}|present|current|now))"
    r"|(?:present|current)\b"
)

_EMAIL_RE = re.compile(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", re.I)
_PHONE_RE = re.compile(r"(?:\+?\d[\d\s().-]{7,}\d)")
_SEP_CANDIDATES = (" \u2022 ", " • ", " | ", " · ", " – ", " — ", " / ", " |")


class _Strict(BaseModel):
    model_config = ConfigDict(extra="forbid")


class ParagraphInfo(_Strict):
    """One body paragraph flattened for the mapping UI."""

    id: int
    text: str
    is_bullet: bool = False
    is_heading_candidate: bool = False
    has_tab: bool = False
    has_hyperlink: bool = False
    run_count: int = 0
    preview: str = ""


class FieldCandidate(_Strict):
    """A suggested character span for one semantic field."""

    field: str
    span: CharSpan
    confidence: float = Field(ge=0.0, le=1.0)
    preview: str = ""
    #: Which detected section (by `SectionCandidate.heading_paragraph_id`) this
    #: candidate belongs to. `None` only for a candidate emitted outside the
    #: per-section loop (there are none left after `_section_field_candidates`
    #: replaced the old kind-wide emission — kept optional so a `FieldCandidate`
    #: constructed elsewhere, e.g. in a test, still validates).
    section_heading_paragraph_id: int | None = None


class SectionCandidate(_Strict):
    """A detected section heading and its body range."""

    key: str
    heading_paragraph_id: int
    heading_text: str
    body_start: int
    body_end: int
    entry_count: int = 0
    bullet_count: int = 0
    confidence: float = Field(ge=0.0, le=1.0)
    aliases_matched: str = ""


class Issue(_Strict):
    """Blocking or non-blocking finding from analysis."""

    code: str
    message: str
    blocking: bool = False


class AnalyzeResult(_Strict):
    """Full preflight report returned by POST /api/template/analyze."""

    source_sha256: str
    paragraphs: list[ParagraphInfo]
    sections: list[SectionCandidate]
    suggested_profile: TemplateProfile | None = None
    field_candidates: list[FieldCandidate] = Field(default_factory=list)
    issues: list[Issue] = Field(default_factory=list)
    ready: bool = False

    @property
    def blockers(self) -> list[Issue]:
        """Issues that prevent installation."""
        return [i for i in self.issues if i.blocking]


@dataclass
class _Para:
    """Internal paragraph view used while scanning."""

    id: int
    paragraph: Paragraph
    text: str
    is_bullet: bool
    has_tab: bool
    has_hyperlink: bool
    runs: list = field(default_factory=list)
    #: `None` for a body-level paragraph; set when this paragraph lives inside a table
    #: cell. See `docx_text.iter_document_paragraphs` — this is the additive sidecar
    #: that lets table-layout detection reason about "does this row have a second
    #: populated cell" without changing what a paragraph id *means*.
    location: docx_text.ParaLocation | None = None


def sha256_bytes(raw: bytes) -> str:
    """Return the hex SHA-256 of `raw`."""
    return hashlib.sha256(raw).hexdigest()


def is_bullet(paragraph: Paragraph) -> bool:
    """True when the paragraph carries Word list numbering properties."""
    return paragraph._p.find(f".//{qn('w:numPr')}") is not None


def has_tab(paragraph: Paragraph) -> bool:
    """True when any run contains a tab element or a literal tab character."""
    for run in paragraph.runs:
        if "\t" in (run.text or ""):
            return True
        if run._r.find(qn("w:tab")) is not None:
            return True
    return False


def has_hyperlink(paragraph: Paragraph) -> bool:
    """True when the paragraph contains a w:hyperlink wrapper."""
    return paragraph._p.find(qn("w:hyperlink")) is not None


def _has_tab_like(p: _Para) -> bool:
    """`p` is laid out as a two-part line: a real tab stop, or — in a table layout — a
    paragraph whose row carries a second populated cell.

    A section heading is never either; an entry header almost always is. This is the
    exact structural role a literal tab plays in a paragraph-layout resume
    ("Company | Location\\tDates"), generalised to a table layout where the same split
    is expressed as two cells instead of text before/after a tab character.
    """
    if p.has_tab:
        return True
    loc = p.location
    return loc is not None and loc.row_content_cells >= 2


def _document_has_tables(doc) -> bool:
    """True when the body contains at least one table."""
    return bool(doc.tables)


def _document_has_textboxes(doc) -> bool:
    """True when any paragraph hosts a drawing/textbox (common multi-column cue).

    Checks paragraphs inside table cells too — a textbox parked in a cell is just as
    strong a multi-column cue as one at body level.
    """
    for paragraph, _location in docx_text.iter_document_paragraphs(doc):
        p = paragraph._p
        if p.find(f".//{qn('w:txbxContent')}") is not None:
            return True
        if p.find(f".//{qn('w:drawing')}") is not None:
            return True
    return False


@dataclass(frozen=True)
class TableShape:
    """The one top-level table a table-layout resume lives in."""

    index: int
    rows: int
    grid_cols: int


def classify_table_layout(
    doc, paras: list[_Para], heading_fp_classes: frozenset[_Fingerprint] = frozenset()
) -> tuple[TableShape | None, list[Issue]]:
    """Decide whether this document's table forms ONE logical reading column (a
    `layout="table"` resume: content laid out top-to-bottom, the table used only to
    right-align dates/locations without tab stops) or a genuine multi-column/sidebar
    layout, which stays unsupported.

    Deliberately not width- or majority-based: a real single-column resume table puts
    a narrow label column beside a wide value column in one row (e.g. "Languages:" at
    1509 twips beside a 8859-twip value cell) and a wide main column beside a narrow
    date column in another, so no monotone width rule holds, and a linear resume is
    routinely *majority* two-cell rows (every entry header) against a minority of
    full-width rows (headings, bullets) — a full-width-fraction rule rejects exactly
    the documents it should accept. What actually holds for a linear table: it never
    runs two content streams in parallel. Bullets and section headings always live in
    the reading column (cell 0); nothing is vertically merged; no row holds three or
    more independently-populated cells.

    Returns `(None, issues)` with a single blocking issue on the first structural
    failure (checked in a fixed order so the message is deterministic), or
    `(TableShape, [])` for a linear table with no issues of its own — the caller still
    runs every other detector afterward, since "this table is linear" says nothing yet
    about whether a usable experience/contact mapping was actually found in it.
    """
    if docx_text.has_nested_tables(doc):
        return None, [
            Issue(
                code="nested_tables",
                message=(
                    "A table inside this document contains another table nested in one "
                    "of its cells. Only a single flat layout table is supported."
                ),
                blocking=True,
            )
        ]

    tables = doc.tables
    if not tables:
        return None, []
    if len(tables) > 1:
        return None, [
            Issue(
                code="multiple_tables",
                message=(
                    f"Document contains {len(tables)} separate tables. Only a single "
                    "table used as an invisible layout grid is supported — merge them "
                    "into one table in Word/Google Docs, or remove the extra one."
                ),
                blocking=True,
            )
        ]

    table = tables[0]
    if table._tbl.find(f".//{qn('w:vMerge')}") is not None:
        return None, [
            Issue(
                code="table_vertical_merge",
                message=(
                    "The document's table vertically merges cells, which is how a "
                    "sidebar layout spans a column across several rows. Only a table "
                    "used purely as an invisible single-column layout grid — no "
                    "vertical merges — is supported."
                ),
                blocking=True,
            )
        ]

    table_paras = [p for p in paras if p.location is not None and p.location.table == 0]
    if not table_paras:
        return None, []

    rows_seen = sorted({p.location.row for p in table_paras})  # type: ignore[union-attr]
    grid_cols = len(table.columns)

    for row in rows_seen:
        row_paras = [p for p in table_paras if p.location.row == row]  # type: ignore[union-attr]
        content_cells = row_paras[0].location.row_content_cells  # type: ignore[union-attr]
        if content_cells >= 3:
            return None, [
                Issue(
                    code="table_parallel_columns",
                    message=(
                        f"Row {row} of the table has {content_cells} independently "
                        "populated cells, so this document reads as more than two "
                        "parallel columns. Only a table used as an invisible "
                        "single-column layout grid is supported."
                    ),
                    blocking=True,
                )
            ]

    for p in table_paras:
        if p.is_bullet and p.location.cell != 0:  # type: ignore[union-attr]
            return None, [
                Issue(
                    code="table_sidebar_bullets",
                    message=(
                        f"Row {p.location.row}'s second cell contains bulleted text, "  # type: ignore[union-attr]
                        "so this document reads as two parallel columns. Only tables "
                        "used as an invisible single-column layout grid are "
                        "supported — a sidebar or two-column resume cannot be mapped."
                    ),
                    blocking=True,
                )
            ]

    found_heading = False
    for p in table_paras:
        stripped = p.text.strip()
        if not stripped or p.is_bullet or _is_chrome(p.text):
            continue
        key, conf, _alias = _classify_heading(stripped)
        if key is None or conf < 1.0:
            # Anything short of an exact alias match (a bare heuristic keyword hit —
            # "Technologies" in a company name matching the "skills" heuristic — or an
            # unaliased structural guess like "LEADERSHIP") needs the same structural
            # corroboration `_analyze_document`'s own loop requires for a sub-1.0 match:
            # not shaped like an entry header or the line right under one
            # (`_has_tab_like`, which — row-based — happens to catch both at once,
            # since a title paragraph shares its entry header's two-populated-cell
            # row), not earlier than the name/contact block, and — when the document
            # has a detectable heading-formatting class — a matching fingerprint.
            # Without this, an entry header ("Kidod Science & Technologies") or a
            # short all-caps field that merely looks heading-shaped by text alone (a
            # state abbreviation like "CA" in a location cell) reads as a sidebar
            # heading purely because it shares its row with the entry it belongs to.
            if _has_tab_like(p):
                continue
            if key is None and (p.id < 2 or not _looks_like_heading(stripped)):
                continue
            if heading_fp_classes and _fingerprint(p) not in heading_fp_classes:
                continue
        found_heading = True
        loc = p.location
        assert loc is not None
        if loc.cell != 0 or loc.row_content_cells != 1:
            return None, [
                Issue(
                    code="table_sidebar_headings",
                    message=(
                        f"{stripped!r} (row {loc.row}) looks like a section heading "
                        "but shares its row with other content, which is how a "
                        "sidebar heading is laid out. Only tables used as an "
                        "invisible single-column layout grid are supported."
                    ),
                    blocking=True,
                )
            ]

    if not found_heading:
        return None, [
            Issue(
                code="table_no_headings",
                message=(
                    "This document's table contains no recognizable section heading, "
                    "so it reads as a data table rather than a resume layout grid."
                ),
                blocking=True,
            )
        ]

    return TableShape(index=0, rows=len(rows_seen), grid_cols=grid_cols), []


def _classify_heading(text: str) -> tuple[str | None, float, str]:
    """Map heading text to a canonical section key with a confidence score."""
    stripped = text.strip()
    upper = stripped.upper()
    if upper in _HEADING_ALIASES:
        return _HEADING_ALIASES[upper], 1.0, upper

    # Short all-caps lines that contain a known keyword.
    if len(stripped) <= 40 and stripped == stripped.upper() and any(c.isalpha() for c in stripped):
        for alias, key in _HEADING_ALIASES.items():
            if alias in upper or any(tok in upper for tok in alias.split() if len(tok) > 3):
                # Prefer stronger tokens.
                if key == "experience" and "EXPERIENCE" in upper:
                    return key, 0.85, alias
                if key == "education" and "EDUCATION" in upper:
                    return key, 0.85, alias
                if key == "projects" and "PROJECT" in upper:
                    return key, 0.85, alias
                if key == "skills" and ("SKILL" in upper or "TECHNOLOG" in upper):
                    return key, 0.85, alias

    lower = stripped.lower()
    heuristics = (
        ("experience", ("experience", "employment", "work history"), 0.6),
        ("education", ("education", "academic"), 0.6),
        ("projects", ("project",), 0.55),
        ("skills", ("skills", "technologies", "tech stack"), 0.55),
    )
    for key, needles, conf in heuristics:
        if any(n in lower for n in needles) and len(stripped) <= 48:
            return key, conf, needles[0]
    return None, 0.0, ""


def _is_chrome(text: str) -> bool:
    """Blank, or a decorative rule/underscore line — never an entry header or content.

    Without this, a horizontal-rule paragraph (`"______________"`) reads as an entry
    header to `_split_entries` (non-blank, non-bullet), and whatever field gets mapped to
    "the first header in the section" lands on the rule instead of the real entry below
    it — a live bug in an installed profile, traced back to exactly this.

    Delegates to `docx_text.is_chrome_text` so `template_build` (which clones exactly the
    chrome paragraphs recorded here as spacer donors) cannot drift from this judgement."""
    return docx_text.is_chrome_text(text)


def _looks_like_heading(text: str) -> bool:
    """Structural heading candidate: short, no tab, no colon, no date, mostly uppercase.

    Catches a section title `_classify_heading`'s alias table has no entry for (e.g.
    "OTHER ACTIVITIES", "CERTIFICATIONS") — a user can name a section anything, so no
    fixed alias list can be complete. Deliberately permissive; the caller still requires
    the paragraph to be followed by entry- or bullet-shaped content before treating it as
    a real heading, so an ALL-CAPS one-line achievement is not mistaken for one.
    """
    stripped = text.strip()
    if not stripped or len(stripped) > 48 or "\t" in text or ":" in stripped:
        return False
    if _DATE_RE.search(stripped):
        return False
    letters = [c for c in stripped if c.isalpha()]
    if not letters:
        return False
    return sum(1 for c in letters if c.isupper()) / len(letters) >= 0.8


@dataclass(frozen=True)
class _Fingerprint:
    """Formatting signature of one paragraph, used to cluster paragraphs that were very
    likely authored with the same intent (e.g. "this is a section heading", "this is an
    entry header") independent of what their text says. Read from python-docx's typed
    paragraph/run formatting API, not raw OOXML — every attribute here already has a
    clean typed accessor, and reading it that way is what keeps this resilient to which
    specific XML shape an export happens to use for "bold" or "12pt"."""

    style: str | None
    bold: bool | None
    size: int | None  # EMU, from the first run's explicit font size
    alignment: int | None  # WD_ALIGN_PARAGRAPH value
    left_indent: int | None  # EMU
    space_before: int | None  # EMU
    space_after: int | None  # EMU
    has_tab: bool
    caps_bucket: str


def _caps_bucket(text: str) -> str:
    """Coarse case-shape bucket: "all_caps" (>=80% of letters upper), "title_case"
    (starts upper, not all), or "mixed" (starts lower, or no letters at all)."""
    stripped = text.strip()
    letters = [c for c in stripped if c.isalpha()]
    if not letters:
        return "mixed"
    if sum(1 for c in letters if c.isupper()) / len(letters) >= 0.8:
        return "all_caps"
    if stripped[0].isupper():
        return "title_case"
    return "mixed"


def _fingerprint(p: _Para) -> _Fingerprint:
    """Compute `p`'s formatting fingerprint."""
    paragraph = p.paragraph
    style = paragraph.style.name if paragraph.style is not None else None
    pf = paragraph.paragraph_format
    first_run = p.runs[0] if p.runs else None
    first_size = first_run.font.size if first_run is not None else None
    return _Fingerprint(
        style=style,
        bold=first_run.bold if first_run is not None else None,
        size=first_size.emu if first_size is not None else None,
        alignment=pf.alignment.value if pf.alignment is not None else None,
        left_indent=pf.left_indent.emu if pf.left_indent is not None else None,
        space_before=pf.space_before.emu if pf.space_before is not None else None,
        space_after=pf.space_after.emu if pf.space_after is not None else None,
        has_tab=_has_tab_like(p),
        caps_bucket=_caps_bucket(p.text),
    )


def _introduces_content(
    p: _Para, paras: list[_Para], heading_fp_classes: frozenset[_Fingerprint] = frozenset()
) -> bool:
    """True when `p`'s own body — everything after it, up to the next *recognizable*
    heading, or the end of the document — contains at least one Word list bullet or one
    tab-aligned line (the shape of an entry header: "Company | Location\\tDates").

    This is the guard `_looks_like_heading`'s docstring has long claimed the caller
    enforces; nothing actually implemented it until now, which is how an unrecognized
    heading like "PROFESSIONAL SUMMARY" — followed only by a paragraph of prose, never
    a bullet or a tab-aligned entry header — got defaulted to `kind="experience"` by
    the structural fallback below, dragging a real job's header/title prototype onto
    the summary sentence and silently dropping its dates.

    A bulletless, tabless single-line education entry (a real, supported shape — see
    `education_bullets_not_list`) is unaffected: its header line has a tab even when
    its degree line does not, so the section still introduces content.

    "Recognizable" deliberately excludes a merely short/plain/no-tab line: an entry
    header without its own trailing date on the same line — a company name on its own
    line, title on the next — is *also* short and plain, and stopping the scan there
    made a real section heading with a two-line first entry (e.g. "WORK EXPERIENCE"
    followed by "AMAZON WEB SERVICES" then a tab+date title line) register as
    introducing nothing, since the scan gave up one line too early. Only two things
    count as an unambiguous next heading: an alias/keyword text match
    (`_classify_heading`), or — once the caller has a corroborating fingerprint class in
    hand (see `_heading_classes`) — membership in that class. Both are exactly the
    signals the caller itself already trusts to call something a heading, so this
    cannot recognize a "heading" the rest of the module would not.
    """
    for q in paras:
        if q.id <= p.id or _is_chrome(q.text):
            continue
        stripped = q.text.strip()
        if not stripped:
            continue
        if q.is_bullet or _has_tab_like(q):
            return True
        if _classify_heading(stripped)[0] is not None:
            return False  # reached an unambiguous next heading with nothing found
        if heading_fp_classes and _fingerprint(q) in heading_fp_classes:
            return False
    return False


def _heading_classes(paras: list[_Para]) -> set[_Fingerprint]:
    """Formatting signatures that recur across the document among short, non-bulleted,
    non-chrome, content-introducing paragraphs — the structural signal that
    corroborates (or fails to corroborate) a text-based heading match.

    A resume's own section headings are typically styled identically to each other and
    to nothing else in the body, so real headings cluster into one shared class of
    several members; a single stray ALL-CAPS line (a product name, an emphatic opener)
    does not repeat and never forms a class of its own.

    Returns the empty set when nothing qualifies (too few candidates, or no fingerprint
    shared by more than one) — callers must degrade to today's text-only behavior in
    that case, not treat an empty result as "nothing is a heading."
    """
    candidates: dict[_Fingerprint, list[_Para]] = {}
    for p in paras:
        if p.is_bullet or _is_chrome(p.text):
            continue
        stripped = p.text.strip()
        if not stripped or len(stripped) > 48 or _has_tab_like(p) or ":" in stripped:
            continue
        if _DATE_RE.search(stripped):
            continue
        if not _introduces_content(p, paras):
            continue
        candidates.setdefault(_fingerprint(p), []).append(p)

    return {fp for fp, members in candidates.items() if len(members) >= 2}


def _immediately_follows_entry_header(p: _Para, paras: list[_Para]) -> bool:
    """True when the nearest preceding non-chrome paragraph has a tab and is not
    itself a bullet — the shape of an entry header ("Company | Location\\tDates").

    A paragraph sitting right after one — a job title line, e.g. "Experience
    Designer" immediately under "Acme Corp | Remote\\tJan 2023 - Present" — is almost
    certainly that entry's title, never a new section heading, regardless of what its
    own text happens to say. This is the case fingerprint corroboration alone cannot
    catch: a title line's formatting can coincidentally land in the same class as the
    document's real headings (both are short, unbulleted, content-introducing text),
    exactly as a real heading styled slightly differently ("SKILLS & Interests" mixing
    case against all-caps siblings) also fails corroboration — confidence and
    formatting alone can't tell the two apart, but position relative to an entry
    header can.
    """
    for q in reversed(paras[: p.id]):
        if _is_chrome(q.text):
            continue
        return _has_tab_like(q) and not q.is_bullet
    return False


def _bootstrap_split_entries(paras: list[_Para]) -> list[list[_Para]]:
    """Group a section body into entries by the bullet-anchored rule alone: a new entry
    starts at any non-bullet paragraph with text, as long as the previous entry (if
    any) has already seen a bullet. Kept as its own function because `_split_entries`
    uses it as a bootstrap pass, then tries to structurally corroborate it — see there.
    """
    entries: list[list[_Para]] = []
    for p in paras:
        if _is_chrome(p.text):
            continue
        starts = (not p.is_bullet) and bool(p.text.strip())
        if starts and (not entries or any(x.is_bullet for x in entries[-1])):
            entries.append([p])
        elif entries:
            entries[-1].append(p)
    return entries


def _split_entries(paras: list[_Para]) -> list[list[_Para]]:
    """Group a section body into entries (an entry-header paragraph, its title/bullets,
    until the next one).

    Two passes. The first (`_bootstrap_split_entries`) is today's contract. The second
    clusters each bootstrap entry's own header paragraph by formatting fingerprint, and
    when a majority share one, re-splits using the *union* of the bootstrap rule
    ("previous entry already saw a bullet") and that class's membership — never the
    fingerprint rule alone. This is what fixes a bulletless pseudo-entry (a stray
    heading-shaped line the classifier could not name, or a genuinely bulletless
    single-line entry) swallowing the real entry after it: the bullet-anchored rule
    alone cannot tell "this entry just doesn't have a bullet yet" from "this line was
    never a real entry start", but a formatting match against the *other* entries' own
    headers can, and critically can find a boundary the bootstrap pass never proposed in
    the first place since it scans every paragraph in the section, not just the ones
    bootstrap already treated as entry starts.

    Union, not replacement: an entry whose header simply lacks the dominant format (an
    otherwise-ordinary entry that happens to have no tab-aligned date, e.g. an
    unfinished/current role) is *already* a correct bootstrap boundary — a
    fingerprint-only re-split would incorrectly re-merge it into its predecessor for no
    reason but the format mismatch, trading one entry-splitting bug for another.
    """
    bootstrap = _bootstrap_split_entries(paras)
    if len(bootstrap) < 2:
        return bootstrap

    header_fps = [_fingerprint(entry[0]) for entry in bootstrap]
    dominant_fp, dominant_count = Counter(header_fps).most_common(1)[0]
    if dominant_count <= len(header_fps) / 2:
        return bootstrap  # no majority formatting agreement; keep the bootstrap split

    entries: list[list[_Para]] = []
    for p in paras:
        if _is_chrome(p.text):
            continue
        plain = (not p.is_bullet) and bool(p.text.strip())
        bootstrap_starts = plain and (not entries or any(x.is_bullet for x in entries[-1]))
        fingerprint_starts = plain and _fingerprint(p) == dominant_fp
        if bootstrap_starts or fingerprint_starts:
            entries.append([p])
        elif entries:
            entries[-1].append(p)
        # else: paragraph before the first matching header is dropped, matching
        # `_bootstrap_split_entries`'s own behavior of ignoring anything before the
        # first detected entry start.
    return entries


def _representative_run(runs: list[list[int]], applicable: int) -> list[int]:
    """Pick the run to reproduce from every run observed at one slot.

    Two filters, in order. A slot must have been observed in a **majority** of the places
    it could apply, so one stray blank does not add a spacer everywhere the mapping is
    later applied. Among the surviving runs the **modal length** wins, so a section that
    happens to carry three blank paragraphs where every other section carries one does
    not inflate the gap document-wide; ties go to the shorter run, since over-spacing
    costs page height and this only ever needs to look consistent.
    """
    if not runs or not applicable or len(runs) <= applicable / 2:
        return []
    lengths = [len(r) for r in runs]
    best = min(set(lengths), key=lambda n: (-lengths.count(n), n))
    return next(r for r in runs if len(r) == best)


def _chrome_runs(body: list[_Para]) -> list[tuple[_Para | None, list[_Para], _Para | None]]:
    """Split `body` into maximal runs of chrome paragraphs, each with the non-chrome
    paragraph immediately before and after it (`None` at a boundary of the list)."""
    runs: list[tuple[_Para | None, list[_Para], _Para | None]] = []
    i = 0
    while i < len(body):
        if not _is_chrome(body[i].text):
            i += 1
            continue
        start = i
        while i < len(body) and _is_chrome(body[i].text):
            i += 1
        before = body[start - 1] if start > 0 else None
        after = body[i] if i < len(body) else None
        runs.append((before, body[start:i], after))
    return runs


def _detect_spacing(
    paras: list[_Para], section_candidates: list[SectionCandidate]
) -> SpacingProfile:
    """Detect the chrome-paragraph runs a document reproduces consistently around
    sections: before each heading, right after each heading, and between entries.

    Runs, not single paragraphs, because real exports put several chrome paragraphs at
    one slot — a horizontal rule plus a blank under every heading is the common case, and
    two consecutive blanks between entries is not rare. Detecting only a single blank
    both dropped the rule (it is not blank, so it matched nothing and was deleted with
    the rest of the body) and missed the two-blank gap entirely (each blank's neighbour
    was the other blank rather than the bullet/header pair the boundary test looked for).
    """
    if not section_candidates:
        return SpacingProfile()

    before_runs: list[list[int]] = []
    after_runs: list[list[int]] = []
    between_runs: list[list[int]] = []
    between_applicable = 0

    for sec in section_candidates:
        # Walk back from the heading while the paragraph above is chrome.
        run: list[int] = []
        pid = sec.heading_paragraph_id - 1
        while pid >= 0 and _is_chrome(paras[pid].text):
            run.insert(0, pid)
            pid -= 1
        if run:
            before_runs.append(run)

        body = [p for p in paras if sec.body_start <= p.id < sec.body_end]
        runs = _chrome_runs(body)

        # The after-heading run is the one that opens the body (nothing non-chrome
        # before it inside this section).
        for before, chrome, _after in runs:
            if before is None:
                after_runs.append([p.id for p in chrome])
            break

        if sec.key in ("experience", "projects", "education") and sec.entry_count >= 2:
            between_applicable += 1
            for before, chrome, after in runs:
                # An entry boundary: the run closes one entry's bullets and opens the
                # next entry's header. A run with no `after` is the trailing gap before
                # the *next section's* heading (counted as `before_heading` instead), and
                # a run whose predecessor is not a bullet sits inside an entry (between a
                # company line and its title line, say), not between two entries.
                if (
                    before is not None
                    and before.is_bullet
                    and after is not None
                    and not after.is_bullet
                    and after.text.strip()
                ):
                    between_runs.append([p.id for p in chrome])
                    break

    n = len(section_candidates)
    return SpacingProfile(
        before_heading=_representative_run(before_runs, n),
        after_heading=_representative_run(after_runs, n),
        between_entries=_representative_run(between_runs, between_applicable),
    )


def _contact_separator(text: str) -> str:
    """Pick the most likely contact-line separator from the prototype text."""
    for sep in _SEP_CANDIDATES:
        if sep in text:
            return sep
    # Fallback: single bullet character with spaces.
    if "\u2022" in text:
        return " \u2022 "
    if "|" in text:
        return " | "
    return " \u2022 "


def _contact_field_order(text: str) -> list[ContactField]:
    """Infer which contact fields appear and in what order from prototype text."""
    order: list[ContactField] = []
    lower = text.lower()
    # Scan left-to-right by finding earliest match positions.
    finds: list[tuple[int, ContactField]] = []
    if _EMAIL_RE.search(text):
        finds.append((_EMAIL_RE.search(text).start(), "email"))  # type: ignore[union-attr]
    phone = _PHONE_RE.search(text)
    # A bare year range ("2021 - 2025") is digits-space-punctuation-digits just like a
    # phone number, and `_PHONE_RE` alone cannot tell them apart — exclude anything
    # that also looks like a date before trusting it as a phone.
    if phone and not _DATE_RE.search(phone.group(0)):
        finds.append((phone.start(), "phone"))
    for label, field in (("linkedin", "linkedin"), ("github", "github")):
        idx = lower.find(label)
        if idx >= 0:
            finds.append((idx, field))  # type: ignore[arg-type]
    # Location: leftover leading chunk before first known token, if any.
    if finds:
        first = min(p for p, _ in finds)
        head = text[:first].strip(" \u2022|·/–—\t")
        if head and "@" not in head and not _PHONE_RE.fullmatch(head.replace(" ", "")):
            finds.append((0, "location"))
    elif text.strip():
        finds.append((0, "location"))

    for _, field in sorted(finds, key=lambda t: t[0]):
        if field not in order:
            order.append(field)
    # Always allow the full set at render time; order is preference for present fields.
    for extra in ("location", "email", "phone", "linkedin", "github"):
        if extra not in order:
            order.append(extra)  # type: ignore[arg-type]
    return order


#: A location paragraph in a table's contact block reads as "City, ST" or "City, ST
#: ZIP" — a comma followed by a two-letter state code. Deliberately narrow: unlike
#: `_contact_field_order`'s single-line fallback (which claims *any* leftover text as
#: `location` because there's nowhere else for it to go), a street-address paragraph
#: ("23320 Arroyo Dr.") must NOT match this, so it can be reported as unmapped rather
#: than guessed at — see `_detect_name_and_contact`.
_LOCATION_LIKE_RE = re.compile(r",\s*[A-Z]{2}\b")


def _contact_fields_present(text: str) -> list[ContactField]:
    """Which contact fields one paragraph's own text plausibly names, in the order
    they appear. Unlike `_contact_field_order`, this never pads out to the full
    five-field set — an empty return means "this paragraph isn't a contact field at
    all" (a street address, say), which is exactly what `_detect_name_and_contact`
    needs to tell a mappable slot from one that must stay a literal.
    """
    finds: list[tuple[int, ContactField]] = []
    email_m = _EMAIL_RE.search(text)
    if email_m:
        finds.append((email_m.start(), "email"))
    phone_m = _PHONE_RE.search(text)
    if phone_m and not _DATE_RE.search(phone_m.group(0)):
        finds.append((phone_m.start(), "phone"))
    lower = text.lower()
    for label, field in (("linkedin", "linkedin"), ("github", "github")):
        idx = lower.find(label)
        if idx >= 0:
            finds.append((idx, field))  # type: ignore[arg-type]
    if not finds and _LOCATION_LIKE_RE.search(text):
        finds.append((0, "location"))
    return [f for _, f in sorted(finds, key=lambda t: t[0])]


def _detect_name_and_contact(
    paras: list[_Para], first_heading_id: int | None
) -> tuple[int, _Para | None, list[ContactSlot], list[_Para]]:
    """Name plus the paragraph(s) holding contact info, classified by regex.

    "First two non-empty non-heading paragraphs" (the rule this replaces) is wrong the
    moment a table layout spreads name/address/email/phone across four separate
    paragraphs in four cells: the second paragraph is the street-address line, so a
    joined contact line would be built from an address, and the email/phone would be
    baked into the template as unchanging literals.

    Scans every non-bullet, non-blank paragraph above the first detected section
    heading (or, when none was found at all, falls back to today's exact
    "non-heading-classified content paragraph" scan, so an unrecognizable document
    still fails the same way it always has). Only attempts per-paragraph slot
    classification when the block actually spans a table (at least one candidate
    paragraph carries a `location`) — a paragraph-layout resume with a stray extra
    line before its first heading keeps today's exact behaviour: the first remaining
    paragraph is the whole contact line, verbatim.

    Returns `(name_id, contact_paragraph, slots, unmapped)`. At most one of
    `contact_paragraph` / `slots` is set — `slots` only when the contact block is
    table-shaped *and* splits into two or more recognizable fields; `unmapped` lists
    any table-cell paragraph in the block that named no contact field at all (an
    address line, typically), reported so it's never silently dropped.
    """
    if first_heading_id is not None:
        scope = [p for p in paras if p.id < first_heading_id and p.text.strip() and not p.is_bullet]
    else:
        scope = [
            p
            for p in paras
            if p.text.strip() and not p.is_bullet and _classify_heading(p.text)[0] is None
        ]

    if not scope:
        return 0, None, [], []

    name = scope[0]
    rest = scope[1:]
    if not rest:
        return name.id, None, [], []

    is_table_shaped = any(p.location is not None for p in rest)
    if len(rest) <= 1 or not is_table_shaped:
        return name.id, rest[0], [], []

    slots: list[ContactSlot] = []
    unmapped: list[_Para] = []
    for p in rest:
        fields = _contact_fields_present(p.text)
        if fields:
            slots.append(ContactSlot(paragraph_id=p.id, fields=fields))
        else:
            unmapped.append(p)

    if not slots:
        return name.id, rest[0], [], []

    return name.id, None, slots, unmapped


def _span(paragraph_id: int, start: int, end: int) -> CharSpan:
    """Build a CharSpan."""
    return CharSpan(paragraph_id=paragraph_id, start=start, end=end)


def _header_fields_from_text(
    para: _Para,
    *,
    primary: str,
    secondary: str | None,
    date_field: str,
    exclude_after: int | None = None,
) -> tuple[HeaderFieldMapping, list[FieldCandidate]]:
    """Heuristic split of a header line into primary | secondary \\t dates.

    `exclude_after` clips the primary/secondary search region to `text[:exclude_after]`,
    used to keep a trailing hyperlink label (mapped separately as `link`) out of the
    `secondary`/`tech` span — without it, a header with a link but no tech
    (``"Name | Github\\tdate"``) reads the link's own label as `secondary`, giving `tech`
    and `link` the same span and making the build reject them as overlapping.
    """
    text = para.text
    candidates: list[FieldCandidate] = []
    fields: dict[str, OptionalSpan] = {}

    tab_idx = text.find("\t")
    before = text if tab_idx < 0 else text[:tab_idx]
    after = "" if tab_idx < 0 else text[tab_idx + 1 :]

    date_alignment = "tab" if tab_idx >= 0 else "inline"
    if after.strip() or (tab_idx < 0 and _DATE_RE.search(text)):
        if after.strip():
            start = tab_idx + 1
            # Skip leading whitespace after tab.
            while start < len(text) and text[start].isspace():
                start += 1
            end = len(text.rstrip())
            fields[date_field] = OptionalSpan(
                present=True, span=_span(para.id, start, end)
            )
            candidates.append(
                FieldCandidate(
                    field=date_field,
                    span=fields[date_field].span,  # type: ignore[arg-type]
                    confidence=0.9 if tab_idx >= 0 else 0.6,
                    preview=text[start:end],
                )
            )
        else:
            m = _DATE_RE.search(text)
            if m:
                fields[date_field] = OptionalSpan(
                    present=True, span=_span(para.id, m.start(), m.end())
                )
                date_alignment = "inline"
                candidates.append(
                    FieldCandidate(
                        field=date_field,
                        span=fields[date_field].span,  # type: ignore[arg-type]
                        confidence=0.7,
                        preview=m.group(0),
                    )
                )
                before = text[: m.start()].rstrip(" |·/–—")
            else:
                fields[date_field] = OptionalSpan(present=False)
    else:
        fields[date_field] = OptionalSpan(present=False)

    if exclude_after is not None and exclude_after < len(before):
        before = before[:exclude_after]

    # Split before-tab on a pipe-like separator. Only the first two segments become
    # primary/secondary — further segments (e.g. " | Github" on project headers) are
    # left for project-specific link detection so they do not inflate `tech` and then
    # overlap when the link span is merged in.
    sep = None
    for candidate in (" | ", "|", " — ", " – ", " - "):
        if candidate in before:
            sep = candidate
            break
    if sep and secondary:
        parts = [p.strip() for p in before.split(sep) if p.strip()]
        left_s = parts[0] if parts else ""
        right_s = parts[1] if len(parts) > 1 else ""
        if not left_s:
            fields[primary] = OptionalSpan(present=False)
            fields[secondary] = OptionalSpan(present=False)
        else:
            # Map back to offsets in `text` (before is a prefix of text).
            left_start = text.find(left_s)
            if left_start < 0:
                fields[primary] = OptionalSpan(present=False)
                fields[secondary] = OptionalSpan(present=False)
            else:
                fields[primary] = OptionalSpan(
                    present=True,
                    span=_span(para.id, left_start, left_start + len(left_s)),
                )
                candidates.append(
                    FieldCandidate(
                        field=primary,
                        span=fields[primary].span,  # type: ignore[arg-type]
                        confidence=0.85,
                        preview=left_s,
                    )
                )
                if right_s:
                    # Search after the primary so a repeated token cannot point left.
                    right_start = text.find(right_s, left_start + len(left_s))
                    if right_start < 0:
                        fields[secondary] = OptionalSpan(present=False)
                    else:
                        fields[secondary] = OptionalSpan(
                            present=True,
                            span=_span(
                                para.id, right_start, right_start + len(right_s)
                            ),
                        )
                        candidates.append(
                            FieldCandidate(
                                field=secondary,
                                span=fields[secondary].span,  # type: ignore[arg-type]
                                confidence=0.8,
                                preview=right_s,
                            )
                        )
                else:
                    fields[secondary] = OptionalSpan(present=False)
    else:
        primary_text = before.strip()
        if primary_text:
            start = text.find(primary_text)
            fields[primary] = OptionalSpan(
                present=True, span=_span(para.id, start, start + len(primary_text))
            )
            candidates.append(
                FieldCandidate(
                    field=primary,
                    span=fields[primary].span,  # type: ignore[arg-type]
                    confidence=0.7,
                    preview=primary_text,
                )
            )
        else:
            fields[primary] = OptionalSpan(present=False)
        if secondary:
            fields[secondary] = OptionalSpan(present=False)

    header = HeaderFieldMapping(
        header_paragraph_id=para.id,
        fields=fields,
        date_alignment=date_alignment,  # type: ignore[arg-type]
    )
    return header, candidates


_BARE_YEAR_RE = re.compile(r"\b(?:19|20)\d{2}\b")


def _entry_main_paragraphs(entry: list[_Para]) -> list[_Para]:
    """`entry` with its header's cross-cell siblings (location/dates in a table
    layout's other cell, sharing the header's row) removed — every other paragraph
    (title lines, bullets, later rows) is untouched.

    Without this, `titles`/`detail_paras`-style scans over `entry[1:]` pick up a table
    layout's location/date cell as if it were a job title or an education detail line,
    since depth-first flattening puts the whole row — both cells — before the next
    row's bullets.
    """
    head = entry[0]
    loc = head.location
    if loc is None:
        return entry
    return [
        p
        for p in entry
        if not (p.location is not None and p.location.row == loc.row and p.location.cell != loc.cell)
    ]


def _header_fields_across_cells(
    entry: list[_Para], *, primary: str, secondary: str | None, date_field: str
) -> tuple[HeaderFieldMapping, list[FieldCandidate]] | None:
    """Cross-cell counterpart to `_header_fields_from_text`, for a table layout where
    an entry header's company/school/name sits in one cell and its location/dates sit
    in the row's other cell (`Company | Location\\tDates` expressed as two cells
    instead of text either side of a tab).

    Returns `None` when the header paragraph isn't in a table row with a second
    populated cell — the caller falls back to the ordinary single-paragraph split.

    Location vs. dates in the side cell is resolved positionally, corroborated by
    `_DATE_RE`, never by `_DATE_RE` alone: it requires a month name or a year range, so
    a degree line's "Class of 2027" matches nothing, and a regex-first rule would wrongly
    leave it unmapped. With two side paragraphs, the second is dates (mirroring
    `Company | Location\\tDates`'s left-to-right order); with one, it's dates only when
    it looks date-shaped, else location — the two swap if `_DATE_RE` disagrees with
    that call for either paragraph.
    """
    head = entry[0]
    loc = head.location
    if loc is None or loc.row_content_cells < 2:
        return None

    main = [
        p for p in entry if p.location is not None and p.location.row == loc.row and p.location.cell == loc.cell
    ]
    side = [
        p
        for p in entry
        if p.location is not None
        and p.location.row == loc.row
        and p.location.cell != loc.cell
        and p.text.strip()
    ]
    if not main or not side:
        return None

    def _looks_like_date(text: str) -> bool:
        return bool(_DATE_RE.search(text)) or bool(_BARE_YEAR_RE.search(text))

    secondary_para: _Para | None
    date_para: _Para | None
    if len(side) >= 2:
        secondary_para, date_para = side[0], side[-1]
    else:
        p0 = side[0]
        if _looks_like_date(p0.text):
            secondary_para, date_para = None, p0
        else:
            secondary_para, date_para = p0, None

    if (
        secondary_para is not None
        and date_para is not None
        and _looks_like_date(secondary_para.text)
        and not _looks_like_date(date_para.text)
    ):
        secondary_para, date_para = date_para, secondary_para

    fields: dict[str, OptionalSpan] = {}
    candidates: list[FieldCandidate] = []

    def _field_span(name: str, para: _Para, confidence: float) -> None:
        text = para.text
        stripped = text.strip()
        start = text.find(stripped)
        span = _span(para.id, start, start + len(stripped))
        fields[name] = OptionalSpan(present=True, span=span)
        candidates.append(FieldCandidate(field=name, span=span, confidence=confidence, preview=stripped))

    _field_span(primary, main[0], 0.85)
    if secondary and secondary_para is not None:
        _field_span(secondary, secondary_para, 0.8)
    elif secondary:
        fields[secondary] = OptionalSpan(present=False)
    if date_para is not None:
        _field_span(date_field, date_para, 0.9)
    else:
        fields[date_field] = OptionalSpan(present=False)

    header = HeaderFieldMapping(
        header_paragraph_id=main[0].id,
        fields=fields,
        date_alignment="separate_paragraph",
        date_paragraph_id=date_para.id if date_para is not None else None,
    )
    return header, candidates


def _entry_header_fields(
    entry: list[_Para],
    *,
    primary: str,
    secondary: str | None,
    date_field: str,
    exclude_after: int | None = None,
) -> tuple[HeaderFieldMapping, list[FieldCandidate]]:
    """Header fields for one entry, cross-cell aware.

    Delegates to `_header_fields_from_text` (one paragraph, split on a tab or a
    pipe-like separator) unless the entry's header paragraph sits in a table row with a
    second populated cell, in which case `_header_fields_across_cells` reads
    location/dates out of that other cell instead. Keyed on "the row has a second
    physical cell", never on a fixed `gridSpan` — a document's WORK EXPERIENCE rows can
    be a 3+1 split while its LEADERSHIP rows are 2+2, and both are the same shape
    logically.
    """
    cross = _header_fields_across_cells(entry, primary=primary, secondary=secondary, date_field=date_field)
    if cross is not None:
        return cross
    return _header_fields_from_text(
        entry[0], primary=primary, secondary=secondary, date_field=date_field, exclude_after=exclude_after
    )


def _reconcile_header_fields(
    entries: list[list[_Para]],
    *,
    primary: str,
    secondary: str | None,
    date_field: str,
) -> tuple[list[list[_Para]], dict[str, bool], dict[str, float]]:
    """Run `_header_fields_from_text` on *every* entry's own header line, not just the
    one prototype the caller will eventually pick, and determine per-field whether a
    majority of entries actually carry it.

    A single-prototype's field detection is only as reliable as whichever entry
    happens to score highest for "looks like the best header" (`_exp_score` and
    friends) — if that one entry's dates genuinely are not present (a header formatted
    slightly differently than the rest, a role with no end date typed), the mapping
    would lose the field for *every* rendered entry, even though the rest of the
    section clearly has it.

    Returns `(candidate_entries, field_majority, field_confidence)`:
    - `candidate_entries`: the subset of `entries` whose own header's field-presence
      signature (which of primary/secondary/date_field it found) matches the modal
      signature across the section. The caller picks its prototype from *this* list
      (via `_exp_score` or similar), not from `entries` at large, so an outlier entry
      can never become the prototype merely by scoring highest on some unrelated
      dimension (more runs, a title line) while disagreeing with the section on which
      fields even exist.
    - `field_majority`: per field name, whether more than half of all entries carry it
      — what a caller uses to decide whether a field should be considered reliably
      present for the section as a whole, independent of what the eventually-chosen
      prototype's own header spans say.
    - `field_confidence`: per field name, the raw presence rate (0.0-1.0) across all
      entries — becomes `FieldCandidate.confidence`, which today is a fixed constant
      regardless of how many entries actually agree a field is there.

    Falls back to treating every entry as a candidate, with 100% confidence on
    whichever fields the single entry has, when there is only one entry — nothing to
    reconcile against.
    """
    if not entries:
        return [], {}, {}
    if len(entries) == 1:
        header, _ = _entry_header_fields(
            entries[0], primary=primary, secondary=secondary, date_field=date_field
        )
        names = [primary] + ([secondary] if secondary else []) + [date_field]
        presence = {
            n: header.fields.get(n, OptionalSpan(present=False)).present for n in names
        }
        return list(entries), presence, {n: 1.0 if presence[n] else 0.0 for n in names}

    per_entry: list[tuple[list[_Para], HeaderFieldMapping]] = []
    for entry in entries:
        header, _ = _entry_header_fields(
            entry, primary=primary, secondary=secondary, date_field=date_field
        )
        per_entry.append((entry, header))

    field_names = [primary] + ([secondary] if secondary else []) + [date_field]

    def _present(header: HeaderFieldMapping, name: str) -> bool:
        return header.fields.get(name, OptionalSpan(present=False)).present

    field_confidence = {
        name: sum(1 for _, h in per_entry if _present(h, name)) / len(per_entry)
        for name in field_names
    }
    field_majority = {name: field_confidence[name] > 0.5 for name in field_names}

    def _signature(header: HeaderFieldMapping) -> tuple[bool, ...]:
        return tuple(_present(header, n) for n in field_names)

    modal_signature = Counter(_signature(h) for _, h in per_entry).most_common(1)[0][0]
    candidate_entries = [entry for entry, h in per_entry if _signature(h) == modal_signature]

    return candidate_entries, field_majority, field_confidence


def _prototype_consistency_issue(
    proto: list[_Para], para_ids: dict[str, int | None], label: str
) -> Issue | None:
    """Blocking issue when any of `para_ids` (e.g. header/title/bullet paragraph ids)
    falls outside `proto`'s own paragraph span.

    A prototype whose pieces come from different entries — most often via a "fall back
    to any bullet in the section" escape hatch, reached when the chosen entry itself
    has none — produces a template whose built loop body mixes fragments from
    unrelated entries: a company from one job, a bullet from another. Never silently
    accepted, the same "never silently truncate" rule CLAUDE.md states for the fit
    loop, applied here to "never silently splice."
    """
    proto_ids = {p.id for p in proto}
    offenders = {
        name: pid for name, pid in para_ids.items() if pid is not None and pid not in proto_ids
    }
    if not offenders:
        return None
    detail = ", ".join(f"{name} is paragraph {pid}" for name, pid in offenders.items())
    return Issue(
        code="prototype_spans_multiple_entries",
        message=(
            f"{label} prototype fields come from different entries: {detail} "
            f"(prototype entry spans paragraphs {sorted(proto_ids)})."
        ),
        blocking=True,
    )


def _exp_score(entry: list[_Para]) -> tuple:
    """Prefer the experience entry with a tab-like header and a title + bullets — the
    same tie-break `_analyze_document`'s installed-prototype selection uses, hoisted to
    module level so `_section_field_candidates` can reuse it per section."""
    header = entry[0]
    bullets = [x for x in entry[1:] if x.is_bullet]
    titles = [x for x in entry[1:] if not x.is_bullet and x.text.strip()]
    return (
        1 if _has_tab_like(header) else 0,
        len(header.runs),
        1 if titles else 0,
        len(bullets),
    )


def _edu_score(entry: list[_Para]) -> tuple:
    """Prefer the education entry with a tab-like header and more runs — mirrors
    `_exp_score`'s role for education's prototype selection."""
    return (1 if _has_tab_like(entry[0]) else 0, len(entry[0].runs))


def _proj_score(entry: list[_Para]) -> int:
    """Fewer header runs first — a project header without a tech/date suffix ("Note
    Engine" alone) usually has fewer runs than one with a full "Name | Tech\\tDate"
    line, and the plainer header is the safer prototype to build a template from."""
    return len(entry[0].runs)


def _detect_project_link(
    proto: list[_Para],
) -> tuple[OptionalSpan, int | None, FieldCandidate | None]:
    """Hyperlink-based `link` field for a project entry's header line, detected before
    name/tech splitting so `exclude_after` can keep the link label out of `tech`.

    Returns `(link_span, exclude_after, candidate)` — `candidate` is `None` when the
    header paragraph carries no hyperlink before its first tab, in which case the other
    two values are the "absent" defaults.
    """
    header = proto[0]
    if not header.has_hyperlink:
        return OptionalSpan(present=False), None, None
    text = header.text
    tab = text.find("\t")
    limit = len(text) if tab < 0 else tab
    in_region = [
        (s, e) for s, e in docx_text.hyperlink_char_spans(header.paragraph) if e <= limit
    ]
    if not in_region:
        return OptionalSpan(present=False), None, None
    start, end = in_region[-1]
    span = _span(header.id, start, end)
    return (
        OptionalSpan(present=True, span=span),
        start,
        FieldCandidate(field="link", span=span, confidence=0.75, preview=text[start:end]),
    )


def _section_field_candidates(
    paras: list[_Para],
    sections: list[SectionCandidate],
    *,
    primary: str,
    secondary: str | None,
    date_field: str,
    pick: Callable[[list[list[_Para]]], list[_Para]],
    include_title: bool = False,
    include_link: bool = False,
) -> list[FieldCandidate]:
    """Header-field candidates for *every* detected section of one kind, not just the
    kind's single installed prototype.

    The profile installs one prototype per kind (`section_by_key`, pooled across every
    same-kind section via `combined_body`) — correct, and unchanged; that's what the
    caller's own `header`/`*_mapping` construction still uses. But the wizard's
    `AnalyzeReport` shows one field-detection row per *section*, so a kind-wide
    candidate set left every section but the pooled prototype's own looking like nothing
    was detected in it — a document with "WORK EXPERIENCE" and "LEADERSHIP" sections
    showed a false "company/dates — not detected" under LEADERSHIP even though both
    sections' entries have both fields. Confidence here is each section's own presence
    rate across its own entries (via `_reconcile_header_fields` scoped to that section),
    which is what the row claims to be reporting, rather than the kind-wide rate.
    """
    out: list[FieldCandidate] = []
    for sec in sections:
        entries = _split_entries(paras[sec.body_start : sec.body_end])
        if not entries:
            continue
        candidate_entries, _field_majority, field_confidence = _reconcile_header_fields(
            entries, primary=primary, secondary=secondary, date_field=date_field
        )
        proto = pick(candidate_entries)

        exclude_after: int | None = None
        if include_link:
            _link, exclude_after, link_cand = _detect_project_link(proto)
            if link_cand is not None:
                out.append(
                    link_cand.model_copy(
                        update={"section_heading_paragraph_id": sec.heading_paragraph_id}
                    )
                )

        _header, hcands = _entry_header_fields(
            proto,
            primary=primary,
            secondary=secondary,
            date_field=date_field,
            exclude_after=exclude_after,
        )
        for cand in hcands:
            out.append(
                cand.model_copy(
                    update={
                        "confidence": field_confidence.get(cand.field, cand.confidence),
                        "section_heading_paragraph_id": sec.heading_paragraph_id,
                    }
                )
            )

        if include_title:
            proto_main = _entry_main_paragraphs(proto)
            titles = [x for x in proto_main[1:] if not x.is_bullet and x.text.strip()]
            if titles:
                stripped = titles[0].text.strip()
                out.append(
                    FieldCandidate(
                        field="title",
                        span=_span(titles[0].id, 0, len(stripped)),
                        confidence=0.9,
                        preview=stripped[:80],
                        section_heading_paragraph_id=sec.heading_paragraph_id,
                    )
                )
    return out


def _skills_spans(para: _Para) -> tuple[CharSpan, CharSpan, str] | None:
    """Split a skills line into label + body on the first colon."""
    text = para.text
    if ":" not in text:
        return None
    idx = text.find(":")
    label = text[:idx].rstrip()
    # Include trailing spaces after colon in the separator; body starts at first non-space.
    sep_end = idx + 1
    while sep_end < len(text) and text[sep_end] == " ":
        sep_end += 1
    body = text[sep_end:]
    if not label or not body.strip():
        return None
    label_start = text.find(label)
    return (
        _span(para.id, label_start, label_start + len(label)),
        _span(para.id, sep_end, sep_end + len(body.rstrip())),
        text[idx:sep_end],
    )


def _skills_rows_across_cells(body: list[_Para]) -> list[tuple[_Para, _Para]] | None:
    """Every (label paragraph, value paragraph) pair in a table layout's skills grid,
    row by row — the shared row-pairing logic behind both `_skills_pair_across_cells`
    (one representative pair, for `SkillsMapping`) and `resume_import`'s own need for
    *every* pair (to actually import each group's items, not just learn the format).

    Depth-first flattening yields every label before every value within such a row
    ("Languages:, Skills:, Interests:, <fluent…>, <pivot tables…>, <piano…>"), which
    `_skills_spans` (one paragraph, split on a colon) cannot read at all — its
    "Languages:" has nothing after the colon in the SAME paragraph and returns `None`.

    Pairs cell *i* of the lower-indexed column with cell *i* of the higher-indexed one,
    within each row `body` touches, requiring every row to have exactly two populated
    cells with equal paragraph counts and non-empty label/value text on every pair —
    any row that doesn't fit that shape (this is not a label/value grid after all)
    makes the whole thing return `None` rather than a partial, unreliable pairing.
    """
    rows: dict[tuple[int, int], dict[int, list[_Para]]] = {}
    order: list[tuple[int, int]] = []
    for p in body:
        if p.location is None:
            return None
        key = (p.location.table, p.location.row)
        if key not in rows:
            rows[key] = {}
            order.append(key)
        rows[key].setdefault(p.location.cell, []).append(p)

    if not order:
        return None

    pairs: list[tuple[_Para, _Para]] = []
    for key in order:
        cells = rows[key]
        cell_ids = sorted(cells)
        if len(cell_ids) != 2:
            return None
        left, right = cells[cell_ids[0]], cells[cell_ids[1]]
        if not left or len(left) != len(right):
            return None
        for lp, rp in zip(left, right):
            if not lp.text.strip() or not rp.text.strip():
                return None
            pairs.append((lp, rp))
    return pairs


def _skills_pair_across_cells(body: list[_Para]) -> tuple[CharSpan, CharSpan, str] | None:
    """Single representative `(label_span, body_span, separator)` pair — same contract
    as `_skills_spans` — since that's all `SkillsMapping` needs: a single prototype
    pair whose formatting the build step clones once per `SkillGroup` at render time.
    See `_skills_rows_across_cells` for the row-pairing this is built on.
    """
    pairs = _skills_rows_across_cells(body)
    if not pairs:
        return None
    lp, rp = pairs[0]
    label = lp.text.strip()
    if label.endswith(":"):
        label = label[:-1].rstrip()
    value = rp.text.strip()
    if not label or not value:
        return None
    label_start = lp.text.find(label)
    body_start = rp.text.find(value)
    return (
        _span(lp.id, label_start, label_start + len(label)),
        _span(rp.id, body_start, body_start + len(value)),
        ": ",
    )
    return first


def _load_paras(doc) -> list[_Para]:
    """Flatten document body paragraphs — including any inside tables — into indexed
    `_Para` records, via `docx_text.iter_document_paragraphs`.

    THE id space: every `CharSpan.paragraph_id` and every bare `*_paragraph_id` field
    in `template_profile.py` is an index into this exact sequence.
    `template_build._para_by_id` must enumerate identically — see that function.
    """
    out: list[_Para] = []
    for i, (paragraph, location) in enumerate(docx_text.iter_document_paragraphs(doc)):
        text = paragraph.text or ""
        out.append(
            _Para(
                id=i,
                paragraph=paragraph,
                text=text,
                is_bullet=is_bullet(paragraph),
                has_tab=has_tab(paragraph),
                has_hyperlink=has_hyperlink(paragraph),
                runs=list(paragraph.runs),
                location=location,
            )
        )
    return out


def analyze_docx(
    path: str | bytes | None = None,
    *,
    raw: bytes | None = None,
    overrides: dict[int, str | None] | None = None,
) -> AnalyzeResult:
    """Analyze a DOCX from a filesystem path or in-memory bytes.

    Prefer `raw=` for uploads (hash is of the exact bytes). Path mode reads the file
    and hashes those bytes.

    `overrides` maps a paragraph id to a user-confirmed kind (`"experience"`,
    `"education"`, `"projects"`, `"skills"`, `"list"`) or `None` to say "this is not a
    section, regardless of what the heuristics think" — see `_analyze_document`.
    """
    import tempfile
    from pathlib import Path

    if raw is not None:
        digest = sha256_bytes(raw)
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(raw)
            tmp_path = Path(tmp.name)
        try:
            doc = docx.Document(str(tmp_path))
        finally:
            tmp_path.unlink(missing_ok=True)
    elif path is not None:
        path_obj = Path(path)
        digest = sha256_bytes(path_obj.read_bytes())
        doc = docx.Document(str(path_obj))
    else:
        raise ValueError("analyze_docx requires path or raw bytes")

    return _analyze_document(doc, digest, overrides=overrides)


def _analyze_document(
    doc, digest: str, *, overrides: dict[int, str | None] | None = None
) -> AnalyzeResult:
    """Run structural analysis on an open Document.

    `overrides` (paragraph id -> forced kind, or `None` for "not a section") lets the
    wizard's remap step correct a specific heading's classification without touching
    any other paragraph's — see the "user-confirmed" branch below. Every heuristic gate
    (fingerprint corroboration, has-tab exclusion, `_introduces_content`, …) is a
    signal for *guessing*; a user override is not a guess, so it bypasses all of them.
    """
    issues: list[Issue] = []
    paras = _load_paras(doc)
    overrides = overrides or {}

    # Computed early (usually this sits right before the heading-detection loop below)
    # specifically so `classify_table_layout` can use the same corroboration signal:
    # a short all-caps paragraph that merely looks heading-shaped (a state abbreviation
    # like "CA" in a location cell, say) must not be mistaken for a sidebar heading
    # just because nothing else disqualifies it — the fingerprint check is what tells
    # the two apart, since only real headings recur with matching formatting.
    heading_fp_classes = _heading_classes(paras)

    table_shape: TableShape | None = None
    if _document_has_tables(doc):
        table_shape, table_issues = classify_table_layout(doc, paras, heading_fp_classes)
        issues.extend(table_issues)
    if _document_has_textboxes(doc):
        issues.append(
            Issue(
                code="textboxes",
                message=(
                    "Document contains drawings or text boxes, which usually means a "
                    "multi-column or sidebar layout. Only single-column body text is supported."
                ),
                blocking=True,
            )
        )

    if len(paras) < 2:
        issues.append(
            Issue(
                code="too_short",
                message="Document needs at least a name line and a contact line.",
                blocking=True,
            )
        )

    paragraph_infos = [
        ParagraphInfo(
            id=p.id,
            text=p.text,
            is_bullet=p.is_bullet,
            is_heading_candidate=bool(_classify_heading(p.text)[0]),
            has_tab=p.has_tab,
            has_hyperlink=p.has_hyperlink,
            run_count=len(p.runs),
            preview=(p.text[:120] + ("…" if len(p.text) > 120 else "")),
        )
        for p in paras
    ]

    # Detect every section heading in the document — alias-matched or structurally
    # inferred — without deduplicating by kind. A resume may have any number of
    # experience-shaped sections (WORK EXPERIENCE, LEADERSHIP EXPERIENCE, OTHER
    # ACTIVITIES, …); each becomes its own candidate, and the code below pools same-kind
    # candidates together when picking that kind's prototype entry.
    # Structure-first corroboration (see `_heading_classes`'s docstring): a formatting
    # signature shared by several short, content-introducing paragraphs is what a
    # resume's own section headings typically look like. `heading_fp_classes` was
    # already computed above (`classify_table_layout` needs it too) — the alias table
    # below only *names* a heading's kind; this is what actually decides whether the
    # structural fallback gets to guess at all, and downgrades a merely-plausible text
    # match that nothing else in the document agrees with. Empty when nothing qualifies
    # (too few candidates, or no repeated formatting) — every use below degrades to
    # today's text-only behavior in that case, exactly as if this feature did not exist.

    raw_headings: list[SectionCandidate] = []
    for p in paras:
        if p.is_bullet or not p.text.strip() or _is_chrome(p.text):
            continue
        if p.id in overrides:
            forced = overrides[p.id]
            if forced is None:
                continue  # user confirmed: not a section, regardless of the heuristics
            raw_headings.append(
                SectionCandidate(
                    key=forced,
                    heading_paragraph_id=p.id,
                    heading_text=p.text.strip(),
                    body_start=p.id + 1,
                    body_end=len(paras),
                    confidence=1.0,
                    aliases_matched="user-confirmed",
                )
            )
            continue
        key, conf, alias = _classify_heading(p.text)
        # `p.id >= 2` keeps the structural fallback off the name/contact lines — both are
        # short, and a name in particular is very often all-caps or Title Case, which
        # would otherwise misclassify paragraph 0 as a heading. Paragraphs 0/1 are name
        # and contact everywhere else in this codebase (`build_name`/`build_contact`
        # legacy mode, `content_paras[0]`/`[1]` below); an alias match is unaffected by
        # this guard since a name or contact line never happens to equal a known alias.
        if key is None and p.id >= 2 and _looks_like_heading(p.text):
            # Structural fallback: no alias matched, but this looks like a heading. A
            # user can name a section anything, so no fixed alias list can be complete.
            #
            # Two hard gates, both required, neither previously enforced despite this
            # function's own docstring claiming the first one: the line must actually
            # introduce something (a bullet, or a tab-aligned entry header) before the
            # next heading-shaped line — a "PROFESSIONAL SUMMARY" followed only by a
            # sentence of prose introduces nothing and is not a section. And when the
            # document has a detectable heading-formatting class at all, an unaliased
            # candidate must belong to it — this is the "guess" path with zero other
            # evidence, so structural corroboration is required here, not just
            # preferred.
            if not _introduces_content(p, paras, frozenset(heading_fp_classes)):
                continue
            if heading_fp_classes and _fingerprint(p) not in heading_fp_classes:
                continue
            # Default to "experience" — the common case for an unnamed
            # achievements-with-employer section — unless it is immediately followed by
            # bullets with no entry header, which is a plain list-shaped section
            # (certifications, awards, …).
            next_content = next(
                (q for q in paras if q.id > p.id and not _is_chrome(q.text)), None
            )
            key = "list" if next_content is not None and next_content.is_bullet else "experience"
            conf, alias = 0.4, "structural"
        if key is None:
            continue
        uncorroborated = bool(heading_fp_classes) and _fingerprint(p) not in heading_fp_classes
        # The <=0.6-confidence tiers — a keyword appearing anywhere in the lowercased
        # text, with no case requirement at all — are weak enough that an ordinary
        # entry header matches them exactly like a real heading would: a job title
        # ("Experience Designer") right under a company/dates line, or a later entry's
        # own tab-aligned "Name\tDate" header (an "Advocate of Sexual Education in
        # School\t2022" activity inside "OTHER ACTIVITIES" matching the "education"
        # keyword). Fingerprint corroboration alone can't screen these out: an entry
        # header's formatting can land in the same class as the document's real
        # headings just as easily as a genuine heading styled slightly differently
        # ("SKILLS & Interests" mixing case) can fail to. Two structural, position-based
        # signals tell the two apart instead: a real section heading is never itself
        # tab-aligned with a trailing date (`_heading_classes`'s own candidate filter
        # already assumes this), and never sits immediately after another entry's own
        # header line. Exact alias matches (conf == 1.0, e.g. literally "EDUCATION")
        # are trusted on text alone regardless of position or formatting.
        if conf < 1.0 and (_has_tab_like(p) or _immediately_follows_entry_header(p, paras)):
            continue
        if conf < 1.0 and uncorroborated:
            conf = min(conf, 0.4)
            issues.append(
                Issue(
                    code="heading_formatting_mismatch",
                    message=(
                        f"{p.text.strip()!r} (paragraph {p.id}) looks like a {key} "
                        "heading by its text, but its formatting does not match the "
                        "document's other section headings. Confirm this is really a "
                        "section heading."
                    ),
                    blocking=False,
                )
            )
        raw_headings.append(
            SectionCandidate(
                key=key,
                heading_paragraph_id=p.id,
                heading_text=p.text.strip(),
                body_start=p.id + 1,
                body_end=len(paras),
                confidence=conf,
                aliases_matched=alias,
            )
        )
    raw_headings.sort(key=lambda s: s.heading_paragraph_id)

    # Resolve body_end from the next heading OF ANY KIND, so an embedded same-kind
    # sub-heading (e.g. "LEADERSHIP EXPERIENCE" after "WORK EXPERIENCE") is never
    # mis-read as an entry header by `_split_entries` — each heading's slice already
    # excludes every other heading paragraph by construction.
    section_candidates: list[SectionCandidate] = []
    by_kind: dict[str, list[SectionCandidate]] = {}
    for i, sec in enumerate(raw_headings):
        end = (
            raw_headings[i + 1].heading_paragraph_id
            if i + 1 < len(raw_headings)
            else len(paras)
        )
        body = paras[sec.body_start : end]
        entries = _split_entries(body)
        bullets = sum(1 for x in body if x.is_bullet)
        resolved = sec.model_copy(
            update={"body_end": end, "entry_count": len(entries), "bullet_count": bullets}
        )
        section_candidates.append(resolved)
        by_kind.setdefault(resolved.key, []).append(resolved)

    # One representative heading per kind — the first found, in document order — whose
    # `heading_paragraph_id`/`heading_text` become that kind's mapping fields (what
    # `template_build` anchors the kind's tagged prototype on). `combined_body` pools
    # every same-kind heading's body for prototype/bullet selection, so the best entry
    # can come from any of them, not only the first.
    section_by_key: dict[str, SectionCandidate] = {
        key: candidates[0] for key, candidates in by_kind.items()
    }
    combined_body: dict[str, list[_Para]] = {
        key: [p for sec in candidates for p in paras[sec.body_start : sec.body_end]]
        for key, candidates in by_kind.items()
    }

    if "experience" not in section_by_key:
        issues.append(
            Issue(
                code="missing_experience",
                message="Could not find an Experience / Work Experience section heading.",
                blocking=True,
            )
        )

    field_candidates: list[FieldCandidate] = []
    suggested: TemplateProfile | None = None

    # Name + contact: everything above the first detected heading, classified by regex
    # — see `_detect_name_and_contact` for why "first two non-heading paragraphs" isn't
    # enough once a table layout spreads the contact block across several paragraphs.
    first_heading_id = section_candidates[0].heading_paragraph_id if section_candidates else None
    name_id, contact_para, contact_slots, unmapped_contact_paras = _detect_name_and_contact(
        paras, first_heading_id
    )

    if contact_para is None and not contact_slots:
        issues.append(
            Issue(
                code="missing_contact",
                message="Could not find a contact line after the name.",
                blocking=True,
            )
        )
    for p in unmapped_contact_paras:
        issues.append(
            Issue(
                code="contact_unmapped_paragraph",
                message=(
                    f"{p.text.strip()!r} (paragraph {p.id}) is part of the contact "
                    "block but doesn't look like an email, phone, location, or "
                    "profile link. It will stay a literal in the template."
                ),
                blocking=False,
            )
        )

    # Manual bullet glyph warning (non-blocking unless no native bullets in experience).
    for p in paras:
        stripped = p.text.lstrip()
        if stripped.startswith(("•", "●", "○", "-", "–", "—")) and not p.is_bullet:
            issues.append(
                Issue(
                    code="manual_bullets",
                    message=(
                        f"Paragraph {p.id} looks like a bullet but is not a Word list item. "
                        "Convert lists to real bullets in Word/Google Docs before uploading."
                    ),
                    blocking=False,
                )
            )
            break

    enabled = EnabledSections(
        education="education" in section_by_key,
        experience="experience" in section_by_key,
        projects="projects" in section_by_key,
        skills="skills" in section_by_key,
        list_section="list" in section_by_key,
    )

    experience_mapping: ExperienceMapping | None = None
    education_mapping: EducationMapping | None = None
    projects_mapping: ProjectsMapping | None = None
    skills_mapping: SkillsMapping | None = None
    list_mapping: ListMapping | None = None

    if "experience" in section_by_key:
        sec = section_by_key["experience"]
        body = combined_body["experience"]
        entries = _split_entries(body)
        if not entries:
            issues.append(
                Issue(
                    code="empty_experience",
                    message="Experience section has no entries to use as a prototype.",
                    blocking=True,
                )
            )
        else:
            # 4d: reconcile field presence across every entry's own header before
            # picking a prototype, so an outlier entry — one that scores well on
            # `_exp_score` (more runs, a title line) but happens to lack a field most
            # other entries have — can never become the prototype and silently lose
            # that field for the whole section.
            candidate_entries, field_majority, field_confidence = _reconcile_header_fields(
                entries, primary="company", secondary="location", date_field="dates"
            )

            proto = max(candidate_entries, key=_exp_score)
            header_para = proto[0]
            header, _hcands = _entry_header_fields(
                proto, primary="company", secondary="location", date_field="dates"
            )
            field_candidates.extend(
                _section_field_candidates(
                    paras,
                    by_kind["experience"],
                    primary="company",
                    secondary="location",
                    date_field="dates",
                    pick=lambda entries: max(entries, key=_exp_score),
                    include_title=True,
                )
            )
            if not field_majority.get("dates", True):
                issues.append(
                    Issue(
                        code="experience_dates_not_detected",
                        message=(
                            "No experience entry's header has a detected date. Every "
                            "rendered job would lose its dates — map the date span "
                            "manually or confirm the header format."
                        ),
                        blocking=True,
                    )
                )
            elif field_confidence.get("dates", 1.0) < 1.0:
                issues.append(
                    Issue(
                        code="experience_dates_partial",
                        message=(
                            f"Only {field_confidence['dates']:.0%} of experience "
                            "entries have a detected date; the rest will render "
                            "without one."
                        ),
                        blocking=False,
                    )
                )
            proto_main = _entry_main_paragraphs(proto)
            titles = [x for x in proto_main[1:] if not x.is_bullet and x.text.strip()]
            bullets = [x for x in proto[1:] if x.is_bullet]
            if not bullets:
                # Fall back to any bullet in the section.
                bullets = [x for x in body if x.is_bullet]
            if not bullets:
                issues.append(
                    Issue(
                        code="no_experience_bullets",
                        message=(
                            "Experience section has no Word list bullets. "
                            "Entries need real list formatting."
                        ),
                        blocking=True,
                    )
                )
            else:
                title_span: OptionalSpan
                title_pid: int | None
                if titles:
                    t = titles[0]
                    title_span = OptionalSpan(
                        present=True, span=_span(t.id, 0, len(t.text.strip()))
                    )
                    title_pid = t.id
                else:
                    title_span = OptionalSpan(present=False)
                    title_pid = None
                    issues.append(
                        Issue(
                            code="inline_title",
                            message=(
                                "No separate job-title paragraph found under the experience "
                                "header. Map a title field or ensure each job has a title line."
                            ),
                            blocking=True,
                        )
                    )
                experience_mapping = ExperienceMapping(
                    heading_paragraph_id=sec.heading_paragraph_id,
                    heading_text=sec.heading_text,
                    prototype_entry_start=header_para.id,
                    header=header,
                    title=title_span,
                    title_paragraph_id=title_pid,
                    bullet_paragraph_id=bullets[0].id,
                )
                consistency_issue = _prototype_consistency_issue(
                    proto,
                    {"header": header_para.id, "title": title_pid, "bullet": bullets[0].id},
                    "Experience",
                )
                if consistency_issue is not None:
                    issues.append(consistency_issue)

    if "education" in section_by_key:
        sec = section_by_key["education"]
        body = combined_body["education"]
        entries = _split_entries(body)
        if entries:
            proto = max(entries, key=_edu_score)
            header, _hcands = _entry_header_fields(
                proto, primary="school", secondary="location", date_field="dates"
            )
            field_candidates.extend(
                _section_field_candidates(
                    paras,
                    by_kind["education"],
                    primary="school",
                    secondary="location",
                    date_field="dates",
                    pick=lambda entries: max(entries, key=_edu_score),
                )
            )
            proto_main = _entry_main_paragraphs(proto)
            plain_lines = [x for x in proto_main[1:] if not x.is_bullet and x.text.strip()]
            real_bullets = [x for x in proto[1:] if x.is_bullet] or [
                x for x in body if x.is_bullet
            ]
            plain_fallback = not real_bullets
            if plain_lines and real_bullets:
                # A prose degree line right under the header ("Bachelor of Arts in...",
                # not itself a Word bullet) followed by separately bulleted detail lines
                # (GPA, Dean's List, coursework) — distinct from the shape below, where
                # the degree line IS the first bullet. `edu.degree_line` is a single
                # field, so the prose line is the only sound choice for it; the real
                # bullets become the `edu.details` loop's prototype and, at render time,
                # its actual items.
                bullets = [plain_lines[0]] + real_bullets
            elif plain_fallback:
                # No real Word-list bullets under this entry. `retarget_bullet` (called
                # at build time) creates a paragraph's numbering properties rather than
                # requiring them to already exist, so a plain degree line still produces
                # a working template — it becomes a real bullet in the output. A warning,
                # not a blocker: the visual result is a reasonable, working outcome.
                bullets = [x for x in proto[1:] if x.text.strip()] or [
                    x for x in body if x.text.strip()
                ]
            else:
                bullets = real_bullets
            if not bullets:
                issues.append(
                    Issue(
                        code="no_education_bullets",
                        message="Education section has no Word list bullets for degree/details.",
                        blocking=True,
                    )
                )
            else:
                if plain_fallback:
                    issues.append(
                        Issue(
                            code="education_bullets_not_list",
                            message=(
                                "Education degree/detail line is not a Word list bullet; "
                                "it will be converted to one in the tagged template."
                            ),
                            blocking=False,
                        )
                    )
                degree = bullets[0]
                detail = bullets[1] if len(bullets) > 1 else bullets[0]
                education_mapping = EducationMapping(
                    heading_paragraph_id=sec.heading_paragraph_id,
                    heading_text=sec.heading_text,
                    prototype_entry_start=proto[0].id,
                    header=header,
                    degree_paragraph_id=degree.id,
                    detail_paragraph_id=detail.id,
                )
                consistency_issue = _prototype_consistency_issue(
                    proto,
                    {"header": proto[0].id, "degree": degree.id, "detail": detail.id},
                    "Education",
                )
                if consistency_issue is not None:
                    issues.append(consistency_issue)
        else:
            issues.append(
                Issue(
                    code="empty_education",
                    message="Education heading found but the section has no entries.",
                    blocking=False,
                )
            )
            enabled = enabled.model_copy(update={"education": False})

    if "projects" in section_by_key:
        sec = section_by_key["projects"]
        body = combined_body["projects"]
        entries = _split_entries(body)
        if entries:
            # 4d: same reconciliation as experience — narrow to the entries agreeing
            # with the section's own modal field-presence signature before applying
            # the existing "prefer fewer runs" tie-break.
            proj_candidate_entries, proj_field_majority, proj_field_confidence = (
                _reconcile_header_fields(
                    entries, primary="name", secondary="tech", date_field="date"
                )
            )
            proto = min(proj_candidate_entries, key=_proj_score)

            # Detect the link BEFORE splitting name/tech: its own span (not a fixed word
            # list like "Github"/"Demo"/"Live") comes from the hyperlink itself, so any
            # label works. `exclude_after` then keeps that label out of `tech` — without
            # it, a project with a link but no tech ("Name | Github\tdate") reads the
            # label as tech and the two fields end up with the same span, which the
            # builder rejects as an overlap.
            link, exclude_after, link_cand = _detect_project_link(proto)
            if link_cand is not None:
                field_candidates.append(link_cand)

            # Cross-cell aware (`_entry_header_fields`, not the plain-text-only
            # `_header_fields_from_text`) so a table-layout Projects section — name/tech
            # in one cell, date in the row's other cell — reconciles and installs
            # identically instead of reconciling fine and then losing the date on the
            # actual installed prototype.
            header, _hcands = _entry_header_fields(
                proto,
                primary="name",
                secondary="tech",
                date_field="date",
                exclude_after=exclude_after,
            )
            field_candidates.extend(
                _section_field_candidates(
                    paras,
                    by_kind["projects"],
                    primary="name",
                    secondary="tech",
                    date_field="date",
                    pick=lambda entries: min(entries, key=_proj_score),
                    include_link=True,
                )
            )
            if not proj_field_majority.get("date", True):
                issues.append(
                    Issue(
                        code="project_dates_not_detected",
                        message=(
                            "No project entry's header has a detected date. Every "
                            "rendered project would lose its date — map the date span "
                            "manually or confirm the header format."
                        ),
                        blocking=True,
                    )
                )
            elif proj_field_confidence.get("date", 1.0) < 1.0:
                issues.append(
                    Issue(
                        code="project_dates_partial",
                        message=(
                            f"Only {proj_field_confidence['date']:.0%} of project "
                            "entries have a detected date; the rest will render "
                            "without one."
                        ),
                        blocking=False,
                    )
                )
            bullets = [x for x in proto[1:] if x.is_bullet] or [
                x for x in body if x.is_bullet
            ]
            if not bullets:
                issues.append(
                    Issue(
                        code="no_project_bullets",
                        message="Projects section has no Word list bullets.",
                        blocking=True,
                    )
                )
            else:
                projects_mapping = ProjectsMapping(
                    heading_paragraph_id=sec.heading_paragraph_id,
                    heading_text=sec.heading_text,
                    prototype_entry_start=proto[0].id,
                    header=header,
                    link=link,
                    bullet_paragraph_id=bullets[0].id,
                )
                consistency_issue = _prototype_consistency_issue(
                    proto,
                    {"header": proto[0].id, "bullet": bullets[0].id},
                    "Project",
                )
                if consistency_issue is not None:
                    issues.append(consistency_issue)
        else:
            enabled = enabled.model_copy(update={"projects": False})

    if "skills" in section_by_key:
        sec = section_by_key["skills"]
        body = [p for p in combined_body["skills"] if p.text.strip()]
        if body:
            proto = next((p for p in body if ":" in p.text), body[0])
            spans = _skills_spans(proto)
            proto_id = proto.id
            if spans is None:
                # Not one paragraph split on a colon — try a table layout's label
                # cell/value cell pairing before giving up.
                cross = _skills_pair_across_cells(body)
                if cross is not None:
                    spans = cross
                    proto_id = cross[0].paragraph_id
            if spans is None:
                issues.append(
                    Issue(
                        code="skills_format",
                        message=(
                            "Skills lines should look like 'Label: item, item'. "
                            "Could not split the prototype line."
                        ),
                        blocking=True,
                    )
                )
            else:
                label_span, body_span, sep = spans
                skills_mapping = SkillsMapping(
                    heading_paragraph_id=sec.heading_paragraph_id,
                    heading_text=sec.heading_text,
                    prototype_paragraph_id=proto_id,
                    label_span=label_span,
                    body_span=body_span,
                    separator=sep,
                )
                field_candidates.append(
                    FieldCandidate(
                        field="skills_label",
                        span=label_span,
                        confidence=0.9,
                        preview=proto.text[label_span.start : label_span.end],
                        section_heading_paragraph_id=sec.heading_paragraph_id,
                    )
                )
        else:
            enabled = enabled.model_copy(update={"skills": False})

    if "list" in section_by_key:
        sec = section_by_key["list"]
        body = [p for p in combined_body["list"] if p.text.strip()]
        bullets = [p for p in body if p.is_bullet]
        if bullets:
            list_mapping = ListMapping(
                heading_paragraph_id=sec.heading_paragraph_id,
                heading_text=sec.heading_text,
                bullet_paragraph_id=bullets[0].id,
            )
        else:
            enabled = enabled.model_copy(update={"list_section": False})

    for key in ("education", "projects", "skills"):
        if key not in section_by_key:
            issues.append(
                Issue(
                    code=f"omit_{key}",
                    message=f"No {key.title()} section detected; it will be omitted from the template.",
                    blocking=False,
                )
            )

    blockers = [i for i in issues if i.blocking]
    if (
        not blockers
        and experience_mapping is not None
        and (contact_para is not None or contact_slots)
        and enabled.experience
    ):
        if contact_slots:
            contact = ContactMapping(
                paragraph_id=contact_slots[0].paragraph_id,
                slots=contact_slots,
            )
        else:
            assert contact_para is not None
            contact = ContactMapping(
                paragraph_id=contact_para.id,
                field_order=_contact_field_order(contact_para.text),
                separator=_contact_separator(contact_para.text),
            )

        # Generic mode is needed the moment fixed mode could not represent what was
        # found: more than one heading of some kind (two experience-shaped sections
        # cannot both keep their own title/position under one hard-coded heading), a
        # `list`-kind section (fixed mode has no such prototype at all), or a table
        # layout (always generic — see `TemplateProfile.layout`'s docstring). Otherwise
        # today's exact single-heading-per-kind case stays on fixed mode, byte-identical
        # to before this existed.
        is_table_layout = table_shape is not None
        needs_generic = (
            is_table_layout or "list" in by_kind or any(len(v) > 1 for v in by_kind.values())
        )
        detected_sections: list[DetectedSection] = []
        heading_prototype: HeadingPrototype | None = None
        spacing = SpacingProfile()
        if needs_generic:
            seen_ids: set[str] = set()
            for sec in section_candidates:
                base = config.slugify(sec.heading_text) or sec.key
                candidate_id = base
                suffix = 2
                while candidate_id in seen_ids:
                    candidate_id = f"{base}-{suffix}"
                    suffix += 1
                seen_ids.add(candidate_id)
                detected_sections.append(
                    DetectedSection(
                        id=candidate_id,
                        title=sec.heading_text,
                        kind=_TO_GENERIC_KIND.get(sec.key, sec.key),  # type: ignore[arg-type]
                        heading_paragraph_id=sec.heading_paragraph_id,
                    )
                )
            heading_prototype = HeadingPrototype(
                paragraph_id=section_candidates[0].heading_paragraph_id
            )
            # A table layout's inter-section gaps come from heading rows' own paragraph
            # spacing and dedicated spacer rows, not counted blank paragraphs —
            # `_detect_spacing`'s chrome-run model doesn't translate, and
            # `TemplateProfile` rejects a table-layout profile carrying spacing donors.
            spacing = SpacingProfile() if is_table_layout else _detect_spacing(paras, section_candidates)

        suggested = TemplateProfile(
            source_sha256=digest,
            name_paragraph_id=name_id,
            contact=contact,
            enabled=enabled,
            experience=experience_mapping,
            education=education_mapping if enabled.education else None,
            projects=projects_mapping if enabled.projects else None,
            skills=skills_mapping if enabled.skills else None,
            list_section=list_mapping if enabled.list_section else None,
            normalization=NormalizationFlags(),
            warnings=[i.message for i in issues if not i.blocking],
            section_mode="generic" if needs_generic else "fixed",
            sections=detected_sections,
            heading_prototype=heading_prototype,
            spacing=spacing,
            layout="table" if is_table_layout else "paragraph",
            paragraph_count=len(paras),
        )

    ready = suggested is not None and not blockers
    return AnalyzeResult(
        source_sha256=digest,
        paragraphs=paragraph_infos,
        sections=section_candidates,
        suggested_profile=suggested,
        field_candidates=field_candidates,
        issues=issues,
        ready=ready,
    )


def validate_profile_against_doc(
    profile: TemplateProfile,
    *,
    raw: bytes,
) -> list[Issue]:
    """Re-check a confirmed profile against the exact upload bytes before install."""
    issues: list[Issue] = []
    digest = sha256_bytes(raw)
    if profile.source_sha256 != digest:
        issues.append(
            Issue(
                code="hash_mismatch",
                message=(
                    "Mapping was confirmed against a different file than the one being "
                    "installed. Re-analyze the upload."
                ),
                blocking=True,
            )
        )
        return issues

    result = analyze_docx(raw=raw)
    # Structural re-validation: required paragraphs must still exist and spans fit.
    para_by_id = {p.id: p for p in result.paragraphs}
    if profile.name_paragraph_id not in para_by_id:
        issues.append(
            Issue(
                code="bad_name",
                message=f"Name paragraph {profile.name_paragraph_id} is out of range.",
                blocking=True,
            )
        )
    if profile.contact.paragraph_id not in para_by_id:
        issues.append(
            Issue(
                code="bad_contact",
                message=f"Contact paragraph {profile.contact.paragraph_id} is out of range.",
                blocking=True,
            )
        )

    def _check_span(span: CharSpan | None, label: str) -> None:
        """Blocking issue when `span` is out of range or straddles a tab.

        The tab check matters because `_tag_mapped_header` treats a tab inside a mapped
        span as a hard error at build time (the tab is what keeps a date right-aligned);
        catching it here turns that failure into a readable install-time rejection
        instead of a build crash — or, before that fix existed, a silently garbled
        template.
        """
        if span is None:
            return
        para = para_by_id.get(span.paragraph_id)
        if para is None:
            issues.append(
                Issue(
                    code="bad_span",
                    message=f"{label}: paragraph {span.paragraph_id} missing.",
                    blocking=True,
                )
            )
            return
        if span.end > len(para.text):
            issues.append(
                Issue(
                    code="bad_span",
                    message=(
                        f"{label}: span [{span.start}:{span.end}] exceeds paragraph "
                        f"length {len(para.text)}."
                    ),
                    blocking=True,
                )
            )
            return
        if "\t" in para.text[span.start : span.end]:
            issues.append(
                Issue(
                    code="span_has_tab",
                    message=(
                        f"{label}: span [{span.start}:{span.end}] contains a tab. "
                        "Map date fields after the tab, not through it."
                    ),
                    blocking=True,
                )
            )

    def _present_spans(fields: dict[str, OptionalSpan]) -> list[tuple[str, CharSpan]]:
        """(field name, span) pairs for mapped fields that are actually present."""
        return [
            (name, field.span)
            for name, field in fields.items()
            if field.present and field.span is not None
        ]

    def _check_no_overlap(spans: list[tuple[str, CharSpan]]) -> None:
        """Blocking issue when two mapped fields on the same paragraph overlap.

        Grouped by paragraph because header fields occasionally live off the header
        paragraph (a `date_paragraph_id` on its own line); those never collide.
        """
        by_paragraph: dict[int, list[tuple[str, CharSpan]]] = {}
        for label, span in spans:
            by_paragraph.setdefault(span.paragraph_id, []).append((label, span))
        for paragraph_id, entries in by_paragraph.items():
            ordered = sorted(entries, key=lambda e: e[1].start)
            for (label_a, span_a), (label_b, span_b) in zip(ordered, ordered[1:]):
                if span_b.start < span_a.end:
                    issues.append(
                        Issue(
                            code="overlapping_spans",
                            message=(
                                f"{label_a!r} and {label_b!r} both claim text in "
                                f"paragraph {paragraph_id}: {label_a} ends at "
                                f"{span_a.end}, {label_b} starts at {span_b.start}."
                            ),
                            blocking=True,
                        )
                    )

    def _check_bullet(paragraph_id: int | None, label: str, *, strict: bool = True) -> None:
        """Issue when a bullet-prototype paragraph is missing or not a list item.

        A missing paragraph always blocks. A non-list paragraph blocks only when
        `strict` — the experience/project bullet loop is the resume's main visual list
        content, where real Word numbering matters. Education's degree/detail role is not
        strict: `template_build.retarget_bullet` creates a paragraph's numbering
        properties rather than requiring them, so a plain degree line still builds fine
        and only needs a non-blocking heads-up (see `template_analyze`'s
        `education_bullets_not_list` issue, raised for the same reason at analyze time).
        """
        if paragraph_id is None:
            return
        para = para_by_id.get(paragraph_id)
        if para is None:
            issues.append(
                Issue(
                    code="bad_bullet",
                    message=f"{label}: paragraph {paragraph_id} is missing.",
                    blocking=True,
                )
            )
        elif not para.is_bullet:
            issues.append(
                Issue(
                    code="bullet_not_list",
                    message=f"{label}: paragraph {paragraph_id} is not a Word list item.",
                    blocking=strict,
                )
            )

    def _check_header_is_entry_start(paragraph_id: int, label: str) -> None:
        """Blocking issue when a header prototype paragraph is a bullet or blank — the
        mapping cannot possibly be pointing at a real entry header in that case."""
        para = para_by_id.get(paragraph_id)
        if para is None:
            return  # already reported by _check_span / the earlier existence checks
        if para.is_bullet or not para.text.strip():
            issues.append(
                Issue(
                    code="header_not_entry_start",
                    message=(
                        f"{label}: paragraph {paragraph_id} is "
                        f"{'a bullet' if para.is_bullet else 'blank'}, not an entry header."
                    ),
                    blocking=True,
                )
            )

    def _check_date_span(span: CharSpan | None, label: str) -> None:
        """Non-blocking warning when a span mapped to a date field does not itself
        look like a date — confirms the mapping actually landed on a date, not some
        other text that happened to survive detection."""
        if span is None:
            return
        para = para_by_id.get(span.paragraph_id)
        if para is None:
            return  # already reported by _check_span
        text = para.text[span.start : span.end]
        if not _DATE_RE.search(text):
            issues.append(
                Issue(
                    code="date_span_not_date_shaped",
                    message=(
                        f"{label}: mapped span {text!r} does not look like a date. "
                        "Confirm this field is mapped correctly."
                    ),
                    blocking=False,
                )
            )

    if not profile.enabled.experience:
        issues.append(
            Issue(
                code="experience_required",
                message="Experience section cannot be disabled.",
                blocking=True,
            )
        )
    else:
        exp = profile.experience
        exp_spans = _present_spans(exp.header.fields)
        if exp.title.present and exp.title.span is not None:
            exp_spans.append(("title", exp.title.span))
        for label, span in exp_spans:
            _check_span(span, label)
        _check_no_overlap(exp_spans)
        _check_bullet(exp.bullet_paragraph_id, "experience bullet prototype")
        _check_header_is_entry_start(
            exp.header.header_paragraph_id, "experience header prototype"
        )
        exp_dates = exp.header.fields.get("dates")
        if exp_dates is not None and exp_dates.present:
            _check_date_span(exp_dates.span, "experience dates")

    if profile.enabled.education and profile.education is not None:
        edu = profile.education
        edu_spans = _present_spans(edu.header.fields)
        for label, span in edu_spans:
            _check_span(span, label)
        _check_no_overlap(edu_spans)
        _check_bullet(edu.degree_paragraph_id, "education degree paragraph", strict=False)
        if edu.detail_paragraph_id is not None:
            _check_bullet(edu.detail_paragraph_id, "education detail paragraph", strict=False)
        _check_header_is_entry_start(
            edu.header.header_paragraph_id, "education header prototype"
        )
        edu_dates = edu.header.fields.get("dates")
        if edu_dates is not None and edu_dates.present:
            _check_date_span(edu_dates.span, "education dates")

    if profile.enabled.projects and profile.projects is not None:
        proj = profile.projects
        proj_spans = _present_spans(proj.header.fields)
        if proj.link.present and proj.link.span is not None:
            proj_spans.append(("link", proj.link.span))
        for label, span in proj_spans:
            _check_span(span, label)
        _check_no_overlap(proj_spans)
        _check_bullet(proj.bullet_paragraph_id, "project bullet prototype")
        _check_header_is_entry_start(
            proj.header.header_paragraph_id, "project header prototype"
        )
        proj_date = proj.header.fields.get("date")
        if proj_date is not None and proj_date.present:
            _check_date_span(proj_date.span, "project date")

    if profile.enabled.skills and profile.skills is not None:
        skl = profile.skills
        skl_spans = [("skills_label", skl.label_span), ("skills_body", skl.body_span)]
        for label, span in skl_spans:
            _check_span(span, label)
        _check_no_overlap(skl_spans)
        if skl.prototype_paragraph_id not in para_by_id:
            issues.append(
                Issue(
                    code="bad_span",
                    message=(
                        f"skills prototype: paragraph {skl.prototype_paragraph_id} "
                        "is missing."
                    ),
                    blocking=True,
                )
            )

    if profile.section_mode == "generic":
        if (
            profile.heading_prototype is not None
            and profile.heading_prototype.paragraph_id not in para_by_id
        ):
            issues.append(
                Issue(
                    code="bad_heading_prototype",
                    message=(
                        f"heading_prototype: paragraph "
                        f"{profile.heading_prototype.paragraph_id} is out of range."
                    ),
                    blocking=True,
                )
            )
        for s in profile.sections:
            if s.heading_paragraph_id not in para_by_id:
                issues.append(
                    Issue(
                        code="bad_detected_section",
                        message=(
                            f"Detected section {s.id!r} ({s.title!r}): heading "
                            f"paragraph {s.heading_paragraph_id} is out of range."
                        ),
                        blocking=True,
                    )
                )
        for label, donor_ids in (
            ("spacing.before_heading", profile.spacing.before_heading),
            ("spacing.after_heading", profile.spacing.after_heading),
            ("spacing.between_entries", profile.spacing.between_entries),
        ):
            for donor_id in donor_ids:
                donor = para_by_id.get(donor_id)
                # Non-blocking: a stale or no-longer-chrome donor should degrade to a
                # narrower gap at build time, not block the whole install — a spacer is a
                # cosmetic fidelity nicety, not load-bearing content.
                if donor is None or not docx_text.is_chrome_text(donor.text):
                    issues.append(
                        Issue(
                            code="spacer_donor_missing",
                            message=(
                                f"{label}: donor paragraph {donor_id} is missing or no "
                                "longer a blank/rule line; it will be omitted."
                            ),
                            blocking=False,
                        )
                    )

    return issues
