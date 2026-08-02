"""Generate `templates/main_template.docx` from `templates/original_export.docx`.

This is the ONLY thing permitted to produce the tagged template (see CLAUDE.md). Run it
again whenever the resume is re-exported from Google Docs:

    python scripts/build_template.py

Why a script rather than hand-tagging in Word: the header lines carry up to 15
separately-formatted runs plus per-project hyperlinks, so tags must land inside specific
runs to preserve bold/tab/link formatting. Doing that by hand is error-prone and has to be
redone on every resume update — and the resume already changed once during development.

What gets tagged, and what deliberately does not:

- Tagged: the name line (plain substitution from `contact.name`), the contact line (as a
  RichText so LinkedIn/GitHub stay hyperlinks), EDUCATION (looped so coursework/GPA come
  from the master resume), WORK EXPERIENCES, PROJECTS, and SKILLS.

The transformation works by treating one entry in each section as a prototype, cloning its
XML (which carries all formatting), and deleting the rest. Formatting is therefore
inherited rather than reconstructed.
"""

from __future__ import annotations

import copy
import sys
from pathlib import Path

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.text.run import Run

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

from resume_tailor import config  # noqa: E402

W = qn("w:p")
NOTO_MARKER_FONT = "Noto Sans Symbols"


# --------------------------------------------------------------------------------------
# Low-level XML helpers
# --------------------------------------------------------------------------------------


def make_para(text: str) -> OxmlElement:
    """Build a bare paragraph carrying a `{%p %}` control tag.

    The `p` suffix is load-bearing, not cosmetic. A plain `{% for %}` is substituted
    inline, so the loop body starts mid-paragraph and each iteration re-emits the
    surrounding `<w:p>` fragments — nesting a paragraph inside a paragraph. The result
    is well-formed XML that violates the OOXML schema, and Word refuses to open it with
    a generic "Word experienced an error trying to open the file". `{%p %}` makes
    docxtpl drop the whole enclosing paragraph instead, so loops span entire paragraphs.

    Deliberately style-less: these paragraphs are deleted at render time, so formatting
    on them would be wasted, and inheriting a list style would leave a stray bullet
    behind if a tag were ever mistyped.
    """
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    p.append(r)
    return p


def delete(paragraph: Paragraph) -> None:
    paragraph._p.getparent().remove(paragraph._p)


def strip_hyperlinks(paragraph: Paragraph) -> list[str]:
    """Remove `w:hyperlink` wrappers, returning the link texts that were dropped.

    Each project carries its own URL, so a hyperlink baked into the template would point
    every rendered project at the prototype's target. The link is reinstated at render
    time as a docxtpl `RichText`, which supplies the correct per-project URL.
    """
    dropped = []
    for link in paragraph._p.findall(qn("w:hyperlink")):
        dropped.append("".join(t.text or "" for t in link.iter(qn("w:t"))))
        paragraph._p.remove(link)
    return dropped


def has_tab_element(run) -> bool:
    """Whether the run's tab is a `<w:tab/>` element rather than a literal character.

    Word writes tab stops as elements, which `run.text` renders as "\\t". Treating that
    rendered form as if it were real text and rewriting it produced a duplicated tab —
    the element survived the rewrite and a literal tab was added alongside it.
    """
    return run._r.find(qn("w:tab")) is not None


def set_run_text(run, text: str, *, keep_tabs: bool = False) -> None:
    """Replace a run's text, preserving significant whitespace.

    With `keep_tabs`, any `<w:tab/>` elements are left in place and `text` is appended
    after them — use it for the run that holds a header's alignment tab.
    """
    for t in run._r.findall(qn("w:t")):
        run._r.remove(t)
    if not keep_tabs:
        for tab in run._r.findall(qn("w:tab")):
            run._r.remove(tab)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run._r.append(t)


def clone_run_after(run, text: str) -> Run:
    """Duplicate a run (inheriting its formatting) with new text, placed just after it."""
    new_r = copy.deepcopy(run._r)
    run._r.addnext(new_r)
    cloned = Run(new_r, run._parent)
    set_run_text(cloned, text)
    return cloned


def strip_bold(run: Run) -> None:
    """Remove bold from a run's character properties."""
    rPr = run._r.find(qn("w:rPr"))
    if rPr is None:
        return
    for tag in ("w:b", "w:bCs"):
        el = rPr.find(qn(tag))
        if el is not None:
            rPr.remove(el)


def tab_run_index(paragraph: Paragraph) -> int:
    """Return the index of the run holding the header's date-alignment tab."""
    for i, run in enumerate(paragraph.runs):
        if "\t" in run.text or has_tab_element(run):
            return i
    return len(paragraph.runs) - 1


def leading_runs_before_tab(paragraph: Paragraph) -> int:
    """Count runs before the header tab — more runs usually means company/location split."""
    return tab_run_index(paragraph)


def set_num_id(paragraph: Paragraph, num_id: str) -> None:
    """Point a list paragraph at a different numbering definition instance."""
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        paragraph._p.insert(0, pPr)
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        numPr = OxmlElement("w:numPr")
        pPr.insert(0, numPr)
    num_id_el = numPr.find(qn("w:numId"))
    if num_id_el is None:
        num_id_el = OxmlElement("w:numId")
        numPr.append(num_id_el)
    num_id_el.set(qn("w:val"), num_id)


def strip_right_indent(paragraph: Paragraph) -> None:
    """Drop a paragraph's right indent — it narrows the column and inflates line count."""
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        return
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        return
    if ind.get(qn("w:right")) is not None:
        del ind.attrib[qn("w:right")]


# Word single spacing under lineRule=auto is 240 (240ths of a line).
SINGLE_LINE = "240"


def set_single_spacing(paragraph: Paragraph, *, exact: bool = False) -> None:
    """Force Word single line spacing on a paragraph, preserving before/after gaps.

    `lineRule=auto` is Word's "Single". Bullets also get `exact` so a taller substitute
    bullet font (common under LibreOffice when Noto is missing) cannot inflate the
    auto line box past single — measured wraps were ~15.7pt for 10pt body text.
    """
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        paragraph._p.insert(0, pPr)
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:line"), SINGLE_LINE)
    # exact locks height in twips (240 = 12pt); auto is Word UI "Single".
    spacing.set(qn("w:lineRule"), "exact" if exact else "auto")


def normalize_single_spacing(doc) -> None:
    """Write single line spacing onto every body paragraph in the document.

    Google Docs exports can leave spacing unset (inherits a looser default) or carry
    Multiple > 1 on some paragraphs. Prototypes alone cannot fix education bullets or
    the name/contact lines, so this runs over the whole document after tagging.
    """
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        # Control-tag paragraphs are dropped at render time; leave them bare.
        if text.startswith("{%p "):
            continue
        set_single_spacing(paragraph, exact=is_bullet(paragraph))


# --------------------------------------------------------------------------------------
# Document structure
# --------------------------------------------------------------------------------------

SECTIONS = ("EDUCATION", "WORK EXPERIENCES", "PROJECTS", "SKILLS")


def is_bullet(p: Paragraph) -> bool:
    """A real list paragraph, i.e. one carrying numbering properties."""
    return p._p.find(f".//{qn('w:numPr')}") is not None


def find_sections(doc) -> dict[str, tuple[int, int]]:
    """Map each section heading to the (start, end) paragraph range of its body."""
    heads: dict[str, int] = {}
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text in SECTIONS:
            heads[text] = i

    bounds: dict[str, tuple[int, int]] = {}
    ordered = sorted(heads.items(), key=lambda kv: kv[1])
    for n, (name, idx) in enumerate(ordered):
        end = ordered[n + 1][1] if n + 1 < len(ordered) else len(doc.paragraphs)
        bounds[name] = (idx + 1, end)
    return bounds


def split_entries(paragraphs: list[Paragraph]) -> list[list[Paragraph]]:
    """Group a section's paragraphs into entries.

    An entry starts at a non-bullet paragraph that has actual text; its bullets and any
    sub-heading (a job title) belong to it until the next such paragraph.
    """
    entries: list[list[Paragraph]] = []
    for p in paragraphs:
        starts_entry = not is_bullet(p) and p.text.strip()
        if starts_entry and (not entries or any(is_bullet(x) for x in entries[-1])):
            entries.append([p])
        elif entries:
            entries[-1].append(p)
    return entries


def header_run_count(entry: list[Paragraph]) -> int:
    return len(entry[0].runs)


def vertical_cost(p: Paragraph) -> tuple[int, int]:
    """Approximate how much vertical space a paragraph's formatting costs.

    Returns (line spacing, right indent); lower is more compact. Word's default line
    value is 240 (single), and a right indent narrows the text column, causing earlier
    wrapping and therefore taller blocks.
    """
    line, right = 240, 0
    pPr = p._p.find(qn("w:pPr"))
    if pPr is not None:
        spacing = pPr.find(qn("w:spacing"))
        if spacing is not None and spacing.get(qn("w:line")):
            line = int(float(spacing.get(qn("w:line"))))
        ind = pPr.find(qn("w:ind"))
        if ind is not None and ind.get(qn("w:right")):
            right = int(float(ind.get(qn("w:right"))))
    return line, right


def _num_to_abstract(numbering_root) -> dict[str, str]:
    """Map numbering instance ids to their abstract definition ids."""
    mapping: dict[str, str] = {}
    for num in numbering_root.findall(qn("w:num")):
        nid = num.get(qn("w:numId"))
        abs_el = num.find(qn("w:abstractNumId"))
        if nid is not None and abs_el is not None:
            mapping[nid] = abs_el.get(qn("w:val"), "")
    return mapping


def _abstract_has_noto_marker(numbering_root, abstract_id: str) -> bool:
    """Return whether an abstract numbering definition pins Noto for its lvl0 bullet."""
    for anum in numbering_root.findall(qn("w:abstractNum")):
        if anum.get(qn("w:abstractNumId")) != abstract_id:
            continue
        for lvl in anum.findall(qn("w:lvl")):
            if lvl.get(qn("w:ilvl")) != "0":
                continue
            num_fmt = lvl.find(qn("w:numFmt"))
            if num_fmt is None or num_fmt.get(qn("w:val")) != "bullet":
                continue
            rPr = lvl.find(qn("w:rPr"))
            rf = rPr.find(qn("w:rFonts")) if rPr is not None else None
            return rf is not None and rf.get(qn("w:ascii")) == NOTO_MARKER_FONT
    return False


def discover_noto_num_id(doc) -> str | None:
    """Find a list instance id whose lvl0 bullet marker uses Noto Sans Symbols.

    Prefer an id used by education bullets so experience/projects match that section.
    """
    numbering = doc.part.numbering_part
    root = numbering.element
    num_to_abs = _num_to_abstract(root)

    bounds = find_sections(doc)
    candidates: list[str] = []
    if "EDUCATION" in bounds:
        start, end = bounds["EDUCATION"]
        for paragraph in doc.paragraphs[start:end]:
            if not is_bullet(paragraph):
                continue
            pPr = paragraph._p.find(qn("w:pPr"))
            numPr = pPr.find(qn("w:numPr")) if pPr is not None else None
            if numPr is None:
                continue
            num_id_el = numPr.find(qn("w:numId"))
            if num_id_el is None:
                continue
            num_id = num_id_el.get(qn("w:val"))
            if num_id and _abstract_has_noto_marker(root, num_to_abs.get(num_id, "")):
                return num_id
            if num_id:
                candidates.append(num_id)

    for num_id, abstract_id in num_to_abs.items():
        if _abstract_has_noto_marker(root, abstract_id):
            return num_id
    return candidates[0] if candidates else None


def normalize_bullet_numbering(doc) -> None:
    """Pin Noto Sans Symbols on every lvl0 bullet marker in numbering.xml."""
    numbering = doc.part.numbering_part
    root = numbering.element

    ref_rfonts = None
    for anum in root.findall(qn("w:abstractNum")):
        for lvl in anum.findall(qn("w:lvl")):
            if lvl.get(qn("w:ilvl")) != "0":
                continue
            num_fmt = lvl.find(qn("w:numFmt"))
            if num_fmt is None or num_fmt.get(qn("w:val")) != "bullet":
                continue
            rPr = lvl.find(qn("w:rPr"))
            rf = rPr.find(qn("w:rFonts")) if rPr is not None else None
            if rf is not None and rf.get(qn("w:ascii")) == NOTO_MARKER_FONT:
                ref_rfonts = copy.deepcopy(rf)
                break
        if ref_rfonts is not None:
            break

    if ref_rfonts is None:
        ref_rfonts = OxmlElement("w:rFonts")
        for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
            ref_rfonts.set(qn(f"w:{attr}"), NOTO_MARKER_FONT)

    for anum in root.findall(qn("w:abstractNum")):
        for lvl in anum.findall(qn("w:lvl")):
            if lvl.get(qn("w:ilvl")) != "0":
                continue
            num_fmt = lvl.find(qn("w:numFmt"))
            if num_fmt is None or num_fmt.get(qn("w:val")) != "bullet":
                continue
            rPr = lvl.find(qn("w:rPr"))
            if rPr is None:
                rPr = OxmlElement("w:rPr")
                lvl.append(rPr)
            old_rf = rPr.find(qn("w:rFonts"))
            if old_rf is not None:
                rPr.remove(old_rf)
            rPr.insert(0, copy.deepcopy(ref_rfonts))


def retarget_bullet(paragraph: Paragraph, num_id: str | None) -> None:
    """Use the canonical small-marker list id and drop any right-indent inflation."""
    if num_id is not None:
        set_num_id(paragraph, num_id)
    strip_right_indent(paragraph)


def pick_bullet_prototype(entries: list[list[Paragraph]]) -> Paragraph:
    """Choose the most compact bullet in a section to serve as the loop body.

    The source resume is internally inconsistent: some entries' bullets carry a 1.15
    line-spacing override and a right indent, others carry neither. Since one prototype
    drives every rendered bullet, that formatting is necessarily normalised — a bullet
    from any entry may appear where another entry's used to, so per-entry spacing cannot
    survive the loop regardless.

    Given that, normalise *downward*. Picking the loosest prototype silently inflated the
    resume by roughly 15% of a line per bullet and pushed a one-page resume onto two.
    """
    bullets = [p for entry in entries for p in entry if is_bullet(p)]
    if not bullets:
        raise RuntimeError("Section contains no bullet paragraphs to use as a prototype.")
    return min(bullets, key=vertical_cost)


# --------------------------------------------------------------------------------------
# Tagging
# --------------------------------------------------------------------------------------


def tag_header(paragraph: Paragraph, fields: list[str], *, tail_field: str) -> None:
    """Tag a header line of the form `<bold> | <plain>\\t<date>`.

    Runs before the tab receive `fields` in order; the run holding the tab keeps it and
    the text after the tab becomes `tail_field`. Working run-by-run rather than rewriting
    the paragraph is what preserves the bold prefix and the right-aligned tab stop.
    """
    tab_idx = tab_run_index(paragraph)
    runs = paragraph.runs

    # Everything before the tab carries the leading fields.
    lead = runs[:tab_idx]
    for i, run in enumerate(lead):
        set_run_text(run, fields[i] if i < len(fields) else "")

    # More fields than runs: insert plain runs instead of merging into the bold run.
    if len(fields) > len(lead) and lead:
        anchor = lead[-1]
        for field in fields[len(lead) :]:
            plain = clone_run_after(anchor, field)
            strip_bold(plain)
            anchor = plain

    runs = paragraph.runs
    tab_idx = tab_run_index(paragraph)

    # The tab run keeps its tab; the date follows it.
    tab_run = runs[tab_idx]
    element_tab = has_tab_element(tab_run)
    after_tab = tab_run.text.split("\t", 1)[1] if "\t" in tab_run.text else ""

    if after_tab or tab_idx == len(runs) - 1:
        # Tab and date share one run.
        if element_tab:
            set_run_text(tab_run, tail_field, keep_tabs=True)
        else:
            set_run_text(tab_run, "\t" + tail_field)
        for extra in runs[tab_idx + 1 :]:
            paragraph._p.remove(extra._r)
    else:
        # Tab and date are in separate runs; keep that split so the date keeps its own
        # formatting (the tab run is sometimes bold, the date never is).
        if element_tab:
            set_run_text(tab_run, "", keep_tabs=True)
        else:
            set_run_text(tab_run, "\t")
        set_run_text(runs[tab_idx + 1], tail_field)
        for extra in runs[tab_idx + 2 :]:
            paragraph._p.remove(extra._r)


def tag_bullet(paragraph: Paragraph, expr: str) -> None:
    """Collapse a bullet to a single tagged run, keeping its list formatting."""
    runs = paragraph.runs
    set_run_text(runs[0], expr)
    for extra in runs[1:]:
        paragraph._p.remove(extra._r)


def build_loop(
    entry: list[Paragraph],
    anchor: Paragraph,
    *,
    outer: str,
    inner: str,
) -> None:
    """Insert a cloned, tagged entry wrapped in for-loops, before `anchor`."""
    anchor._p.addprevious(make_para(outer))
    for para in entry:
        anchor._p.addprevious(copy.deepcopy(para._p))
    anchor._p.addprevious(make_para("{%p endfor %}"))
    _ = inner


# --------------------------------------------------------------------------------------
# Section builders
# --------------------------------------------------------------------------------------


def build_name(doc) -> None:
    """Replace the name line with a plain `{{ }}` tag so it comes from the master resume.

    Plain substitution is fine here (unlike the contact line): the name carries no
    hyperlink, so there is no risk of a run nesting inside `<w:t>`.
    """
    if not doc.paragraphs:
        raise RuntimeError("Document is missing the name line (expected paragraph 0).")
    paragraph = doc.paragraphs[0]
    runs = paragraph.runs
    if not runs:
        raise RuntimeError("Name line has no runs to tag.")
    # Keep the first run's display-name formatting; collapse everything else into it.
    # Deliberately a distinct top-level key, not `contact.name` — `contact` in the render
    # context is the RichText contact *line* (`{{r contact }}`), which has no attributes.
    set_run_text(runs[0], "{{ name }}")
    for extra in runs[1:]:
        paragraph._p.remove(extra._r)


def build_contact(doc) -> None:
    """Replace the contact line with a single RichText tag.

    The line carries LinkedIn (and optionally GitHub) as real hyperlinks, which differ
    per person and must not be baked into the template. `{{r contact }}` is required —
    a plain `{{ }}` nests the hyperlink's `<w:r>` inside a `<w:t>` and the link vanishes
    on read-back. The baked-in hyperlink is stripped first for the same reason project
    URLs are.
    """
    # Contact is the second body paragraph in the export: name, then location/email/…
    if len(doc.paragraphs) < 2:
        raise RuntimeError("Document is missing the contact line (expected paragraph 1).")
    paragraph = doc.paragraphs[1]
    strip_hyperlinks(paragraph)
    runs = paragraph.runs
    if not runs:
        raise RuntimeError("Contact line has no runs to tag.")
    # Keep the first run's Spectral formatting; collapse everything else into it.
    set_run_text(runs[0], "{{r contact }}")
    for extra in runs[1:]:
        paragraph._p.remove(extra._r)


def build_education(doc, entries: list[list[Paragraph]], *, noto_num_id: str | None) -> None:
    """Replace all education entries with one tagged, looped prototype.

    Degree text and detail bullets (coursework joined upstream, plus free-text details)
    come from the master resume so the editor can toggle GPA and edit coursework without
    re-exporting the Google Doc.
    """
    if not entries:
        raise RuntimeError("EDUCATION section has no entries to prototype.")
    # Prefer a header whose school and location already live in separate runs so
    # school stays bold and location stays plain after tagging.
    prototype = max(entries, key=lambda entry: leading_runs_before_tab(entry[0]))
    header, *rest = prototype
    bullets = [p for p in rest if is_bullet(p)]
    if not bullets:
        raise RuntimeError("Education entry is missing a degree/detail bullet.")
    degree = bullets[0]
    # Prefer a different bullet for the detail loop body so tagging one does not
    # overwrite the other when they happen to be the same paragraph object.
    section_bullets = [p for entry in entries for p in entry if is_bullet(p)]
    others = [p for p in section_bullets if p._p is not degree._p]
    detail = min(others, key=vertical_cost) if others else degree
    # When the section has only a degree bullet, clone it so the two tags stay independent.
    if detail._p is degree._p:
        detail_p = copy.deepcopy(degree._p)
        degree._p.addnext(detail_p)
        detail = Paragraph(detail_p, degree._parent)
    retarget_bullet(degree, noto_num_id)
    retarget_bullet(detail, noto_num_id)

    tag_header(
        header,
        ["{{ edu.school }} | ", "{{ edu.location }}"],
        tail_field="{{ edu.dates }}",
    )
    tag_bullet(degree, "{{ edu.degree_line }}")
    tag_bullet(detail, "{{ detail }}")

    anchor = entries[0][0]
    anchor._p.addprevious(make_para("{%p for edu in education %}"))
    anchor._p.addprevious(copy.deepcopy(header._p))
    anchor._p.addprevious(copy.deepcopy(degree._p))
    anchor._p.addprevious(make_para("{%p for detail in edu.details %}"))
    anchor._p.addprevious(copy.deepcopy(detail._p))
    anchor._p.addprevious(make_para("{%p endfor %}"))
    anchor._p.addprevious(make_para("{%p endfor %}"))

    # Drop the temporary clone if we inserted one next to the degree bullet.
    if detail._p.getparent() is not None and detail not in [
        p for entry in entries for p in entry
    ]:
        delete(detail)

    for entry in entries:
        for para in entry:
            delete(para)


def build_experience(doc, entries: list[list[Paragraph]], *, noto_num_id: str | None) -> None:
    """Replace all experience entries with one tagged, looped prototype."""
    # Prefer a header whose company and location already live in separate runs so
    # company stays bold and location stays plain after tagging.
    prototype = max(entries, key=lambda entry: leading_runs_before_tab(entry[0]))

    header, *rest = prototype
    title = next((p for p in rest if not is_bullet(p) and p.text.strip()), None)
    if title is None:
        raise RuntimeError("Experience entry is missing a title paragraph.")
    # The bullet prototype is chosen across the whole section, independently of the
    # header, so spacing is normalised to the tightest variant present.
    bullet = pick_bullet_prototype(entries)
    retarget_bullet(bullet, noto_num_id)

    tag_header(header, ["{{ job.company }} | ", "{{ job.location }}"], tail_field="{{ job.dates }}")
    set_run_text(title.runs[0], "{{ job.title }}")
    for extra in title.runs[1:]:
        title._p.remove(extra._r)
    tag_bullet(bullet, "{{ bullet }}")

    anchor = entries[0][0]
    anchor._p.addprevious(make_para("{%p for job in experience %}"))
    anchor._p.addprevious(copy.deepcopy(header._p))
    anchor._p.addprevious(copy.deepcopy(title._p))
    anchor._p.addprevious(make_para("{%p for bullet in job.bullets %}"))
    anchor._p.addprevious(copy.deepcopy(bullet._p))
    anchor._p.addprevious(make_para("{%p endfor %}"))
    anchor._p.addprevious(make_para("{%p endfor %}"))

    for entry in entries:
        for para in entry:
            delete(para)


def build_projects(doc, entries: list[list[Paragraph]], *, noto_num_id: str | None) -> None:
    """Replace all project entries with one tagged, looped prototype."""
    prototype = min(entries, key=header_run_count)
    header, *rest = prototype
    bullet = pick_bullet_prototype(entries)
    retarget_bullet(bullet, noto_num_id)

    # The per-project URL differs, so the baked-in hyperlink must go; render.py
    # reinstates it as a RichText carrying the right target.
    strip_hyperlinks(header)
    tag_header(
        header,
        # `{{r ... }}` is docxtpl's RichText tag — the inline counterpart of `{%p %}`.
        # A plain `{{ }}` substitutes the value as text, which nests the hyperlink's
        # `<w:r>` inside a `<w:t>`; `w:t` may only contain characters, so the link is
        # silently dropped when the document is read back.
        # Separator before the link lives in the RichText (render.py), not here —
        # otherwise suppressing the link leaves a dangling " | ".
        ["{{ proj.name }} | ", "{{ proj.tech }}", "{{r proj.link }}"],
        tail_field="{{ proj.date }}",
    )
    tag_bullet(bullet, "{{ bullet }}")

    anchor = entries[0][0]
    anchor._p.addprevious(make_para("{%p for proj in projects %}"))
    anchor._p.addprevious(copy.deepcopy(header._p))
    anchor._p.addprevious(make_para("{%p for bullet in proj.bullets %}"))
    anchor._p.addprevious(copy.deepcopy(bullet._p))
    anchor._p.addprevious(make_para("{%p endfor %}"))
    anchor._p.addprevious(make_para("{%p endfor %}"))

    for entry in entries:
        for para in entry:
            delete(para)


def build_skills(doc, paragraphs: list[Paragraph]) -> None:
    """Replace the skills lines with a single looped, tagged line."""
    lines = [p for p in paragraphs if p.text.strip()]
    if not lines:
        return
    # The prototype must keep the export's two-run split: a bold label run and a plain
    # body run. Writing both tags into the label run and dropping the rest discards the
    # plain run's formatting, so the entire rendered line comes out bold.
    prototype = min(
        (p for p in lines if len(p.runs) >= 2), key=lambda p: len(p.runs), default=lines[0]
    )

    # `entries`, not `items`: Jinja resolves `x.items` to the dict's built-in method
    # before the key, and the method's repr gets mangled into bogus XML that Word
    # refuses to open. Avoid any name that collides with a dict attribute.
    #
    # The colon lives in the label run because that is how the export bolds it; the
    # separating space moves to the body run, where `set_run_text` preserves it.
    set_run_text(prototype.runs[0], "{{ group.label }}:")
    if len(prototype.runs) >= 2:
        body = prototype.runs[1]
    else:
        body = clone_run_after(prototype.runs[0], "")
        body.bold = False  # the clone inherits the label's bold
    set_run_text(body, " {{ group.entries }}")
    for extra in list(prototype.runs[2:]):
        prototype._p.remove(extra._r)

    anchor = lines[0]
    anchor._p.addprevious(make_para("{%p for group in skills %}"))
    anchor._p.addprevious(copy.deepcopy(prototype._p))
    anchor._p.addprevious(make_para("{%p endfor %}"))

    for p in lines:
        delete(p)


# --------------------------------------------------------------------------------------
# Entry point
# --------------------------------------------------------------------------------------


def main() -> int:
    src = config.BASELINE_TEMPLATE_PATH
    dst = config.DEFAULT_TEMPLATE_PATH

    if not src.exists():
        print(
            f"ERROR: {src} not found.\n"
            "Export your resume from Google Docs (File > Download > Microsoft Word) "
            "and save it there.",
            file=sys.stderr,
        )
        return 1

    doc = docx.Document(src)
    bounds = find_sections(doc)

    missing = [s for s in ("WORK EXPERIENCES", "PROJECTS", "SKILLS") if s not in bounds]
    if missing:
        print(
            f"ERROR: could not find section heading(s): {', '.join(missing)}.\n"
            "The build script locates sections by their exact all-caps heading text. "
            "If you renamed a heading in the Google Doc, update SECTIONS in this script.",
            file=sys.stderr,
        )
        return 1

    paras = doc.paragraphs

    if "EDUCATION" not in bounds:
        print(
            "ERROR: could not find section heading: EDUCATION.\n"
            "The build script locates sections by their exact all-caps heading text.",
            file=sys.stderr,
        )
        return 1

    edu_start, edu_end = bounds["EDUCATION"]
    exp_start, exp_end = bounds["WORK EXPERIENCES"]
    proj_start, proj_end = bounds["PROJECTS"]
    skill_start, skill_end = bounds["SKILLS"]

    education_entries = split_entries(paras[edu_start:edu_end])
    experience_entries = split_entries(paras[exp_start:exp_end])
    project_entries = split_entries(paras[proj_start:proj_end])
    skill_paras = paras[skill_start:skill_end]

    print(
        f"found {len(education_entries)} education, "
        f"{len(experience_entries)} experience entries, "
        f"{len(project_entries)} projects"
    )

    # Discover Noto list id from the *source* education bullets before they are replaced.
    noto_num_id = discover_noto_num_id(doc)

    build_name(doc)
    build_contact(doc)
    build_education(doc, education_entries, noto_num_id=noto_num_id)
    build_experience(doc, experience_entries, noto_num_id=noto_num_id)
    build_projects(doc, project_entries, noto_num_id=noto_num_id)
    build_skills(doc, skill_paras)
    normalize_bullet_numbering(doc)
    # After tagging: prototypes may still carry Multiple > 1 from the export, and
    # header lines are never re-prototyped — force single everywhere.
    normalize_single_spacing(doc)

    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(dst)
    print(f"wrote {dst}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
